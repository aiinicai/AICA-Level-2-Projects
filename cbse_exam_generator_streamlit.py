
import streamlit as st
import random
from io import BytesIO
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.enums import TA_CENTER
from docx import Document

st.set_page_config(page_title="CBSE Exam Generator", page_icon="📚", layout="wide")

DATA = {
    "Mathematics": {
        "Chapter 1 - Real Numbers": [
            ("Fill in the Blank", "The decimal expansion of a rational number is either terminating or ________.", "non-terminating recurring", "Easy", 1),
            ("MCQ", "Which of the following is irrational?", "A. 2/3   B. √2   C. 0.25   D. 5", "Medium", 1),
            ("Short Answer", "Use Euclid's division algorithm to find the HCF of 135 and 225.", "45", "Medium", 3),
            ("Long Answer", "Prove that √5 is irrational.", "Assume √5 = p/q in lowest terms and derive a contradiction using divisibility by 5.", "Hard", 5),
        ],
        "Chapter 2 - Polynomials": [
            ("Fill in the Blank", "The zeroes of a polynomial are the values of x for which the polynomial becomes ________.", "zero", "Easy", 1),
            ("MCQ", "If α and β are zeroes of x² - 5x + 6, then α+β is:", "A. 6   B. 5   C. -5   D. -6", "Easy", 1),
            ("Short Answer", "Find the zeroes of x² - 7x + 12.", "3 and 4", "Medium", 3),
            ("Long Answer", "Explain the relationship between zeroes and coefficients of a quadratic polynomial.", "For ax²+bx+c, α+β=-b/a and αβ=c/a.", "Hard", 5),
        ],
        "Chapter 3 - Pair of Linear Equations": [
            ("Fill in the Blank", "Two linear equations in two variables have a unique solution when their graphs are ________.", "intersecting", "Easy", 1),
            ("MCQ", "The pair x+y=5 and x-y=1 has solution:", "A. (2,3)   B. (3,2)   C. (4,1)   D. (1,4)", "Easy", 1),
            ("Short Answer", "Solve 2x+y=7 and x-y=2.", "x=3, y=1", "Medium", 3),
            ("Long Answer", "Solve a pair of linear equations by elimination and explain the steps.", "Eliminate one variable by multiplying equations appropriately, add/subtract, then back-substitute.", "Hard", 5),
        ],
    },
    "Science": {
        "Chapter 1 - Chemical Reactions and Equations": [
            ("Fill in the Blank", "A reaction in which a substance gains oxygen is called ________.", "oxidation", "Easy", 1),
            ("MCQ", "Zn + CuSO4 → ZnSO4 + Cu is a:", "A. Combination   B. Displacement   C. Decomposition   D. Neutralisation", "Easy", 1),
            ("Short Answer", "Why is magnesium ribbon cleaned before burning?", "To remove the magnesium oxide layer so that it burns readily.", "Medium", 3),
            ("Long Answer", "Explain combination, decomposition, displacement and double-displacement reactions with examples.", "Define each type and provide a balanced chemical equation for each.", "Hard", 5),
        ],
        "Chapter 2 - Acids, Bases and Salts": [
            ("Fill in the Blank", "Acids turn blue litmus paper ________.", "red", "Easy", 1),
            ("MCQ", "The pH of a neutral solution at room temperature is approximately:", "A. 0   B. 5   C. 7   D. 14", "Easy", 1),
            ("Short Answer", "What is a neutralisation reaction? Give one example.", "Reaction of an acid with a base to form salt and water; e.g. HCl + NaOH → NaCl + H2O.", "Medium", 3),
            ("Long Answer", "Explain the importance of pH in everyday life with suitable examples.", "Discuss digestion, tooth decay, soil treatment and self-defence of plants/animals as appropriate.", "Hard", 5),
        ],
        "Chapter 3 - Metals and Non-metals": [
            ("Fill in the Blank", "Metals generally form ________ ions by losing electrons.", "positive", "Easy", 1),
            ("MCQ", "Which metal is liquid at room temperature?", "A. Iron   B. Mercury   C. Copper   D. Aluminium", "Easy", 1),
            ("Short Answer", "Why do ionic compounds have high melting points?", "Strong electrostatic forces hold oppositely charged ions together.", "Medium", 3),
            ("Long Answer", "Explain the formation of ionic compounds using electron transfer.", "Describe loss/gain of electrons, ion formation and electrostatic attraction with an example such as NaCl.", "Hard", 5),
        ],
    },
    "Social Science": {
        "Chapter 1 - The Rise of Nationalism in Europe": [
            ("Fill in the Blank", "The French Revolution began in the year ________.", "1789", "Easy", 1),
            ("MCQ", "Who hosted the Congress of Vienna?", "A. Metternich   B. Napoleon   C. Garibaldi   D. Bismarck", "Easy", 1),
            ("Short Answer", "What did liberal nationalism mean in nineteenth-century Europe?", "It emphasised freedom for the individual and equality before law, along with representative government.", "Medium", 3),
            ("Long Answer", "Explain the process of German unification.", "Discuss the role of Prussia, Bismarck, wars and the proclamation of the German Empire.", "Hard", 5),
        ],
        "Chapter 2 - Nationalism in India": [
            ("Fill in the Blank", "The Non-Cooperation Movement was adopted in the year ________.", "1920", "Easy", 1),
            ("MCQ", "The Salt March ended at:", "A. Dandi   B. Delhi   C. Lahore   D. Bombay", "Easy", 1),
            ("Short Answer", "Why did Mahatma Gandhi launch the Civil Disobedience Movement?", "To challenge colonial laws and demand swaraj, beginning notably with the salt law.", "Medium", 3),
            ("Long Answer", "Describe the major features of the Civil Disobedience Movement.", "Explain the salt march, participation of different groups, spread of the movement and its limitations.", "Hard", 5),
        ],
        "Chapter 3 - The Making of a Global World": [
            ("Fill in the Blank", "The movement of people, goods and capital across countries is associated with ________.", "globalisation", "Easy", 1),
            ("MCQ", "Indentured labour was especially associated with:", "A. Contract labour migration   B. Factory ownership   C. Military service   D. Land reform", "Medium", 1),
            ("Short Answer", "What was the impact of technology on nineteenth-century global trade?", "Steamships, railways and communication reduced travel time and transport costs.", "Medium", 3),
            ("Long Answer", "Explain how the nineteenth century created a global economy.", "Discuss trade, migration, capital flows, technology and colonial expansion.", "Hard", 5),
        ],
    }
}

OFFICIAL_LINKS = {
    "NCERT Textbooks": "https://ncert.nic.in/textbook.php",
    "ePathshala": "https://epathshala.nic.in/",
    "CBSE Academic": "https://cbseacademic.nic.in/"
}

def collect_questions(subject, chapters, types, difficulty, total_marks):
    pool = []
    for ch in chapters:
        for q in DATA[subject][ch]:
            if q[0] in types and (difficulty == "Mixed" or q[3] == difficulty):
                pool.append((ch, *q))
    random.shuffle(pool)
    selected, marks = [], 0
    # Prefer a balanced selection and never exceed requested marks.
    for q in pool:
        if marks + q[5] <= total_marks:
            selected.append(q)
            marks += q[5]
        if marks == total_marks:
            break
    return selected, marks

def make_pdf(title, selected, total):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    title_style = styles["Title"]; title_style.alignment = TA_CENTER
    story = [Paragraph(title, title_style), Spacer(1, 10),
             Paragraph(f"Maximum Marks: {total}", styles["Heading2"]), Spacer(1, 12)]
    section_names = {"Fill in the Blank":"SECTION A - FILL IN THE BLANKS",
                     "MCQ":"SECTION B - MULTIPLE CHOICE QUESTIONS",
                     "Short Answer":"SECTION C - SHORT ANSWER QUESTIONS",
                     "Long Answer":"SECTION D - LONG ANSWER QUESTIONS"}
    n = 1
    for typ in ["Fill in the Blank","MCQ","Short Answer","Long Answer"]:
        qs = [x for x in selected if x[1] == typ]
        if not qs: continue
        story += [Paragraph(section_names[typ], styles["Heading2"])]
        for x in qs:
            story += [Paragraph(f"{n}. {x[2]}  [{x[5]} mark(s)]", styles["BodyText"]), Spacer(1, 6)]
            n += 1
        story.append(Spacer(1, 8))
    doc.build(story)
    return buf.getvalue()

def make_answer_pdf(title, selected):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, rightMargin=40, leftMargin=40, topMargin=40, bottomMargin=40)
    styles = getSampleStyleSheet()
    story = [Paragraph(title + " - ANSWER KEY", styles["Title"]), Spacer(1, 12)]
    for i, x in enumerate(selected, 1):
        story += [Paragraph(f"{i}. {x[2]}", styles["BodyText"]),
                  Paragraph(f"<b>Answer:</b> {x[3]}", styles["BodyText"]), Spacer(1, 8)]
    doc.build(story)
    return buf.getvalue()

def make_docx(title, selected, answer_key=False):
    d = Document()
    d.add_heading(title, 0)
    d.add_paragraph("Maximum Marks: " + str(sum(x[5] for x in selected)))
    for i, x in enumerate(selected, 1):
        d.add_heading(x[1], level=2) if i == 1 or selected[i-2][1] != x[1] else None
        d.add_paragraph(f"{i}. {x[2]} [{x[5]} mark(s)]")
        if answer_key:
            d.add_paragraph("Answer: " + x[3])
    out = BytesIO(); d.save(out); return out.getvalue()

st.title("📚 CBSE Exam Paper Generator")
st.caption("Version 1 • Grade 10 sample • Teacher-focused")

with st.sidebar:
    st.header("⚙️ Paper Settings")
    subject = st.selectbox("Subject", list(DATA.keys()))
    chapters = st.multiselect("Chapter(s)", list(DATA[subject].keys()), default=[list(DATA[subject].keys())[0]])
    types = st.multiselect("Question Types", ["Fill in the Blank","MCQ","Short Answer","Long Answer"],
                           default=["Fill in the Blank","MCQ","Short Answer","Long Answer"])
    difficulty = st.selectbox("Difficulty", ["Mixed","Easy","Medium","Hard"])
    total_marks = st.selectbox("Total Marks", [10,20,30,40,50], index=3)
    generate = st.button("📝 Generate Paper", use_container_width=True)

st.info("📌 Sample data contains the first three chapters for Grade 10 Mathematics, Science and Social Science. Replace/extend the question bank for full syllabus use.")

if generate:
    if not chapters or not types:
        st.error("Please select at least one chapter and one question type.")
    else:
        selected, achieved = collect_questions(subject, chapters, types, difficulty, total_marks)
        st.session_state["paper"] = selected
        st.session_state["subject"] = subject
        st.session_state["achieved"] = achieved

if "paper" in st.session_state:
    selected = st.session_state["paper"]
    achieved = st.session_state["achieved"]
    st.subheader("📄 Paper Preview")
    st.write(f"**Subject:** {st.session_state['subject']}  |  **Generated Marks:** {achieved}")
    if achieved < total_marks:
        st.warning("The sample question bank does not contain enough questions to reach the selected marks. Add more questions to the database.")
    for i, x in enumerate(selected, 1):
        st.markdown(f"**{i}. [{x[1]} | {x[3]} | {x[5]} mark(s)]**")
        st.write(x[2])
        if x[1] == "MCQ":
            st.write(x[3])
    title = f"CBSE Class 10 {st.session_state['subject']} Examination"
    c1, c2, c3 = st.columns(3)
    with c1:
        st.download_button("📄 Download Question Paper PDF", make_pdf(title, selected, achieved),
                           "CBSE_Question_Paper.pdf", "application/pdf")
    with c2:
        st.download_button("✅ Download Answer Key PDF", make_answer_pdf(title, selected),
                           "CBSE_Answer_Key.pdf", "application/pdf")
    with c3:
        st.download_button("📝 Download Word Paper", make_docx(title, selected),
                           "CBSE_Question_Paper.docx",
                           "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

st.divider()
st.subheader("📚 Official Textbook Resources")
for name, url in OFFICIAL_LINKS.items():
    st.link_button(name, url)
st.caption("Use official NCERT/CBSE resources for textbook downloads. The app does not redistribute copyrighted textbooks.")
