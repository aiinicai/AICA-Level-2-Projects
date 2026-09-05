"""
SECTION 3 - INPUT NORMALIZATION ENGINE

Converts arbitrary, heterogeneous input files into canonical, structured,
machine-verifiable representation.

Core Principles:
- Lossless extraction
- Schema-first, typed fields
- Explicit uncertainty
- Source traceability
- NO business logic validation
- NO semantic interpretation

Lifecycle: LOAD -> PARSE -> EXTRACT -> NORMALIZE -> VALIDATE -> EMIT
"""

import os
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

from models.entities import Evidence, StandardEntity
from services.entity_utils import make_entity_id
from services.file_processor import extract_pdf_text

# ==================== STEP 3: STANDARD ENTITY COERCION ====================

def _coerce_to_standard_entities(
    file_name: str,
    source_path: str,
    raw_entities: List[Any],
    doc_type: str,
) -> List[Dict[str, Any]]:
    """
    Converts task-specific entity shapes into StandardEntity dicts.

    Handles three cases:
    1. Already StandardEntity objects -> convert to dict
    2. Dict with entity_type/fields/evidence -> ensure complete
    3. Plain dicts -> wrap as GENERIC entity with minimal evidence

    Args:
        file_name: Source filename
        source_path: Full path to source file
        raw_entities: List of entities from task pack extractor
        doc_type: Document type (e.g., "pdf_challan")

    Returns:
        List of StandardEntity dicts
    """
    out: List[StandardEntity] = []

    for i, ent in enumerate(raw_entities or []):
        # Case 1: Already StandardEntity
        if isinstance(ent, StandardEntity):
            out.append(ent)
            continue

        # Case 2: Dict-like entity from task pack extractor
        if isinstance(ent, dict):
            # If task pack already uses entity_type/fields/evidence, keep it
            if "entity_type" in ent and "fields" in ent:
                # Ensure evidence exists
                ev = ent.get("evidence") or []
                if not ev:
                    ev = [Evidence(source_file=file_name, source_path=source_path, location="unknown").to_dict()]
                ent["evidence"] = ev

                # Ensure entity_id exists
                if not ent.get("entity_id"):
                    ent["entity_id"] = f"{ent['entity_type']}:{doc_type}:{i}"

                out.append(
                    StandardEntity(
                        entity_type=ent["entity_type"],
                        entity_id=ent["entity_id"],
                        fields=ent["fields"],
                        evidence=[
                            Evidence(**e) if isinstance(e, dict) else Evidence(source_file=file_name)
                            for e in ent["evidence"]
                        ],
                    )
                )
                continue

            # Otherwise wrap unknown dict into GENERIC entity
            fields = ent
            entity_type = f"{doc_type.upper()}_ENTITY"
            # Fallback: use common generic keys for unknown entities
            generic_keys = ["id", "number", "name"]
            entity_id = make_entity_id(entity_type, fields, generic_keys)
            out.append(
                StandardEntity(
                    entity_type=entity_type,
                    entity_id=entity_id,
                    fields=fields,
                    evidence=[Evidence(source_file=file_name, source_path=source_path, location="unknown")],
                )
            )
            continue

        # Case 3: Unknown type -> wrap as string
        out.append(
            StandardEntity(
                entity_type=f"{doc_type.upper()}_ENTITY",
                entity_id=f"{doc_type}:{i}",
                fields={"value": str(ent)},
                evidence=[Evidence(source_file=file_name, source_path=source_path, location="unknown")],
            )
        )

    return [e.to_dict() for e in out]


# ==================== 3.3 FILE CLASSIFICATION ====================
# NOTE: Document classification is now delegated to task packs.
# Core only handles generic file type detection (PDF, Excel, etc.)
# Task-specific classification is handled by task pack.


# ==================== 3.13 AMBIGUITY LOGGING ====================

class AmbiguityLogger:
    """Logs ambiguities encountered during normalization"""

    def __init__(self):
        self.ambiguities = []

    def log(self, ambiguity_type: str, file: str, details: str, context: Any = None):
        """
        Log an ambiguity.

        Args:
            ambiguity_type: Type of ambiguity (HEADER_AMBIGUITY, TYPE_UNCERTAIN, etc.)
            file: File where ambiguity occurred
            details: Human-readable description
            context: Additional context data
        """
        self.ambiguities.append({
            'type': ambiguity_type,
            'file': file,
            'details': details,
            'context': context,
            'timestamp': datetime.utcnow().isoformat()
        })

    def get_ambiguities(self) -> List[Dict[str, Any]]:
        """Get all logged ambiguities"""
        return self.ambiguities


# ==================== 3.8.4 TYPE NORMALIZATION ====================
# NOTE: Task-specific normalization (like section codes) is now delegated to task packs.
# This allows different tasks to have different normalization rules.


def normalize_number(raw_value: Any) -> Dict[str, Any]:
    """
    Normalize number with source traceability.

    Examples:
        34,102 -> 34102
        ₹34,102 -> {amount: 34102, currency: INR}

    Returns:
        {
            "raw": "34,102",
            "value": 34102,
            "type": "number"
        }
    """
    if pd.isna(raw_value) or raw_value == '':
        return {"raw": raw_value, "value": None, "type": "null"}

    raw_str = str(raw_value)

    # Handle currency
    currency = None
    if '₹' in raw_str:
        currency = 'INR'
        raw_str = raw_str.replace('₹', '')
    elif '$' in raw_str:
        currency = 'USD'
        raw_str = raw_str.replace('$', '')
    elif '€' in raw_str:
        currency = 'EUR'
        raw_str = raw_str.replace('€', '')

    # Remove commas, spaces
    cleaned = raw_str.replace(',', '').replace(' ', '').strip()

    # Try to convert to number
    try:
        value = float(cleaned)
        result = {
            "raw": raw_value,
            "value": value,
            "type": "currency" if currency else "number"
        }
        if currency:
            result['currency'] = currency
        return result
    except ValueError:
        return {"raw": raw_value, "value": None, "type": "invalid_number"}


def normalize_date(raw_value: Any) -> Dict[str, Any]:
    """
    Normalize date to ISO 8601 (YYYY-MM-DD).

    Examples:
        07/10/25 -> 2025-10-07
        06-Nov-2025 -> 2025-11-06

    Returns:
        {
            "raw": "07/10/25",
            "value": "2025-10-07",
            "type": "date"
        }
    """
    if pd.isna(raw_value) or raw_value == '':
        return {"raw": raw_value, "value": None, "type": "null"}

    # If already datetime object
    if isinstance(raw_value, datetime):
        return {
            "raw": str(raw_value),
            "value": raw_value.strftime('%Y-%m-%d'),
            "type": "date"
        }

    raw_str = str(raw_value).strip()

    # Try common date formats
    date_formats = [
        '%d/%m/%Y',   # 07/10/2025
        '%d/%m/%y',   # 07/10/25
        '%m/%d/%Y',   # 10/07/2025
        '%m/%d/%y',   # 10/07/25
        '%Y-%m-%d',   # 2025-10-07
        '%d-%b-%Y',   # 07-Oct-2025
        '%d-%B-%Y',   # 07-October-2025
        '%d.%m.%Y',   # 07.10.2025
    ]

    for fmt in date_formats:
        try:
            dt = datetime.strptime(raw_str, fmt)
            return {
                "raw": raw_value,
                "value": dt.strftime('%Y-%m-%d'),
                "type": "date"
            }
        except ValueError:
            continue

    # Failed to parse
    return {"raw": raw_value, "value": None, "type": "invalid_date"}


def normalize_text(raw_value: Any) -> Dict[str, Any]:
    """Normalize text field"""
    if pd.isna(raw_value) or raw_value == '':
        return {"raw": raw_value, "value": None, "type": "null"}

    return {
        "raw": raw_value,
        "value": str(raw_value).strip(),
        "type": "text"
    }


# ==================== 3.5 TEXT FILE EXTRACTOR ====================

def extract_text_file(file_path: str, file_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract text file (.txt, .md) with line numbers.

    Output:
        {
            "document_type": "TEXT",
            "lines": [
                {"line_no": 1, "content": "..."}
            ]
        }
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        lines_content = f.readlines()

    lines = [
        {"line_no": i + 1, "content": line.rstrip('\n')}
        for i, line in enumerate(lines_content)
    ]

    return {
        "document_type": "TEXT",
        "file": file_metadata.get('name'),
        "lines": lines
    }


# ==================== 3.6 DOCX EXTRACTOR ====================

def extract_docx_file(file_path: str, file_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract DOCX file with structure preservation.

    Output:
        {
            "document_type": "DOCX",
            "blocks": [
                {"type": "heading", "level": 1, "text": "..."},
                {"type": "paragraph", "text": "..."},
                {"type": "table", "rows": [...]}
            ]
        }
    """
    from docx import Document

    doc = Document(file_path)
    blocks = []

    # Extract paragraphs (including headings)
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            level = int(para.style.name.replace('Heading ', ''))
            blocks.append({
                "type": "heading",
                "level": level,
                "text": para.text
            })
        else:
            blocks.append({
                "type": "paragraph",
                "text": para.text
            })

    # Extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            rows.append(row_data)

        blocks.append({
            "type": "table",
            "rows": rows
        })

    return {
        "document_type": "DOCX",
        "file": file_metadata.get('name'),
        "blocks": blocks
    }


# ==================== 3.7 PDF EXTRACTOR ====================

def extract_pdf_file(file_path: str, file_metadata: Dict[str, Any]) -> Dict[str, Any]:
    """
    Extract PDF with page-wise text.

    Output:
        {
            "document_type": "PDF",
            "pages": [
                {"page_no": 1, "text": "...", "ocr_used": false}
            ]
        }
    """
    # Use existing PDF extractor
    full_text = extract_pdf_text(file_path)

    # Simple page splitting (not perfect, but reasonable)
    pages = []
    page_no = 1
    pages.append({
        "page_no": page_no,
        "text": full_text,
        "ocr_used": False  # TODO: Implement OCR fallback
    })

    return {
        "document_type": "PDF",
        "file": file_metadata.get('name'),
        "pages": pages
    }


# ==================== 3.8 SPREADSHEET EXTRACTOR (CRITICAL) ====================

def detect_header_row(df: pd.DataFrame) -> Tuple[int, List[str], Optional[str]]:
    """
    Detect header row in spreadsheet.

    Rules:
    - First non-empty row
    - Majority non-numeric
    - Keyword presence

    Returns:
        (header_row_index, column_headers, ambiguity_message)
    """
    for i in range(min(10, len(df))):  # Check first 10 rows
        row = df.iloc[i]

        # Check if row is mostly text (not numeric)
        non_numeric_count = sum(1 for val in row if isinstance(val, str) and val.strip())

        if non_numeric_count > len(row) / 2:
            headers = [str(val).strip() for val in row]
            return (i, headers, None)

    # Ambiguous - no clear header found
    return (0, list(df.columns), "No clear header row detected - using default column names")


def normalize_spreadsheet_row(row_index: int, row_data: pd.Series, headers: List[str],
                               file_metadata: Dict[str, Any], sheet_name: str) -> Dict[str, Any]:
    """
    Normalize a single spreadsheet row with typed fields and source traceability.

    Output:
        {
            "row_index": 14,
            "fields": {
                "document_no": {"raw": "34102", "value": 34102, "type": "number"},
                "date": {"raw": "07/10/25", "value": "2025-10-07", "type": "date"},
                "amount": {"raw": "34,102", "value": 34102, "type": "currency"}
            },
            "source": {
                "file": "data.xlsx",
                "sheet": "Sheet1",
                "row": 14
            }
        }
    """
    fields = {}

    for col_idx, header in enumerate(headers):
        if col_idx >= len(row_data):
            continue

        raw_value = row_data.iloc[col_idx]

        # Infer type and normalize
        # Check if looks like a date
        if any(keyword in header.lower() for keyword in ['date', 'time', 'when']):
            normalized = normalize_date(raw_value)
        # Check if looks like a number/amount
        elif any(keyword in header.lower() for keyword in ['amount', 'total', 'sum', 'value', 'no', 'number']):
            normalized = normalize_number(raw_value)
        # Default to text
        else:
            normalized = normalize_text(raw_value)

        # Add source traceability to each field
        normalized['source'] = {
            'file': file_metadata.get('name'),
            'sheet': sheet_name,
            'row': row_index + 1,  # 1-indexed
            'column': header
        }

        fields[header] = normalized

    return {
        "row_index": row_index,
        "fields": fields,
        "source": {
            "file": file_metadata.get('name'),
            "sheet": sheet_name,
            "row": row_index + 1
        }
    }


def extract_spreadsheet_file(file_path: str, file_metadata: Dict[str, Any],
                              ambiguity_logger: AmbiguityLogger) -> Dict[str, Any]:
    """
    Extract spreadsheet (.xlsx, .csv) with full normalization.

    Output:
        {
            "document_type": "SPREADSHEET",
            "sheets": [
                {
                    "sheet_name": "Sheet1",
                    "headers": [...],
                    "rows": [...]
                }
            ]
        }
    """
    ext = os.path.splitext(file_path)[1].lower()

    sheets_data = []

    if ext == '.csv':
        # CSV has single sheet
        df = pd.read_csv(file_path, header=None)
        header_row_idx, headers, ambiguity = detect_header_row(df)

        if ambiguity:
            ambiguity_logger.log('HEADER_AMBIGUITY', file_metadata.get('name'), ambiguity)

        # Skip header row in data
        data_df = df.iloc[header_row_idx + 1:]

        rows = []
        for idx, row in data_df.iterrows():
            normalized_row = normalize_spreadsheet_row(idx, row, headers, file_metadata, 'Sheet1')
            rows.append(normalized_row)

        sheets_data.append({
            "sheet_name": "Sheet1",
            "headers": headers,
            "rows": rows
        })

    elif ext in ['.xlsx', '.xlsm', '.xls']:
        # Excel can have multiple sheets
        xls = pd.ExcelFile(file_path)

        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=None)

            header_row_idx, headers, ambiguity = detect_header_row(df)

            if ambiguity:
                ambiguity_logger.log('HEADER_AMBIGUITY', f"{file_metadata.get('name')}:{sheet_name}", ambiguity)

            # Skip header row in data
            data_df = df.iloc[header_row_idx + 1:]

            rows = []
            for idx, row in data_df.iterrows():
                normalized_row = normalize_spreadsheet_row(idx, row, headers, file_metadata, sheet_name)
                rows.append(normalized_row)

            sheets_data.append({
                "sheet_name": sheet_name,
                "headers": headers,
                "rows": rows
            })

    return {
        "document_type": "SPREADSHEET",
        "file": file_metadata.get('name'),
        "sheets": sheets_data
    }


# ==================== 3.4 FILE TYPE DISPATCH ====================

def extract_generic(file_path: str, file_metadata: Dict[str, Any],
                    ambiguity_logger: AmbiguityLogger) -> Dict[str, Any]:
    """
    Generic, task-agnostic extraction by file type. Every common format is
    supported so any task can be normalized without a specialized extractor.
    """
    filename = file_metadata.get('name', '')
    ext = os.path.splitext(filename)[1].lower()

    if ext in ['.txt', '.md']:
        return extract_text_file(file_path, file_metadata)
    if ext in ['.docx', '.doc']:
        return extract_docx_file(file_path, file_metadata)
    if ext in ['.csv', '.xlsx', '.xlsm', '.xls']:
        return extract_spreadsheet_file(file_path, file_metadata, ambiguity_logger)
    if ext == '.pdf':
        return extract_pdf_file(file_path, file_metadata)

    return {
        "document_type": "UNKNOWN",
        "file": filename,
        "handled": False,
        "reason": f"Unsupported format: {ext}",
    }


def dispatch_extractor(file_path: str, file_metadata: Dict[str, Any],
                       ambiguity_logger: AmbiguityLogger, task=None) -> Dict[str, Any]:
    """
    Extract a file into a normalized artifact.

    Order of precedence:
    1. Task pack first refusal — if a task pack is loaded and classifies the file
       as a known type, use its specialized extractor (e.g. TDS challan/BSCT/26Q).
    2. Generic fallback — otherwise extract by file type (text/docx/spreadsheet/pdf).

    Every returned artifact carries a uniform 'text' + 'tables' view (see
    _attach_generic_view) so downstream criteria/checks/AI can consume any file
    type the same way.
    """
    filename = file_metadata.get('name', '')

    try:
        # 1. Task pack first refusal (any extension)
        if task is not None:
            try:
                doc_type = task.classify(file_path)
            except Exception:
                doc_type = "unknown"

            if doc_type and doc_type != "unknown":
                raw_entities = task.extract(doc_type, file_path)
                standard_entities = _coerce_to_standard_entities(
                    file_name=filename,
                    source_path=file_path,
                    raw_entities=raw_entities,
                    doc_type=doc_type,
                )
                return _attach_generic_view({
                    "document_type": doc_type,
                    "file": filename,
                    "source_path": file_path,
                    "handled": True,
                    "entities": standard_entities,
                    "extractor": f"{task.task_id}:{doc_type}",
                    "confidence": 0.8,
                })

        # 2. Generic fallback
        return _attach_generic_view(extract_generic(file_path, file_metadata, ambiguity_logger))

    except Exception as e:
        ambiguity_logger.log('EXTRACTION_ERROR', filename, str(e))
        return {
            "document_type": "ERROR",
            "file": filename,
            "handled": False,
            "reason": str(e),
        }


# ==================== 3.4b UNIFORM GENERIC VIEW ====================

def _cell_value(value: Any) -> Any:
    """Unwrap a normalized spreadsheet cell ({raw,value,type}) to its plain value."""
    if isinstance(value, dict) and 'value' in value and 'type' in value:
        return value.get('value')
    return value


def _attach_generic_view(artifact: Dict[str, Any], max_text_chars: int = 20000) -> Dict[str, Any]:
    """
    Attach a uniform {'text', 'tables'} view to any extracted artifact so that
    the criteria engine, generic deterministic checks, and the AI validator can
    consume PDFs, docs, spreadsheets, and task-pack entities identically.

    - text:   a flattened textual rendering (truncated to max_text_chars)
    - tables: list of {'name', 'headers', 'rows'} where rows are lists of cell values
    """
    if not isinstance(artifact, dict):
        return artifact

    doc_type = artifact.get('document_type')
    text_parts: List[str] = []
    tables: List[Dict[str, Any]] = []

    if doc_type == 'TEXT':
        text_parts = [ln.get('content', '') for ln in artifact.get('lines', [])]

    elif doc_type == 'DOCX':
        for block in artifact.get('blocks', []):
            if block.get('type') in ('paragraph', 'heading'):
                if block.get('text'):
                    text_parts.append(block['text'])
            elif block.get('type') == 'table':
                rows = block.get('rows', [])
                headers = rows[0] if rows else []
                tables.append({'name': 'table', 'headers': headers, 'rows': rows[1:] if rows else []})

    elif doc_type == 'PDF':
        text_parts = [pg.get('text', '') for pg in artifact.get('pages', [])]

    elif doc_type == 'SPREADSHEET':
        for sheet in artifact.get('sheets', []):
            headers = sheet.get('headers', [])
            rows = []
            for row in sheet.get('rows', []):
                fields = row.get('fields', {})
                rows.append([_cell_value(fields.get(h)) for h in headers])
            tables.append({'name': sheet.get('sheet_name', 'Sheet1'), 'headers': headers, 'rows': rows})
            text_parts.append(f"# {sheet.get('sheet_name', 'Sheet1')}")
            text_parts.append(", ".join(str(h) for h in headers))
            for r in rows[:200]:
                text_parts.append(", ".join("" if c is None else str(c) for c in r))

    elif artifact.get('entities'):
        # Task-pack entities (or coerced generic entities)
        entities = artifact.get('entities', [])
        header_keys: List[str] = []
        for ent in entities:
            for k in (ent.get('fields', {}) if isinstance(ent, dict) else {}):
                if k not in header_keys:
                    header_keys.append(k)
        rows = []
        for ent in entities:
            fields = ent.get('fields', {}) if isinstance(ent, dict) else {}
            rows.append([fields.get(k) for k in header_keys])
            text_parts.append("; ".join(f"{k}={fields.get(k)}" for k in header_keys))
        if header_keys:
            tables.append({'name': doc_type or 'entities', 'headers': header_keys, 'rows': rows})

    text = "\n".join(str(p) for p in text_parts if p is not None)
    if max_text_chars and len(text) > max_text_chars:
        text = text[:max_text_chars] + "\n... [truncated]"

    artifact['text'] = text
    artifact['tables'] = tables
    return artifact


# ==================== 3.16 SECTION 3 ORCHESTRATOR ====================

class Section3Normalization:
    """
    Input Normalization Engine - Section 3 orchestrator.

    Takes Section 2 manifest, produces normalized inputs.
    NOW GENERIC: Uses task packs for task-specific extraction.
    """

    def __init__(self, section2_manifest: Dict[str, Any], task_id: str):
        """
        Args:
            section2_manifest: Output from Section 2 ingestion
            task_id: Task identifier (e.g., "tds_26q")
        """
        self.manifest = section2_manifest
        self.task_id = task_id
        self.ambiguity_logger = AmbiguityLogger()
        self.extraction_logs = []

        # Load task pack
        from services.task_loader import load_task
        self.task = load_task(task_id)

    def normalize(self) -> Dict[str, Any]:
        """
        Execute full normalization pipeline.

        Returns:
            {
                "normalized_inputs": [...],
                "extraction_logs": [...],
                "ambiguities": [...],
                "source_map": {...}
            }
        """
        print("=" * 80)
        print("SECTION 3: INPUT NORMALIZATION START")
        print("=" * 80)

        normalized_inputs = []
        source_map = {}

        # Process ONLY files in INPUTS role
        input_files = self.manifest.get('role_index', {}).get('INPUTS', [])

        print(f"\n[1/3] Processing {len(input_files)} input files...")

        for file_meta in input_files:
            file_path = file_meta.get('local_path')
            filename = file_meta.get('name')

            if not file_path or not os.path.exists(file_path):
                self.ambiguity_logger.log('FILE_NOT_FOUND', filename, f"Local file not found: {file_path}")
                continue

            print(f"   -> Extracting: {filename}")

            # Dispatch to extractor (with task pack)
            extracted = dispatch_extractor(file_path, file_meta, self.ambiguity_logger, task=self.task)

            # Use extracted document type directly (task pack handles classification)
            doc_type = extracted.get('document_type', 'UNKNOWN')
            confidence = extracted.get('confidence', 0.0)

            extracted['classified_as'] = doc_type
            extracted['classification_confidence'] = confidence

            normalized_inputs.append(extracted)

            # Update source map
            source_map[filename] = {
                'onedrive_file_id': file_meta.get('onedrive_file_id'),
                'path': file_meta.get('path'),
                'document_type': doc_type
            }

            self.extraction_logs.append({
                'file': filename,
                'document_type': doc_type,
                'confidence': confidence,
                'status': 'success' if extracted.get('document_type') != 'ERROR' else 'error'
            })

        print(f"   [OK] Extracted {len(normalized_inputs)} files")

        # Process OUTPUT files (ADDED - was missing!)
        normalized_outputs = []
        output_files = self.manifest.get('role_index', {}).get('OUTPUTS', [])

        if output_files:
            print(f"\n[2/5] Processing {len(output_files)} output files...")

            for file_meta in output_files:
                file_path = file_meta.get('local_path')
                filename = file_meta.get('name')

                if not file_path or not os.path.exists(file_path):
                    self.ambiguity_logger.log('FILE_NOT_FOUND', filename, f"Local file not found: {file_path}")
                    continue

                print(f"   -> Extracting: {filename}")

                # Dispatch to extractor (with task pack)
                extracted = dispatch_extractor(file_path, file_meta, self.ambiguity_logger, task=self.task)

                # Use extracted document type directly (task pack handles classification)
                doc_type = extracted.get('document_type', 'UNKNOWN')
                confidence = extracted.get('confidence', 0.0)

                extracted['classified_as'] = doc_type
                extracted['classification_confidence'] = confidence

                normalized_outputs.append(extracted)

                # Update source map
                source_map[filename] = {
                    'onedrive_file_id': file_meta.get('onedrive_file_id'),
                    'path': file_meta.get('path'),
                    'document_type': doc_type
                }

                self.extraction_logs.append({
                    'file': filename,
                    'document_type': doc_type,
                    'confidence': confidence,
                    'status': 'success' if extracted.get('document_type') != 'ERROR' else 'error'
                })

            print(f"   [OK] Extracted {len(normalized_outputs)} output files")
        else:
            print("\n[2/5] No output files to process")

        # Get ambiguities
        ambiguities = self.ambiguity_logger.get_ambiguities()

        if ambiguities:
            print(f"\n[3/5] Logged {len(ambiguities)} ambiguities")
        else:
            print("\n[3/5] No ambiguities detected")

        print("\n[4/5] Building output...")

        result = {
            "normalized_inputs": normalized_inputs,
            "normalized_outputs": normalized_outputs,  # ADDED - was missing!
            "extraction_logs": self.extraction_logs,
            "ambiguities": ambiguities,
            "source_map": source_map
        }

        print("\n[5/5] Complete!\n")

        print("=" * 80)
        print("SECTION 3: INPUT NORMALIZATION COMPLETE")
        print("=" * 80)

        return result


# ==================== CONVENIENCE FUNCTION ====================

def normalize_inputs(section2_manifest: Dict[str, Any], task_id: str) -> Dict[str, Any]:
    """
    Convenience function for Section 3 normalization.

    Args:
        section2_manifest: Output from Section 2
        task_id: Task identifier (e.g., "tds_26q")

    Returns:
        Normalized inputs
    """
    normalizer = Section3Normalization(section2_manifest, task_id=task_id)
    return normalizer.normalize()
