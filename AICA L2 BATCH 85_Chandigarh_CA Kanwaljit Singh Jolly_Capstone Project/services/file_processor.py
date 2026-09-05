"""
File Processor Service
Intelligently processes and summarizes different file types
"""
import csv
import json
import os

import fitz  # PyMuPDF
from docx import Document
from openpyxl import load_workbook

DEFAULT_TEXT_PREVIEW_CHARS = 1200
DEFAULT_PDF_MAX_PAGES = 8
DEFAULT_PDF_MAX_CHARS = 12000
DEFAULT_DOCX_MAX_CHARS = 8000
DEFAULT_EXCEL_MAX_ROWS = 100
DEFAULT_EXCEL_MAX_CHARS = 15000
DEFAULT_CSV_SAMPLE_ROWS = 8
DEFAULT_CSV_STAT_ROWS = 200
DEFAULT_EXCEL_SAMPLE_ROWS = 25
DEFAULT_EXCEL_STAT_ROWS = 200
DEFAULT_JSON_SAMPLE_KEYS = 25
DEFAULT_CHUNK_MAX_CHARS = 2000
DEFAULT_CHUNK_MAX_LINES = 200
DEFAULT_CHUNK_CSV_ROWS = 200
DEFAULT_CHUNK_JSON_ITEMS = 200


def _normalize_rel_path(path):
    if not path:
        return ""
    return path.replace(os.sep, '/')


def _chunk_lines(lines, max_chars=DEFAULT_CHUNK_MAX_CHARS, max_lines=DEFAULT_CHUNK_MAX_LINES):
    chunks = []
    start_line = 1
    current = []
    current_len = 0

    for idx, line in enumerate(lines, 1):
        line_len = len(line)
        if (current and (current_len + line_len > max_chars or len(current) >= max_lines)):
            chunks.append({
                'start_line': start_line,
                'end_line': start_line + len(current) - 1,
                'text': "".join(current)
            })
            current = []
            current_len = 0
            start_line = idx

        current.append(line)
        current_len += line_len

    if current:
        chunks.append({
            'start_line': start_line,
            'end_line': start_line + len(current) - 1,
            'text': "".join(current)
        })

    return chunks


def _chunk_paragraphs(paragraphs, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    chunks = []
    start_idx = 1
    current = []
    current_len = 0

    for idx, text in enumerate(paragraphs, 1):
        line = text + "\n"
        line_len = len(line)
        if current and (current_len + line_len > max_chars):
            chunks.append({
                'start_para': start_idx,
                'end_para': start_idx + len(current) - 1,
                'text': "".join(current)
            })
            current = []
            current_len = 0
            start_idx = idx

        current.append(line)
        current_len += line_len

    if current:
        chunks.append({
            'start_para': start_idx,
            'end_para': start_idx + len(current) - 1,
            'text': "".join(current)
        })

    return chunks


def _chunk_csv_rows(headers, rows, max_rows=DEFAULT_CHUNK_CSV_ROWS, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    chunks = []
    start_idx = 1
    current = []
    current_len = 0

    header_line = "Headers: " + " | ".join(headers) + "\n"
    for idx, row in enumerate(rows, 1):
        sanitized = []
        for cell in row:
            cell_text = str(cell)
            if len(cell_text) > 200:
                cell_text = cell_text[:200] + "..."
            sanitized.append(cell_text)
        row_line = f"Row {idx}: " + " | ".join(sanitized) + "\n"

        if current and (len(current) >= max_rows or current_len + len(row_line) > max_chars):
            chunk_text = header_line + "".join(current)
            chunks.append({
                'start_row': start_idx,
                'end_row': start_idx + len(current) - 1,
                'text': chunk_text
            })
            current = []
            current_len = 0
            start_idx = idx

        current.append(row_line)
        current_len += len(row_line)

    if current:
        chunk_text = header_line + "".join(current)
        chunks.append({
            'start_row': start_idx,
            'end_row': start_idx + len(current) - 1,
            'text': chunk_text
        })

    return chunks


def _chunk_json(data, max_items=DEFAULT_CHUNK_JSON_ITEMS, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    chunks = []
    if isinstance(data, dict):
        keys = list(data.keys())
        for start in range(0, len(keys), max_items):
            slice_keys = keys[start:start + max_items]
            payload = {key: data[key] for key in slice_keys}
            text = json.dumps(payload, ensure_ascii=True, sort_keys=True)
            if len(text) > max_chars:
                text = text[:max_chars] + "..."
            chunks.append({
                'location': f"keys {start + 1}-{start + len(slice_keys)}",
                'text': text
            })
    elif isinstance(data, list):
        for start in range(0, len(data), max_items):
            items = data[start:start + max_items]
            text = json.dumps(items, ensure_ascii=True)
            if len(text) > max_chars:
                text = text[:max_chars] + "..."
            chunks.append({
                'location': f"items {start + 1}-{start + len(items)}",
                'text': text
            })
    else:
        chunks.append({
            'location': 'value',
            'text': json.dumps(data, ensure_ascii=True)
        })

    return chunks


def extract_file_chunks(path, rel_name=None, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    """
    Extract full-content chunks with stable references.
    Returns list of chunk dicts with file, location, and text.
    """
    rel_name = _normalize_rel_path(rel_name or os.path.basename(path))
    file_lower = rel_name.lower()
    chunks = []

    if file_lower.endswith(('.txt', '.md', '.log')):
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        for entry in _chunk_lines(lines, max_chars=max_chars):
            chunks.append({
                'file': rel_name,
                'location': f"lines {entry['start_line']}-{entry['end_line']}",
                'text': entry['text']
            })
        return chunks

    if file_lower.endswith('.docx'):
        doc = Document(path)
        paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        for entry in _chunk_paragraphs(paragraphs, max_chars=max_chars):
            chunks.append({
                'file': rel_name,
                'location': f"paragraphs {entry['start_para']}-{entry['end_para']}",
                'text': entry['text']
            })
        return chunks

    if file_lower.endswith('.pdf'):
        try:
            with fitz.open(path) as pdf:
                for page_index, page in enumerate(pdf, 1):
                    page_text = page.get_text() or ""
                    if not page_text.strip():
                        continue
                    page_lines = [line + "\n" for line in page_text.splitlines()]
                    for entry in _chunk_lines(page_lines, max_chars=max_chars):
                        chunks.append({
                            'file': rel_name,
                            'location': f"page {page_index}, lines {entry['start_line']}-{entry['end_line']}",
                            'text': entry['text']
                        })
            return chunks
        except Exception as exc:
            return [{
                'file': rel_name,
                'location': 'pdf',
                'text': f"Error reading PDF: {str(exc)}"
            }]

    if file_lower.endswith('.csv'):
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            rows = [row for row in reader]
        for entry in _chunk_csv_rows(headers, rows, max_chars=max_chars):
            chunks.append({
                'file': rel_name,
                'location': f"rows {entry['start_row']}-{entry['end_row']}",
                'text': entry['text']
            })
        return chunks

    if file_lower.endswith(('.xlsx', '.xls', '.xlsm', '.xlsb')):
        try:
            import pandas as pd
        except Exception:
            return [{
                'file': rel_name,
                'location': 'sheet',
                'text': 'Excel content extraction not available (do not treat as error).'
            }]

        try:
            sheets = pd.read_excel(path, sheet_name=None)
        except Exception as exc:
            return [{
                'file': rel_name,
                'location': 'sheet',
                'text': f"Excel content extraction failed: {str(exc)}"
            }]
        for sheet_name, df in sheets.items():
            headers = [str(col) for col in df.columns.tolist()]
            rows = df.astype(str).values.tolist()
            for entry in _chunk_csv_rows(headers, rows, max_chars=max_chars):
                chunks.append({
                    'file': rel_name,
                    'location': f"sheet {sheet_name}, rows {entry['start_row']}-{entry['end_row']}",
                    'text': entry['text']
                })
        return chunks

    if file_lower.endswith('.json'):
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        data = json.loads(content) if content.strip() else {}
        for entry in _chunk_json(data, max_chars=max_chars):
            location = entry.get('location', 'json')
            chunks.append({
                'file': rel_name,
                'location': location,
                'text': entry.get('text', '')
            })
        return chunks

    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            lines = f.readlines()
        for entry in _chunk_lines(lines, max_chars=max_chars):
            chunks.append({
                'file': rel_name,
                'location': f"lines {entry['start_line']}-{entry['end_line']}",
                'text': entry['text']
            })
    except Exception as exc:
        chunks.append({
            'file': rel_name,
            'location': 'unknown',
            'text': f"Error reading file: {str(exc)}"
        })

    return chunks


def extract_file_chunks_alt(path, rel_name=None, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    """
    Alternate extractor for conflict detection.
    """
    rel_name = _normalize_rel_path(rel_name or os.path.basename(path))
    file_lower = rel_name.lower()
    chunks = []

    if file_lower.endswith('.pdf'):
        try:
            text = extract_pdf_text(path, max_pages=15, max_chars=max_chars * 2)
            lines = [line + "\n" for line in text.splitlines()]
            for entry in _chunk_lines(lines, max_chars=max_chars):
                chunks.append({
                    'file': rel_name,
                    'location': f"alt lines {entry['start_line']}-{entry['end_line']}",
                    'text': entry['text']
                })
            return chunks
        except Exception as exc:
            return [{
                'file': rel_name,
                'location': 'pdf',
                'text': f"Alt PDF error: {str(exc)}"
            }]

    if file_lower.endswith('.csv'):
        try:
            import pandas as pd
            df = pd.read_csv(path)
            headers = [str(col) for col in df.columns.tolist()]
            rows = df.astype(str).values.tolist()
            for entry in _chunk_csv_rows(headers, rows, max_chars=max_chars):
                chunks.append({
                    'file': rel_name,
                    'location': f"alt rows {entry['start_row']}-{entry['end_row']}",
                    'text': entry['text']
                })
            return chunks
        except Exception as exc:
            return [{
                'file': rel_name,
                'location': 'csv',
                'text': f"Alt CSV error: {str(exc)}"
            }]

    if file_lower.endswith(('.log', '.txt', '.md')):
        try:
            content = read_text_file(path, max_chars=max_chars * 4)
            lines = [line + "\n" for line in content.splitlines()]
            for entry in _chunk_lines(lines, max_chars=max_chars):
                chunks.append({
                    'file': rel_name,
                    'location': f"alt lines {entry['start_line']}-{entry['end_line']}",
                    'text': entry['text']
                })
            return chunks
        except Exception as exc:
            return [{
                'file': rel_name,
                'location': 'log',
                'text': f"Alt log error: {str(exc)}"
            }]

    return extract_file_chunks(path, rel_name=rel_name, max_chars=max_chars)


def detect_extraction_conflicts(primary_chunks, alt_chunks, min_ratio=0.6):
    if not primary_chunks or not alt_chunks:
        return []

    def text_len(chunks):
        return sum(len(chunk.get('text', '') or '') for chunk in chunks)

    primary_len = text_len(primary_chunks)
    alt_len = text_len(alt_chunks)
    smaller = min(primary_len, alt_len)
    larger = max(primary_len, alt_len)
    ratio = (smaller / larger) if larger else 1.0

    if ratio >= min_ratio:
        return []

    return [{
        'primary_len': primary_len,
        'alt_len': alt_len,
        'ratio': round(ratio, 3)
    }]


def extract_folder_chunks_dual(folder_path, prefix=None, recursive=False, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return [], []

    files = list_folder_files(folder_path, recursive=recursive)
    combined = []
    conflicts = []

    for file_info in files:
        rel_name = _normalize_rel_path(file_info.get('name'))
        if prefix:
            rel_name = f"{prefix}/{rel_name}"
        full_path = os.path.join(folder_path, *file_info.get('name', '').split('/'))
        primary = extract_file_chunks(full_path, rel_name=rel_name, max_chars=max_chars)
        alternate = extract_file_chunks_alt(full_path, rel_name=rel_name, max_chars=max_chars)
        combined.extend(primary)

        conflict = detect_extraction_conflicts(primary, alternate)
        if conflict:
            conflicts.append({
                'file': rel_name,
                'details': conflict
            })

    return combined, conflicts


def extract_folder_chunks(folder_path, prefix=None, recursive=False, max_chars=DEFAULT_CHUNK_MAX_CHARS):
    """
    Extract chunks for all files in a folder.
    """
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return []

    files = list_folder_files(folder_path, recursive=recursive)
    chunks = []
    for file_info in files:
        rel_name = _normalize_rel_path(file_info.get('name'))
        if prefix:
            rel_name = f"{prefix}/{rel_name}"
        full_path = os.path.join(folder_path, *file_info.get('name', '').split('/'))
        file_chunks = extract_file_chunks(full_path, rel_name=rel_name, max_chars=max_chars)
        chunks.extend(file_chunks)

    return chunks


def read_workflow_file(path):
    """
    Read workflow file (supports .txt, .md, .docx, .pdf)

    Args:
        path: Path to workflow file

    Returns:
        str: Workflow content
    """
    if not os.path.exists(path):
        raise FileNotFoundError(f"Workflow file not found at: {path}")

    # Get file extension
    file_lower = os.path.basename(path).lower()

    # Extract content based on file type
    if file_lower.endswith('.docx'):
        content = extract_docx_text(path, max_chars=DEFAULT_DOCX_MAX_CHARS)
    elif file_lower.endswith('.pdf'):
        content = extract_pdf_text(path, max_pages=DEFAULT_PDF_MAX_PAGES, max_chars=DEFAULT_PDF_MAX_CHARS)
    else:
        # Plain text file
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

    if not content.strip():
        raise ValueError("Workflow file is empty or could not be read")

    return content


def truncate_text(text, max_chars):
    """Return a head/tail preview to keep context while limiting size."""
    if max_chars is None or max_chars <= 0:
        return text

    text = text or ""
    if len(text) <= max_chars:
        return text

    head = int(max_chars * 0.65)
    tail = max_chars - head - 40
    if tail < 0:
        tail = 0

    return (
        text[:head]
        + "\n... [truncated to fit prompt budget] ...\n"
        + (text[-tail:] if tail else "")
    )


def list_folder_files(folder_path, recursive=False):
    """
    List file metadata for a folder.

    Args:
        folder_path: Path to folder containing files

    Returns:
        list[dict]: [{name, ext, size_kb}]
    """
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return []

    files = []
    if recursive:
        for root, _, filenames in os.walk(folder_path):
            for filename in filenames:
                full_path = os.path.join(root, filename)
                if not os.path.isfile(full_path):
                    continue
                rel_path = os.path.relpath(full_path, folder_path)
                files.append(rel_path)
    else:
        files = [
            f for f in os.listdir(folder_path)
            if os.path.isfile(os.path.join(folder_path, f))
        ]

    result = []
    for filename in sorted(files):
        filepath = os.path.join(folder_path, *filename.split('/'))
        ext = os.path.splitext(filename)[1].lower() or 'none'
        size_kb = os.path.getsize(filepath) / 1024
        result.append({
            'name': filename.replace(os.sep, '/'),
            'ext': ext,
            'size_kb': round(size_kb, 1)
        })

    return result


def format_file_manifest(files):
    if not files:
        return "No files found."

    lines = ["File manifest:"]
    for file_info in files:
        lines.append(
            f"- {file_info['name']} ({file_info['ext']}, {file_info['size_kb']} KB)"
        )
    return "\n".join(lines)


def process_folder_files(
    folder_path,
    preview_chars=DEFAULT_TEXT_PREVIEW_CHARS,
    csv_sample_rows=DEFAULT_CSV_SAMPLE_ROWS,
    csv_stat_rows=DEFAULT_CSV_STAT_ROWS,
    recursive=False
):
    """
    Intelligently process all files in a folder and return summary

    Args:
        folder_path: Path to folder containing files
        preview_chars: Maximum chars per file preview
        csv_sample_rows: Sample rows to include for CSVs
        csv_stat_rows: Max rows to scan for CSV stats

    Returns:
        str: Formatted summary of all files
    """
    if not os.path.exists(folder_path):
        return f"Folder not found: {folder_path}"

    if not os.path.isdir(folder_path):
        return f"Not a directory: {folder_path}"

    files = list_folder_files(folder_path)

    if not files:
        return f"Empty folder: {folder_path}"

    summaries = []
    summaries.append(f"Folder: {os.path.basename(folder_path)}")
    summaries.append(format_file_manifest(files))
    summaries.append("")

    for file_info in files:
        filename = file_info['name']
        filepath = os.path.join(folder_path, filename)

        try:
            file_lower = filename.lower()

            if file_lower.endswith(('.txt', '.md')):
                content = read_text_file(filepath, max_chars=preview_chars * 2)
                summaries.append(f"File: {filename}")
                summaries.append("Type: Text")
                summaries.append(f"Preview:\n{truncate_text(content, preview_chars)}")
                summaries.append("")

            elif file_lower.endswith('.csv'):
                summary = summarize_csv(
                    filepath,
                    max_rows=csv_sample_rows,
                    max_stat_rows=csv_stat_rows
                )
                summaries.append(f"File: {filename}")
                summaries.append("Type: CSV")
                summaries.append(summary)
                summaries.append("")

            elif file_lower.endswith('.docx'):
                content = extract_docx_text(filepath, max_chars=DEFAULT_DOCX_MAX_CHARS)
                summaries.append(f"File: {filename}")
                summaries.append("Type: DOCX")
                summaries.append(f"Preview:\n{truncate_text(content, preview_chars)}")
                summaries.append("")

            elif file_lower.endswith('.pdf'):
                content = extract_pdf_text(
                    filepath,
                    max_pages=DEFAULT_PDF_MAX_PAGES,
                    max_chars=DEFAULT_PDF_MAX_CHARS
                )
                summaries.append(f"File: {filename}")
                summaries.append("Type: PDF")
                summaries.append(f"Preview:\n{truncate_text(content, preview_chars)}")
                summaries.append("")

            elif file_lower.endswith(('.xlsx', '.xls', '.xlsm', '.xlsb')):
                content = extract_excel_text(filepath, max_rows=DEFAULT_EXCEL_MAX_ROWS, max_chars=DEFAULT_EXCEL_MAX_CHARS)
                summaries.append(f"File: {filename}")
                summaries.append("Type: Excel")
                summaries.append(f"Preview:\n{content}")
                summaries.append("")

            elif file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                summaries.append(f"File: {filename}")
                summaries.append("Type: Image")
                summaries.append("Preview: Image content requires OCR.")
                summaries.append("")

            else:
                try:
                    content = read_text_file(filepath, max_chars=preview_chars)
                    summaries.append(f"File: {filename}")
                    summaries.append("Type: Unknown (text preview)")
                    summaries.append(f"Preview:\n{truncate_text(content, preview_chars)}")
                    summaries.append("")
                except Exception:
                    summaries.append(f"File: {filename}")
                    summaries.append("Type: Binary/Unsupported")
                    summaries.append("Preview: Not available.")
                    summaries.append("")

        except Exception as e:
            summaries.append(f"File: {filename}")
            summaries.append(f"Error: {str(e)}")
            summaries.append("")

    return "\n".join(summaries).strip()


def process_folder_files_structured(
    folder_path,
    preview_chars=DEFAULT_TEXT_PREVIEW_CHARS,
    csv_sample_rows=DEFAULT_CSV_SAMPLE_ROWS,
    csv_stat_rows=DEFAULT_CSV_STAT_ROWS,
    excel_max_rows=DEFAULT_EXCEL_MAX_ROWS,
    excel_max_chars=DEFAULT_EXCEL_MAX_CHARS,
    recursive=False
):
    """
    Process all files in a folder and return structured summaries.
    """
    if not os.path.exists(folder_path):
        return []

    if not os.path.isdir(folder_path):
        return []

    files = list_folder_files(folder_path, recursive=recursive)
    if not files:
        return []

    summaries = []

    for file_info in files:
        filename = file_info['name']
        filepath = os.path.join(folder_path, *filename.split('/'))
        file_lower = filename.lower()

        entry = {
            'name': filename,
            'ext': file_info.get('ext'),
            'size_kb': file_info.get('size_kb'),
            'type': 'Unknown',
            'summary': '',
            'stats': {}
        }

        try:
            if file_lower.endswith(('.txt', '.md')):
                content = read_text_file(filepath, max_chars=preview_chars * 2)
                entry['type'] = 'Text'
                entry['summary'] = truncate_text(content, preview_chars)

            elif file_lower.endswith('.csv'):
                csv_structured = summarize_csv_structured(
                    filepath,
                    max_rows=csv_sample_rows,
                    max_stat_rows=csv_stat_rows
                )
                entry['type'] = 'CSV'
                entry['summary'] = csv_structured.get('summary', '')
                entry['stats'] = csv_structured.get('stats', {})

            elif file_lower.endswith('.json'):
                json_structured = summarize_json_structured(filepath, max_keys=DEFAULT_JSON_SAMPLE_KEYS)
                entry['type'] = 'JSON'
                entry['summary'] = json_structured.get('summary', '')
                entry['stats'] = json_structured.get('stats', {})

            elif file_lower.endswith('.docx'):
                content = extract_docx_text(filepath, max_chars=DEFAULT_DOCX_MAX_CHARS)
                entry['type'] = 'DOCX'
                entry['summary'] = truncate_text(content, preview_chars)

            elif file_lower.endswith('.pdf'):
                content = extract_pdf_text(
                    filepath,
                    max_pages=DEFAULT_PDF_MAX_PAGES,
                    max_chars=DEFAULT_PDF_MAX_CHARS
                )
                entry['type'] = 'PDF'
                entry['summary'] = truncate_text(content, preview_chars)

            elif file_lower.endswith(('.xlsx', '.xls', '.xlsm', '.xlsb')):
                entry['type'] = 'Excel'
                content = extract_excel_text(filepath, max_rows=excel_max_rows, max_chars=excel_max_chars)
                entry['summary'] = content

            elif file_lower.endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp')):
                entry['type'] = 'Image'
                entry['summary'] = 'Preview: Image content requires OCR.'

            else:
                try:
                    content = read_text_file(filepath, max_chars=preview_chars)
                    entry['type'] = 'Unknown (text preview)'
                    entry['summary'] = truncate_text(content, preview_chars)
                except Exception:
                    entry['type'] = 'Binary/Unsupported'
                    entry['summary'] = 'Preview: Not available.'

        except Exception as e:
            entry['type'] = 'Error'
            entry['summary'] = f"Error: {str(e)}"

        summaries.append(entry)

    return summaries


def format_folder_summary(folder_name, structured_summaries, include_manifest=True):
    if not structured_summaries:
        return f"Folder: {folder_name}\nNo files found."

    lines = [f"Folder: {folder_name}"]

    if include_manifest:
        manifest = [
            {
                'name': entry.get('name'),
                'ext': entry.get('ext'),
                'size_kb': entry.get('size_kb')
            }
            for entry in structured_summaries
        ]
        lines.append(format_file_manifest(manifest))
        lines.append("")

    for entry in structured_summaries:
        lines.append(f"File: {entry.get('name')}")
        lines.append(f"Type: {entry.get('type')}")
        summary = entry.get('summary') or ''
        if summary:
            lines.append(f"Preview:\n{summary}")
        stats = entry.get('stats') or {}
        if stats:
            lines.append(f"Stats: {json.dumps(stats, ensure_ascii=True)}")
        lines.append("")

    return "\n".join(lines).strip()


def read_text_file(path, max_chars=10000):
    """Read plain text file"""
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read(max_chars)
    return content


def summarize_csv(path, max_rows=DEFAULT_CSV_SAMPLE_ROWS, max_stat_rows=DEFAULT_CSV_STAT_ROWS):
    """
    Summarize CSV file with structure, samples, and basic numeric stats.
    """
    structured = summarize_csv_structured(path, max_rows=max_rows, max_stat_rows=max_stat_rows)
    return structured.get('summary', '')


def summarize_csv_structured(path, max_rows=DEFAULT_CSV_SAMPLE_ROWS, max_stat_rows=DEFAULT_CSV_STAT_ROWS):
    """
    Summarize CSV file and return structured stats + summary text.
    """
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            headers = next(reader, [])
            data_rows = []
            row_count = 0

            numeric_stats = []
            for _ in headers:
                numeric_stats.append({
                    'count': 0,
                    'min': None,
                    'max': None
                })

            for row in reader:
                row_count += 1
                if len(data_rows) < max_rows:
                    data_rows.append(row)

                if row_count <= max_stat_rows:
                    for idx, cell in enumerate(row[:len(headers)]):
                        value = cell.strip()
                        if not value:
                            continue
                        try:
                            num = float(value.replace(',', ''))
                        except ValueError:
                            continue

                        stats = numeric_stats[idx]
                        stats['count'] += 1
                        stats['min'] = num if stats['min'] is None else min(stats['min'], num)
                        stats['max'] = num if stats['max'] is None else max(stats['max'], num)

        if not headers:
            return {
                'summary': "Empty CSV file",
                'stats': {'row_count': 0, 'headers': []}
            }

        summary_parts = []
        summary_parts.append("CSV structure:")
        summary_parts.append(f"- Rows: {row_count}")
        summary_parts.append(f"- Columns: {len(headers)}")
        summary_parts.append(f"- Headers: {', '.join(headers)}")

        numeric_lines = []
        numeric_ranges = []
        for idx, stats in enumerate(numeric_stats):
            if stats['count']:
                header_name = headers[idx] if idx < len(headers) else f"Column {idx + 1}"
                numeric_lines.append(
                    f"  - {header_name}: min={stats['min']}, max={stats['max']} (n={stats['count']})"
                )
                numeric_ranges.append({
                    'column': header_name,
                    'min': stats['min'],
                    'max': stats['max'],
                    'count': stats['count']
                })

        if numeric_lines:
            summary_parts.append("Numeric column ranges:")
            summary_parts.extend(numeric_lines)

        sample_rows = []
        if data_rows:
            summary_parts.append(f"Sample rows (first {min(max_rows, len(data_rows))}):")
            for i, row in enumerate(data_rows[:max_rows], 1):
                truncated_row = [cell[:200] + '...' if len(cell) > 200 else cell for cell in row]
                summary_parts.append(f"  Row {i}: {' | '.join(truncated_row)}")
                sample_rows.append(truncated_row)

        return {
            'summary': "\n".join(summary_parts),
            'stats': {
                'row_count': row_count,
                'headers': headers,
                'numeric_ranges': numeric_ranges,
                'sample_rows': sample_rows
            }
        }

    except Exception as e:
        return {
            'summary': f"Error reading CSV: {str(e)}",
            'stats': {'error': str(e)}
        }


def summarize_excel_structured(path, max_rows=DEFAULT_EXCEL_SAMPLE_ROWS, max_stat_rows=DEFAULT_EXCEL_STAT_ROWS):
    """
    Summarize Excel file with per-sheet structure and sample rows.
    """
    try:
        import pandas as pd
    except Exception as e:
        return {
            'summary': f"Excel content extraction not available: {str(e)}",
            'stats': {'error': str(e)}
        }

    try:
        sheets = pd.read_excel(path, sheet_name=None)
    except Exception as e:
        return {
            'summary': f"Error reading Excel: {str(e)}",
            'stats': {'error': str(e)}
        }

    if not sheets:
        return {
            'summary': "Empty Excel file",
            'stats': {'sheets': {}}
        }

    summaries = []
    sheet_stats = {}

    for sheet_name, df in sheets.items():
        headers = [str(col) for col in df.columns.tolist()]
        row_count = len(df.index)
        sample_rows = []

        if row_count:
            sample_df = df.head(max_rows).astype(str)
            sample_rows = sample_df.values.tolist()

        numeric_ranges = []
        if row_count:
            stat_df = df.head(max_stat_rows)
            for col in headers:
                series = pd.to_numeric(stat_df.get(col), errors='coerce')
                series = series.dropna()
                if series.empty:
                    continue
                numeric_ranges.append({
                    'column': col,
                    'min': float(series.min()),
                    'max': float(series.max()),
                    'count': int(series.count())
                })

        summaries.append(f"Sheet: {sheet_name}")
        summaries.append(f"- Rows: {row_count}")
        summaries.append(f"- Columns: {', '.join(headers) if headers else 'None'}")

        if numeric_ranges:
            summaries.append("Numeric column ranges:")
            for rng in numeric_ranges:
                summaries.append(
                    f"  - {rng['column']}: min={rng['min']}, max={rng['max']} (n={rng['count']})"
                )

        if sample_rows:
            summaries.append(f"Sample rows (first {min(max_rows, row_count)}):")
            for idx, row in enumerate(sample_rows, 1):
                truncated_row = [cell[:50] + '...' if len(cell) > 50 else cell for cell in row]
                summaries.append(f"  Row {idx}: {' | '.join(truncated_row)}")

        summaries.append("")

        sheet_stats[sheet_name] = {
            'row_count': row_count,
            'headers': headers,
            'numeric_ranges': numeric_ranges,
            'sample_rows': sample_rows
        }

    return {
        'summary': "\n".join(summaries).strip(),
        'stats': {
            'sheets': sheet_stats
        }
    }


def summarize_json_structured(path, max_keys=DEFAULT_JSON_SAMPLE_KEYS):
    """
    Summarize JSON file with top-level keys and basic shape info.
    """
    try:
        with open(path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        data = json.loads(content)

        stats = {'type': type(data).__name__}
        if isinstance(data, dict):
            keys = list(data.keys())
            stats['keys'] = keys[:max_keys]
            summary = f"JSON object with {len(keys)} keys. Sample keys: {', '.join(keys[:max_keys])}"
        elif isinstance(data, list):
            stats['length'] = len(data)
            sample_keys = []
            for item in data[: min(len(data), 5)]:
                if isinstance(item, dict):
                    for key in item.keys():
                        if key not in sample_keys:
                            sample_keys.append(key)
                        if len(sample_keys) >= max_keys:
                            break
                if len(sample_keys) >= max_keys:
                    break
            stats['sample_keys'] = sample_keys
            summary = f"JSON array with {len(data)} items. Sample keys: {', '.join(sample_keys)}"
        else:
            summary = f"JSON value type: {type(data).__name__}"

        return {'summary': summary, 'stats': stats}
    except Exception as e:
        return {'summary': f"Error reading JSON: {str(e)}", 'stats': {'error': str(e)}}


def extract_docx_text(path, max_chars=DEFAULT_DOCX_MAX_CHARS):
    """
    Extract text from DOCX file

    Args:
        path: Path to DOCX file
        max_chars: Maximum characters to extract

    Returns:
        str: Extracted text
    """
    try:
        doc = Document(path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

            current_length = sum(len(p) for p in paragraphs)
            if current_length > max_chars:
                break

        content = "\n".join(paragraphs)
        return content[:max_chars]

    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def extract_pdf_text(path, max_pages=DEFAULT_PDF_MAX_PAGES, max_chars=DEFAULT_PDF_MAX_CHARS):
    """
    Extract text from PDF file

    Args:
        path: Path to PDF file
        max_pages: Maximum number of pages to process
        max_chars: Maximum characters to extract

    Returns:
        str: Extracted text
    """
    try:
        text_chunks = []
        total_chars = 0

        with fitz.open(path) as pdf:
            total_pages = len(pdf)

            for page_num, page in enumerate(pdf, 1):
                if page_num > max_pages:
                    text_chunks.append(f"\n... (Skipped remaining {total_pages - max_pages} pages)")
                    break

                text = page.get_text()
                if text.strip():
                    page_text = f"[Page {page_num}]\n{text}\n"
                    text_chunks.append(page_text)
                    total_chars += len(page_text)

                    if total_chars > max_chars:
                        text_chunks.append(f"\n... (Content truncated at {max_chars} characters)")
                        break

        result = "".join(text_chunks)
        return result[:max_chars]

    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_excel_text(path, max_rows=DEFAULT_EXCEL_MAX_ROWS, max_chars=DEFAULT_EXCEL_MAX_CHARS):
    """
    Extract text content from Excel file (.xlsx, .xlsm, .xls)

    Args:
        path: Path to Excel file
        max_rows: Maximum rows to extract per worksheet
        max_chars: Maximum characters to extract total

    Returns:
        str: Formatted text representation of Excel content
    """
    try:
        # Load workbook (read-only for better performance)
        wb = load_workbook(path, read_only=True, data_only=True)

        content_parts = []
        total_chars = 0

        # Add workbook summary
        sheet_names = wb.sheetnames
        content_parts.append(f"Excel Workbook: {os.path.basename(path)}")
        content_parts.append(f"Worksheets ({len(sheet_names)}): {', '.join(sheet_names)}")
        content_parts.append("")

        # Process each worksheet
        for sheet_name in sheet_names:
            ws = wb[sheet_name]

            # Get sheet dimensions
            max_row = ws.max_row
            max_col = ws.max_column

            content_parts.append(f"[Worksheet: {sheet_name}]")
            content_parts.append(f"Dimensions: {max_row} rows × {max_col} columns")
            content_parts.append("")

            # Extract column headers (first row)
            headers = []
            if max_row > 0:
                for col in range(1, min(max_col + 1, 200)):  # Limit to 200 columns
                    cell_value = ws.cell(row=1, column=col).value
                    if cell_value is not None:
                        headers.append(str(cell_value))
                    else:
                        headers.append(f"Col{col}")

                if headers:
                    content_parts.append(f"Headers: {' | '.join(headers)}")
                    content_parts.append("")

            # Extract sample data rows
            rows_to_show = min(max_rows, max_row)
            data_rows = max(0, rows_to_show - 1)  # Exclude header row
            if data_rows > 0:
                content_parts.append(
                    f"Sample Data (showing {data_rows} of {max(0, max_row - 1)} data rows):"
                )

                for row_num in range(2, 2 + data_rows):
                    pairs = []
                    for col in range(1, min(max_col + 1, 200)):
                        header = headers[col - 1] if col - 1 < len(headers) else f"Col{col}"
                        cell_value = ws.cell(row=row_num, column=col).value
                        value = str(cell_value)[:200] if cell_value is not None else ""
                        pairs.append(f"{header}={value}")

                    content_parts.append(f"  Row {row_num}: {' | '.join(pairs)}")

                content_parts.append("")

            # Check if we've exceeded max_chars
            current_content = "\n".join(content_parts)
            total_chars = len(current_content)
            if total_chars > max_chars:
                content_parts.append(f"\n... (Content truncated at {max_chars} characters)")
                break

        wb.close()

        result = "\n".join(content_parts)
        return result[:max_chars]

    except Exception as e:
        return f"Error reading Excel file: {str(e)}"


def get_file_stats(folder_path):
    """
    Get statistics about files in a folder

    Args:
        folder_path: Path to folder

    Returns:
        dict: File statistics
    """
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return {'error': 'Folder not found'}

    files = [
        f for f in os.listdir(folder_path)
        if os.path.isfile(os.path.join(folder_path, f))
    ]

    file_types = {}
    total_size = 0

    for filename in files:
        filepath = os.path.join(folder_path, filename)
        ext = os.path.splitext(filename)[1].lower()
        size = os.path.getsize(filepath)

        file_types[ext] = file_types.get(ext, 0) + 1
        total_size += size

    return {
        'total_files': len(files),
        'total_size_mb': total_size / (1024 * 1024),
        'file_types': file_types
    }


def load_kb_files(folder_path, max_chars=8000, recursive=True):
    """
    Load KB files for retrieval (text-focused).
    """
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return []

    files = list_folder_files(folder_path, recursive=recursive)
    results = []

    for file_info in files:
        filename = file_info['name']
        filepath = os.path.join(folder_path, *filename.split('/'))
        file_lower = filename.lower()

        try:
            if file_lower.endswith(('.txt', '.md', '.log')):
                content = read_text_file(filepath, max_chars=max_chars)
            elif file_lower.endswith('.csv'):
                content = summarize_csv(filepath, max_rows=30, max_stat_rows=200)
            elif file_lower.endswith('.pdf'):
                content = extract_pdf_text(filepath, max_pages=10, max_chars=max_chars)
            elif file_lower.endswith('.docx'):
                content = extract_docx_text(filepath, max_chars=max_chars)
            elif file_lower.endswith(('.xlsx', '.xls', '.xlsm', '.xlsb')):
                content = extract_excel_text(filepath, max_rows=DEFAULT_EXCEL_MAX_ROWS, max_chars=max_chars)
            else:
                content = read_text_file(filepath, max_chars=max_chars)

            results.append({
                'name': filename,
                'content': truncate_text(content, max_chars)
            })
        except Exception as e:
            results.append({
                'name': filename,
                'content': f"(Error loading file: {str(e)})"
            })

    return results


def extract_text_from_file(path, max_chars=10000):
    """
    Generic text extractor for any supported file type.
    Returns plain text content.

    Args:
        path: Path to file
        max_chars: Maximum characters to extract

    Returns:
        str: Extracted text content
    """
    if not os.path.exists(path):
        return ""

    file_lower = os.path.basename(path).lower()

    try:
        if file_lower.endswith('.pdf'):
            return extract_pdf_text(path, max_pages=10, max_chars=max_chars)
        elif file_lower.endswith('.docx'):
            return extract_docx_text(path, max_chars=max_chars)
        elif file_lower.endswith(('.xlsx', '.xls', '.xlsm', '.xlsb')):
            return extract_excel_text(path, max_rows=50, max_chars=max_chars)
        elif file_lower.endswith('.csv'):
            return summarize_csv(path, max_rows=30, max_stat_rows=200)
        elif file_lower.endswith(('.txt', '.md', '.log')):
            return read_text_file(path, max_chars=max_chars)
        else:
            # Try reading as text
            return read_text_file(path, max_chars=max_chars)
    except Exception as e:
        return f"Error extracting text from {os.path.basename(path)}: {str(e)}"
