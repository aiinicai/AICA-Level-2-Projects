"""
Universal Multi-Format Ingestion Engine for Forensic Audit & Benford's Law Suite.

Supports:
- Excel (.xlsx, .xls, .xlsm)
- Word Documents (.docx - tabular & structured financial text)
- PDF Documents (.pdf - vector tables & digital text extraction)
- Delimited Text & Ledger Dumps (.csv, .tsv, .psv, .txt, .log, .dat)
- Semi-Structured Data (.json, .jsonl, .xml)
- High-Performance & Database Formats (.parquet, .sqlite, .db)
- Local Files, Folders, and Network Server Paths (UNC \\server\share\... and drive letters).

Features:
- Diagnostic warnings for password-protected or image-only scanned files with actionable suggestions.
- Cryptographic SHA-256 dataset fingerprinting.
- Intelligent financial column auto-mapping (Amount, Date, Vendor, Invoice, Narration).
"""

import os
import re
import io
import json
import sqlite3
import hashlib
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path

import pandas as pd
import numpy as np

# Optional / Specialized parsers
try:
    import docx
except ImportError:
    docx = None

try:
    import pymupdf  # PyMuPDF
except ImportError:
    pymupdf = None

try:
    import pypdf
except ImportError:
    pypdf = None

try:
    from lxml import etree
except ImportError:
    etree = None


# ============================================================================
# COLUMN MAPPING HEURISTICS FOR FORENSIC FINANCIAL DATA
# ============================================================================

COLUMN_ALIASES = {
    'amount': [
        'amount', 'amt', 'txn_amt', 'transaction_amount', 'gross_amount', 'net_amount',
        'invoice_amount', 'bill_amount', 'debit', 'credit', 'payment_amount', 'total',
        'val', 'value', 'dr_amt', 'cr_amt', 'line_total', 'paid_amount', 'voucher_amount',
        'transaction_val', 'trans_amount', 'ledger_balance', 'closing_balance'
    ],
    'date': [
        'date', 'txn_date', 'transaction_date', 'posting_date', 'invoice_date',
        'bill_date', 'voucher_date', 'entry_date', 'val_date', 'value_date',
        'trans_date', 'created_at', 'timestamp', 'doc_date'
    ],
    'vendor': [
        'vendor', 'vendor_name', 'party', 'party_name', 'payee', 'beneficiary',
        'supplier', 'supplier_name', 'merchant', 'entity', 'account_name',
        'customer', 'client', 'remitter', 'contractor', 'creditor', 'debtor'
    ],
    'invoice_no': [
        'invoice', 'invoice_no', 'invoice_num', 'inv_no', 'bill_no', 'voucher_no',
        'voucher_num', 'ref_no', 'reference', 'trans_id', 'transaction_id',
        'cheque_no', 'utr_no', 'utr', 'doc_no', 'document_number'
    ],
    'description': [
        'description', 'narration', 'particulars', 'remarks', 'details',
        'purpose', 'item_name', 'memo', 'line_desc', 'comments', 'notes'
    ]
}


class DataIngestionResult:
    """Encapsulates dataset load status, metadata, records, and diagnostic warnings."""
    def __init__(
        self,
        success: bool,
        file_path: str,
        file_name: str,
        records: List[Dict[str, Any]],
        columns: List[str],
        column_mapping: Dict[str, str],
        dataset_hash: str,
        error_message: Optional[str] = None,
        limitation_warning: Optional[str] = None,
        recommendation: Optional[str] = None
    ):
        self.success = success
        self.file_path = file_path
        self.file_name = file_name
        self.records = records
        self.columns = columns
        self.column_mapping = column_mapping
        self.dataset_hash = dataset_hash
        self.error_message = error_message
        self.limitation_warning = limitation_warning
        self.recommendation = recommendation
        self.row_count = len(records)

    def to_dict(self) -> Dict[str, Any]:
        return {
            "success": self.success,
            "file_path": self.file_path,
            "file_name": self.file_name,
            "row_count": self.row_count,
            "columns": self.columns,
            "column_mapping": self.column_mapping,
            "dataset_hash": self.dataset_hash,
            "error_message": self.error_message,
            "limitation_warning": self.limitation_warning,
            "recommendation": self.recommendation,
            "sample_records": self.records[:50] if self.records else []
        }


# ============================================================================
# UNIVERSAL DATA LOADER CLASS
# ============================================================================

class UniversalDataLoader:
    """Universal parser for multi-format local and network financial files."""

    @staticmethod
    def compute_file_hash(file_bytes: bytes) -> str:
        """Computes SHA-256 hash for dataset fingerprinting."""
        return hashlib.sha256(file_bytes).hexdigest()

    @staticmethod
    def auto_detect_columns(columns: List[str]) -> Dict[str, str]:
        """Maps dataset column headers to standardized forensic financial fields."""
        mapped = {
            'amount': None,
            'date': None,
            'vendor': None,
            'invoice_no': None,
            'description': None
        }

        col_lower_map = {c.lower().strip().replace(' ', '_').replace('-', '_'): c for c in columns}

        for target_field, aliases in COLUMN_ALIASES.items():
            for alias in aliases:
                # Exact match
                if alias in col_lower_map:
                    mapped[target_field] = col_lower_map[alias]
                    break
            # Fuzzy containment match if not found
            if not mapped[target_field]:
                for col_key, original_col in col_lower_map.items():
                    if any(alias in col_key for alias in aliases[:5]):
                        mapped[target_field] = original_col
                        break

        # Fallback for amount if not detected: find first numeric column
        return {k: v for k, v in mapped.items() if v is not None}

    @classmethod
    def load_from_path(cls, file_path_str: str) -> DataIngestionResult:
        """Loads dataset from local file path or network UNC path."""
        path = Path(file_path_str.strip())
        if not path.exists():
            return DataIngestionResult(
                success=False,
                file_path=file_path_str,
                file_name=path.name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash="",
                error_message=f"File path does not exist or network path is unreachable: {file_path_str}",
                recommendation="Please verify the file path or network permissions and ensure read access."
            )

        if path.is_dir():
            return cls.load_from_directory(str(path))

        try:
            with open(path, "rb") as f:
                content = f.read()
            return cls.load_from_bytes(content, path.name, str(path))
        except PermissionError:
            return DataIngestionResult(
                success=False,
                file_path=str(path),
                file_name=path.name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash="",
                error_message="Access Denied / Permission Error reading file.",
                limitation_warning="The file is currently locked by another application (e.g. Excel) or lacks OS read privileges.",
                recommendation="Please close any open instances of the file in Excel or grant read permissions."
            )
        except Exception as e:
            return DataIngestionResult(
                success=False,
                file_path=str(path),
                file_name=path.name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash="",
                error_message=f"Error reading file: {str(e)}"
            )

    @classmethod
    def load_from_directory(cls, dir_path_str: str) -> DataIngestionResult:
        """Scans directory and loads all supported tabular files, merging records."""
        path = Path(dir_path_str)
        supported_exts = {'.csv', '.tsv', '.xlsx', '.xls', '.json', '.parquet', '.txt', '.xml'}
        files = [p for p in path.iterdir() if p.is_file() and p.suffix.lower() in supported_exts]

        if not files:
            return DataIngestionResult(
                success=False,
                file_path=dir_path_str,
                file_name=path.name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash="",
                error_message=f"No supported tabular data files found in directory: {dir_path_str}",
                recommendation="Ensure the directory contains .csv, .xlsx, .json, .parquet, or .xml files."
            )

        # Load first valid file or merge compatible files
        combined_records = []
        combined_columns = []
        file_hashes = []

        for f in files:
            res = cls.load_from_path(str(f))
            if res.success and res.records:
                combined_records.extend(res.records)
                file_hashes.append(res.dataset_hash)
                if not combined_columns:
                    combined_columns = res.columns

        if not combined_records:
            return DataIngestionResult(
                success=False,
                file_path=dir_path_str,
                file_name=path.name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash="",
                error_message="Could not extract tabular records from files in directory.",
                limitation_warning="Files in directory were unreadable or contained unsupported formats."
            )

        agg_hash = hashlib.sha256("".join(file_hashes).encode()).hexdigest()
        col_mapping = cls.auto_detect_columns(combined_columns)

        return DataIngestionResult(
            success=True,
            file_path=dir_path_str,
            file_name=f"Directory Ingest ({len(files)} files)",
            records=combined_records,
            columns=combined_columns,
            column_mapping=col_mapping,
            dataset_hash=agg_hash
        )

    @classmethod
    def load_from_bytes(cls, content: bytes, file_name: str, file_path: str = "") -> DataIngestionResult:
        """Parses byte buffer into structured records based on file extension and signature."""
        dataset_hash = cls.compute_file_hash(content)
        ext = Path(file_name).suffix.lower()

        # 1. CSV / TSV / PSV / TXT
        if ext in ('.csv', '.tsv', '.psv', '.txt', '.log', '.dat'):
            return cls._parse_delimited(content, file_name, file_path, dataset_hash)

        # 2. Excel (.xlsx, .xls, .xlsm)
        elif ext in ('.xlsx', '.xls', '.xlsm'):
            return cls._parse_excel(content, file_name, file_path, dataset_hash, ext)

        # 3. Word Documents (.docx)
        elif ext == '.docx':
            return cls._parse_docx(content, file_name, file_path, dataset_hash)

        # 4. PDF Documents (.pdf)
        elif ext == '.pdf':
            return cls._parse_pdf(content, file_name, file_path, dataset_hash)

        # 5. JSON / JSONL
        elif ext in ('.json', '.jsonl'):
            return cls._parse_json(content, file_name, file_path, dataset_hash)

        # 6. XML
        elif ext == '.xml':
            return cls._parse_xml(content, file_name, file_path, dataset_hash)

        # 7. Parquet
        elif ext == '.parquet':
            return cls._parse_parquet(content, file_name, file_path, dataset_hash)

        # 8. SQLite Database (.sqlite, .db)
        elif ext in ('.sqlite', '.db'):
            return cls._parse_sqlite(content, file_name, file_path, dataset_hash)

        # Unsupported format
        else:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"Unsupported file format: '{ext}'.",
                recommendation="Please provide standard forensic file formats: CSV, Excel (.xlsx/.xls), Word (.docx), PDF, JSON, XML, Parquet, or SQLite."
            )

    # ------------------------------------------------------------------------
    # INDIVIDUAL FORMAT PARSERS
    # ------------------------------------------------------------------------

    @classmethod
    def _parse_delimited(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Parses CSV, TSV, and delimited text with automatic encoding & delimiter detection."""
        encodings = ['utf-8', 'utf-8-sig', 'latin1', 'cp1252', 'iso-8859-1', 'utf-16']
        df = None
        
        for enc in encodings:
            try:
                text_stream = io.BytesIO(content)
                # Try auto-detect separator
                df = pd.read_csv(text_stream, encoding=enc, sep=None, engine='python', on_bad_lines='skip')
                if df is not None and not df.empty:
                    break
            except Exception:
                continue

        if df is None or df.empty:
            # Try fallback with comma or tab
            for sep in [',', '\t', '|', ';']:
                try:
                    df = pd.read_csv(io.BytesIO(content), sep=sep, encoding='latin1', on_bad_lines='skip')
                    if df is not None and not df.empty and len(df.columns) > 1:
                        break
                except Exception:
                    pass

        if df is None or df.empty:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message="Unable to parse delimited text file. Empty or invalid structure.",
                recommendation="Ensure the file has header columns and valid delimiter-separated rows."
            )

        return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)

    @classmethod
    def _parse_excel(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str, ext: str) -> DataIngestionResult:
        """Parses Excel spreadsheets (.xlsx, .xls, .xlsm) with password check."""
        try:
            engine = 'openpyxl' if ext in ('.xlsx', '.xlsm') else 'xlrd'
            df = pd.read_excel(io.BytesIO(content), engine=engine)
            if df is None or df.empty:
                return DataIngestionResult(
                    success=False,
                    file_path=file_path,
                    file_name=file_name,
                    records=[],
                    columns=[],
                    column_mapping={},
                    dataset_hash=dataset_hash,
                    error_message="Excel workbook contains no data rows in the default active sheet.",
                    recommendation="Ensure the primary sheet contains financial transaction records with column headers."
                )
            return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)

        except Exception as e:
            err_str = str(e).lower()
            if 'password' in err_str or 'encrypted' in err_str or 'protected' in err_str:
                return DataIngestionResult(
                    success=False,
                    file_path=file_path,
                    file_name=file_name,
                    records=[],
                    columns=[],
                    column_mapping={},
                    dataset_hash=dataset_hash,
                    error_message="Excel file is Password-Protected or Encrypted.",
                    limitation_warning="The forensic ingestion engine cannot decrypt password-protected spreadsheets without credentials.",
                    recommendation="Please remove the password protection in Excel (File > Info > Protect Workbook > Unprotect) and re-upload."
                )
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"Failed to read Excel file: {str(e)}",
                recommendation="Check if the file is corrupted or save as a standard .xlsx / .csv format."
            )

    @classmethod
    def _parse_docx(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Extracts tables or delimited financial text from Word .docx documents."""
        if not docx:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message="python-docx library is required for Word document parsing."
            )

        try:
            doc = docx.Document(io.BytesIO(content))
            all_rows = []
            headers = []

            # 1. Look for tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    if any(cells):
                        table_rows.append(cells)

                if len(table_rows) >= 2:
                    headers = table_rows[0]
                    for r in table_rows[1:]:
                        if len(r) == len(headers):
                            all_rows.append(dict(zip(headers, r)))
                        else:
                            padded = r + [''] * (len(headers) - len(r))
                            all_rows.append(dict(zip(headers, padded[:len(headers)])))

            # 2. If no table found, look for tab/comma separated paragraphs
            if not all_rows:
                lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                parsed_lines = []
                for line in lines:
                    if '\t' in line:
                        parsed_lines.append(line.split('\t'))
                    elif ',' in line:
                        parsed_lines.append(line.split(','))
                
                if len(parsed_lines) >= 2:
                    headers = parsed_lines[0]
                    for r in parsed_lines[1:]:
                        all_rows.append(dict(zip(headers, r + [''] * (len(headers) - len(r)))))

            if not all_rows:
                return DataIngestionResult(
                    success=False,
                    file_path=file_path,
                    file_name=file_name,
                    records=[],
                    columns=[],
                    column_mapping={},
                    dataset_hash=dataset_hash,
                    error_message="No structured tabular data found in Word document.",
                    limitation_warning="Word document contains paragraphs without structured tables or delimited columns.",
                    recommendation="Convert tabular audit data to Excel (.xlsx) or CSV format for optimal ingestion."
                )

            df = pd.DataFrame(all_rows)
            return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)

        except Exception as e:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"Error reading Word document: {str(e)}"
            )

    @classmethod
    def _parse_pdf(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """
        Extracts structured tables and financial text from PDF.
        Detects encrypted PDFs and scanned image-only PDFs.
        """
        # 1. Check password encryption with pypdf
        if pypdf:
            try:
                reader = pypdf.PdfReader(io.BytesIO(content))
                if reader.is_encrypted:
                    return DataIngestionResult(
                        success=False,
                        file_path=file_path,
                        file_name=file_name,
                        records=[],
                        columns=[],
                        column_mapping={},
                        dataset_hash=dataset_hash,
                        error_message="PDF is Password-Protected.",
                        limitation_warning="The PDF file is encrypted with a password and cannot be read.",
                        recommendation="Please remove the password security from the PDF before uploading."
                    )
            except Exception:
                pass

        # 2. Extract text & tables using PyMuPDF (fitz)
        if pymupdf:
            try:
                doc = pymupdf.open(stream=content, filetype="pdf")
                total_text = ""
                all_table_data = []

                for page in doc:
                    text = page.get_text()
                    total_text += text
                    
                    # Try table extraction if available in pymupdf
                    try:
                        tabs = page.find_tables()
                        for tab in tabs:
                            tab_df = tab.extract()
                            if tab_df and len(tab_df) >= 2:
                                all_table_data.extend(tab_df)
                    except Exception:
                        pass

                # Check if scanned image-only
                if len(total_text.strip()) < 50 and not all_table_data:
                    return DataIngestionResult(
                        success=False,
                        file_path=file_path,
                        file_name=file_name,
                        records=[],
                        columns=[],
                        column_mapping={},
                        dataset_hash=dataset_hash,
                        error_message="PDF contains scanned images without digital text.",
                        limitation_warning="The PDF appears to be a scanned image or photo without selectable digital text or OCR layer.",
                        recommendation="Please run OCR on the PDF document or export the financial data directly from your accounting software as Excel/CSV."
                    )

                if all_table_data:
                    headers = [str(h) if h else f"Col_{i}" for i, h in enumerate(all_table_data[0])]
                    rows = []
                    for r in all_table_data[1:]:
                        if any(r):
                            rows.append(dict(zip(headers, r + [''] * (len(headers) - len(r)))))
                    if rows:
                        df = pd.DataFrame(rows)
                        return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)

                # Fallback: parse lines with regex
                lines = [l.strip() for l in total_text.split('\n') if l.strip()]
                records = []
                for line in lines:
                    # Look for amounts in line
                    amounts = re.findall(r'[\d,]+\.\d{2}|\b\d{3,}\b', line)
                    if amounts:
                        records.append({
                            "Line_Text": line,
                            "Extracted_Amount": amounts[-1].replace(',', '')
                        })

                if records:
                    df = pd.DataFrame(records)
                    return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)

            except Exception as e:
                return DataIngestionResult(
                    success=False,
                    file_path=file_path,
                    file_name=file_name,
                    records=[],
                    columns=[],
                    column_mapping={},
                    dataset_hash=dataset_hash,
                    error_message=f"Failed to extract records from PDF: {str(e)}",
                    recommendation="Export ledger directly to CSV or Excel (.xlsx) for guaranteed precision."
                )

        return DataIngestionResult(
            success=False,
            file_path=file_path,
            file_name=file_name,
            records=[],
            columns=[],
            column_mapping={},
            dataset_hash=dataset_hash,
            error_message="PDF parser library is not configured.",
            recommendation="Please use Excel (.xlsx) or CSV file."
        )

    @classmethod
    def _parse_json(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Parses JSON or JSON Lines (JSONL)."""
        try:
            text = content.decode('utf-8')
        except UnicodeDecodeError:
            text = content.decode('latin1')

        try:
            # Try standard JSON
            data = json.loads(text)
            if isinstance(data, dict):
                # Look for root list inside dictionary
                for k, v in data.items():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                        data = v
                        break
            if isinstance(data, list) and len(data) > 0:
                df = pd.json_normalize(data)
                return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)
        except Exception:
            # Try JSON Lines
            try:
                records = [json.loads(line) for line in text.split('\n') if line.strip()]
                if records:
                    df = pd.json_normalize(records)
                    return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)
            except Exception:
                pass

        return DataIngestionResult(
            success=False,
            file_path=file_path,
            file_name=file_name,
            records=[],
            columns=[],
            column_mapping={},
            dataset_hash=dataset_hash,
            error_message="Invalid JSON / JSONL structure.",
            recommendation="Ensure JSON contains an array of transaction objects."
        )

    @classmethod
    def _parse_xml(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Parses XML financial ledger exports."""
        try:
            df = pd.read_xml(io.BytesIO(content))
            if df is not None and not df.empty:
                return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)
        except Exception as e:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"XML Parsing error: {str(e)}",
                recommendation="Check XML schema or provide standard CSV/Excel format."
            )

    @classmethod
    def _parse_parquet(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Parses columnar Parquet datasets."""
        try:
            df = pd.read_parquet(io.BytesIO(content))
            return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)
        except Exception as e:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"Parquet load error: {str(e)}"
            )

    @classmethod
    def _parse_sqlite(cls, content: bytes, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Parses SQLite database files."""
        try:
            # Write to temp in-memory SQLite
            conn = sqlite3.connect(":memory:")
            # Restore bytes
            conn.deserialize(content)
            cursor = conn.cursor()
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'")
            tables = cursor.fetchall()
            if not tables:
                return DataIngestionResult(
                    success=False,
                    file_path=file_path,
                    file_name=file_name,
                    records=[],
                    columns=[],
                    column_mapping={},
                    dataset_hash=dataset_hash,
                    error_message="SQLite database contains no accessible user tables."
                )
            
            table_name = tables[0][0]
            df = pd.read_sql_query(f"SELECT * FROM [{table_name}]", conn)
            return cls._dataframe_to_result(df, file_name, file_path, dataset_hash)
        except Exception as e:
            return DataIngestionResult(
                success=False,
                file_path=file_path,
                file_name=file_name,
                records=[],
                columns=[],
                column_mapping={},
                dataset_hash=dataset_hash,
                error_message=f"SQLite reading error: {str(e)}"
            )

    # ------------------------------------------------------------------------
    # HELPERS
    # ------------------------------------------------------------------------

    @classmethod
    def _dataframe_to_result(cls, df: pd.DataFrame, file_name: str, file_path: str, dataset_hash: str) -> DataIngestionResult:
        """Converts cleaned pandas DataFrame to DataIngestionResult."""
        # Clean column names (strip whitespace)
        df.columns = [str(c).strip() for c in df.columns]
        
        # Replace NaN with empty string or None for JSON serialization
        records = df.replace({np.nan: None}).to_dict(orient='records')
        columns = list(df.columns)
        mapping = cls.auto_detect_columns(columns)

        return DataIngestionResult(
            success=True,
            file_path=file_path,
            file_name=file_name,
            records=records,
            columns=columns,
            column_mapping=mapping,
            dataset_hash=dataset_hash
        )
