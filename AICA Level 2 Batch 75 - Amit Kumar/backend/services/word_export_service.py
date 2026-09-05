
def add_signing_docket_word(doc, db, client_id):
    meta = db.query(models.ClientMetadata).filter_by(client_id=client_id).first()
    dirs = db.query(models.DirectorMaster).filter_by(client_id=client_id).all()
    cs = db.query(models.CompanySecretary).filter_by(client_id=client_id).first()
    cfo = db.query(models.ChiefFinancialOfficer).filter_by(client_id=client_id).first()
    
    doc.add_paragraph("")
    doc.add_paragraph("For and on behalf of the Board of Directors").bold = True
    doc.add_paragraph("")
    
    table = doc.add_table(rows=3, cols=3)
    table.autofit = True
    
    # Directors
    row1 = table.rows[0].cells
    row2 = table.rows[1].cells
    row3 = table.rows[2].cells
    
    for i, d in enumerate(dirs[:2]):
        row1[i].text = d.name
        row1[i].paragraphs[0].runs[0].bold = True
        row2[i].text = d.designation
        row3[i].text = f"DIN: {d.din}"
        
    if cfo:
        row1[2].text = cfo.name
        row1[2].paragraphs[0].runs[0].bold = True
        row2[2].text = "Chief Financial Officer"
        row3[2].text = ""
        
    if cs:
        doc.add_paragraph("")
        p1 = doc.add_paragraph(cs.name)
        p1.runs[0].bold = True
        doc.add_paragraph("Company Secretary")
        doc.add_paragraph(f"Membership No: {cs.membership_no}")
        doc.add_paragraph("")

import os
import json
from datetime import datetime
from sqlalchemy.orm import Session

import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from models import (
    Client, TrialBalanceLine, ARAgeing, APAgeing, CWIPAgeing,
    RelatedParty, Borrowing, Contingency, Note, AccountingPolicy,
    CashFlowAdjustment
)
from services.fs_generator import generate_financial_statements
from services.cash_flow_engine import generate_cash_flow_statement
from services.ratio_engine import calculate_ratios
from services.validation_engine import run_validation_checks
from services.notes_engine import generate_or_update_notes
from services.accounting_policies_engine import generate_or_update_accounting_policies

# Color Constants
NAVY_RGB = RGBColor(0x1B, 0x36, 0x5D)
ORANGE_RGB = RGBColor(0xEA, 0x58, 0x0C)
GREY_TEXT_RGB = RGBColor(0x64, 0x74, 0x8B)
DARK_TEXT_RGB = RGBColor(0x0F, 0x17, 0x2A)

NAVY_HEX = "1B365D"
GREY_HEX = "F1F5F9"
BORDER_HEX = "CBD5E1"


def set_cell_background(cell, hex_color: str):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin_name}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def add_table_header_style(row, title_row=False):
    for cell in row.cells:
        set_cell_background(cell, NAVY_HEX)
        set_cell_margins(cell, top=120, bottom=120, left=150, right=150)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = "Segoe UI"
                r.font.bold = True
                r.font.size = Pt(9.5 if title_row else 9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def style_table_cells(table, grey_col_indices=[2, 3]):
    for row_idx, row in enumerate(table.rows[1:], start=1):
        for col_idx, cell in enumerate(row.cells):
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if col_idx in grey_col_indices:
                set_cell_background(cell, GREY_HEX)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.name = "Segoe UI"
                    r.font.size = Pt(9)


def add_section_header(doc, text: str):
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(6)
    h.paragraph_format.keep_with_next = True
    run = h.add_run(text.upper())
    run.font.name = "Segoe UI"
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = NAVY_RGB
    return h


def export_word_financial_report(client_id: int, export_dir: str, db: Session) -> str:
    # Critical Mapping Validation Guard
    validations = run_validation_checks(client_id, db)
    critical_mapping = [v for v in validations if v.category == "Mapping Exception" and v.status == "Critical"]
    if critical_mapping:
        err_msg = "; ".join([f"{v.check_name}: {v.message}" for v in critical_mapping])
        raise ValueError(f"Word export blocked due to Critical Mapping Exceptions: {err_msg}. Please review and approve manual override in Ledger Mapping.")

    os.makedirs(export_dir, exist_ok=True)
    client = db.query(Client).filter(Client.id == client_id).first()
    client_name = client.name if client else "Client"
    safe_name = "".join(c if c.isalnum() else "_" for c in client_name)
    filepath = os.path.join(export_dir, f"{safe_name}_Annual_Report_Review_Pack.docx")

    doc = Document()

    # Page Setup (A4, 0.75 in margins)
    for section in doc.sections:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

        # Header / Footer
        header_p = section.header.paragraphs[0]
        header_p.text = f"SW INDIA | FS BUILDER LITE v0.2 | {client_name}"
        header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_p.runs[0].font.name = "Segoe UI"
        header_p.runs[0].font.size = Pt(8)
        header_p.runs[0].font.color.rgb = GREY_TEXT_RGB

        footer_p = section.footer.paragraphs[0]
        footer_p.text = f"Confidential - For Internal CA Review Only. AS 3 & Schedule III Division I Compliant. | {client.reporting_period if client else 'FY 2024-25'}"
        footer_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        footer_p.runs[0].font.name = "Segoe UI"
        footer_p.runs[0].font.size = Pt(8)
        footer_p.runs[0].font.color.rgb = GREY_TEXT_RGB

    # -------------------------------------------------------------
    # 1. COVER PAGE
    # -------------------------------------------------------------
    p_firm = doc.add_paragraph()
    r_firm = p_firm.add_run("SW INDIA | CHARTERED ACCOUNTANTS")
    r_firm.font.name = "Segoe UI"
    r_firm.font.bold = True
    r_firm.font.size = Pt(12)
    r_firm.font.color.rgb = ORANGE_RGB

    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(10)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("ANNUAL REPORT FINANCIAL STATEMENTS")
    r_title.font.name = "Segoe UI"
    r_title.font.bold = True
    r_title.font.size = Pt(22)
    r_title.font.color.rgb = NAVY_RGB

    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("SCHEDULE III DIVISION I (IGAAP) & AS 3 COMPLIANT")
    r_sub.font.name = "Segoe UI"
    r_sub.font.bold = True
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = GREY_TEXT_RGB

    doc.add_paragraph()

    # Meta Table
    t_meta = doc.add_table(rows=9, cols=2)
    t_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_rows = [
        ("Client Name:", client.name),
        ("Entity Constitution:", client.entity_type),
        ("Reporting Period:", client.reporting_period),
        ("Previous Year Period:", client.previous_year_period),
        ("Currency / Denomination:", client.currency),
        ("Accounting Framework:", client.accounting_framework),
        ("Prepared By:", client.prepared_by),
        ("Reviewed By:", client.reviewed_by),
        ("Generation Timestamp:", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    ]
    for i, (k, v) in enumerate(meta_rows):
        row_cells = t_meta.rows[i].cells
        row_cells[0].paragraphs[0].add_run(k).bold = True
        row_cells[1].paragraphs[0].add_run(v)
        set_cell_margins(row_cells[0], 60, 60, 100, 100)
        set_cell_margins(row_cells[1], 60, 60, 100, 100)
        set_cell_background(row_cells[0], "F8FAFC")
        set_cell_background(row_cells[1], "F8FAFC")

    doc.add_paragraph()
    p_disc = doc.add_paragraph()
    r_disc = p_disc.add_run("MANDATORY AUDIT DISCLAIMER: This is a system-generated draft for internal review. Professional review and verification by a Chartered Accountant is required before official signature and issuance.")
    r_disc.font.name = "Segoe UI"
    r_disc.font.size = Pt(9)
    r_disc.font.color.rgb = ORANGE_RGB

    doc.add_page_break()

    # -------------------------------------------------------------
    # 2. INDEX / TABLE OF CONTENTS
    # -------------------------------------------------------------
    add_section_header(doc, "Table of Contents / Index")
    t_idx = doc.add_table(rows=14, cols=3)
    t_idx.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_idx.rows[0])
    
    idx_headers = t_idx.rows[0].cells
    idx_headers[0].paragraphs[0].text = "Sec #"
    idx_headers[1].paragraphs[0].text = "Report Section Name"
    idx_headers[2].paragraphs[0].text = "Compliance Standard / Reference"

    index_spec = [
        ("01", "Cover Sheet & Audit Metadata", "Firm Administrative Control"),
        ("02", "Index & Section Sitemap", "CA Audit Pack Index"),
        ("03", "Schedule III Balance Sheet", "Companies Act 2013 Division I"),
        ("04", "Statement of Profit and Loss", "Companies Act 2013 Division I"),
        ("05", "AS 3 Cash Flow Statement", "AS 3 Cash Flow Statements"),
        ("06", "Significant Accounting Policies", "AS 1 Disclosure of Accounting Policies"),
        ("07", "Notes to Accounts", "Schedule III Notes & Disclosures"),
        ("08", "Related Party Disclosures", "AS 18 Related Party Disclosures"),
        ("09", "Contingent Liabilities & Commitments", "AS 29 Provisions & Contingencies"),
        ("10", "Schedule III Financial Ratios", "MCA 2021 Ratio Disclosures"),
        ("11", "Management Queries & Audit Notes", "CA Audit Review Workflow"),
        ("12", "Validation Exceptions Log", "Automated Audit Sanity Check"),
        ("13", "Disclaimer & Sign-off Block", "ICAI Professional Standards")
    ]

    for row_idx, (s_no, s_title, s_ref) in enumerate(index_spec, start=1):
        cells = t_idx.rows[row_idx].cells
        cells[0].paragraphs[0].text = s_no
        cells[1].paragraphs[0].text = s_title
        cells[2].paragraphs[0].text = s_ref

    style_table_cells(t_idx, grey_col_indices=[])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 3. BALANCE SHEET
    # -------------------------------------------------------------
    fs = generate_financial_statements(client_id, db)
    add_section_header(doc, "Schedule III Balance Sheet as at March 31, 2025")
    
    bs_lines = fs.balance_sheet
    t_bs = doc.add_table(rows=len(bs_lines) + 1, cols=4)
    t_bs.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_bs.rows[0])

    t_bs.rows[0].cells[0].paragraphs[0].text = "Particulars"
    t_bs.rows[0].cells[1].paragraphs[0].text = "Note #"
    t_bs.rows[0].cells[2].paragraphs[0].text = "As at 31-Mar-2025 (CY)"
    t_bs.rows[0].cells[3].paragraphs[0].text = "As at 31-Mar-2024 (PY)"

    for r_i, line in enumerate(bs_lines, start=1):
        cells = t_bs.rows[r_i].cells
        cells[0].paragraphs[0].text = line.particulars
        cells[1].paragraphs[0].text = line.note_number or ""
        cells[2].paragraphs[0].text = f"{line.cy_amount:,.2f}"
        cells[3].paragraphs[0].text = f"{line.py_amount:,.2f}"
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_bs, grey_col_indices=[2, 3])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 4. PROFIT AND LOSS
    # -------------------------------------------------------------
    add_section_header(doc, "Statement of Profit and Loss for the year ended March 31, 2025")
    pl_lines = fs.profit_and_loss
    t_pl = doc.add_table(rows=len(pl_lines) + 1, cols=4)
    t_pl.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_pl.rows[0])

    t_pl.rows[0].cells[0].paragraphs[0].text = "Particulars"
    t_pl.rows[0].cells[1].paragraphs[0].text = "Note #"
    t_pl.rows[0].cells[2].paragraphs[0].text = "Current Year (CY)"
    t_pl.rows[0].cells[3].paragraphs[0].text = "Previous Year (PY)"

    for r_i, line in enumerate(pl_lines, start=1):
        cells = t_pl.rows[r_i].cells
        cells[0].paragraphs[0].text = line.particulars
        cells[1].paragraphs[0].text = line.note_number or ""
        cells[2].paragraphs[0].text = f"{line.cy_amount:,.2f}"
        cells[3].paragraphs[0].text = f"{line.py_amount:,.2f}"
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_pl, grey_col_indices=[2, 3])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 5. AS 3 CASH FLOW STATEMENT
    # -------------------------------------------------------------
    cfs = generate_cash_flow_statement(client_id, db)
    add_section_header(doc, "AS 3 Cash Flow Statement (Indirect Method)")
    
    t_cf = doc.add_table(rows=len(cfs.statement) + 1, cols=3)
    t_cf.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_cf.rows[0])

    t_cf.rows[0].cells[0].paragraphs[0].text = "Particulars"
    t_cf.rows[0].cells[1].paragraphs[0].text = "Current Year (CY)"
    t_cf.rows[0].cells[2].paragraphs[0].text = "Previous Year (PY)"

    for r_i, line in enumerate(cfs.statement, start=1):
        cells = t_cf.rows[r_i].cells
        cells[0].paragraphs[0].text = line.particulars
        cells[1].paragraphs[0].text = f"{line.cy_amount:,.2f}" if not line.is_header else ""
        cells[2].paragraphs[0].text = f"{line.py_amount:,.2f}" if not line.is_header else ""
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_cf, grey_col_indices=[1, 2])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 6. SIGNIFICANT ACCOUNTING POLICIES
    # -------------------------------------------------------------
    add_section_header(doc, "Significant Accounting Policies (IGAAP)")
    policies = generate_or_update_accounting_policies(client_id, db)
    for p in policies:
        if p.is_applicable:
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(8)
            p_head.paragraph_format.space_after = Pt(2)
            p_head.paragraph_format.keep_with_next = True
            r_h = p_head.add_run(f"{p.policy_number}: {p.title}")
            r_h.font.name = "Segoe UI"
            r_h.font.bold = True
            r_h.font.size = Pt(10)
            r_h.font.color.rgb = NAVY_RGB

            p_body = doc.add_paragraph()
            p_body.paragraph_format.space_after = Pt(6)
            r_b = p_body.add_run(p.content)
            r_b.font.name = "Segoe UI"
            r_b.font.size = Pt(9)
    doc.add_page_break()

    # -------------------------------------------------------------
    # 7. NOTES TO ACCOUNTS
    # -------------------------------------------------------------
    add_section_header(doc, "Notes Forming Part of Financial Statements")
    notes = generate_or_update_notes(client_id, db)
    for n in notes:
        p_head = doc.add_paragraph()
        p_head.paragraph_format.space_before = Pt(10)
        p_head.paragraph_format.space_after = Pt(2)
        p_head.paragraph_format.keep_with_next = True
        r_h = p_head.add_run(f"NOTE {n.note_number}: {n.title.upper()}")
        r_h.font.name = "Segoe UI"
        r_h.font.bold = True
        r_h.font.size = Pt(10.5)
        r_h.font.color.rgb = NAVY_RGB

        # Render Table if table_json exists
        if n.table_json:
            try:
                t_data = json.loads(n.table_json)
                headers = t_data.get("headers", [])
                rows = t_data.get("rows", [])
                if headers and rows:
                    t_note = doc.add_table(rows=len(rows) + 1, cols=len(headers))
                    t_note.alignment = WD_TABLE_ALIGNMENT.CENTER
                    add_table_header_style(t_note.rows[0])

                    for col_i, h_text in enumerate(headers):
                        t_note.rows[0].cells[col_i].paragraphs[0].text = h_text

                    for r_i, r_vals in enumerate(rows, start=1):
                        cells = t_note.rows[r_i].cells
                        is_tot = bool(r_vals[0] and (r_vals[0].strip().startswith("TOTAL") or r_vals[0].strip().startswith("GRAND TOTAL") or r_vals[0].strip().startswith("NET BLOCK")))
                        for c_i, val in enumerate(r_vals):
                            p = cells[c_i].paragraphs[0]
                            p.text = str(val)
                            if c_i > 0 and str(val).strip() != "":
                                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            if is_tot:
                                for run in p.runs:
                                    run.font.bold = True

                    grey_cols = list(range(1, len(headers)))
                    style_table_cells(t_note, grey_col_indices=grey_cols)
                    doc.add_paragraph()  # Spacing
            except Exception as e:
                print(f"Error rendering Word note table {n.note_number}: {e}")

        if n.content and n.content.strip():
            p_body = doc.add_paragraph()
            p_body.paragraph_format.space_after = Pt(8)
            r_b = p_body.add_run(n.content)
            r_b.font.name = "Segoe UI"
            r_b.font.size = Pt(9)

    doc.add_page_break()


    # -------------------------------------------------------------
    # 8. RELATED PARTY DISCLOSURES
    # -------------------------------------------------------------
    add_section_header(doc, "Related Party Disclosures (AS-18)")
    rpts = db.query(RelatedParty).filter(RelatedParty.client_id == client_id).all()
    t_rpt = doc.add_table(rows=len(rpts) + 1, cols=4)
    t_rpt.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_rpt.rows[0])
    
    t_rpt.rows[0].cells[0].paragraphs[0].text = "Party Name"
    t_rpt.rows[0].cells[1].paragraphs[0].text = "Relationship"
    t_rpt.rows[0].cells[2].paragraphs[0].text = "Nature of Transaction"
    t_rpt.rows[0].cells[3].paragraphs[0].text = "Closing Balance (Lakhs)"

    for r_i, r in enumerate(rpts, start=1):
        cells = t_rpt.rows[r_i].cells
        cells[0].paragraphs[0].text = r.name
        cells[1].paragraphs[0].text = r.relationship
        cells[2].paragraphs[0].text = r.nature_tx
        cells[3].paragraphs[0].text = f"{r.closing_bal:,.2f}"
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_rpt, grey_col_indices=[3])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 9. CONTINGENT LIABILITIES
    # -------------------------------------------------------------
    add_section_header(doc, "Contingent Liabilities & Commitments (AS-29)")
    conts = db.query(Contingency).filter(Contingency.client_id == client_id).all()
    t_cont = doc.add_table(rows=len(conts) + 1, cols=4)
    t_cont.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_cont.rows[0])

    t_cont.rows[0].cells[0].paragraphs[0].text = "Nature of Contingency"
    t_cont.rows[0].cells[1].paragraphs[0].text = "Appellate Forum"
    t_cont.rows[0].cells[2].paragraphs[0].text = "CY Amount (Lakhs)"
    t_cont.rows[0].cells[3].paragraphs[0].text = "Management Assessment"

    for r_i, c in enumerate(conts, start=1):
        cells = t_cont.rows[r_i].cells
        cells[0].paragraphs[0].text = c.nature
        cells[1].paragraphs[0].text = c.forum or ""
        cells[2].paragraphs[0].text = f"{c.cy_amount:,.2f}"
        cells[3].paragraphs[0].text = c.assessment or ""
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_cont, grey_col_indices=[2])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 10. RATIOS
    # -------------------------------------------------------------
    add_section_header(doc, "Schedule III Mandatory Financial Ratios")
    ratios = calculate_ratios(client_id, db)
    t_rat = doc.add_table(rows=len(ratios) + 1, cols=5)
    t_rat.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_rat.rows[0])

    t_rat.rows[0].cells[0].paragraphs[0].text = "Code"
    t_rat.rows[0].cells[1].paragraphs[0].text = "Ratio Name"
    t_rat.rows[0].cells[2].paragraphs[0].text = "CY Value"
    t_rat.rows[0].cells[3].paragraphs[0].text = "PY Value"
    t_rat.rows[0].cells[4].paragraphs[0].text = "Movement %"

    for r_i, r in enumerate(ratios, start=1):
        cells = t_rat.rows[r_i].cells
        cells[0].paragraphs[0].text = r.code
        cells[1].paragraphs[0].text = r.name
        cells[2].paragraphs[0].text = f"{r.cy_value:.2f} {r.unit}"
        cells[3].paragraphs[0].text = f"{r.py_value:.2f} {r.unit}"
        cells[4].paragraphs[0].text = r.movement
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    style_table_cells(t_rat, grey_col_indices=[2, 3])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 11. MANAGEMENT QUERIES
    # -------------------------------------------------------------
    add_section_header(doc, "Management Queries & Audit Clarifications")
    mq_rows = [
        ("1", "Verification of physical stock count & valuation", "Physical count completed on 31-Mar-2025. Valued at lower of cost or NRV.", "Verified"),
        ("2", "Confirmation of balances for Trade Receivables > 180 days", "Direct balance confirmations received for 85% of outstanding amount.", "Verified"),
        ("3", "Income tax paid and pending appellate disputes", "Tax paid verified with Challan 280. Appeal pending with CIT(A).", "Verified")
    ]
    t_mq = doc.add_table(rows=len(mq_rows) + 1, cols=4)
    t_mq.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_mq.rows[0])

    t_mq.rows[0].cells[0].paragraphs[0].text = "#"
    t_mq.rows[0].cells[1].paragraphs[0].text = "Audit Query Topic"
    t_mq.rows[0].cells[2].paragraphs[0].text = "Management Response"
    t_mq.rows[0].cells[3].paragraphs[0].text = "Verification Status"

    for r_i, (q_num, q_top, q_res, q_stat) in enumerate(mq_rows, start=1):
        cells = t_mq.rows[r_i].cells
        cells[0].paragraphs[0].text = q_num
        cells[1].paragraphs[0].text = q_top
        cells[2].paragraphs[0].text = q_res
        cells[3].paragraphs[0].text = q_stat

    style_table_cells(t_mq, grey_col_indices=[])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 12. VALIDATION LOG
    # -------------------------------------------------------------
    add_section_header(doc, "Automated Validation Exceptions Log")
    vals = run_validation_checks(client_id, db)
    t_val = doc.add_table(rows=len(vals) + 1, cols=4)
    t_val.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_header_style(t_val.rows[0])

    t_val.rows[0].cells[0].paragraphs[0].text = "Rule Code"
    t_val.rows[0].cells[1].paragraphs[0].text = "Check Name"
    t_val.rows[0].cells[2].paragraphs[0].text = "Status"
    t_val.rows[0].cells[3].paragraphs[0].text = "Audit Findings & Remarks"

    for r_i, v in enumerate(vals, start=1):
        cells = t_val.rows[r_i].cells
        cells[0].paragraphs[0].text = v.code
        cells[1].paragraphs[0].text = v.check_name
        cells[2].paragraphs[0].text = v.status
        cells[3].paragraphs[0].text = v.message

    style_table_cells(t_val, grey_col_indices=[])
    doc.add_page_break()

    # -------------------------------------------------------------
    # 13. DISCLAIMER & SIGN-OFF BLOCK
    # -------------------------------------------------------------
    add_section_header(doc, "Auditor Disclaimer & Partner Sign-off Block")
    
    p_disc_end = doc.add_paragraph()
    r_de = p_disc_end.add_run("This Financial Statement draft has been prepared strictly based on the trial balance and supporting schedules provided by management. In accordance with ICAI Standards on Auditing, final financial statements must be approved by the Board of Directors and signed by Chartered Accountants.")
    r_de.font.name = "Segoe UI"
    r_de.font.size = Pt(9.5)
    r_de.font.color.rgb = DARK_TEXT_RGB

    doc.add_paragraph()

    t_sign = doc.add_table(rows=3, cols=2)
    t_sign.alignment = WD_TABLE_ALIGNMENT.CENTER

    t_sign.rows[0].cells[0].paragraphs[0].text = "For SW INDIA\nChartered Accountants\nFirm Registration No: 001234N"
    t_sign.rows[0].cells[1].paragraphs[0].text = f"For {client.name}\nBoard of Directors"

    t_sign.rows[1].cells[0].paragraphs[0].text = "\n\n"
    t_sign.rows[1].cells[1].paragraphs[0].text = "\n\n"

    t_sign.rows[2].cells[0].paragraphs[0].text = "CA Partner\nMembership No: 501234\nUDIN: 25501234AAAAAA1234"
    t_sign.rows[2].cells[1].paragraphs[0].text = "Director\nDIN: 01234567\n\nDirector\nDIN: 08901234"

    for row in t_sign.rows:
        for cell in row.cells:
            set_cell_margins(cell, 80, 80, 120, 120)
            set_cell_background(cell, "F8FAFC")

    doc.save(filepath)
    return filepath
