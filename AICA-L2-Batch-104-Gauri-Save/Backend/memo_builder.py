"""
Generates a Word summary memo for a batch invoice reconciliation run.

CHANGE: added a "Revenue Support Schedules Referenced" section so the memo
shows, per month, the actual Operating Revenue and Extraordinary/Non-Operating
Income figures the reconciliation was run against — without this, findings
like "Fee Computed on Total Revenue" reference numbers the reader can't see
anywhere else in the document. revenue_schedules is optional and backward
compatible: existing callers that don't pass it still get a working memo,
just without this section.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
from reconciliation import InvoiceReconciliation
from extraction import ExtractedDoc

NAVY = RGBColor(0x0B, 0x1F, 0x3A)
BRASS = RGBColor(0x9C, 0x7A, 0x2E)
RED = RGBColor(0xA3, 0x24, 0x1F)
AMBER = RGBColor(0xA8, 0x71, 0x0D)
GREEN = RGBColor(0x1E, 0x7A, 0x3A)
GREY = RGBColor(0x66, 0x66, 0x66)

SEVERITY_COLOR = {"CRITICAL": RED, "FLAG": AMBER, "REVIEW": AMBER, "OK": GREEN}


def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_col_widths(table, widths_inches):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_inches):
            cell.width = Inches(w)


def _fmt_amount(n):
    if n is None:
        return "\u2013"
    return f"{n:,.0f}"


def build_batch_memo(agreement: ExtractedDoc, reconciliations: list[InvoiceReconciliation],
                      narrative: dict, output_path: str,
                      revenue_schedules: "list[ExtractedDoc] | None" = None):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # --- Header --------------------------------------------------------------
    title = doc.add_paragraph()
    run = title.add_run("INTERCOMPANY INVOICE CONSISTENCY REVIEW")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Internal Working Paper — Batch Reconciliation Summary")
    sub_run.italic = True
    sub_run.font.color.rgb = BRASS

    doc.add_paragraph(f"Generated: {date.today().isoformat()}")
    doc.add_paragraph(
        f"Reference agreement: {agreement.filename} "
        f"({agreement.provider_name} — {agreement.recipient_name})"
    )
    if agreement.fee_percentage is not None:
        doc.add_paragraph(
            f"Fee formula per agreement: {agreement.fee_percentage:.0f}% of Operating Revenue "
            f"(Extraordinary/Non-Operating Income excluded per Clause 2)."
        )
    doc.add_paragraph(f"Invoices reviewed: {len(reconciliations)}")

    doc.add_paragraph()

    # --- Status Summary Table --------------------------------------------------
    doc.add_heading("1. Batch Summary", level=1)
    status_counts = {}
    for r in reconciliations:
        status_counts[r.status] = status_counts.get(r.status, 0) + 1

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    hdr = summary_table.rows[0].cells
    hdr[0].text = "Status"
    hdr[1].text = "Count"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True
        _shade_cell(c, "0B1F3A")
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for status in ["CRITICAL", "FLAG", "REVIEW", "OK"]:
        if status in status_counts:
            row = summary_table.add_row()
            row.cells[0].text = status
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = SEVERITY_COLOR.get(status)
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = str(status_counts[status])
    _set_col_widths(summary_table, [2.5, 1.5])

    doc.add_paragraph()

    # --- Revenue Support Schedules Referenced -----------------------------------
    # Shows the actual monthly figures the reconciliation was computed against,
    # so any finding referencing "Operating Revenue" or "Total Revenue" for a
    # given month is traceable to a number the reader can see in this memo.
    if revenue_schedules:
        doc.add_heading("2. Revenue Support Schedules Referenced", level=1)
        doc.add_paragraph(
            "The figures below were certified by finance in monthly Revenue Support Schedules "
            "supplied separately from the agreement, and are the basis for the expected-fee "
            "calculations in Section 5."
        )
        rev_table = doc.add_table(rows=1, cols=4)
        rev_table.style = "Table Grid"
        hdr = rev_table.rows[0].cells
        for i, label in enumerate(["Month", "Operating Revenue", "Extraordinary Income", "Total Revenue"]):
            hdr[i].text = label
            hdr[i].paragraphs[0].runs[0].bold = True
            _shade_cell(hdr[i], "0B1F3A")
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        from datetime import datetime as _dt

        def _month_sort_key(s):
            try:
                return _dt.strptime(s.schedule_month, "%B %Y")
            except (ValueError, TypeError):
                return _dt.max  # unparseable months sort last rather than crashing

        sorted_schedules = sorted(
            [s for s in revenue_schedules if s.schedule_month],
            key=_month_sort_key
        )
        for s in sorted_schedules:
            row = rev_table.add_row()
            row.cells[0].text = s.schedule_month
            row.cells[1].text = _fmt_amount(s.operating_revenue)
            row.cells[2].text = _fmt_amount(s.extraordinary_income)
            row.cells[3].text = _fmt_amount(s.total_revenue)
            if s.extraordinary_income:
                row.cells[2].paragraphs[0].runs[0].font.color.rgb = AMBER
                row.cells[2].paragraphs[0].runs[0].bold = True
        _set_col_widths(rev_table, [1.5, 1.8, 1.8, 1.5])
        for row in rev_table.rows[1:]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run_ in para.runs:
                        run_.font.size = Pt(9)
        doc.add_paragraph()
        section_offset = 1
    else:
        section_offset = 0

    # --- Executive Summary -------------------------------------------------------
    doc.add_heading(f"{2 + section_offset}. Executive Summary", level=1)
    doc.add_paragraph(narrative.get("executive_summary", "Not available."))

    # --- Prioritized Items ------------------------------------------------------
    doc.add_heading(f"{3 + section_offset}. Prioritized Items", level=1)
    for i, item in enumerate(narrative.get("prioritized_items", []), 1):
        doc.add_paragraph(f"{i}. {item}")

    doc.add_paragraph()

    # --- Per-Invoice Detailed Findings --------------------------------------------
    doc.add_heading(f"{4 + section_offset}. Per-Invoice Detailed Findings", level=1)
    for r in reconciliations:
        p = doc.add_paragraph()
        run = p.add_run(f"{r.invoice.filename}  —  {r.status}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = SEVERITY_COLOR.get(r.status, RGBColor(0, 0, 0))

        method_note = doc.add_paragraph()
        method_run = method_note.add_run(
            f"Extraction method: {r.invoice.extraction_method.upper()}"
            + (" — OCR-sourced, verify flagged fields manually." if r.invoice.extraction_method == "ocr" else "")
        )
        method_run.italic = True
        method_run.font.size = Pt(8.5)
        method_run.font.color.rgb = GREY

        findings_table = doc.add_table(rows=0, cols=2)
        findings_table.style = "Table Grid"
        for f in r.findings:
            row = findings_table.add_row()
            row.cells[0].text = f.severity
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = SEVERITY_COLOR.get(f.severity, RGBColor(0,0,0))
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = f"{f.check}: {f.detail}"
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run_ in para.runs:
                        run_.font.size = Pt(9)
        _set_col_widths(findings_table, [0.9, 5.5])
        doc.add_paragraph()

    # --- Recommended Next Steps --------------------------------------------------
    doc.add_heading(f"{5 + section_offset}. Recommended Next Steps", level=1)
    for step in narrative.get("recommended_next_steps", []):
        doc.add_paragraph(step, style="List Bullet")

    # --- Footer / Disclaimer --------------------------------------------------
    doc.add_paragraph()
    disclaimer = doc.add_paragraph()
    disclaimer_run = disclaimer.add_run(
        "Disclaimer: This memo is generated by an AI-assisted document reconciliation tool for "
        "internal review purposes. Fields extracted via OCR from scanned or image documents carry "
        "inherent uncertainty and are flagged accordingly — verify these manually before relying on "
        "them. This tool does not substitute professional judgment or a full audit of underlying "
        "documentation."
    )
    disclaimer_run.italic = True
    disclaimer_run.font.size = Pt(8.5)
    disclaimer_run.font.color.rgb = GREY

    doc.save(output_path)
    return output_path
