"""
app.py â€” R K Muley & Co | Tax Notice Litigation Assistant v8.1
Main Streamlit entry point.

Run: streamlit run app.py

v8.1 changes: dead imports removed, inner duplicate imports resolved,
preliminary block is now user-confirmed, SDK updated to google.genai v1.x.
See README.md for the full change log.
"""
from __future__ import annotations

# â”€â”€ Path fix â€” must be first, before any local imports â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
# Guarantees data/, engines/, services/ are importable regardless of how
# Streamlit is launched (Windows, Linux, any working directory).
import sys
import os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
# â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€

import io
import base64
import html
import logging
import re
import uuid
from datetime import date, datetime, timedelta
from pathlib import Path

import streamlit as st

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    Document = None
    DOCX_AVAILABLE = False

# â”€â”€ Internal modules â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
from config import (
    APP_NAME, APP_VERSION, FIRM_NAME, FIRM_SUBTITLE,
    APP_STEPS, FEATURES, PORTAL_TEXTBOX_LIMIT, MAX_DRAFT_CHARS,
    GEMINI_MODELS,
)
from database import (
    DatabaseMigrationEngine, write_audit_trail,
    persist_notice_store, get_system_health, smoke_test,
    save_session, load_session,
)
from security import require_auth, rbac_check, logout
from data.case_laws import get_relevant_case_laws, VERIFIED_CASE_LAWS
from data.prompts import (
    EXTRACTION_PROMPT, DRAFTING_PROMPT, COVER_NOTE_PROMPT,
    FORM35_APPEAL_PROMPT, ITAT_NOTICE_RESPONSE_PROMPT,
    PENALTY_PROCEEDINGS_PROMPT,
)
from services.gemini_service import call_gemini, get_api_key, APICallError
from services.knowledge_vault import CAKnowledgeVault
from engines.hallucination_guard import HallucinationGuard
from engines.draft_risk_checker import DraftRiskChecker, detect_third_person, detect_markdown_residue
from engines.reply_scorer import ReplySuccessScorer, NoticeProbabilityPredictor

# â”€â”€ Conditional imports â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False


logger = logging.getLogger("RKMuley.App.v9")

# â”€â”€ Module-level singletons (created once, not in every rerun) â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
_db_engine = DatabaseMigrationEngine()
_vault     = CAKnowledgeVault()
LOGO_PATH  = Path(__file__).resolve().parent / "assets" / "rk_muley_logo.png"


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# PDF EXTRACTION
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def extract_text_from_pdf(pdf_bytes: bytes) -> tuple[str, str]:
    """Returns (text, method_name). Raises ValueError if all methods fail."""
    # Try pypdf
    if PYPDF_AVAILABLE:
        try:
            reader = pypdf.PdfReader(io.BytesIO(pdf_bytes))
            text = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            if text:
                return text, "pypdf"
        except Exception as exc:
            logger.debug("pypdf failed: %s", exc)

    # Try pdfplumber
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                text = "\n".join(p.extract_text() or "" for p in pdf.pages).strip()
            if text:
                return text, "pdfplumber"
        except Exception as exc:
            logger.debug("pdfplumber failed: %s", exc)

    raise ValueError(
        "All PDF extraction methods failed. "
        "The PDF may be scanned/image-based. Please use OCR (Adobe Acrobat / Google Drive) before uploading."
    )


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TEXT UTILITIES
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def clean_markdown_from_draft(text: str) -> str:
    """Strip markdown artefacts that appear as garbage on the portal."""
    text = re.sub(r"\*{1,3}", "", text)
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"_{1,2}([^_]+)_{1,2}", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"^\s*[-â€¢]\s", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def split_for_portal(text: str, max_chars: int = MAX_DRAFT_CHARS) -> list[str]:
    """Split a long response into portal-uploadable chunks at paragraph boundaries."""
    if len(text) <= max_chars:
        return [text]
    parts: list[str] = []
    remaining = text
    part_num = 1
    while len(remaining) > max_chars:
        cut = remaining[:max_chars]
        split_at = cut.rfind("\n\n")
        if split_at < int(max_chars * 0.75):
            split_at = cut.rfind(".")
        if split_at <= 0:
            split_at = max_chars
        part_text = remaining[:split_at + 1]
        header = f"[PART {part_num} OF RESPONSE]\n\n"
        parts.append(header + part_text)
        remaining = remaining[split_at + 1:].strip()
        part_num += 1
    if remaining:
        parts.append(f"[PART {part_num} â€” FINAL]\n\n" + remaining)
    return parts


def redact_sensitive_text_for_ai(text: str) -> tuple[str, dict[str, int]]:
    """Redact direct client identifiers before any text is sent to Gemini."""
    redacted = text
    counts: dict[str, int] = {}

    patterns = {
        "PAN": r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",
        "TAN": r"\b[A-Z]{4}[0-9]{5}[A-Z]\b",
        "GSTIN": r"\b[0-3][0-9][A-Z]{5}[0-9]{4}[A-Z][1-9A-Z]Z[0-9A-Z]\b",
        "Aadhaar": r"\b\d{4}\s?\d{4}\s?\d{4}\b",
        "Email": r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "Mobile": r"\b(?:\+91[-\s]?)?[6-9]\d{9}\b",
        "IFSC": r"\b[A-Z]{4}0[A-Z0-9]{6}\b",
        "CIN": r"\b[LU][0-9]{5}[A-Z]{2}[0-9]{4}[A-Z]{3}[0-9]{6}\b",
        "BankAccount": r"\b(?:A/c|Account|Acct|Bank Account)\s*(?:No\.?|Number)?\s*[:\-]?\s*\d{6,18}\b",
    }
    for label, pattern in patterns.items():
        redacted, count = re.subn(pattern, f"[{label.upper()}_REDACTED]", redacted, flags=re.IGNORECASE)
        counts[label] = count

    field_patterns = {
        "TaxpayerName": r"(?im)^(\s*(?:taxpayer|assessee|applicant)\s+name\s*[:\-]\s*).+$",
        "FatherName": r"(?im)^(\s*father'?s?\s+name\s*[:\-]\s*).+$",
        "Address": r"(?im)^(\s*(?:address|registered address|communication address)\s*[:\-]\s*).+$",
    }
    for label, pattern in field_patterns.items():
        redacted, count = re.subn(pattern, rf"\1[{label.upper()}_REDACTED]", redacted)
        counts[label] = count

    try:
        from presidio_analyzer import AnalyzerEngine
        from presidio_anonymizer import AnonymizerEngine

        analyzer = AnalyzerEngine()
        anonymizer = AnonymizerEngine()
        results = analyzer.analyze(text=redacted, language="en", entities=["PERSON", "LOCATION"], score_threshold=0.75)
        if results:
            redacted = anonymizer.anonymize(text=redacted, analyzer_results=results).text
            counts["PresidioNER"] = len(results)
    except Exception:
        pass

    return redacted, {k: v for k, v in counts.items() if v}


def anonymise_prompt_identity(prompt: str, assessee_name: str, pan: str) -> str:
    """Keep draft prompts useful while preventing client name/PAN from going to Gemini."""
    if assessee_name and assessee_name not in {"[ASSESSEE NAME]", "[ASSESSEE_NAME]"}:
        prompt = prompt.replace(assessee_name, "[ASSESSEE_NAME]")
    if pan and pan not in {"[PAN]", "[PAN_REDACTED]"}:
        prompt = prompt.replace(pan, "[PAN]")
    prompt, _ = redact_sensitive_text_for_ai(prompt)
    return prompt


def restore_local_identity(text: str, assessee_name: str, pan: str) -> str:
    return (
        text.replace("[ASSESSEE_NAME]", assessee_name or "[ASSESSEE NAME]")
        .replace("[TAXPAYERNAME_REDACTED]", assessee_name or "[ASSESSEE NAME]")
        .replace("[PAN_REDACTED]", pan or "[PAN]")
        .replace("[PAN]", pan or "[PAN]")
    )


def extract_issue_blocks(extraction: str) -> list[dict[str, str]]:
    """Parse issue blocks from the extraction output for Step 3 structured inputs."""
    part_match = re.search(
        r"PART\s*3\s*:\s*ISSUES IDENTIFIED(.*?)(?=PART\s*4\s*:|$)",
        extraction,
        re.IGNORECASE | re.DOTALL,
    )
    source = part_match.group(1) if part_match else extraction
    starts = list(re.finditer(r"(?m)^\s*(?:ISSUE\s*)?(\d+)[\).:\-]\s*(.*)$", source, re.IGNORECASE))
    issues: list[dict[str, str]] = []
    if starts:
        for idx, match in enumerate(starts):
            end = starts[idx + 1].start() if idx + 1 < len(starts) else len(source)
            number = match.group(1)
            body = source[match.start():end].strip()
            title = match.group(2).strip() or f"Issue {number}"
            issues.append({"number": number, "title": title[:120], "body": body})
    else:
        allegation_lines = re.split(r"(?=\n\s*\(a\)\s*ALLEGATION|\n\s*ALLEGATION\s*:)", source, flags=re.IGNORECASE)
        for idx, body in enumerate([b.strip() for b in allegation_lines if "ALLEGATION" in b.upper()], 1):
            issues.append({"number": str(idx), "title": f"Issue {idx}", "body": body})
    return issues[:12]


def make_portal_safe_text(text: str) -> str:
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\u2022": "-",
        "\u20b9": "Rs.",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"[^\x09\x0A\x0D\x20-\x7E]", "", text)
    return text.strip()


def build_word_package(draft: str, cover: str, extraction: str) -> bytes:
    if not DOCX_AVAILABLE or Document is None:
        raise RuntimeError("python-docx is not installed.")
    doc = Document()
    doc.add_heading("R K Muley & Co - Tax Notice Litigation Assistant", level=1)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_heading("Portal Cover Note", level=2)
    doc.add_paragraph(make_portal_safe_text(cover or "No cover note generated."))
    doc.add_heading("Full Response", level=2)
    for para in make_portal_safe_text(draft).split("\n\n"):
        doc.add_paragraph(para)
    doc.add_heading("Extracted Issues", level=2)
    for para in make_portal_safe_text(extraction).split("\n\n"):
        doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_annexure_schedule(user_inputs: str) -> str:
    """Extract document mentions from user inputs and build a numbered schedule."""
    doc_pattern = re.compile(
        r"MY DOCUMENTS\s*:?\s*(.*?)(?=\n(?:ISSUE|MY POSITION|MY FACTS|CASE LAW|$))",
        re.IGNORECASE | re.DOTALL,
    )
    all_docs: list[str] = []
    for match in doc_pattern.finditer(user_inputs):
        block = match.group(1).strip()
        for item in re.split(r"[,;\n]", block):
            cleaned = item.strip().lstrip("0123456789.-) ")
            if len(cleaned) > 4:
                all_docs.append(cleaned)

    if not all_docs:
        all_docs = [
            "ITR Acknowledgement for the relevant Assessment Year",
            "Form 26AS (downloaded as on date of submission)",
            "Bank account statement for the relevant period",
            "Form 16 / Form 16A (as applicable)",
        ]

    seen: set[str] = set()
    unique_docs: list[str] = []
    for doc in all_docs:
        key = doc.lower().strip()
        if key not in seen:
            seen.add(key)
            unique_docs.append(doc)

    lines = [
        "\n\nANNEXURE SCHEDULE",
        "(Documents Submitted in Support of this Response)",
        "-" * 55,
    ]
    for idx, doc in enumerate(unique_docs, 1):
        lines.append(f"Annexure {idx:<3} â€” {doc}")
    lines.append(
        "\nNote: All documents are self-attested copies. "
        "Originals will be produced for verification if called upon."
    )
    return "\n".join(lines)


# v9 helper layer. These functions are intentionally local to app.py so the
# monolith remains deployable while the new workflows are still beta.
def feature_enabled(name: str, default: bool = False) -> bool:
    return bool(FEATURES.get(name, default))


def extract_annexure_items(user_inputs: str) -> list[str]:
    doc_pattern = re.compile(
        r"MY DOCUMENTS\s*:?\s*(.*?)(?=\n(?:ISSUE|MY POSITION|MY FACTS|CASE LAW|$))",
        re.IGNORECASE | re.DOTALL,
    )
    all_docs: list[str] = []
    for match in doc_pattern.finditer(user_inputs):
        block = match.group(1).strip()
        for item in re.split(r"[,;\n]", block):
            cleaned = item.strip().lstrip("0123456789.-) ")
            if len(cleaned) > 4:
                all_docs.append(cleaned)
    if not all_docs:
        all_docs = [
            "ITR Acknowledgement for the relevant Assessment Year",
            "Form 26AS (downloaded as on date of submission)",
            "Bank account statement for the relevant period",
            "Form 16 / Form 16A (as applicable)",
        ]
    seen: set[str] = set()
    unique_docs: list[str] = []
    for doc in all_docs:
        key = doc.lower().strip()
        if key not in seen:
            seen.add(key)
            unique_docs.append(doc)
    return unique_docs


def build_annexure_schedule(user_inputs: str, evidence_rows: list[dict] | None = None) -> str:
    unique_docs = extract_annexure_items(user_inputs)
    status_by_doc = {
        str(row.get("Document", "")).strip().lower(): str(row.get("Status", "")).strip()
        for row in (evidence_rows or [])
        if row.get("Document")
    }
    lines = [
        "\n\nANNEXURE SCHEDULE",
        "(Documents Submitted in Support of this Response)",
        "-" * 55,
    ]
    for idx, doc in enumerate(unique_docs, 1):
        status = status_by_doc.get(doc.lower().strip())
        suffix = f" [{status}]" if status else ""
        lines.append(f"Annexure {idx:<3} - {doc}{suffix}")
    lines.append(
        "\nNote: All documents are self-attested copies. "
        "Originals will be produced for verification if called upon."
    )
    return "\n".join(lines)


def extract_notice_quantum(extraction: str) -> float | None:
    candidates: list[float] = []
    for line in extraction.splitlines():
        if not re.search(r"(quantum|demand|tax effect|amount|addition|income|penalty)", line, re.I):
            continue
        for raw in re.findall(r"(?:rs\.?|inr)?\s*([0-9][0-9,]*(?:\.\d+)?)", line, re.I):
            try:
                candidates.append(float(raw.replace(",", "")))
            except ValueError:
                continue
    return max(candidates) if candidates else None


def reconcile_ais_tis(rows: list[dict], notice_quantum: float | None) -> dict:
    clean_rows: list[dict] = []
    total_delta = 0.0
    for row in rows:
        label = str(row.get("Category", "")).strip() or "Uncategorised"
        try:
            ais_amount = float(row.get("AIS Amount", 0) or 0)
            tis_amount = float(row.get("TIS Amount", 0) or 0)
        except (TypeError, ValueError):
            ais_amount, tis_amount = 0.0, 0.0
        delta = ais_amount - tis_amount
        total_delta += delta
        clean_rows.append({
            "Category": label,
            "AIS Amount": ais_amount,
            "TIS Amount": tis_amount,
            "Delta": delta,
            "Remarks": str(row.get("Remarks", "")).strip(),
        })
    flags: list[str] = []
    abs_delta = abs(total_delta)
    if notice_quantum is not None:
        tolerance = max(100.0, notice_quantum * 0.01)
        if abs(abs_delta - notice_quantum) <= tolerance:
            flags.append("AIS/TIS delta broadly matches the notice quantum.")
        elif abs_delta > tolerance:
            flags.append("AIS/TIS delta does not match the extracted notice quantum.")
    elif abs_delta:
        flags.append("Notice quantum could not be extracted automatically; verify manually.")
    return {
        "rows": clean_rows,
        "total_delta": total_delta,
        "notice_quantum": notice_quantum,
        "flags": flags,
    }


def parse_due_date(extraction: str) -> date | None:
    due_m = re.search(r"Due Date.*?:\s*(.+)", extraction, re.IGNORECASE)
    if not due_m:
        return None
    raw = due_m.group(1).strip().splitlines()[0]
    raw = re.sub(r"\b(on or before|by|not found|n/a)\b", "", raw, flags=re.I).strip(" :-")
    for fmt in ("%d-%m-%Y", "%d/%m/%Y", "%Y-%m-%d", "%d.%m.%Y", "%d %B %Y", "%d %b %Y"):
        try:
            return datetime.strptime(raw, fmt).date()
        except ValueError:
            continue
    return None


def _ics_escape(value: str) -> str:
    return value.replace("\\", "\\\\").replace("\n", "\\n").replace(",", "\\,").replace(";", "\\;")


def generate_ics_invite(due_date: date, summary: str, description: str) -> str:
    dtstamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    start = due_date.strftime("%Y%m%d")
    end = (due_date + timedelta(days=1)).strftime("%Y%m%d")
    uid = f"{uuid.uuid4()}@rkmuley-tax-notice-v9"
    return "\r\n".join([
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//R K Muley & Co//Tax Notice Litigation Assistant v9//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
        "BEGIN:VEVENT",
        f"UID:{uid}",
        f"DTSTAMP:{dtstamp}",
        f"DTSTART;VALUE=DATE:{start}",
        f"DTEND;VALUE=DATE:{end}",
        f"SUMMARY:{_ics_escape(summary)}",
        f"DESCRIPTION:{_ics_escape(description)}",
        "BEGIN:VALARM",
        "TRIGGER:-P1D",
        "ACTION:DISPLAY",
        "DESCRIPTION:Tax notice response due tomorrow",
        "END:VALARM",
        "END:VEVENT",
        "END:VCALENDAR",
    ]) + "\r\n"


def logo_data_uri() -> str:
    if not LOGO_PATH.exists():
        return ""
    encoded = base64.b64encode(LOGO_PATH.read_bytes()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def section_270a_detected(extraction: str) -> bool:
    return bool(re.search(r"\b270A\b", extraction, re.IGNORECASE))


def penalty_proceeding_detected(extraction: str) -> bool:
    return bool(re.search(
        r"\b(270A|270AA|271(?:\([^)]+\))?|271AAB|271AAC|271AAD|271B|271C|271D|271E|271H|272A|273B|274|275|penalty|show cause.*penalt)",
        extraction,
        re.IGNORECASE,
    ))


def build_form_68_immunity_draft(extraction: str, assessee: str, pan: str) -> str:
    return (
        "FORM 68 / SECTION 270AA IMMUNITY REQUEST - WORKING DRAFT\n\n"
        f"Applicant: {assessee}\nPAN: {pan}\n\n"
        "The assessee proposes to seek immunity under section 270AA against penalty "
        "proceedings under section 270A, subject to statutory eligibility. The final "
        "Form 68 filing should be made only after confirming that the tax and interest "
        "payable as per the relevant assessment/reassessment order have been paid within "
        "the prescribed time and that no appeal has been filed against the said order.\n\n"
        "Prayer:\n"
        "The assessee respectfully requests immunity from imposition of penalty under "
        "section 270A and from prosecution, in accordance with section 270AA of the "
        "Income-tax Act, 1961. The assessee undertakes to comply with all requirements "
        "and to furnish proof of payment and non-filing of appeal wherever required.\n\n"
        "Extraction basis:\n"
        f"{extraction[:2500]}"
    )


def compute_form35_fee(assessed_income: float, other_matter: bool = False) -> int:
    """Return first-appeal fee under section 249/Form 35 slabs."""
    if other_matter:
        return 250
    if assessed_income <= 100000:
        return 250
    if assessed_income <= 200000:
        return 500
    return 1000


def count_words(text: str) -> int:
    return len(re.findall(r"\b[\w()/.%-]+\b", text or ""))


def build_form35_metadata_text(
    forum: str,
    order_section: str,
    order_date_value: date,
    service_date_value: date,
    assessed_income: float,
    demand_amount: str,
    fee_amount: int,
    delay_days: int,
    wants_rule46a: bool,
    wants_stay: bool,
    wants_personal_hearing: bool,
) -> str:
    return "\n".join([
        f"Target authority: {forum}",
        f"Order appealed against: Section {order_section}",
        f"Order date: {order_date_value.strftime('%d-%m-%Y')}",
        f"Date of service: {service_date_value.strftime('%d-%m-%Y')}",
        f"Assessed income as per order: Rs. {assessed_income:,.0f}",
        f"Disputed demand: {demand_amount or 'Not provided'}",
        f"Computed appeal fee: Rs. {fee_amount}",
        "Payment route: e-Pay Tax / Other Receipts / Minor Head 500 / Appeal Fees",
        f"Delay beyond 30 days: {delay_days} day(s)",
        f"Condonation required: {'Yes' if delay_days > 0 else 'No'}",
        f"Rule 46A additional evidence application required: {'Yes' if wants_rule46a else 'No'}",
        f"Stay of demand request required: {'Yes' if wants_stay else 'No'}",
        f"Personal hearing requested: {'Yes' if wants_personal_hearing else 'No'}",
        "Portal constraints: Statement of Facts within 1000 words; each Ground of Appeal within 100 words; use plain portal-safe text.",
    ])


def build_penalty_metadata_text(
    penalty_section: str,
    notice_date_value: date,
    reply_due_value: date,
    quantum_appeal_pending: bool,
    form35_date: str,
    form35_ack: str,
    requested_abeyance: bool,
    defective_notice: bool,
    delayed_reply: bool,
    delay_reason: str,
    wants_273b: bool,
    wants_270aa: bool,
) -> str:
    delay_days = max(0, (date.today() - reply_due_value).days)
    return "\n".join([
        f"Penalty section: {penalty_section or 'Not specified'}",
        f"Penalty notice date: {notice_date_value.strftime('%d-%m-%Y')}",
        f"Reply due date: {reply_due_value.strftime('%d-%m-%Y')}",
        f"Delay in response as on today: {delay_days} day(s)",
        f"Delayed reply / condonation requested: {'Yes' if delayed_reply or delay_days > 0 else 'No'}",
        f"Reason for delay: {delay_reason or 'Not provided'}",
        f"Quantum appeal pending: {'Yes' if quantum_appeal_pending else 'No'}",
        f"Form 35 date: {form35_date or 'Not provided'}",
        f"Form 35 acknowledgment: {form35_ack or 'Not provided'}",
        f"Request abeyance of penalty proceedings: {'Yes' if requested_abeyance else 'No'}",
        f"Defective/vague notice objection: {'Yes' if defective_notice else 'No'}",
        f"Section 273B reasonable cause submission: {'Yes' if wants_273b else 'No'}",
        f"Section 270AA/Form 68 immunity pathway considered: {'Yes' if wants_270aa else 'No'}",
        "Penalty guide principles: penalty is separate from assessment; addition alone is not automatic penalty; check DIN, limitation, satisfaction, charge clarity, natural justice, and evidence.",
        "Section 270A handling: distinguish under-reporting from misreporting; 200 percent misreporting requires a specific section 270A(9) basis.",
    ])


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# PROCEDURAL VALIDATOR (inline â€” from engines would be imported in full version)
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def run_procedural_checks(extraction: str) -> list[str]:
    """
    Run all deterministic procedural checks and return a list of flag strings.
    Imports ProceduralValidator from engines/procedural_validator.py in production.
    Inline here for single-file deployability.
    """
    from config import DIN_PATTERN
    flags: list[str] = []

    # Extract DIN
    din_m = re.search(r"DIN.*?:\s*(.+)", extraction, re.IGNORECASE)
    din_val = din_m.group(1).strip() if din_m else ""

    if not din_val or din_val.upper() in ("NOT FOUND IN NOTICE", "N/A", ""):
        flags.append(
            "DIN ABSENT - NOTICE POTENTIALLY VOID\n"
            "CBDT Circular No. 19/2019 mandates a computer-generated DIN on all communications. "
            "A communication without a valid DIN is void and not acted upon. "
            "Raise as preliminary threshold objection before engaging on merits."
        )
    elif not DIN_PATTERN.search(din_val):
        flags.append(
            f"DIN FORMAT NEEDS MANUAL VERIFICATION: '{din_val}'\n"
            f"The DIN does not match the app's current ITBA pattern library. "
            f"Verify against the original notice and ITBA portal before raising this as an objection."
        )

    # Extract primary section
    sec_m = re.search(r"Primary Section Invoked:\s*(.+)", extraction, re.IGNORECASE)
    prim_sec = sec_m.group(1).strip() if sec_m else ""

    if any(s in prim_sec for s in ["148", "147"]):
        flags.append(
            "REASSESSMENT NOTICE - SECTION 148A COMPLIANCE REQUIRED\n"
            "Finance Act 2021 mandates the four-step 148A procedure before a Sec 148 notice. "
            "Verify: (1) 148A(b) SCN issued with min 7 days? "
            "(2) Reasoned 148A(d) order passed considering your reply? "
            "(3) PCIT/CCIT sanction obtained u/s 151? "
            "(4) Tangible material (not mere change of opinion) cited in 148A(d) order?"
        )

    # Check AY for time-bar
    ay_m = re.search(r"Assessment Year.*?:\s*([\d\-]+)", extraction, re.IGNORECASE)
    if ay_m:
        ay_str = ay_m.group(1).strip()
        year_m = re.search(r"(\d{4})", ay_str)
        if year_m:
            ay_start = int(year_m.group(1))
            today = date.today()
            ay_end = date(ay_start + 1, 3, 31)
            years_elapsed = (today - ay_end).days / 365.25

            if ay_start < 2018:
                flags.append(
                    f"TIME-BARRED NOTICE - AY {ay_str}\n"
                    f"AY {ay_str} is more than 6 years old. Notice is prima facie time-barred "
                    f"under Section 149 as amended by Finance Act 2021. "
                    f"The 10-year exception applies only if escaped income >= Rs. 50 lakhs "
                    f"AND PCIT sanction was obtained. Raise as threshold objection."
                )
            elif 2018 <= ay_start <= 2020:
                flags.append(
                    f"SAVINGS-CLAUSE PERIOD - AY {ay_str}\n"
                    f"AY {ay_str} falls in the transitional savings-clause window (CBDT Circular 6/2024). "
                    f"Verify: (1) Is the notice within 6 years from end of AY? "
                    f"(2) Has AO cited tangible material? (3) PCIT sanction obtained u/s 151?"
                )

    return flags


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# CSS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def inject_css() -> None:
    """Inject custom CSS â€” MUST be called inside main(), not at module level."""
    st.markdown("""
<style>
.main { padding: 1.5rem 2rem; }
.block-container { max-width: 1200px; }

.firm-header {
    background: linear-gradient(135deg, #1a237e 0%, #283593 60%, #3949ab 100%);
    color: white; padding: 1.3rem 1.8rem; border-radius: 10px;
    margin-bottom: 1rem; display: flex; align-items: center;
    justify-content: space-between; gap: 1rem;
}
.firm-header .brand-left { display: flex; align-items: center; gap: 1rem; min-width: 0; }
.firm-header .logo-wrap {
    background: white; border-radius: 8px; padding: 0.45rem 0.7rem;
    box-shadow: 0 6px 16px rgba(0,0,0,0.14); flex: 0 0 auto;
}
.firm-header .firm-logo { height: 54px; width: auto; display: block; }
.firm-header .firm-name { font-size: 1.75rem; font-weight: 800; }
.firm-header .firm-sub  { font-size: 0.88rem; opacity: 0.85; margin-top: 4px; }
.firm-header .app-lbl   { text-align: right; font-size: 0.82rem; opacity: 0.82; flex: 0 0 auto; }
@media (max-width: 760px) {
    .firm-header { align-items: flex-start; flex-direction: column; }
    .firm-header .brand-left { align-items: flex-start; flex-direction: column; }
    .firm-header .app-lbl { text-align: left; }
    .firm-header .firm-logo { height: 44px; }
}

.output-box {
    background: #fafafa; border: 1px solid #ddd;
    border-left: 5px solid #3949ab; border-radius: 6px;
    padding: 1.3rem; margin: 0.8rem 0; font-size: 0.9rem; line-height: 1.85;
}
.proc-defect-box {
    background: #fff8e1; border: 2px solid #f9a825;
    border-left: 6px solid #e65100; border-radius: 8px;
    padding: 1rem 1.2rem; margin: 0.8rem 0; font-size: 0.88rem;
}
.section-header {
    background: #e8eaf6; border-left: 4px solid #3949ab;
    padding: 0.5rem 1rem; border-radius: 0 4px 4px 0;
    font-weight: 700; font-size: 0.95rem; color: #1a237e;
    margin: 1rem 0 0.5rem;
}
.portal-note {
    background: #e3f2fd; border: 1px solid #90caf9;
    border-radius: 6px; padding: 0.75rem 1rem;
    font-size: 0.84rem; color: #1565c0; margin: 0.5rem 0;
}
.warning-box {
    background: #fff3e0; border: 1px solid #ffb74d;
    border-left: 4px solid #f57c00; border-radius: 6px;
    padding: 0.75rem 1rem; font-size: 0.85rem;
}
</style>
""", unsafe_allow_html=True)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# SIDEBAR
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def render_sidebar() -> tuple[str, str]:
    """Render sidebar. Returns (model_choice, api_key_override)."""
    with st.sidebar:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), use_container_width=True)
        st.markdown(f"### {FIRM_NAME}")
        st.caption(f"{FIRM_SUBTITLE} | v{APP_VERSION}")

        if st.session_state.get("auth_username"):
            display = st.session_state.get("auth_display", "User")
            role    = st.session_state.get("auth_role", "article")
            st.markdown(f"**{display}** `{role}`")
            if st.button("Logout", use_container_width=True):
                logout()

        st.divider()

        # API key â€” sidebar entry (legacy support for local dev without secrets)
        api_key_override = ""
        if not get_api_key():
            api_key_override = st.text_input(
                "Gemini API Key",
                type="password",
                placeholder="AIzaSyAVrRmGLPMw_8kFQe6EJx7WOhnbP_Bg-iE",
                help="Free key at aistudio.google.com",
            )
            if api_key_override:
                # Store temporarily in session â€” will be picked up by get_api_key()
                # This is the legacy path; prefer secrets.toml
                st.session_state["_api_key_"] = api_key_override
        else:
            st.success("API key configured")

        st.markdown("**AI Model**")
        model_options = [m["id"] for m in GEMINI_MODELS]
        model_labels  = [m["label"] for m in GEMINI_MODELS]
        model_idx     = st.selectbox(
            "Select Model",
            options=range(len(model_options)),
            format_func=lambda i: model_labels[i],
            index=0,
        )
        model_choice = model_options[model_idx]

        st.divider()
        st.caption("Session Progress")
        progress_keys = [
            ("extraction_result", "Step 1: Extraction"),
            ("user_inputs_text",  "Step 3: Inputs"),
            ("draft_response",    "Step 4: Draft"),
        ]
        for key, label in progress_keys:
            status = "Done" if st.session_state.get(key) else "Pending"
            st.caption(f"{status}: {label}")

    return model_choice, api_key_override


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# FIRM HEADER
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def render_firm_header() -> None:
    logo_uri = logo_data_uri()
    logo_html = (
        f'<div class="logo-wrap"><img class="firm-logo" src="{logo_uri}" alt="{FIRM_NAME} logo"></div>'
        if logo_uri else ""
    )
    st.markdown(f"""
<div class="firm-header">
  <div class="brand-left">
    {logo_html}
    <div>
      <div class="firm-name">{FIRM_NAME}</div>
      <div class="firm-sub">{FIRM_SUBTITLE}</div>
    </div>
  </div>
  <div class="app-lbl">
    {APP_NAME}<br>
    Version {APP_VERSION}<br>
    AI-Augmented Legal Workflow
  </div>
</div>
""", unsafe_allow_html=True)


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# TAB RENDERERS
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def render_tab_upload(model_choice: str) -> None:
    st.markdown('<div class="section-header">Upload Notice & Extract</div>',
                unsafe_allow_html=True)

    upload_method = st.radio(
        "Notice input method",
        ["Upload PDF", "Paste Text"],
        horizontal=True,
    )

    notice_text = ""

    if upload_method == "Upload PDF":
        uploaded = st.file_uploader(
            "Upload the Income Tax notice (PDF)",
            type=["pdf"],
            help="Text-based PDFs only. For scanned PDFs, use OCR first.",
        )
        if uploaded:
            with st.spinner("Extracting text from PDF..."):
                try:
                    text, method = extract_text_from_pdf(uploaded.read())
                    st.success(f"Extracted {len(text):,} characters via {method}.")
                    notice_text = text
                except ValueError as exc:
                    st.error(str(exc))
    else:
        notice_text = st.text_area(
            "Paste notice text",
            height=300,
            placeholder="Paste the full text of the Income Tax notice here...",
        )

    if notice_text and st.button("Extract & Analyse Notice", type="primary"):
        with st.spinner("Running AI extraction and procedural scan..."):
            try:
                ai_notice_text, redaction_counts = redact_sensitive_text_for_ai(notice_text[:150_000])
                if redaction_counts:
                    st.info(
                        "Privacy guard active: direct identifiers were redacted before API processing "
                        f"({', '.join(f'{k}: {v}' for k, v in redaction_counts.items())})."
                    )
                extraction = call_gemini(
                    model_name=model_choice,
                    prompt=EXTRACTION_PROMPT.format(notice_text=ai_notice_text),
                    temperature=0.1,
                    max_tokens=4096,
                    step="extraction",
                    username=st.session_state.get("auth_username", "anonymous"),
                )
                st.session_state["extraction_result"]  = extraction
                st.session_state["notice_text_stored"] = ai_notice_text
                st.session_state["pii_mask_counts"] = redaction_counts

                # Procedural scan (deterministic â€” no LLM)
                proc_flags = run_procedural_checks(extraction)
                st.session_state["procedural_flags"] = proc_flags

                # Case law match (deterministic)
                matched = get_relevant_case_laws(extraction)
                st.session_state["matched_laws"] = matched

                write_audit_trail(
                    "EXTRACTION_COMPLETE",
                    "notice",
                    f"proc_flags={len(proc_flags)}",
                    username=st.session_state.get("auth_username", "anonymous"),
                )

                st.success("Extraction complete. Proceed to Review Issues.")

                if proc_flags:
                    st.warning(
                        f"{len(proc_flags)} procedural defect(s) detected. "
                        "Review in Procedural Audit."
                    )

            except APICallError as exc:
                st.error(f"API Error: {exc}")

    if st.session_state.get("extraction_result"):
        with st.expander("Extraction Result", expanded=True):
            st.markdown(
                f'<div class="output-box"><pre style="white-space:pre-wrap;">'
                f'{st.session_state["extraction_result"]}</pre></div>',
                unsafe_allow_html=True,
            )
            st.download_button(
                "Download Extracted Issues",
                data=st.session_state["extraction_result"],
                file_name=f"RKMuley_v9_Extraction_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True,
            )


def render_tab_review() -> None:
    st.markdown('<div class="section-header">Review Extracted Issues</div>',
                unsafe_allow_html=True)
    if "extraction_result" not in st.session_state:
        st.info("Complete Step 1 first.")
        return

    extraction = st.session_state["extraction_result"]

    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**Matched Case Laws**")
        matched = st.session_state.get("matched_laws", "No matches")
        st.markdown(
            f'<div class="output-box" style="max-height:400px;overflow-y:auto;">'
            f'<pre style="white-space:pre-wrap;font-size:0.82rem;">{matched}</pre></div>',
            unsafe_allow_html=True,
        )
    with col2:
        st.markdown("**Procedural Flags**")
        flags = st.session_state.get("procedural_flags", [])
        if flags:
            for flag in flags:
                st.markdown(
                    f'<div class="proc-defect-box">{flag}</div>',
                    unsafe_allow_html=True,
                )
        else:
            st.success("No procedural defects detected.")


def render_tab_inputs() -> None:
    st.markdown('<div class="section-header">Your Inputs</div>',
                unsafe_allow_html=True)
    if "extraction_result" not in st.session_state:
        st.info("Complete Step 1 first.")
        return

    st.markdown("""<div class="portal-note">
For each issue extracted in Step 1, provide your factual position using this template:<br><br>
<strong>ISSUE 1:</strong><br>
MY POSITION: Deny / Admit / Partially Admit<br>
MY FACTS: [Your explanation in 2-3 sentences]<br>
MY DOCUMENTS: [Comma-separated list of documents you have]<br>
CASE LAW: [Any specific case or circular you want cited]
</div>""", unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    with c1:
        assessee_name = st.text_input("Assessee Name", value=st.session_state.get("assessee_name", ""))
    with c2:
        raw_pan = st.text_input("PAN", value=st.session_state.get("assessee_pan", ""),
                                placeholder="ABCDE1234F").upper().strip()
    with c3:
        city = st.text_input("City / Place", value=st.session_state.get("assessee_city", ""))

    from config import PAN_PATTERN
    if raw_pan and not PAN_PATTERN.match(raw_pan):
        st.error(f"PAN format invalid: '{raw_pan}'. Expected format: ABCDE1234F")
        return

    ar_name = st.text_input(
        "Authorised Representative (CA / AR) Name",
        value=st.session_state.get("ar_name", ""),
        placeholder="Leave blank if filing directly",
    )

    issues = extract_issue_blocks(st.session_state["extraction_result"])
    structured_inputs: list[str] = []
    if issues:
        st.markdown("**Issue-wise inputs**")
        st.caption("Enter your position, facts, documents, and case law separately for each extracted issue.")
        for issue in issues:
            num = issue["number"]
            with st.expander(f"Issue {num}: {issue['title']}", expanded=(num == "1")):
                st.text_area(
                    "Extracted issue",
                    value=issue["body"],
                    height=120,
                    disabled=True,
                    key=f"issue_body_{num}",
                )
                position = st.selectbox(
                    "My Position",
                    ["Deny", "Partially Admit", "Admit", "Need More Facts"],
                    key=f"issue_position_{num}",
                )
                facts = st.text_area(
                    "My Facts",
                    value=st.session_state.get(f"issue_facts_{num}", ""),
                    height=120,
                    placeholder="State the factual explanation for this issue.",
                    key=f"issue_facts_{num}",
                )
                docs = st.text_area(
                    "My Documents",
                    value=st.session_state.get(f"issue_docs_{num}", ""),
                    height=80,
                    placeholder="List supporting documents for this issue.",
                    key=f"issue_docs_{num}",
                )
                case_law = st.text_area(
                    "Case Law / Circulars",
                    value=st.session_state.get(f"issue_case_law_{num}", ""),
                    height=80,
                    placeholder="Optional. Add only citations you have verified.",
                    key=f"issue_case_law_{num}",
                )
                structured_inputs.append(
                    f"ISSUE {num}:\n"
                    f"MY POSITION: {position}\n"
                    f"MY FACTS: {facts.strip()}\n"
                    f"MY DOCUMENTS: {docs.strip()}\n"
                    f"CASE LAW: {case_law.strip()}"
                )
        legacy_inputs = st.text_area(
            "Additional general inputs",
            height=140,
            value=st.session_state.get("user_inputs_extra", ""),
            placeholder="Any common background facts, portal notes, or drafting instructions.",
        )
        user_inputs = "\n\n".join(structured_inputs + ([legacy_inputs] if legacy_inputs.strip() else []))
    else:
        user_inputs = st.text_area(
            "Issue-wise inputs (use template above)",
            height=350,
            value=st.session_state.get("user_inputs_text", ""),
            placeholder="ISSUE 1:\nMY POSITION: Deny\nMY FACTS: ...\nMY DOCUMENTS: ...\nCASE LAW: ...",
        )

    ais_tis_rows = st.session_state.get("ais_tis_rows", [
        {"Category": "Salary / TDS", "AIS Amount": 0.0, "TIS Amount": 0.0, "Remarks": ""},
        {"Category": "Interest", "AIS Amount": 0.0, "TIS Amount": 0.0, "Remarks": ""},
        {"Category": "SFT / High-value transaction", "AIS Amount": 0.0, "TIS Amount": 0.0, "Remarks": ""},
    ])
    ais_recon_result = st.session_state.get("ais_recon_result")
    if feature_enabled("AIS_RECON"):
        with st.expander("AIS/TIS Reconciliation Engine (Beta)", expanded=False):
            st.caption("Enter source-wise AIS and TIS figures. The app compares the delta with extracted notice quantum.")
            ais_tis_rows = st.data_editor(
                ais_tis_rows,
                num_rows="dynamic",
                use_container_width=True,
                key="ais_tis_editor",
                column_config={
                    "Category": st.column_config.TextColumn("Category"),
                    "AIS Amount": st.column_config.NumberColumn("AIS Amount", min_value=0.0, step=100.0),
                    "TIS Amount": st.column_config.NumberColumn("TIS Amount", min_value=0.0, step=100.0),
                    "Remarks": st.column_config.TextColumn("Remarks"),
                },
            )
            notice_quantum = extract_notice_quantum(st.session_state["extraction_result"])
            ais_recon_result = reconcile_ais_tis(ais_tis_rows, notice_quantum)
            c_ais, c_tis = st.columns(2)
            with c_ais:
                st.metric("Extracted Notice Quantum", "Not found" if notice_quantum is None else f"Rs. {notice_quantum:,.0f}")
            with c_tis:
                st.metric("AIS - TIS Delta", f"Rs. {ais_recon_result['total_delta']:,.0f}")
            for flag in ais_recon_result["flags"]:
                st.warning(flag)

    evidence_rows = st.session_state.get("evidence_tracker_rows")
    if feature_enabled("EVIDENCE_TRACKER"):
        with st.expander("Evidence Tracker", expanded=False):
            docs = extract_annexure_items(user_inputs)
            status_lookup = {
                str(row.get("Document", "")).strip().lower(): row
                for row in (evidence_rows or [])
                if row.get("Document")
            }
            tracker_seed = []
            for doc in docs:
                existing = status_lookup.get(doc.lower().strip(), {})
                tracker_seed.append({
                    "Document": doc,
                    "Status": existing.get("Status", "Pending Client"),
                    "Notes": existing.get("Notes", ""),
                })
            evidence_rows = st.data_editor(
                tracker_seed,
                num_rows="dynamic",
                use_container_width=True,
                key="evidence_tracker_editor",
                column_config={
                    "Document": st.column_config.TextColumn("Document", disabled=True),
                    "Status": st.column_config.SelectboxColumn(
                        "Status",
                        options=["Ready", "Pending Client", "Unavailable"],
                        required=True,
                    ),
                    "Notes": st.column_config.TextColumn("Notes"),
                },
            )

    if st.button("Save Inputs", type="primary"):
        st.session_state["assessee_name"]   = assessee_name
        st.session_state["assessee_pan"]    = raw_pan
        st.session_state["assessee_city"]   = city
        st.session_state["ar_name"]         = ar_name
        st.session_state["user_inputs_text"] = user_inputs
        if issues:
            st.session_state["user_inputs_extra"] = legacy_inputs
        st.session_state["ais_tis_rows"] = ais_tis_rows
        st.session_state["ais_recon_result"] = ais_recon_result
        st.session_state["evidence_tracker_rows"] = evidence_rows or []
        save_session("assessee_name", assessee_name)
        save_session("user_inputs_text", user_inputs)
        save_session("ais_tis_rows", ais_tis_rows)
        save_session("evidence_tracker_rows", evidence_rows or [])
        write_audit_trail("INPUTS_SAVED", "inputs", f"pan={raw_pan}",
                          username=st.session_state.get("auth_username", "anonymous"))
        st.success("Inputs saved. Proceed to Draft & Download.")


def render_tab_draft(model_choice: str) -> None:
    st.markdown('<div class="section-header">Generate Draft & Download</div>',
                unsafe_allow_html=True)

    if "extraction_result" not in st.session_state:
        st.info("Complete Steps 1 to 3 first.")
        return
    if "user_inputs_text" not in st.session_state:
        st.info("Save your inputs in Step 3 first.")
        return

    extraction   = st.session_state["extraction_result"]
    user_inputs  = st.session_state.get("user_inputs_text", "")
    proc_flags   = st.session_state.get("procedural_flags", [])
    matched_laws = st.session_state.get("matched_laws", "")
    assessee     = st.session_state.get("assessee_name", "[ASSESSEE NAME]")
    pan          = st.session_state.get("assessee_pan", "[PAN]")
    city         = st.session_state.get("assessee_city", "[CITY]")
    ar_name      = st.session_state.get("ar_name", "")

    # Build procedural defect block
    proc_block = ""
    if proc_flags:
        proc_block = "\n\n".join(
            f for f in proc_flags if "void" in f.lower() or "time-barred" in f.lower()
        )

    # Detect faceless proceedings
    faceless = "nfac" in extraction.lower() or "144b" in extraction.lower() or \
               "faceless" in extraction.lower()
    st.session_state["faceless_mode"] = faceless

    if faceless:
        st.info("Faceless Assessment detected (Section 144B / NFAC). "
                "Draft will address the National Faceless Assessment Centre.")
        authority_address = (
            "The National Faceless Assessment Centre,\n"
            "Income Tax Department,\n"
            "New Delhi."
        )
    else:
        authority_address = (
            "The Assessing Officer,\n"
            "[Ward / Circle / Range],\n"
            "Income Tax Department,\n"
            f"{city}."
        )

    ar_line = f"Authorised Representative: {ar_name}" if ar_name else ""

    evidence_rows = st.session_state.get("evidence_tracker_rows", [])
    annexure = build_annexure_schedule(user_inputs, evidence_rows if feature_enabled("EVIDENCE_TRACKER") else None)
    ais_recon = st.session_state.get("ais_recon_result") or {}
    ais_recon_block = ""
    if feature_enabled("AIS_RECON") and ais_recon:
        flags = "; ".join(ais_recon.get("flags", [])) or "No AIS/TIS discrepancy flags."
        ais_recon_block = (
            "\n\nAIS/TIS RECONCILIATION SUMMARY (Beta):\n"
            f"Extracted notice quantum: {ais_recon.get('notice_quantum')}\n"
            f"Total AIS - TIS delta: {ais_recon.get('total_delta')}\n"
            f"Flags: {flags}\n"
            "Use only where the figures have been verified from the assessee portal."
        )

    form_68_enabled = False
    form_68_draft = ""
    if feature_enabled("FORM_68_PATHWAY") and section_270a_detected(extraction):
        st.markdown("**Section 270AA Immunity Workflow**")
        st.info("Section 270A appears in the extraction. Enable this only after checking payment, appeal status, and statutory Form 68 eligibility.")
        form_68_enabled = st.toggle(
            "Prepare Form 68 / Section 270AA immunity draft",
            value=st.session_state.get("form_68_enabled", False),
            key="form_68_toggle",
        )
        if form_68_enabled:
            form_68_draft = build_form_68_immunity_draft(extraction, assessee, pan)
            st.session_state["form_68_enabled"] = True
            st.session_state["form_68_draft"] = form_68_draft
            with st.expander("Preview Form 68 pathway text", expanded=False):
                st.text_area("Immunity request draft", value=form_68_draft, height=220)
        else:
            st.session_state["form_68_enabled"] = False
            st.session_state.pop("form_68_draft", None)

    if feature_enabled("PENALTY_WORKFLOW") and penalty_proceeding_detected(extraction):
        st.markdown("**Penalty Proceedings Workflow**")
        st.info(
            "Penalty proceeding detected. Use this to prepare abeyance request, "
            "defective notice objections, merits reply, and condonation / reasonable-cause submissions."
        )
        with st.expander("Penalty and Condonation Inputs", expanded=False):
            p1, p2, p3 = st.columns(3)
            with p1:
                penalty_section = st.text_input("Penalty section", value="270A", key="penalty_section")
                penalty_notice_date = st.date_input("Penalty notice date", value=date.today(), key="penalty_notice_date")
                penalty_due_date = st.date_input("Penalty reply due date", value=date.today(), key="penalty_due_date")
            with p2:
                quantum_appeal_pending = st.checkbox("Quantum Form 35 appeal is pending", value=False, key="penalty_quantum_pending")
                form35_date = st.text_input("Form 35 filing date", value="", key="penalty_form35_date")
                form35_ack = st.text_input("Form 35 acknowledgment no.", value="", key="penalty_form35_ack")
            with p3:
                requested_abeyance = st.checkbox("Request penalty proceedings be kept in abeyance", value=quantum_appeal_pending, key="penalty_abeyance")
                defective_notice = st.checkbox("Raise defective/vague notice objection", value=True, key="penalty_defective_notice")
                wants_273b = st.checkbox("Include section 273B reasonable cause", value=False, key="penalty_273b")
                delayed_reply = st.checkbox("Include condonation for delayed response", value=False, key="penalty_condonation")

            delay_reason = st.text_area(
                "Reason for delay / reasonable cause",
                placeholder="Example: portal issue, illness, records not available, accountant resignation, consultant change, genuine confusion, NFAC communication issue.",
                key="penalty_delay_reason",
                height=90,
            )
            penalty_metadata = build_penalty_metadata_text(
                penalty_section=penalty_section,
                notice_date_value=penalty_notice_date,
                reply_due_value=penalty_due_date,
                quantum_appeal_pending=quantum_appeal_pending,
                form35_date=form35_date,
                form35_ack=form35_ack,
                requested_abeyance=requested_abeyance,
                defective_notice=defective_notice,
                delayed_reply=delayed_reply,
                delay_reason=delay_reason,
                wants_273b=wants_273b,
                wants_270aa=form_68_enabled,
            )
            st.session_state["penalty_metadata"] = penalty_metadata
            st.caption("Penalty reply will still include a merits response even when abeyance is requested.")

            if st.button("Generate Penalty Response Package", use_container_width=True):
                with st.spinner("Drafting penalty response package..."):
                    try:
                        penalty_prompt = PENALTY_PROCEEDINGS_PROMPT.format(
                            penalty_metadata=penalty_metadata,
                            extraction=redact_sensitive_text_for_ai(extraction[:9000])[0],
                            user_inputs=redact_sensitive_text_for_ai(user_inputs[:7000])[0],
                            verified_laws=get_relevant_case_laws(
                                extraction + "\n" + user_inputs + "\n" + penalty_metadata
                            ),
                        )
                        penalty_prompt = anonymise_prompt_identity(penalty_prompt, assessee, pan)
                        penalty_response = call_gemini(
                            model_name=model_choice,
                            prompt=penalty_prompt,
                            temperature=0.12,
                            max_tokens=8192,
                            step="penalty_response",
                            username=st.session_state.get("auth_username", "anonymous"),
                        )
                        penalty_response = make_portal_safe_text(
                            restore_local_identity(penalty_response, assessee, pan)
                        )
                        st.session_state["penalty_response_draft"] = penalty_response
                        write_audit_trail(
                            "PENALTY_RESPONSE_GENERATED", "penalty",
                            f"chars={len(penalty_response)} section={penalty_section}",
                            username=st.session_state.get("auth_username", "anonymous"),
                        )
                        st.success("Penalty response package generated.")
                    except APICallError as exc:
                        st.error(f"API Error: {exc}")

            penalty_response = st.session_state.get("penalty_response_draft", "")
            if penalty_response:
                st.text_area("Penalty response preview", value=penalty_response, height=280)
                pdl1, pdl2 = st.columns(2)
                with pdl1:
                    st.download_button(
                        "Download Penalty Response (.txt)",
                        data=penalty_response,
                        file_name=f"RKMuley_v9_Penalty_Response_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with pdl2:
                    if DOCX_AVAILABLE:
                        st.download_button(
                            "Download Penalty Response (.docx)",
                            data=build_word_package(penalty_response, "", extraction),
                            file_name=f"RKMuley_v9_Penalty_Response_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

    # Preliminary block options (litigation-safe)
    st.markdown("**Preliminary Submissions - Compliance History**")
    st.markdown(
        '<div class="warning-box">Only check items you can independently confirm. '
        'An incorrect statement in preliminary submissions can be used adversely against the assessee.</div>',
        unsafe_allow_html=True,
    )
    confirm_timely_itr = st.checkbox(
        "Assessee has filed all ITRs within due date (or before this notice)",
        value=False,
        key="prelim_itr",
    )
    confirm_taxes_paid = st.checkbox(
        "All advance tax and self-assessment tax has been paid",
        value=False,
        key="prelim_tax",
    )
    confirm_cooperative = st.checkbox(
        "Assessee is fully cooperative and will produce all documents",
        value=True,
        key="prelim_coop",
    )

    prelim_lines = []
    if confirm_timely_itr:
        prelim_lines.append(
            "I have been filing my Income Tax Returns within the prescribed due dates for all Assessment Years."
        )
    if confirm_taxes_paid:
        prelim_lines.append(
            "I have duly paid all advance tax and self-assessment tax as applicable."
        )
    if confirm_cooperative:
        prelim_lines.append(
            "I submit all information and documents in a spirit of full cooperation with these proceedings "
            "and undertake to produce any further documents called for."
        )

    preliminary_block = " ".join(prelim_lines) if prelim_lines else (
        "I submit this response within the due date and remain committed to full cooperation with these proceedings."
    )

    if st.button("Generate Draft Response", type="primary"):
        with st.spinner("Drafting portal-ready response..."):
            try:
                # Extract key fields for prompt
                din_m    = re.search(r"DIN.*?:\s*(.+)", extraction, re.IGNORECASE)
                date_m   = re.search(r"Notice Date.*?:\s*(.+)", extraction, re.IGNORECASE)
                ay_m     = re.search(r"Assessment Year.*?:\s*(.+)", extraction, re.IGNORECASE)
                sec_m    = re.search(r"Primary Section Invoked.*?:\s*(.+)", extraction, re.IGNORECASE)
                ntype_m  = re.search(r"Notice Type.*?:\s*(.+)", extraction, re.IGNORECASE)
                due_m    = re.search(r"Due Date.*?:\s*(.+)", extraction, re.IGNORECASE)

                prompt = DRAFTING_PROMPT.format(
                    notice_type       = ntype_m.group(1).strip() if ntype_m else "Notice",
                    din               = din_m.group(1).strip() if din_m else "N/A",
                    notice_date       = date_m.group(1).strip() if date_m else "N/A",
                    ay                = ay_m.group(1).strip() if ay_m else "N/A",
                    sections          = sec_m.group(1).strip() if sec_m else "N/A",
                    due_date          = due_m.group(1).strip() if due_m else "N/A",
                    authority_address = authority_address,
                    assessee_name     = assessee,
                    pan               = pan,
                    city              = city,
                    ar_line           = ar_line,
                    preliminary_block = preliminary_block,
                    procedural_defect_block = proc_block,
                    issues_block      = "[AI will generate based on extraction and user inputs below]",
                    prayer_items      = "[AI will generate based on issues]",
                    current_date      = datetime.now().strftime("%d-%m-%Y"),
                    verified_laws     = matched_laws,
                    extracted_issues  = extraction,
                    user_inputs       = user_inputs + ais_recon_block + (
                        "\n\nSECTION 270AA / FORM 68 PATHWAY REQUESTED:\n"
                        "Include a separate, eligibility-qualified immunity pathway note. "
                        "Do not state that immunity is automatic; mention payment and no-appeal conditions."
                        if form_68_enabled else ""
                    ),
                    annexure_schedule = annexure,
                )

                prompt = anonymise_prompt_identity(prompt, assessee, pan)
                draft = call_gemini(
                    model_name=model_choice,
                    prompt=prompt,
                    temperature=0.15,
                    max_tokens=8192,
                    step="drafting",
                    username=st.session_state.get("auth_username", "anonymous"),
                )
                draft = clean_markdown_from_draft(draft)
                draft = restore_local_identity(draft, assessee, pan)
                st.session_state["draft_response"] = draft

                # Generate cover note
                if FEATURES["COVER_NOTE_GEN"]:
                    cover_extraction, _ = redact_sensitive_text_for_ai(extraction[:3000])
                    cover_inputs, _ = redact_sensitive_text_for_ai(user_inputs[:2000])
                    cover = call_gemini(
                        model_name=model_choice,
                        prompt=COVER_NOTE_PROMPT.format(
                            extracted_issues=cover_extraction,
                            user_inputs=cover_inputs,
                            current_date=datetime.now().strftime("%d-%m-%Y"),
                        ),
                        temperature=0.1,
                        max_tokens=1500,
                        step="cover_note",
                        username=st.session_state.get("auth_username", "anonymous"),
                    )
                    cover = restore_local_identity(cover, assessee, pan)
                    st.session_state["cover_note"] = cover

                write_audit_trail("DRAFT_GENERATED", "draft", f"chars={len(draft)}",
                                  username=st.session_state.get("auth_username", "anonymous"))

                # Run Pass D admission scan immediately after draft generation
                # DraftRiskChecker already imported at module top â€” no inner import needed
                adm_findings_raw = DraftRiskChecker.run_passes_a_to_d(draft, extraction)
                admission_findings = [
                    {"phrase": f["issue"], "risk_label": f["issue"].split(":")[0], "context": f["context"]}
                    for f in adm_findings_raw if f["pass"] == "D"
                ]
                st.session_state["admission_findings"] = admission_findings

                if admission_findings:
                    st.warning(
                        f"Pass D flagged {len(admission_findings)} admission risk(s) in draft. "
                        "Review in Risk Checker before submission."
                    )

                st.success(f"Draft generated ({len(draft):,} chars). Review below.")

                if feature_enabled("NOTICE_STORE") and rbac_check("can_notice_store"):
                    masked_extraction, _ = redact_sensitive_text_for_ai(extraction)
                    masked_draft, _ = redact_sensitive_text_for_ai(draft)
                    masked_cover, _ = redact_sensitive_text_for_ai(st.session_state.get("cover_note", ""))
                    row_id = persist_notice_store(
                        pan=redact_sensitive_text_for_ai(pan)[0],
                        ay=ay_m.group(1).strip() if ay_m else "",
                        notice_type=ntype_m.group(1).strip() if ntype_m else "Notice",
                        extraction=masked_extraction,
                        draft=masked_draft,
                        cover=masked_cover,
                        proc_flags=proc_flags,
                        risk_score=0,
                        success_score=0,
                        username=st.session_state.get("auth_username", "anonymous"),
                    )
                    if row_id > 0:
                        st.session_state["notice_store_id"] = row_id

            except APICallError as exc:
                st.error(f"API Error: {exc}")
                return

    if feature_enabled("FORM35_APPEALS"):
        st.divider()
        with st.expander("Form 35 / First Appeal Package", expanded=False):
            st.caption(
                "Prepare Statement of Facts, Grounds of Appeal, personal hearing request, "
                "and optional Rule 46A / stay / condonation drafts. Form 35 is for CIT(A) "
                "or JCIT(A); ITAT appeal filing generally uses Form 36."
            )
            c1, c2, c3 = st.columns(3)
            with c1:
                appeal_forum = st.selectbox("Appeal forum", ["CIT(A)", "JCIT(A)"], key="form35_forum")
                order_section = st.text_input("Order section", value="143(3)", key="form35_order_section")
                assessed_income = st.number_input(
                    "Assessed income as per order",
                    min_value=0.0,
                    value=0.0,
                    step=10000.0,
                    key="form35_assessed_income",
                )
            with c2:
                order_date_value = st.date_input("Order date", value=date.today(), key="form35_order_date")
                service_date_value = st.date_input("Date of service", value=date.today(), key="form35_service_date")
                demand_amount = st.text_input("Disputed demand", value="", key="form35_demand")
            with c3:
                other_matter = st.checkbox("Other matter appeal", value=False, key="form35_other_matter")
                wants_rule46a = st.checkbox("Include Rule 46A application", value=False, key="form35_rule46a")
                wants_stay = st.checkbox("Include stay of demand request", value=True, key="form35_stay")
                wants_personal_hearing = st.checkbox("Request personal hearing", value=True, key="form35_hearing")

            fee_amount = compute_form35_fee(float(assessed_income), other_matter)
            appeal_deadline = service_date_value + timedelta(days=30)
            delay_days = max(0, (date.today() - appeal_deadline).days)
            st.info(
                f"Computed Form 35 fee: Rs. {fee_amount}. "
                f"Indicative 30-day appeal deadline from service date: {appeal_deadline.strftime('%d-%m-%Y')}. "
                f"Delay today: {delay_days} day(s)."
            )

            appeal_metadata = build_form35_metadata_text(
                forum=appeal_forum,
                order_section=order_section,
                order_date_value=order_date_value,
                service_date_value=service_date_value,
                assessed_income=float(assessed_income),
                demand_amount=demand_amount,
                fee_amount=fee_amount,
                delay_days=delay_days,
                wants_rule46a=wants_rule46a,
                wants_stay=wants_stay,
                wants_personal_hearing=wants_personal_hearing,
            )
            st.session_state["form35_metadata"] = appeal_metadata

            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("Generate Form 35 Package", use_container_width=True):
                    with st.spinner("Drafting Form 35 appeal package..."):
                        try:
                            form35_prompt = FORM35_APPEAL_PROMPT.format(
                                appeal_metadata=appeal_metadata,
                                extraction=redact_sensitive_text_for_ai(extraction[:8000])[0],
                                user_inputs=redact_sensitive_text_for_ai(user_inputs[:6000])[0],
                                verified_laws=matched_laws,
                            )
                            form35_prompt = anonymise_prompt_identity(form35_prompt, assessee, pan)
                            package = call_gemini(
                                model_name=model_choice,
                                prompt=form35_prompt,
                                temperature=0.12,
                                max_tokens=8192,
                                step="form35_appeal",
                                username=st.session_state.get("auth_username", "anonymous"),
                            )
                            package = make_portal_safe_text(restore_local_identity(package, assessee, pan))
                            st.session_state["form35_package"] = package
                            write_audit_trail(
                                "FORM35_GENERATED", "appeal",
                                f"chars={len(package)} fee={fee_amount} delay={delay_days}",
                                username=st.session_state.get("auth_username", "anonymous"),
                            )
                            st.success("Form 35 package generated.")
                        except APICallError as exc:
                            st.error(f"API Error: {exc}")
            with col_b:
                if feature_enabled("ITAT_NOTICE_HELPER") and st.button(
                    "Generate ITAT Notice Response", use_container_width=True
                ):
                    with st.spinner("Drafting ITAT-stage notice response..."):
                        try:
                            itat_prompt = ITAT_NOTICE_RESPONSE_PROMPT.format(
                                appeal_metadata=appeal_metadata,
                                extraction=redact_sensitive_text_for_ai(extraction[:8000])[0],
                                user_inputs=redact_sensitive_text_for_ai(user_inputs[:6000])[0],
                                verified_laws=matched_laws,
                            )
                            itat_prompt = anonymise_prompt_identity(itat_prompt, assessee, pan)
                            itat_response = call_gemini(
                                model_name=model_choice,
                                prompt=itat_prompt,
                                temperature=0.12,
                                max_tokens=6000,
                                step="itat_notice_response",
                                username=st.session_state.get("auth_username", "anonymous"),
                            )
                            itat_response = make_portal_safe_text(
                                restore_local_identity(itat_response, assessee, pan)
                            )
                            st.session_state["itat_notice_response"] = itat_response
                            write_audit_trail(
                                "ITAT_NOTICE_RESPONSE_GENERATED", "appeal",
                                f"chars={len(itat_response)}",
                                username=st.session_state.get("auth_username", "anonymous"),
                            )
                            st.success("ITAT-stage response generated.")
                        except APICallError as exc:
                            st.error(f"API Error: {exc}")

            form35_package = st.session_state.get("form35_package", "")
            if form35_package:
                sof_match = re.search(
                    r"STATEMENT OF FACTS.*?(?=GROUNDS OF APPEAL|PERSONAL HEARING|$)",
                    form35_package,
                    flags=re.IGNORECASE | re.DOTALL,
                )
                sof_words = count_words(sof_match.group(0)) if sof_match else 0
                st.caption(f"Statement of Facts word count check: {sof_words} / 1000 approx.")
                st.text_area("Form 35 package preview", value=form35_package, height=280)
                d1, d2 = st.columns(2)
                with d1:
                    st.download_button(
                        "Download Form 35 Package (.txt)",
                        data=form35_package,
                        file_name=f"RKMuley_v9_Form35_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                        mime="text/plain",
                        use_container_width=True,
                    )
                with d2:
                    if DOCX_AVAILABLE:
                        st.download_button(
                            "Download Form 35 Package (.docx)",
                            data=build_word_package(form35_package, "", extraction),
                            file_name=f"RKMuley_v9_Form35_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                        )

            itat_response = st.session_state.get("itat_notice_response", "")
            if itat_response:
                st.text_area("ITAT notice response preview", value=itat_response, height=220)
                st.download_button(
                    "Download ITAT Notice Response (.txt)",
                    data=itat_response,
                    file_name=f"RKMuley_v9_ITAT_Response_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )

    # Display draft
    if "draft_response" in st.session_state:
        draft = st.session_state["draft_response"]
        st.markdown(
            f'<div class="output-box"><pre style="white-space:pre-wrap;font-size:0.88rem;">'
            f'{draft}</pre></div>',
            unsafe_allow_html=True,
        )

        # Downloads
        dl1, dl2, dl3, dl4 = st.columns(4)
        with dl1:
            st.download_button(
                "Draft (.txt)",
                data=make_portal_safe_text(draft),
                file_name=f"RKMuley_v9_Draft_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                mime="text/plain",
                use_container_width=True,
            )
        with dl2:
            cover = st.session_state.get("cover_note", "")
            if cover:
                cover_len = len(cover)
                color = "green" if cover_len <= PORTAL_TEXTBOX_LIMIT else "red"
                st.markdown(
                    f'<div style="background:#f8f9fa;border-radius:6px;padding:0.5rem;'
                    f'text-align:center;font-size:0.82rem;">'
                    f'Cover Note: <strong style="color:{color};">{cover_len:,} chars</strong> '
                    f'(limit: {PORTAL_TEXTBOX_LIMIT:,})</div>',
                    unsafe_allow_html=True,
                )
                st.download_button(
                    "Cover Note (.txt)",
                    data=make_portal_safe_text(cover),
                    file_name=f"RKMuley_v9_CoverNote_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
        with dl3:
            parts = split_for_portal(draft)
            if len(parts) > 1:
                st.warning(f"Draft is long â€” split into {len(parts)} parts for portal.")
                for i, part in enumerate(parts, 1):
                    st.download_button(
                        f"Part {i} of {len(parts)}",
                        data=make_portal_safe_text(part),
                        file_name=f"RKMuley_v9_Part{i}_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                        mime="text/plain",
                        key=f"part_dl_{i}",
                        use_container_width=True,
                    )
            else:
                st.caption("Response fits in one PDF - no split needed.")
        with dl4:
            due_date = parse_due_date(extraction)
            if feature_enabled("ICS_EXPORT", feature_enabled("DEADLINE_CALENDAR")) and due_date:
                ics_data = generate_ics_invite(
                    due_date=due_date,
                    summary=f"Tax notice response due - {pan}",
                    description=f"{FIRM_NAME}: response deadline for {assessee}. Verify portal submission before due date.",
                )
                st.download_button(
                    "Download Calendar Invite",
                    data=ics_data,
                    file_name=f"RKMuley_v9_Deadline_{due_date.strftime('%Y%m%d')}.ics",
                    mime="text/calendar",
                    use_container_width=True,
                )
            else:
                st.caption("Calendar invite available after due date is extracted.")

            form68 = st.session_state.get("form_68_draft", "")
            if feature_enabled("FORM_68_PATHWAY") and form68:
                st.download_button(
                    "Form 68 Pathway (.txt)",
                    data=make_portal_safe_text(form68),
                    file_name=f"RKMuley_v9_Form68_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
            if DOCX_AVAILABLE:
                cover = st.session_state.get("cover_note", "")
                st.download_button(
                    "Portal Word Package (.docx)",
                    data=build_word_package(draft, cover, extraction),
                    file_name=f"RKMuley_v9_PortalPackage_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )

        # Pre-submission checklist
        st.divider()
        st.markdown("**Pre-Submission Checklist**")
        col1, col2, col3 = st.columns(3)
        checks = [
            ["All issues from notice addressed", "Response in first person throughout",
             "No fabricated case laws in draft", "Procedural defects raised as preliminary submissions"],
            [f"Cover note under {PORTAL_TEXTBOX_LIMIT:,} chars",
             "No markdown symbols (*, #, **)", "Assessee name and PAN verified",
             "Without Prejudice wrapper for partial admissions"],
            ["Annexure Schedule complete and numbered", "Documents ready for PDF scan",
             "Response deadline date confirmed", "Personal hearing request included"],
        ]
        for col, items in zip([col1, col2, col3], checks):
            with col:
                for item in items:
                    st.checkbox(item, key=f"chk_{hash(item)}")


def render_tab_procedural() -> None:
    st.markdown('<div class="section-header">Procedural Validity Audit</div>',
                unsafe_allow_html=True)
    st.markdown("""<div class="portal-note">
Deterministic checks â€” zero LLM, zero hallucination risk. Identifies threshold-level defects
that can get a notice quashed before engaging on merits.</div>""", unsafe_allow_html=True)

    if "extraction_result" not in st.session_state:
        st.info("Complete Step 1 first.")
        return

    flags = st.session_state.get("procedural_flags", [])

    if not flags:
        st.success("No major procedural defects detected. Respond on merits.")
    else:
        st.error(f"{len(flags)} procedural defect(s) identified.")
        for i, flag in enumerate(flags, 1):
            severity = "Critical" if "TIME-BARRED" in flag or "ABSENT" in flag else "Requires Verification"
            with st.expander(f"Defect {i}: {severity}", expanded=True):
                st.markdown(f'<div class="proc-defect-box">{flag}</div>',
                            unsafe_allow_html=True)

    # Export report
    report = (
        "R K MULEY & CO - PROCEDURAL AUDIT REPORT v9.0\n"
        f"Generated: {datetime.now().strftime('%d-%m-%Y %H:%M')}\n"
        + "=" * 60 + "\n\n"
        + ("DEFECTS:\n\n" + "\n\n".join(flags) if flags else "No defects detected.")
        + "\n\n" + "=" * 60 + "\n\n"
        + "APPLICABLE CASE LAWS:\n\n"
        + "\n".join(VERIFIED_CASE_LAWS.get("din_circular", []))
        + "\n\n" + "\n".join(VERIFIED_CASE_LAWS.get("sec_148a_procedural", []))
        + "\n\n" + "\n".join(VERIFIED_CASE_LAWS.get("sec_149_limitation", []))
    )
    st.download_button(
        "Download Procedural Audit Report",
        data=report,
        file_name=f"RKMuley_v9_ProcAudit_{datetime.now().strftime('%Y%m%d_%H%M')}.txt",
        mime="text/plain",
    )


def render_tab_risk(model_choice: str) -> None:
    st.markdown('<div class="section-header">Risk Checker & Intelligence Scores</div>',
                unsafe_allow_html=True)

    if not rbac_check("can_risk_checker"):
        st.warning("CA / Manager access required for Risk Checker.")
        return
    if "draft_response" not in st.session_state or "extraction_result" not in st.session_state:
        st.info("Complete Steps 1â€“4 first.")
        return

    draft      = st.session_state["draft_response"]
    extraction = st.session_state["extraction_result"]
    proc_flags = st.session_state.get("procedural_flags", [])
    adm_finds  = st.session_state.get("admission_findings", [])

    # â”€â”€ Run hallucination check â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    user_cits = []
    if st.session_state.get("user_inputs_text"):
        for line in st.session_state["user_inputs_text"].splitlines():
            if "v." in line and re.search(r"\d{4}", line):
                user_cits.append(line.strip())
    hall_report = HallucinationGuard.check(draft, user_cits)

    # â”€â”€ Probability Scores â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    vault_similar = _vault.get_similar(
        keywords=["68", "148", "tds", "69", "penalty", "reopening"][:3]
    )

    npp = st.session_state.get("npp_result") or NoticeProbabilityPredictor.predict(
        extraction, proc_flags, [], adm_finds, vault_similar
    )
    st.session_state["npp_result"] = npp

    rss = st.session_state.get("rss_result") or ReplySuccessScorer.score(
        draft, proc_flags, adm_finds, hall_report, [],
        vault_similar, st.session_state.get("matched_laws", "")
    )
    st.session_state["rss_result"] = rss

    # Display scores â€” using FIXED key names (score, color, grade, verdict)
    st.markdown("### Probability Scores")
    sc1, sc2 = st.columns(2)
    with sc1:
        npp_color = npp["color"]
        st.markdown(
            f'<div style="background:{npp_color}18;border:2px solid {npp_color};'
            f'border-radius:12px;padding:1.2rem;text-align:center;">'
            f'<div style="font-size:2.6rem;font-weight:900;color:{npp_color};">{npp["score"]}%</div>'
            f'<div style="font-weight:700;color:{npp_color};">{npp["verdict"]}</div>'  # FIX: was "band"
            f'<div style="font-size:0.72rem;color:#666;margin-top:6px;">Notice Threat Score</div>'
            f'</div>',
            unsafe_allow_html=True,
        )
    with sc2:
        rss_color = rss["color"]
        st.markdown(
            f'<div style="background:{rss_color}18;border:2px solid {rss_color};'
            f'border-radius:12px;padding:1.2rem;text-align:center;">'
            f'<div style="font-size:2.6rem;font-weight:900;color:{rss_color};">{rss["score"]}%</div>'  # FIX: was "total"
            f'<div style="font-weight:700;color:{rss_color};">{rss["grade"]}</div>'   # FIX: was "band"
            f'<div style="font-size:0.72rem;color:#666;margin-top:6px;">Reply Success Score</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    # Score breakdown
    with st.expander("Reply Score Breakdown", expanded=False):
        for dim, data in rss["breakdown"].items():
            bar_c = "#2e7d32" if data["raw"] >= 70 else ("#e65100" if data["raw"] >= 40 else "#c62828")
            st.markdown(
                f'<div style="margin:3px 0;padding:5px 8px;background:#f8f9fa;border-radius:5px;">'
                f'<span style="min-width:220px;display:inline-block;font-size:0.82rem;">'
                f'{dim}</span>'
                f'<span style="background:{bar_c}22;border:1px solid {bar_c};border-radius:4px;'
                f'padding:1px 10px;font-weight:700;color:{bar_c};font-size:0.82rem;">{data["raw"]}%</span>'
                f' <span style="font-size:0.74rem;color:#666;">{data["note"]}</span></div>',
                unsafe_allow_html=True,
            )

    st.divider()

    # â”€â”€ Passes Aâ€“D â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    st.markdown("### Passes Aâ€“D: Deterministic Checks")
    # run_passes_a_to_d() NOW EXISTS (was missing in v7)
    det_findings = DraftRiskChecker.run_passes_a_to_d(draft, extraction)

    pass_labels = {
        "A": "Section Number Validation",
        "B": "Citation Format Check",
        "C": "Quantum Consistency Check",
        "D": "Admission Language Scan",
    }
    for pid, plabel in pass_labels.items():
        pf = [f for f in det_findings if f["pass"] == pid]
        with st.expander(
            f"Pass {pid}: {plabel} - {'%d issue(s)' % len(pf) if pf else 'CLEAN'}",
            expanded=bool(pf),
        ):
            if not pf:
                st.success(f"Pass {pid} passed â€” no issues found.")
            else:
                for finding in pf:
                    sc_col = "#c62828" if finding["severity"] == "High" else "#e65100"
                    st.markdown(
                        f'<div style="background:{sc_col}0d;border-left:4px solid {sc_col};'
                        f'padding:0.6rem 0.9rem;border-radius:4px;margin:4px 0;">'
                        f'<strong style="color:{sc_col};">{finding["severity"]}</strong>: '
                        f'{finding["issue"]}<br>'
                        f'<code style="font-size:0.78rem;">{finding["context"][:150]}</code><br>'
                        f'<span style="font-size:0.78rem;color:#1565c0;">{finding["action"]}</span>'
                        f'</div>',
                        unsafe_allow_html=True,
                    )

    st.divider()

    # â”€â”€ Pass E: LLM Adversarial Review â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
    st.markdown("### Pass E: LLM Adversarial Self-Review")
    st.markdown('<div class="portal-note">Asks the AI to audit its own draft as an adversarial '
                'reviewer. Catches fabricated sections, invented citations, and logical '
                'contradictions that deterministic passes cannot detect.</div>',
                unsafe_allow_html=True)

    if st.button("Run Pass E - Adversarial Review", type="primary", key="pass_e_btn"):
        with st.spinner("Running adversarial self-review..."):
            # run_pass_e() NOW EXISTS (was missing in v7)
            pe = HallucinationGuard.run_pass_e(
                api_key=get_api_key(),
                model_name=model_choice,
                draft=redact_sensitive_text_for_ai(draft)[0],
                extraction=redact_sensitive_text_for_ai(extraction)[0],
            )
            st.session_state["pass_e_result"] = pe
            write_audit_trail("PASS_E_RUN", "draft",
                              f"risk={pe.get('hallucination_risk', '?')}",
                              username=st.session_state.get("auth_username", "anonymous"))

    pe = st.session_state.get("pass_e_result")
    if pe:
        if not pe.get("layer3_available") or pe.get("error"):
            st.error(f"Pass E failed: {pe.get('error', 'Unknown error')}")
        else:
            if pe.get("parse_warning"):
                st.warning("Pass E returned malformed JSON, so the app used a manual-review fallback summary.")
            risk = pe.get("hallucination_risk", "Unknown")
            rc = {"Low": "#2e7d32", "Medium": "#e65100", "High": "#c62828"}.get(risk, "#666")
            st.markdown(
                f'<div style="background:{rc}18;border:2px solid {rc};border-radius:10px;'
                f'padding:0.9rem 1.2rem;margin:0.8rem 0;">'
                f'<strong style="color:{rc};font-size:1.05rem;">Hallucination Risk: {risk}</strong><br>'
                f'<span style="font-size:0.88rem;">{pe.get("overall_verdict", "")}</span></div>',
                unsafe_allow_html=True,
            )
            pe_issues = pe.get("issues", [])
            if not pe_issues:
                st.success("Pass E: No hallucinations detected by adversarial reviewer.")
            else:
                st.error(f"Pass E flagged {len(pe_issues)} issue(s):")
                for issue in pe_issues:
                    st.markdown(
                        f'<div class="warning-box"><strong>{issue.get("type","")}</strong> '
                        f'({issue.get("severity","")}) - {issue.get("detail","")}</div>',
                        unsafe_allow_html=True,
                    )

            positives = pe.get("positives", [])
            if positives:
                with st.expander("What the reviewer found well-done", expanded=False):
                    for p in positives:
                        st.markdown(f"- {p}")

            if pe.get("recommendation"):
                st.info(f"Recommendation: {pe['recommendation']}")


def render_tab_vault() -> None:
    st.markdown('<div class="section-header">CA Knowledge Vault</div>',
                unsafe_allow_html=True)
    if not rbac_check("can_vault"):
        st.warning("CA / Manager access required.")
        return

    tab_log, tab_view, tab_analytics = st.tabs(["Log Outcome", "View Records", "Analytics"])

    with tab_log:
        st.markdown("**Log a new notice outcome**")
        with st.form("vault_add_form"):
            c1, c2 = st.columns(2)
            with c1:
                assessee = st.text_input("Assessee (anonymous if preferred)", placeholder="Firm A / Individual B")
                ay       = st.text_input("Assessment Year", placeholder="2022-23")
                ntype    = st.selectbox("Notice Type", ["143(1)", "143(2)", "148", "148A", "271", "270A", "Other"])
                sections = st.text_input("Sections involved", placeholder="68, 69, 56(2)(viib)")
            with c2:
                issue_type = st.text_input("Issue Type", placeholder="Cash credit, Unexplained investment...")
                quantum    = st.number_input("Demand Amount (Rs. Lakhs)", min_value=0.0, step=0.5)
                outcome    = st.selectbox("Outcome", ["Pending", "Win", "Loss", "Settle"])
                forum      = st.selectbox("Forum", ["AO Level", "CIT(A)", "ITAT", "HC", "SC"])
            strategy = st.text_area("Defence Strategy Used", height=80)
            lessons  = st.text_area("Lessons / Key Learning", height=80)
            tags     = st.text_input("Tags (comma-separated)", placeholder="faceless, cash, penalty")

            if st.form_submit_button("Save to Vault", type="primary"):
                try:
                    row_id = _vault.add_entry({
                        "assessee": assessee, "ay": ay, "notice_type": ntype,
                        "sections": sections, "issue_type": issue_type,
                        "quantum_lakh": quantum, "outcome": outcome, "forum": forum,
                        "strategy": strategy, "lessons": lessons, "tags": tags,
                        "created_by": st.session_state.get("auth_username", ""),
                    })
                    write_audit_trail("VAULT_ENTRY_ADDED", "vault", f"id={row_id}",
                                      username=st.session_state.get("auth_username", "anonymous"))
                    st.success(f"Saved to vault (record #{row_id}).")
                except Exception as exc:
                    st.error(f"Error saving: {exc}")

    with tab_view:
        entries = _vault.all_entries(limit=200)
        if not entries:
            st.info("No vault records yet. Use 'Log Outcome' to add your first entry.")
        else:
            st.caption(f"{len(entries)} records")
            import pandas as pd
            df = pd.DataFrame(entries)
            display_cols = ["id", "ts_added", "ay", "notice_type", "sections",
                            "issue_type", "quantum_lakh", "outcome", "forum"]
            st.dataframe(df[[c for c in display_cols if c in df.columns]], use_container_width=True)

            csv_data = _vault.export_csv()
            st.download_button("Export CSV", data=csv_data,
                               file_name=f"RKMuley_vault_{datetime.now().strftime('%Y%m%d')}.csv",
                               mime="text/csv")

            if rbac_check("can_delete_vault"):
                del_id = st.number_input("Delete record ID", min_value=1, step=1, value=1)
                if st.button("Delete Record", type="secondary"):
                    try:
                        _vault.delete(int(del_id))
                        write_audit_trail("VAULT_ENTRY_DELETED", "vault", f"id={del_id}",
                                          username=st.session_state.get("auth_username", "anonymous"))
                        st.success(f"Record #{del_id} deleted.")
                        st.rerun()
                    except Exception as exc:
                        st.error(f"Delete failed: {exc}")

    with tab_analytics:
        if not FEATURES["VAULT_ANALYTICS"]:
            st.info("Analytics disabled.")
            return
        stats = _vault.stats()
        c1, c2, c3, c4 = st.columns(4)
        with c1: st.metric("Total Records", stats["total"])
        with c2: st.metric("Wins", stats["wins"])
        with c3: st.metric("Losses", stats["losses"])
        with c4: st.metric("Win Rate", f"{stats['win_rate_pct']}%")


def render_tab_health() -> None:
    st.markdown('<div class="section-header">System Health</div>',
                unsafe_allow_html=True)
    if not rbac_check("can_see_health"):
        st.warning("Admin access required.")
        return

    col1, col2 = st.columns(2)
    with col1:
        if st.button("Run System Health Check", type="primary"):
            health = get_system_health(_db_engine)
            st.json(health)
    with col2:
        if st.button("Run Smoke Tests"):
            results = smoke_test()
            for test, passed in results.items():
                st.markdown(f"**{test}**: {'PASS' if passed else 'FAIL'}")

    if st.button("Re-run DB Migrations"):
        applied = _db_engine.run()
        if applied:
            st.success(f"Applied migrations: {applied}")
        else:
            st.info("All migrations already applied.")


# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•
# MAIN
# â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•â•

def main() -> None:
    # Page config â€” must be first Streamlit call
    st.set_page_config(
        page_title=f"{FIRM_NAME} - {APP_NAME}",
        page_icon=":material/gavel:",
        layout="wide",
        initial_sidebar_state="expanded",
    )

    # Inject CSS (inside main â€” not at module level)
    inject_css()

    # DB migrations (idempotent â€” safe every run)
    _db_engine.run()

    # Auth gate
    if FEATURES["RBAC_LOGIN"]:
        if not require_auth():
            return

    render_firm_header()

    model_choice, _ = render_sidebar()

    # Navigation tabs
    tab_labels = [s["label"] for s in APP_STEPS]
    tabs = st.tabs(tab_labels)

    tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = tabs

    with tab1: render_tab_upload(model_choice)
    with tab2: render_tab_review()
    with tab3: render_tab_inputs()
    with tab4: render_tab_draft(model_choice)
    with tab5: render_tab_procedural()
    with tab6: render_tab_risk(model_choice)
    with tab7: render_tab_vault()
    with tab8: render_tab_health()


if __name__ == "__main__":
    main()
