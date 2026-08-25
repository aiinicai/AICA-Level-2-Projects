"""DOCX export - Module 14.

Converts the markdown produced by generator.py into a formatted Word
document. This is deliberately NOT a general-purpose markdown-to-docx
converter - it handles exactly the subset of markdown constructs
generator.py/templates.py actually produce (headings, bold, italic,
inline code spans, bullet lists, checklist items, and pipe tables), so
its correctness can be verified against this project's own real output
rather than against arbitrary markdown someone might feed it.

Uses python-docx directly rather than a Node/docx-js-based one-off
document workflow, since this needs to run as part of the shipped
Python application (triggered from the Streamlit "Download Report"
button) rather than being authored once and handed over.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from app.reports.generator import ReportContext, generate_report

_INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\*[^*]+?\*)")


def _add_inline_runs(paragraph, text: str) -> None:
    """Split `text` on **bold**, `code`, and *italic* markers and add
    correspondingly-styled runs to `paragraph`. Plain text between
    markers is added as a normal run."""
    pos = 0
    for match in _INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        else:  # *italic*
            paragraph.add_run(token[1:-1]).italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|").replace("|", "")
    return bool(stripped) and all(c in "-: " for c in stripped)


def _parse_table_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _render_table(document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = len(rows[0])
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Wide financial tables (many years of data as columns) need a
    # smaller font to fit without excessive cell-wrapping, even in
    # landscape orientation — narrower tables (e.g. the 5-column
    # Investment Score component table) read fine at the default size
    # and don't need shrinking.
    cell_font_size = Pt(8) if n_cols > 6 else None

    for r, row_data in enumerate(rows):
        for c, cell_text in enumerate(row_data):
            if c >= n_cols:
                continue
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            _add_inline_runs(p, cell_text)
            if cell_font_size is not None:
                for run in p.runs:
                    run.font.size = cell_font_size
            if r == 0:
                for run in p.runs:
                    run.bold = True
    document.add_paragraph()


def markdown_to_docx(markdown_text: str, output_path: Path) -> Path:
    """Convert generator.py's markdown output into a .docx file at
    output_path. Returns output_path for convenience chaining."""
    document = Document()

    # This project's reports include wide financial tables (e.g. 10
    # years + a CAGR column = 11 columns) — portrait orientation caused
    # real, visually-confirmed cell-wrapping ("Metric" wrapping to
    # "Metr/ic") on a real generated report. Landscape gives meaningfully
    # more horizontal room; applied to the whole document (not just
    # table pages) since python-docx's default single-section document
    # doesn't support per-page orientation without inserting section
    # breaks, and a financial report's narrative sections read fine in
    # landscape too (just wider lines, not a readability problem the way
    # a wrapped table cell is).
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("# "):
            document.add_heading(stripped[2:].strip(), level=0)
            i += 1
            continue

        if stripped.startswith("## "):
            document.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue

        if stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            _add_inline_runs(p, stripped)
            i += 1
            continue

        if _is_table_row(stripped):
            table_rows: list[list[str]] = []
            while i < n and _is_table_row(lines[i].strip()):
                row_line = lines[i].strip()
                if not _is_table_separator(row_line):
                    table_rows.append(_parse_table_row(row_line))
                i += 1
            _render_table(document, table_rows)
            continue

        if stripped.startswith("- [x]") or stripped.startswith("- [ ]"):
            checked = stripped.startswith("- [x]")
            text = stripped[5:].strip()
            p = document.add_paragraph(style="List Bullet")
            prefix = "\u2611 " if checked else "\u2610 "
            p.add_run(prefix)
            _add_inline_runs(p, text)
            i += 1
            continue

        if stripped.startswith("- "):
            indent = len(line) - len(line.lstrip(" "))
            style_name = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = document.add_paragraph(style=style_name)
            _add_inline_runs(p, stripped[2:].strip())
            i += 1
            continue

        p = document.add_paragraph()
        _add_inline_runs(p, stripped)
        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    return output_path


def generate_docx_report(ctx: ReportContext, output_path: Path) -> Path:
    """Generate the full report as markdown (generator.generate_report)
    and convert it to a .docx file at output_path. This is the single
    entry point callers (e.g. the Streamlit Final Thesis page) should
    use rather than calling markdown_to_docx directly."""
    markdown_text = generate_report(ctx)
    return markdown_to_docx(markdown_text, output_path)
