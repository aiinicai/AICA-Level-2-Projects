"""Tests for app/reports/docx_export.py.

Beyond just "did a file get written," these tests open the generated
.docx with python-docx and inspect its actual paragraph/table/heading
structure, so a regression in the markdown parsing (like the table-
separator bug found during development - a literal "---" row leaking
into the rendered table) would be caught automatically rather than
requiring a manual visual check every time.
"""

from __future__ import annotations

from docx import Document

from app.core.enums import ConfidenceLevel, DataStatus, ExchangeCode, RiskCategory, RiskSeverity, UnitOfMeasure
from app.core.models import AIInterpretation, Company, HumanReview, MetricResult, RiskItem
from app.reports.docx_export import (
    _is_table_row,
    _is_table_separator,
    generate_docx_report,
    markdown_to_docx,
)
from app.reports.generator import ReportContext


class TestTableLineDetection:
    def test_separator_row_detected(self):
        assert _is_table_separator("|---|---|---|") is True

    def test_separator_row_with_alignment_colons_detected(self):
        assert _is_table_separator("|:---|:---:|---:|") is True

    def test_data_row_not_detected_as_separator(self):
        assert _is_table_separator("| Component | Score |") is False

    def test_separator_with_many_columns_detected(self):
        assert _is_table_separator("|---|---|---|---|---|") is True

    def test_table_row_detection(self):
        assert _is_table_row("| a | b |") is True
        assert _is_table_row("not a table row") is False


class TestMarkdownToDocxStructure:
    def test_title_becomes_heading_level_0(self, tmp_path):
        md = "# My Report Title\n\nSome content.\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        title_paras = [p for p in doc.paragraphs if p.style.name == "Title"]
        assert len(title_paras) == 1
        assert title_paras[0].text == "My Report Title"

    def test_section_header_becomes_heading_level_1(self, tmp_path):
        md = "## 1. Executive Summary\n\nContent here.\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        h1_paras = [p for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert len(h1_paras) == 1
        assert h1_paras[0].text == "1. Executive Summary"

    def test_bold_marker_produces_bold_run(self, tmp_path):
        md = "This has **bold text** in it.\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        bold_runs = [r for p in doc.paragraphs for r in p.runs if r.bold]
        assert any(r.text == "bold text" for r in bold_runs)

    def test_inline_code_produces_monospace_run(self, tmp_path):
        md = "Status: `[LEVEL 1 - Verified]`\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        code_runs = [r for p in doc.paragraphs for r in p.runs if r.font.name == "Consolas"]
        assert any("LEVEL 1" in r.text for r in code_runs)

    def test_bullet_line_becomes_list_bullet_paragraph(self, tmp_path):
        md = "- First item\n- Second item\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullet_paras) == 2

    def test_checked_checklist_item_includes_checkmark(self, tmp_path):
        md = "- [x] Reviewed item\n- [ ] Unreviewed item\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        texts = [p.text for p in doc.paragraphs]
        assert any("\u2611" in t and "Reviewed item" in t for t in texts)
        assert any("\u2610" in t and "Unreviewed item" in t for t in texts)

    def test_table_renders_correct_row_and_column_count(self, tmp_path):
        md = (
            "| Component | Score | Weight |\n"
            "|---|---|---|\n"
            "| Fundamentals | 84.0 | 30% |\n"
            "| Valuation | 45.0 | 20% |\n"
        )
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert len(table.rows) == 3
        assert len(table.columns) == 3

    def test_table_separator_row_never_appears_as_data(self, tmp_path):
        md = (
            "| Component | Score |\n"
            "|---|---|\n"
            "| Fundamentals | 84.0 |\n"
        )
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        table = doc.tables[0]
        all_cell_text = [cell.text for row in table.rows for cell in row.cells]
        assert "---" not in all_cell_text
        assert not any(set(text) <= {"-"} for text in all_cell_text if text)

    def test_table_header_row_is_bold(self, tmp_path):
        md = "| Component | Score |\n|---|---|\n| Fundamentals | 84.0 |\n"
        out = markdown_to_docx(md, tmp_path / "test.docx")
        doc = Document(str(out))
        header_cell = doc.tables[0].cell(0, 0)
        assert all(r.bold for r in header_cell.paragraphs[0].runs)

    def test_empty_markdown_does_not_crash(self, tmp_path):
        out = markdown_to_docx("", tmp_path / "test.docx")
        assert out.exists()

    def test_output_directory_created_if_missing(self, tmp_path):
        nested_path = tmp_path / "nested" / "dir" / "report.docx"
        out = markdown_to_docx("# Title\n", nested_path)
        assert out.exists()


class TestGenerateDocxReportEndToEnd:
    def _context(self):
        company = Company(name="Sona BLW Precision Forgings Ltd", ticker="SONACOMS",
                           exchange=ExchangeCode.NSE, sector="Auto Ancillary")
        metric = MetricResult(metric_name="EBITDA Margin", formula="f", inputs={}, value=0.2524,
                               unit=UnitOfMeasure.PERCENT, period="FY2026", status=DataStatus.OK)
        risk = RiskItem(category=RiskCategory.FINANCIAL, description="Negative FCF", severity=RiskSeverity.LOW)
        interp = AIInterpretation(claim="Management guided for capex expansion",
                                   confidence=ConfidenceLevel.HIGH, model_name="test-model")
        review = HumanReview(target_id=interp.interpretation_id, reviewer_name="Surendra", accepted=True)
        return ReportContext(
            company=company, fundamental_metrics=[metric], risks=[risk],
            management_interpretations=[interp], human_reviews=[review],
        )

    def test_produces_valid_docx_file(self, tmp_path):
        out_path = generate_docx_report(self._context(), tmp_path / "report.docx")
        assert out_path.exists()
        doc = Document(str(out_path))
        assert len(doc.paragraphs) > 0

    def test_all_19_sections_present_as_headings(self, tmp_path):
        out_path = generate_docx_report(self._context(), tmp_path / "report.docx")
        doc = Document(str(out_path))
        h1_headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
        assert len(h1_headings) == 19
        assert h1_headings[0].startswith("1. Executive Summary")
        assert h1_headings[-1].startswith("19. Final Decision-Support Conclusion")

    def test_company_name_in_title(self, tmp_path):
        out_path = generate_docx_report(self._context(), tmp_path / "report.docx")
        doc = Document(str(out_path))
        title = next(p for p in doc.paragraphs if p.style.name == "Title")
        assert "Sona BLW" in title.text

    def test_ai_interpretation_shows_level2_label(self, tmp_path):
        out_path = generate_docx_report(self._context(), tmp_path / "report.docx")
        doc = Document(str(out_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "LEVEL 2" in all_text
        assert "Management guided for capex expansion" in all_text

    def test_checked_review_item_rendered(self, tmp_path):
        out_path = generate_docx_report(self._context(), tmp_path / "report.docx")
        doc = Document(str(out_path))
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "\u2611" in all_text
        assert "Surendra" in all_text
