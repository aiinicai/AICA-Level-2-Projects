"""Tests for app/reports/metric_tables.py (shared pivot/table logic used
by both the Financial Dashboard UI and the report generator), the
report's now-tabular Historical Financial Analysis section, and the
docx landscape-orientation fix for wide tables.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT

from app.core.enums import DataSourceType, DataStatus, ExchangeCode, UnitOfMeasure
from app.core.models import Company, FinancialStatement, MetricResult, SourceMetadata
from app.reports.docx_export import generate_docx_report
from app.reports.generator import ReportContext, generate_report
from app.reports.metric_tables import (
    build_key_financials_metrics,
    dataframe_to_markdown_table,
    format_metric_value,
    pivot_metrics_to_wide_table,
)

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
SAMPLE_EXCEL = PROJECT_ROOT / "data" / "sample" / "Sona_BLW_Precis_screener_export.xlsx"


def _metric(name, period, value, unit=UnitOfMeasure.PERCENT, status=DataStatus.OK):
    return MetricResult(
        metric_name=name, formula="f", inputs={}, value=value, unit=unit,
        period=period, status=status,
    )


def _statement(period, sales, net_profit, total_assets):
    return FinancialStatement(
        company="Test Co", period=period, sales=sales, net_profit=net_profit,
        total_assets=total_assets,
        source=SourceMetadata(
            company="Test Co", source="test", source_type=DataSourceType.MANUAL_ENTRY,
            unit=UnitOfMeasure.INR_CRORE,
        ),
    )


class TestFormatMetricValue:
    def test_inr_crore_formatted_with_rupee_and_cr_suffix(self):
        m = _metric("Sales", "FY2026", 4123.67, unit=UnitOfMeasure.INR_CRORE)
        assert format_metric_value(m) == "\u20b94,123.67 cr"

    def test_per_share_formatted_with_rupee_prefix(self):
        m = _metric("EPS", "FY2026", 10.40, unit=UnitOfMeasure.PER_SHARE)
        assert format_metric_value(m) == "\u20b910.40"

    def test_percent_and_ratio_unchanged_from_prior_behavior(self):
        assert format_metric_value(_metric("X", "FY26", 0.25, unit=UnitOfMeasure.PERCENT)) == "25.00%"
        assert format_metric_value(_metric("X", "FY26", 0.06, unit=UnitOfMeasure.RATIO)) == "0.06x"


class TestBuildKeyFinancialsMetrics:
    def test_produces_three_metrics_per_statement(self):
        statements = [_statement("FY2025", 3226.30, 579.69, 6154.02), _statement("FY2026", 4123.67, 646.42, 6976.44)]
        results = build_key_financials_metrics(statements)
        assert len(results) == 6

    def test_sales_revenue_metric_present_and_correct(self):
        statements = [_statement("FY2026", 4123.67, 646.42, 6976.44)]
        results = build_key_financials_metrics(statements)
        sales_metric = next(m for m in results if m.metric_name == "Sales / Revenue")
        assert sales_metric.value == 4123.67
        assert sales_metric.status == DataStatus.OK

    def test_missing_sales_value_marked_missing_input_not_fabricated(self):
        statements = [_statement("FY2026", None, 646.42, 6976.44)]
        results = build_key_financials_metrics(statements)
        sales_metric = next(m for m in results if m.metric_name == "Sales / Revenue")
        assert sales_metric.value is None
        assert sales_metric.status == DataStatus.MISSING_INPUT

    def test_real_sona_blw_data_produces_correct_revenue_series(self):
        from app.data.loaders import load_screener_excel
        from app.data.financial_data import build_canonical_statements

        raw = load_screener_excel(SAMPLE_EXCEL, company_name="Sona BLW Precision Forgings Ltd")
        statements = build_canonical_statements(raw)
        results = build_key_financials_metrics(statements)

        fy2026_sales = next(m for m in results if m.metric_name == "Sales / Revenue" and m.period == "FY2026")
        assert fy2026_sales.value == 4123.67


class TestDataframeToMarkdownTable:
    def test_empty_dataframe_returns_no_data_message(self):
        import pandas as pd
        assert dataframe_to_markdown_table(pd.DataFrame()) == "*No data available.*"

    def test_basic_table_structure(self):
        df = pivot_metrics_to_wide_table([_metric("ROE", "FY2026", 0.11)])
        table = dataframe_to_markdown_table(df)
        lines = table.split("\n")
        assert lines[0] == "| Metric | FY2026 |"
        assert lines[1] == "|---|---|"
        assert lines[2] == "| ROE | 11.00% |"

    def test_blank_cells_render_as_em_dash(self):
        metrics = [
            _metric("Revenue CAGR (3yr)", "FY2023-FY2026", 0.1899),
            _metric("ROE", "FY2026", 0.11),
        ]
        df = pivot_metrics_to_wide_table(metrics)
        table = dataframe_to_markdown_table(df)
        assert "\u2014" in table

    def test_separator_row_has_correct_column_count(self):
        df = pivot_metrics_to_wide_table([
            _metric("A", "FY2025", 0.1), _metric("A", "FY2026", 0.2),
        ])
        table = dataframe_to_markdown_table(df)
        separator_line = table.split("\n")[1]
        assert separator_line.count("|") == 4


class TestHistoricalFinancialAnalysisIsTabular:
    def test_report_contains_a_markdown_table_not_bullets(self):
        statements = [_statement("FY2026", 4123.67, 646.42, 6976.44)]
        company = Company(name="Test Co", ticker="TEST", exchange=ExchangeCode.NSE)
        ctx = ReportContext(
            company=company, statements=statements,
            fundamental_metrics=[_metric("EBITDA Margin", "FY2026", 0.2524)],
        )
        report = generate_report(ctx)
        idx = report.find("Historical Financial Analysis")
        section = report[idx:idx + 2000]
        assert "| Metric | FY2026 |" in section
        assert "|---|---|" in section
        assert "- **EBITDA Margin**" not in section

    def test_key_financials_table_present_with_real_revenue(self):
        statements = [_statement("FY2026", 4123.67, 646.42, 6976.44)]
        company = Company(name="Test Co", ticker="TEST", exchange=ExchangeCode.NSE)
        ctx = ReportContext(company=company, statements=statements)
        report = generate_report(ctx)
        assert "Sales / Revenue" in report
        assert "4,123.67" in report

    def test_level1_label_present_once_per_table_not_per_cell(self):
        statements = [_statement("FY2026", 4123.67, 646.42, 6976.44)]
        company = Company(name="Test Co", ticker="TEST", exchange=ExchangeCode.NSE)
        ctx = ReportContext(
            company=company, statements=statements,
            fundamental_metrics=[
                _metric("EBITDA Margin", "FY2026", 0.2524),
                _metric("ROE", "FY2026", 0.11),
            ],
        )
        report = generate_report(ctx)
        idx = report.find("Historical Financial Analysis")
        section = report[idx:idx + 2000]
        assert section.count("LEVEL 1") == 2

    def test_no_statements_and_no_metrics_shows_honest_placeholder(self):
        company = Company(name="Test Co", ticker="TEST", exchange=ExchangeCode.NSE)
        ctx = ReportContext(company=company)
        report = generate_report(ctx)
        assert "No financial metrics have been computed for this run." in report


class TestDocxLandscapeOrientation:
    def test_document_is_landscape(self, tmp_path):
        company = Company(name="Test Co", ticker="TEST", exchange=ExchangeCode.NSE)
        statements = [_statement("FY2026", 4123.67, 646.42, 6976.44)]
        ctx = ReportContext(company=company, statements=statements)
        out_path = generate_docx_report(ctx, tmp_path / "report.docx")

        doc = Document(str(out_path))
        section = doc.sections[0]
        assert section.orientation == WD_ORIENT.LANDSCAPE
        assert section.page_width > section.page_height

    def test_wide_table_cells_use_smaller_font(self, tmp_path):
        from app.reports.docx_export import markdown_to_docx

        md = (
            "| Metric | FY19 | FY20 | FY21 | FY22 | FY23 | FY24 | FY25 |\n"
            "|---|---|---|---|---|---|---|---|\n"
            "| ROE | 1 | 2 | 3 | 4 | 5 | 6 | 7 |\n"
        )
        out_path = markdown_to_docx(md, tmp_path / "wide.docx")
        doc = Document(str(out_path))
        table = doc.tables[0]
        data_cell_run = table.cell(1, 1).paragraphs[0].runs[0]
        assert data_cell_run.font.size is not None
        assert data_cell_run.font.size.pt == 8

    def test_narrow_table_keeps_default_font(self, tmp_path):
        from app.reports.docx_export import markdown_to_docx

        md = "| A | B | C |\n|---|---|---|\n| 1 | 2 | 3 |\n"
        out_path = markdown_to_docx(md, tmp_path / "narrow.docx")
        doc = Document(str(out_path))
        table = doc.tables[0]
        data_cell_run = table.cell(1, 0).paragraphs[0].runs[0]
        assert data_cell_run.font.size is None
