"""
============================================================
 PDF STUDIO PRO  (v2)
 A professional, all-in-one PDF toolkit built with PyQt5.
============================================================

FEATURES
  1. Merge PDFs
  2. Split PDF (by page groups)
  3. Extract Pages (into one PDF)
  4. Compress PDF (3 levels)
  5. Password Protect / Remove Password
  6. Add Watermark
  7. Rotate / Delete Pages
  8. Extract Tables from PDF -> Excel
  9. Batch Processing (Compress / Watermark / Protect a whole folder)
 10. Compare PDFs (text diff report + visual side-by-side diff)
 11. Digital Signature (image e-sign stamp, and PKI/PAdES signing with
     a software certificate file)

------------------------------------------------------------
HOW TO RUN (Windows / Mac / Linux)
------------------------------------------------------------
1) Install Python 3.9+ from https://python.org (tick "Add to PATH")
2) Open Command Prompt / Terminal and run this ONE line:

   pip install PyQt5 pypdf PyMuPDF pdfplumber openpyxl Pillow pyHanko python-docx numpy scipy cryptography

3) Save this file as pdf_studio_pro_v2.py
4) Run it:

   python pdf_studio_pro_v2.py

   (You can also open it in IDLE and press F5)

------------------------------------------------------------
IMPORTANT NOTE ON DIGITAL SIGNATURES
------------------------------------------------------------
The "PKI Digital Signature" feature signs using a SOFTWARE certificate
file (.pfx / .p12) - the same format many DSCs can be exported to, and
the format used for cloud/HSM-issued signing certificates.

It does NOT talk to a USB DSC token (e.g. a typical eMudhra/Sify token).
USB tokens require vendor-specific PKCS#11 driver software (the kind
that tools like emSigner use) and are a much larger, hardware-dependent
integration - out of scope for a simple desktop script. If your firm's
signing workflow uses a USB token, that stamp is best done through the
vendor's signer tool; this feature is for certificate FILES.
------------------------------------------------------------
"""

import os
import sys
import io
import difflib

from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtWidgets import (
    QApplication,
    QMainWindow,
    QTabWidget,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QFormLayout,
    QPushButton,
    QListWidget,
    QAbstractItemView,
    QLabel,
    QLineEdit,
    QFileDialog,
    QMessageBox,
    QComboBox,
    QSpinBox,
    QGroupBox,
    QFrame,
    QStackedWidget,
    QProgressBar,
    QCheckBox,
)

from pypdf import PdfReader, PdfWriter
import pymupdf as fitz  # PyMuPDF (using the new recommended import name)
import pdfplumber
import openpyxl
from PIL import Image, ImageDraw
import numpy as np
from scipy import ndimage
from docx import Document
from docx.shared import RGBColor, Pt

from pyhanko.sign import signers, fields as pyhanko_fields
from pyhanko.sign.fields import SigFieldSpec
from pyhanko.pdf_utils.incremental_writer import IncrementalPdfFileWriter


# ============================================================
# SHARED HELPERS - page range parsing
# ============================================================

def parse_page_groups(text, max_page):
    """
    Parses '1,3-5,8-10' into a list of groups, each group being a list of
    1-based page numbers. Each comma-separated token is ONE group.
    """
    text = (text or "").strip()
    if not text:
        raise ValueError("Please enter at least one page number or range.")

    tokens = [t.strip() for t in text.split(",") if t.strip()]
    if not tokens:
        raise ValueError("Please enter at least one page number or range.")

    groups = []
    for token in tokens:
        if "-" in token:
            parts = token.split("-")
            if len(parts) != 2:
                raise ValueError(f"Invalid range: '{token}'")
            start_s, end_s = parts[0].strip(), parts[1].strip()
            if not start_s.isdigit() or not end_s.isdigit():
                raise ValueError(f"Invalid range: '{token}'")
            start, end = int(start_s), int(end_s)
            if start > end:
                start, end = end, start
            pages = list(range(start, end + 1))
        else:
            if not token.isdigit():
                raise ValueError(f"Invalid page number: '{token}'")
            pages = [int(token)]

        for p in pages:
            if p < 1 or p > max_page:
                raise ValueError(
                    f"Page {p} is out of range (this document has {max_page} pages)."
                )
        groups.append(pages)

    return groups


def parse_page_set(text, max_page):
    """Parses '1,3-5,8-10' or 'all' into a flat set of 1-based page numbers."""
    text = (text or "").strip()
    if text.lower() == "all":
        return set(range(1, max_page + 1))
    groups = parse_page_groups(text, max_page)
    flat = set()
    for g in groups:
        flat.update(g)
    return flat


def section_title(text, emoji=""):
    label = QLabel(f"{emoji}  {text}".strip())
    label.setObjectName("tabTitle")
    return label


def hint_label(text):
    label = QLabel(text)
    label.setWordWrap(True)
    label.setObjectName("hintLabel")
    return label


def divider():
    line = QFrame()
    line.setFrameShape(QFrame.HLine)
    line.setFrameShadow(QFrame.Sunken)
    line.setObjectName("divider")
    return line


# ============================================================
# SHARED HELPERS - core PDF operations
# (used both by single-file tabs AND the batch-processing tab)
# ============================================================

def op_compress(input_path, output_path, level_idx):
    """level_idx: 0=Low, 1=Medium, 2=High (rasterize)."""
    doc = fitz.open(input_path)
    if level_idx == 2:
        new_doc = fitz.open()
        for page in doc:
            pix = page.get_pixmap(dpi=110)
            img_bytes = pix.tobytes("jpeg", jpg_quality=55)
            rect = page.rect
            new_page = new_doc.new_page(width=rect.width, height=rect.height)
            new_page.insert_image(rect, stream=img_bytes)
        new_doc.save(output_path, deflate=True, garbage=4)
        new_doc.close()
    elif level_idx == 1:
        doc.save(output_path, deflate=True, garbage=4, clean=True)
    else:
        doc.save(output_path, deflate=True, garbage=3)
    doc.close()


def op_watermark(input_path, output_path, text, fontsize=45):
    doc = fitz.open(input_path)
    for page in doc:
        rect = page.rect
        point = fitz.Point(rect.width * 0.15, rect.height * 0.6)
        # insert_text's own `rotate` only accepts 0/90/180/270, so a true
        # diagonal (45 degree) watermark needs a morph transform instead.
        morph_matrix = fitz.Matrix(1, 1).prerotate(45)
        page.insert_text(
            point, text, fontsize=fontsize, rotate=0,
            color=(0.75, 0.75, 0.75), fontname="helv", overlay=True,
            morph=(point, morph_matrix),
        )
    doc.save(output_path)
    doc.close()


def op_protect(input_path, output_path, password):
    reader = PdfReader(input_path)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(password)
    with open(output_path, "wb") as f:
        writer.write(f)


def op_unlock(input_path, output_path, password):
    """Returns True on success, False if the password was wrong."""
    reader = PdfReader(input_path)
    if reader.is_encrypted:
        result = reader.decrypt(password)
        if result == 0:
            return False
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    with open(output_path, "wb") as f:
        writer.write(f)
    return True


def _position_box(page_rect, box_w, box_h, position, margin=30):
    """Returns a fitz.Rect for a stamp/signature box at the given corner."""
    if position == "Bottom Right":
        x1, y1 = page_rect.width - margin, page_rect.height - margin
        return fitz.Rect(x1 - box_w, y1 - box_h, x1, y1)
    if position == "Bottom Left":
        x0, y1 = margin, page_rect.height - margin
        return fitz.Rect(x0, y1 - box_h, x0 + box_w, y1)
    if position == "Top Right":
        x1, y0 = page_rect.width - margin, margin
        return fitz.Rect(x1 - box_w, y0, x1, y0 + box_h)
    if position == "Top Left":
        x0, y0 = margin, margin
        return fitz.Rect(x0, y0, x0 + box_w, y0 + box_h)
    # default: center
    cx, cy = page_rect.width / 2, page_rect.height / 2
    return fitz.Rect(cx - box_w / 2, cy - box_h / 2, cx + box_w / 2, cy + box_h / 2)


def op_stamp_signature(input_path, output_path, image_path, pages_set, position, box_width=150):
    doc = fitz.open(input_path)
    with Image.open(image_path) as im:
        aspect = im.height / im.width
    box_height = box_width * aspect

    for i, page in enumerate(doc, start=1):
        if i in pages_set:
            box = _position_box(page.rect, box_width, box_height, position)
            page.insert_image(box, filename=image_path)
    doc.save(output_path)
    doc.close()


def op_pki_sign(input_path, output_path, pfx_path, pfx_password, reason, location,
                 page_idx0, position, field_name="Signature1", box_width=180, box_height=70):
    """
    Signs a PDF using a software PKCS#12 certificate (.pfx/.p12), producing
    a visible PAdES-style signature. page_idx0 is 0-based.
    """
    signer = signers.SimpleSigner.load_pkcs12(
        pfx_file=pfx_path,
        passphrase=pfx_password.encode("utf-8"),
    )

    with fitz.open(input_path) as probe:
        if page_idx0 >= len(probe):
            raise ValueError(f"This PDF only has {len(probe)} page(s).")
        rect = probe[page_idx0].rect
        box_rect = _position_box(rect, box_width, box_height, position)
        # pyHanko uses a bottom-left-origin coordinate box (x1,y1,x2,y2)
        pdf_box = (
            box_rect.x0,
            rect.height - box_rect.y1,
            box_rect.x1,
            rect.height - box_rect.y0,
        )

    with open(input_path, "rb") as inf:
        w = IncrementalPdfFileWriter(inf)
        meta = signers.PdfSignatureMetadata(
            field_name=field_name,
            reason=reason or None,
            location=location or None,
        )
        field_spec = SigFieldSpec(
            sig_field_name=field_name,
            on_page=page_idx0,
            box=pdf_box,
        )
        pyhanko_fields.append_signature_field(w, field_spec)
        with open(output_path, "wb") as outf:
            signers.sign_pdf(w, meta, signer=signer, output=outf)


def extract_page_text(path, page_idx0):
    reader = PdfReader(path)
    if page_idx0 < len(reader.pages):
        return reader.pages[page_idx0].extract_text() or ""
    return ""


def render_page_image(path, page_idx0, dpi=130):
    doc = fitz.open(path)
    if page_idx0 >= len(doc):
        doc.close()
        return None
    pix = doc[page_idx0].get_pixmap(dpi=dpi)
    img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
    doc.close()
    return img


def boxed_diff_image(img_a, img_b, threshold=30, dilation=6, padding=4):
    """Returns img_b with red boxes drawn around regions that differ from img_a."""
    if img_a.size != img_b.size:
        img_b_resized = img_b.resize(img_a.size)
    else:
        img_b_resized = img_b

    gray_a = np.array(img_a.convert("L"))
    gray_b = np.array(img_b_resized.convert("L"))
    diff_mask = np.abs(gray_a.astype(int) - gray_b.astype(int)) > threshold

    if not diff_mask.any():
        return img_b_resized.copy(), False

    dilated = ndimage.binary_dilation(diff_mask, iterations=dilation)
    labeled, num_features = ndimage.label(dilated)

    result = img_b_resized.copy()
    draw = ImageDraw.Draw(result)
    for region_id in range(1, num_features + 1):
        ys, xs = np.where(labeled == region_id)
        if len(xs) == 0:
            continue
        x0, x1 = max(xs.min() - padding, 0), min(xs.max() + padding, result.width)
        y0, y1 = max(ys.min() - padding, 0), min(ys.max() + padding, result.height)
        draw.rectangle([x0, y0, x1, y1], outline=(230, 30, 30), width=3)

    return result, True


def generate_text_diff_docx(path_a, path_b, save_path):
    """Returns True if any textual differences were found."""
    reader_a = PdfReader(path_a)
    reader_b = PdfReader(path_b)
    max_pages = max(len(reader_a.pages), len(reader_b.pages))

    doc = Document()
    doc.add_heading("PDF Comparison Report", level=1)
    doc.add_paragraph(f"File A (before): {os.path.basename(path_a)}")
    doc.add_paragraph(f"File B (after): {os.path.basename(path_b)}")
    doc.add_paragraph(f"Pages: {len(reader_a.pages)} vs {len(reader_b.pages)}")
    doc.add_paragraph("")

    any_diff = False
    for i in range(max_pages):
        text_a = reader_a.pages[i].extract_text() if i < len(reader_a.pages) else ""
        text_b = reader_b.pages[i].extract_text() if i < len(reader_b.pages) else ""
        lines_a = (text_a or "").splitlines()
        lines_b = (text_b or "").splitlines()
        diff = list(difflib.unified_diff(lines_a, lines_b, lineterm="", n=0))
        changes = [l for l in diff if (l.startswith("+") or l.startswith("-"))
                   and not l.startswith("+++") and not l.startswith("---")]
        if changes:
            any_diff = True
            doc.add_heading(f"Page {i + 1}", level=2)
            for line in changes:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(10)
                if line.startswith("+"):
                    run.font.color.rgb = RGBColor(0, 128, 0)
                elif line.startswith("-"):
                    run.font.color.rgb = RGBColor(200, 0, 0)

    if not any_diff:
        doc.add_paragraph("No textual differences were found between the two files.")

    doc.save(save_path)
    return any_diff


def generate_visual_diff_pdf(path_a, path_b, save_path, dpi=130):
    """Builds a PDF with one side-by-side (before | after-with-boxes) page per document page."""
    reader_a = PdfReader(path_a)
    reader_b = PdfReader(path_b)
    max_pages = max(len(reader_a.pages), len(reader_b.pages))

    out_doc = fitz.open()
    any_diff = False

    for i in range(max_pages):
        img_a = render_page_image(path_a, i, dpi=dpi)
        img_b = render_page_image(path_b, i, dpi=dpi)

        if img_a is None:
            note_page = out_doc.new_page(width=img_b.width, height=img_b.height + 30)
            note_page.insert_text((10, 20), f"Page {i+1}: only present in File B (after)", fontsize=12)
            buf = io.BytesIO(); img_b.save(buf, format="PNG")
            note_page.insert_image(fitz.Rect(0, 30, img_b.width, 30 + img_b.height), stream=buf.getvalue())
            any_diff = True
            continue
        if img_b is None:
            note_page = out_doc.new_page(width=img_a.width, height=img_a.height + 30)
            note_page.insert_text((10, 20), f"Page {i+1}: only present in File A (before)", fontsize=12)
            buf = io.BytesIO(); img_a.save(buf, format="PNG")
            note_page.insert_image(fitz.Rect(0, 30, img_a.width, 30 + img_a.height), stream=buf.getvalue())
            any_diff = True
            continue

        boxed_b, has_diff = boxed_diff_image(img_a, img_b)
        if has_diff:
            any_diff = True

        gap = 20
        label_h = 30
        total_w = img_a.width + boxed_b.width + gap
        total_h = max(img_a.height, boxed_b.height) + label_h

        page = out_doc.new_page(width=total_w, height=total_h)
        page.insert_text((10, 20), f"Page {i+1} - BEFORE", fontsize=12)
        page.insert_text((img_a.width + gap + 10, 20), "AFTER (differences boxed in red)", fontsize=12)

        buf_a = io.BytesIO(); img_a.save(buf_a, format="PNG")
        buf_b = io.BytesIO(); boxed_b.save(buf_b, format="PNG")
        page.insert_image(fitz.Rect(0, label_h, img_a.width, label_h + img_a.height), stream=buf_a.getvalue())
        page.insert_image(
            fitz.Rect(img_a.width + gap, label_h, img_a.width + gap + boxed_b.width, label_h + boxed_b.height),
            stream=buf_b.getvalue(),
        )

    out_doc.save(save_path)
    out_doc.close()
    return any_diff


# ============================================================
# TAB 1 - MERGE
# ============================================================

class MergeTab(QWidget):
    def __init__(self):
        super().__init__()
        layout = QVBoxLayout()
        layout.addWidget(section_title("Merge PDFs", "🔗"))
        layout.addWidget(hint_label(
            "Add two or more PDF files below. Use Move Up / Move Down to "
            "set the order they will appear in the final document."
        ))

        self.list_widget = QListWidget()
        self.list_widget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        layout.addWidget(self.list_widget)

        btn_row = QHBoxLayout()
        add_btn = QPushButton("➕ Add PDFs")
        add_btn.clicked.connect(self.add_files)
        remove_btn = QPushButton("🗑 Remove Selected")
        remove_btn.clicked.connect(self.remove_files)
        up_btn = QPushButton("⬆ Move Up")
        up_btn.clicked.connect(self.move_up)
        down_btn = QPushButton("⬇ Move Down")
        down_btn.clicked.connect(self.move_down)
        for b in (add_btn, remove_btn, up_btn, down_btn):
            btn_row.addWidget(b)
        layout.addLayout(btn_row)

        merge_btn = QPushButton("🔗 Merge && Save")
        merge_btn.setObjectName("primaryBtn")
        merge_btn.clicked.connect(self.merge_pdfs)
        layout.addWidget(merge_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def add_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "Select PDF files", "", "PDF Files (*.pdf)")
        for f in files:
            self.list_widget.addItem(f)

    def remove_files(self):
        for item in self.list_widget.selectedItems():
            self.list_widget.takeItem(self.list_widget.row(item))

    def move_up(self):
        row = self.list_widget.currentRow()
        if row > 0:
            item = self.list_widget.takeItem(row)
            self.list_widget.insertItem(row - 1, item)
            self.list_widget.setCurrentRow(row - 1)

    def move_down(self):
        row = self.list_widget.currentRow()
        if row != -1 and row < self.list_widget.count() - 1:
            item = self.list_widget.takeItem(row)
            self.list_widget.insertItem(row + 1, item)
            self.list_widget.setCurrentRow(row + 1)

    def merge_pdfs(self):
        count = self.list_widget.count()
        if count < 2:
            QMessageBox.warning(self, "Not enough files", "Please add at least two PDF files to merge.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Merged PDF", "merged.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            writer = PdfWriter()
            for i in range(count):
                path = self.list_widget.item(i).text()
                reader = PdfReader(path)
                for page in reader.pages:
                    writer.add_page(page)
            with open(save_path, "wb") as f:
                writer.write(f)
            self.status_label.setText(f"✅ Merged {count} files → {save_path}")
            QMessageBox.information(self, "Success", f"Merged PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to merge PDFs:\n{e}")


# ============================================================
# TAB 2 - SPLIT
# ============================================================

class SplitTab(QWidget):
    def __init__(self):
        super().__init__()
        self.pdf_path = None
        self.total_pages = 0

        layout = QVBoxLayout()
        layout.addWidget(section_title("Split PDF", "✂️"))
        layout.addWidget(hint_label(
            "Upload a PDF, then enter page numbers/ranges separated by commas. "
            "Each comma-separated group becomes its own PDF file.\n"
            "Example: 1,3-5,8-10  →  creates 3 separate PDFs."
        ))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        self.range_input = QLineEdit()
        self.range_input.setPlaceholderText("e.g. 1,3-5,8-10")
        layout.addWidget(self.range_input)

        split_btn = QPushButton("✂️ Split && Save")
        split_btn.setObjectName("primaryBtn")
        split_btn.clicked.connect(self.split_pdf)
        layout.addWidget(split_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            try:
                reader = PdfReader(path)
                self.total_pages = len(reader.pages)
                self.pdf_path = path
                self.file_label.setText(f"{os.path.basename(path)}  ({self.total_pages} pages)")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not read PDF:\n{e}")

    def split_pdf(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        try:
            groups = parse_page_groups(self.range_input.text(), self.total_pages)
        except ValueError as e:
            QMessageBox.warning(self, "Invalid input", str(e))
            return

        out_dir = QFileDialog.getExistingDirectory(self, "Select output folder")
        if not out_dir:
            return

        base_name = os.path.splitext(os.path.basename(self.pdf_path))[0]
        try:
            reader = PdfReader(self.pdf_path)
            created = []
            for idx, group in enumerate(groups, start=1):
                writer = PdfWriter()
                for p in group:
                    writer.add_page(reader.pages[p - 1])
                suffix = f"page{group[0]}" if len(group) == 1 else f"pages{group[0]}-{group[-1]}"
                out_path = os.path.join(out_dir, f"{base_name}_{idx}_{suffix}.pdf")
                with open(out_path, "wb") as f:
                    writer.write(f)
                created.append(os.path.basename(out_path))

            self.status_label.setText("✅ Created:\n" + "\n".join(created))
            QMessageBox.information(self, "Success", f"Created {len(created)} PDF file(s) in:\n{out_dir}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to split PDF:\n{e}")


# ============================================================
# TAB 3 - EXTRACT PAGES
# ============================================================

class ExtractTab(QWidget):
    def __init__(self):
        super().__init__()
        self.pdf_path = None
        self.total_pages = 0

        layout = QVBoxLayout()
        layout.addWidget(section_title("Extract Pages", "📑"))
        layout.addWidget(hint_label(
            "Upload a PDF, then enter the pages/ranges you want. All of them "
            "will be combined into ONE new PDF, in the order you type them.\n"
            "Example: 1,3-5,8-10  →  one PDF with pages 1,3,4,5,8,9,10."
        ))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        self.range_input = QLineEdit()
        self.range_input.setPlaceholderText("e.g. 1,3-5,8-10")
        layout.addWidget(self.range_input)

        extract_btn = QPushButton("📑 Extract && Save")
        extract_btn.setObjectName("primaryBtn")
        extract_btn.clicked.connect(self.extract_pdf)
        layout.addWidget(extract_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            try:
                reader = PdfReader(path)
                self.total_pages = len(reader.pages)
                self.pdf_path = path
                self.file_label.setText(f"{os.path.basename(path)}  ({self.total_pages} pages)")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not read PDF:\n{e}")

    def extract_pdf(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        try:
            groups = parse_page_groups(self.range_input.text(), self.total_pages)
        except ValueError as e:
            QMessageBox.warning(self, "Invalid input", str(e))
            return

        flat_pages = []
        seen = set()
        for group in groups:
            for p in group:
                if p not in seen:
                    flat_pages.append(p)
                    seen.add(p)

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Extracted PDF", "extracted.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return

        try:
            reader = PdfReader(self.pdf_path)
            writer = PdfWriter()
            for p in flat_pages:
                writer.add_page(reader.pages[p - 1])
            with open(save_path, "wb") as f:
                writer.write(f)
            self.status_label.setText(f"✅ Extracted {len(flat_pages)} pages → {save_path}")
            QMessageBox.information(self, "Success", f"Extracted PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to extract pages:\n{e}")


# ============================================================
# TAB 4 - COMPRESS
# ============================================================

class CompressTab(QWidget):
    LEVELS = [
        "Low (best quality, smaller size cut)",
        "Medium (balanced)",
        "High (smallest size, rasterizes pages)",
    ]

    def __init__(self):
        super().__init__()
        self.pdf_path = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Compress PDF", "🗜️"))
        layout.addWidget(hint_label(
            "Reduce file size for emailing or uploading to portals.\n"
            "'High' rasterizes each page as an image at reduced resolution - "
            "best for scanned documents, but text will no longer be selectable."
        ))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        form = QFormLayout()
        self.level_combo = QComboBox()
        self.level_combo.addItems(self.LEVELS)
        self.level_combo.setCurrentIndex(1)
        form.addRow("Compression level:", self.level_combo)
        layout.addLayout(form)

        compress_btn = QPushButton("🗜️ Compress && Save")
        compress_btn.setObjectName("primaryBtn")
        compress_btn.clicked.connect(self.compress_pdf)
        layout.addWidget(compress_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.pdf_path = path
            self.file_label.setText(os.path.basename(path))

    def compress_pdf(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Compressed PDF", "compressed.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            original_size = os.path.getsize(self.pdf_path)
            op_compress(self.pdf_path, save_path, self.level_combo.currentIndex())
            new_size = os.path.getsize(save_path)
            saved_pct = round((1 - new_size / original_size) * 100, 1) if original_size else 0
            self.status_label.setText(
                f"✅ {original_size/1024:.0f} KB → {new_size/1024:.0f} KB  ({saved_pct}% smaller)\nSaved to: {save_path}"
            )
            QMessageBox.information(self, "Success", f"Compressed PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to compress PDF:\n{e}")


# ============================================================
# TAB 5 - PASSWORD PROTECT / REMOVE
# ============================================================

class PasswordTab(QWidget):
    def __init__(self):
        super().__init__()
        self.protect_path = None
        self.unlock_path = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Password Protect / Unlock", "🔒"))

        protect_box = QGroupBox("Add Password")
        protect_layout = QVBoxLayout()
        prow = QHBoxLayout()
        pbtn = QPushButton("📂 Select PDF")
        pbtn.clicked.connect(self.select_protect_file)
        self.protect_label = QLabel("No file selected")
        prow.addWidget(pbtn)
        prow.addWidget(self.protect_label, 1)
        protect_layout.addLayout(prow)

        self.protect_pw = QLineEdit()
        self.protect_pw.setPlaceholderText("Enter a password to lock this PDF")
        self.protect_pw.setEchoMode(QLineEdit.Password)
        protect_layout.addWidget(self.protect_pw)

        protect_btn = QPushButton("🔒 Protect && Save")
        protect_btn.setObjectName("primaryBtn")
        protect_btn.clicked.connect(self.protect_pdf)
        protect_layout.addWidget(protect_btn)
        protect_box.setLayout(protect_layout)
        layout.addWidget(protect_box)

        layout.addWidget(divider())

        unlock_box = QGroupBox("Remove Password")
        unlock_layout = QVBoxLayout()
        urow = QHBoxLayout()
        ubtn = QPushButton("📂 Select Protected PDF")
        ubtn.clicked.connect(self.select_unlock_file)
        self.unlock_label = QLabel("No file selected")
        urow.addWidget(ubtn)
        urow.addWidget(self.unlock_label, 1)
        unlock_layout.addLayout(urow)

        self.unlock_pw = QLineEdit()
        self.unlock_pw.setPlaceholderText("Enter the current password")
        self.unlock_pw.setEchoMode(QLineEdit.Password)
        unlock_layout.addWidget(self.unlock_pw)

        unlock_btn = QPushButton("🔓 Remove Password && Save")
        unlock_btn.setObjectName("primaryBtn")
        unlock_btn.clicked.connect(self.unlock_pdf)
        unlock_layout.addWidget(unlock_btn)
        unlock_box.setLayout(unlock_layout)
        layout.addWidget(unlock_box)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_protect_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.protect_path = path
            self.protect_label.setText(os.path.basename(path))

    def select_unlock_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.unlock_path = path
            self.unlock_label.setText(os.path.basename(path))

    def protect_pdf(self):
        if not self.protect_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        password = self.protect_pw.text()
        if not password:
            QMessageBox.warning(self, "No password", "Please enter a password.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Protected PDF", "protected.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            op_protect(self.protect_path, save_path, password)
            self.status_label.setText(f"✅ Password-protected PDF saved → {save_path}")
            QMessageBox.information(self, "Success", f"Protected PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to protect PDF:\n{e}")

    def unlock_pdf(self):
        if not self.unlock_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        password = self.unlock_pw.text()
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Unlocked PDF", "unlocked.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            ok = op_unlock(self.unlock_path, save_path, password)
            if not ok:
                QMessageBox.critical(self, "Wrong password", "The password you entered is incorrect.")
                return
            self.status_label.setText(f"✅ Unlocked PDF saved → {save_path}")
            QMessageBox.information(self, "Success", f"Unlocked PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to remove password:\n{e}")


# ============================================================
# TAB 6 - WATERMARK
# ============================================================

class WatermarkTab(QWidget):
    def __init__(self):
        super().__init__()
        self.pdf_path = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Add Watermark", "💧"))
        layout.addWidget(hint_label(
            "Stamp text like 'DRAFT' or 'CONFIDENTIAL' diagonally across every page."
        ))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        form = QFormLayout()
        self.text_input = QLineEdit()
        self.text_input.setPlaceholderText("e.g. DRAFT, CONFIDENTIAL, FOR REVIEW")
        form.addRow("Watermark text:", self.text_input)

        self.size_spin = QSpinBox()
        self.size_spin.setRange(10, 120)
        self.size_spin.setValue(45)
        form.addRow("Font size:", self.size_spin)
        layout.addLayout(form)

        watermark_btn = QPushButton("💧 Add Watermark && Save")
        watermark_btn.setObjectName("primaryBtn")
        watermark_btn.clicked.connect(self.add_watermark)
        layout.addWidget(watermark_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.pdf_path = path
            self.file_label.setText(os.path.basename(path))

    def add_watermark(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        text = self.text_input.text().strip()
        if not text:
            QMessageBox.warning(self, "No text", "Please enter watermark text.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Watermarked PDF", "watermarked.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            op_watermark(self.pdf_path, save_path, text, self.size_spin.value())
            self.status_label.setText(f"✅ Watermarked PDF saved → {save_path}")
            QMessageBox.information(self, "Success", f"Watermarked PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to add watermark:\n{e}")


# ============================================================
# TAB 7 - ROTATE / DELETE PAGES
# ============================================================

class OrganizeTab(QWidget):
    def __init__(self):
        super().__init__()
        self.pdf_path = None
        self.total_pages = 0

        layout = QVBoxLayout()
        layout.addWidget(section_title("Rotate / Delete Pages", "🔄"))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        rotate_box = QGroupBox("Rotate Pages")
        rotate_layout = QFormLayout()
        self.rotate_pages_input = QLineEdit()
        self.rotate_pages_input.setPlaceholderText("e.g. 1,3-5 or 'all'")
        rotate_layout.addRow("Pages:", self.rotate_pages_input)
        self.angle_combo = QComboBox()
        self.angle_combo.addItems(["90° clockwise", "180°", "270° clockwise"])
        rotate_layout.addRow("Rotate by:", self.angle_combo)
        rotate_btn = QPushButton("🔄 Rotate && Save")
        rotate_btn.setObjectName("primaryBtn")
        rotate_btn.clicked.connect(self.rotate_pages)
        rotate_layout.addRow(rotate_btn)
        rotate_box.setLayout(rotate_layout)
        layout.addWidget(rotate_box)

        layout.addWidget(divider())

        delete_box = QGroupBox("Delete Pages")
        delete_layout = QFormLayout()
        self.delete_pages_input = QLineEdit()
        self.delete_pages_input.setPlaceholderText("e.g. 2,6-7")
        delete_layout.addRow("Pages to delete:", self.delete_pages_input)
        delete_btn = QPushButton("🗑 Delete && Save")
        delete_btn.setObjectName("primaryBtn")
        delete_btn.clicked.connect(self.delete_pages)
        delete_layout.addRow(delete_btn)
        delete_box.setLayout(delete_layout)
        layout.addWidget(delete_box)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            try:
                reader = PdfReader(path)
                self.total_pages = len(reader.pages)
                self.pdf_path = path
                self.file_label.setText(f"{os.path.basename(path)}  ({self.total_pages} pages)")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Could not read PDF:\n{e}")

    def rotate_pages(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        try:
            pages_to_rotate = parse_page_set(self.rotate_pages_input.text(), self.total_pages)
        except ValueError as e:
            QMessageBox.warning(self, "Invalid input", str(e))
            return

        angle_map = {0: 90, 1: 180, 2: 270}
        angle = angle_map[self.angle_combo.currentIndex()]

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Rotated PDF", "rotated.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            reader = PdfReader(self.pdf_path)
            writer = PdfWriter()
            for i, page in enumerate(reader.pages, start=1):
                if i in pages_to_rotate:
                    page.rotate(angle)
                writer.add_page(page)
            with open(save_path, "wb") as f:
                writer.write(f)
            self.status_label.setText(f"✅ Rotated {len(pages_to_rotate)} page(s) → {save_path}")
            QMessageBox.information(self, "Success", f"Rotated PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to rotate pages:\n{e}")

    def delete_pages(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        try:
            pages_to_delete = parse_page_set(self.delete_pages_input.text(), self.total_pages)
        except ValueError as e:
            QMessageBox.warning(self, "Invalid input", str(e))
            return

        if len(pages_to_delete) >= self.total_pages:
            QMessageBox.warning(self, "Cannot delete all pages", "You can't delete every page in the PDF.")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save PDF", "pages_deleted.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            reader = PdfReader(self.pdf_path)
            writer = PdfWriter()
            for i, page in enumerate(reader.pages, start=1):
                if i not in pages_to_delete:
                    writer.add_page(page)
            with open(save_path, "wb") as f:
                writer.write(f)
            self.status_label.setText(f"✅ Deleted {len(pages_to_delete)} page(s) → {save_path}")
            QMessageBox.information(self, "Success", f"PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to delete pages:\n{e}")


# ============================================================
# TAB 8 - PDF TO EXCEL (TABLE EXTRACTION)
# ============================================================

class PdfToExcelTab(QWidget):
    def __init__(self):
        super().__init__()
        self.pdf_path = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Extract Tables to Excel", "📊"))
        layout.addWidget(hint_label(
            "Detects tables in the PDF (e.g. financial statements, schedules, "
            "RBI circular tables) and exports each one to its own Excel sheet."
        ))

        file_row = QHBoxLayout()
        select_btn = QPushButton("📂 Select PDF")
        select_btn.clicked.connect(self.select_file)
        self.file_label = QLabel("No file selected")
        self.file_label.setObjectName("fileLabel")
        file_row.addWidget(select_btn)
        file_row.addWidget(self.file_label, 1)
        layout.addLayout(file_row)

        extract_btn = QPushButton("📊 Extract Tables && Save")
        extract_btn.setObjectName("primaryBtn")
        extract_btn.clicked.connect(self.extract_tables)
        layout.addWidget(extract_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_file(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.pdf_path = path
            self.file_label.setText(os.path.basename(path))

    def extract_tables(self):
        if not self.pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        save_path, _ = QFileDialog.getSaveFileName(self, "Save Excel File", "extracted_tables.xlsx", "Excel Files (*.xlsx)")
        if not save_path:
            return

        try:
            wb = openpyxl.Workbook()
            wb.remove(wb.active)
            table_count = 0

            with pdfplumber.open(self.pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    for t_idx, table in enumerate(tables, start=1):
                        table_count += 1
                        sheet_name = f"Page{page_num}_T{t_idx}"[:31]
                        ws = wb.create_sheet(sheet_name)
                        for row in table:
                            clean_row = [cell if cell is not None else "" for cell in row]
                            ws.append(clean_row)

            if table_count == 0:
                QMessageBox.warning(self, "No tables found", "No tables were detected in this PDF.")
                return

            wb.save(save_path)
            self.status_label.setText(f"✅ Extracted {table_count} table(s) → {save_path}")
            QMessageBox.information(self, "Success", f"Excel file saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to extract tables:\n{e}")


# ============================================================
# TAB 9 - BATCH PROCESSING
# ============================================================

class BatchWorker(QThread):
    progress = pyqtSignal(int, int)
    file_result = pyqtSignal(str, bool, str)
    finished_all = pyqtSignal()

    def __init__(self, files, operation, params, out_dir):
        super().__init__()
        self.files = files
        self.operation = operation
        self.params = params
        self.out_dir = out_dir
        self._cancel = False

    def cancel(self):
        self._cancel = True

    def run(self):
        total = len(self.files)
        for i, path in enumerate(self.files, start=1):
            if self._cancel:
                break
            out_path = os.path.join(self.out_dir, os.path.basename(path))
            try:
                if self.operation == "compress":
                    op_compress(path, out_path, self.params["level"])
                elif self.operation == "watermark":
                    op_watermark(path, out_path, self.params["text"], self.params["fontsize"])
                elif self.operation == "protect":
                    op_protect(path, out_path, self.params["password"])
                self.file_result.emit(os.path.basename(path), True, "Done")
            except Exception as e:
                self.file_result.emit(os.path.basename(path), False, str(e))
            self.progress.emit(i, total)
        self.finished_all.emit()


class BatchTab(QWidget):
    def __init__(self):
        super().__init__()
        self.files = []
        self.out_dir = None
        self.worker = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Batch Processing", "📦"))
        layout.addWidget(hint_label(
            "Select a folder of PDFs and apply the same operation to all of them "
            "in one go. Output files keep their original names in the chosen output folder."
        ))

        folder_row = QHBoxLayout()
        folder_btn = QPushButton("📂 Select Folder of PDFs")
        folder_btn.clicked.connect(self.select_folder)
        self.folder_label = QLabel("No folder selected")
        folder_row.addWidget(folder_btn)
        folder_row.addWidget(self.folder_label, 1)
        layout.addLayout(folder_row)

        self.file_list = QListWidget()
        self.file_list.setMaximumHeight(110)
        layout.addWidget(self.file_list)

        op_row = QFormLayout()
        self.operation_combo = QComboBox()
        self.operation_combo.addItems(["Compress", "Watermark", "Password Protect"])
        self.operation_combo.currentIndexChanged.connect(self.on_operation_changed)
        op_row.addRow("Operation:", self.operation_combo)
        layout.addLayout(op_row)

        self.param_stack = QStackedWidget()

        compress_page = QWidget()
        compress_form = QFormLayout()
        self.batch_level_combo = QComboBox()
        self.batch_level_combo.addItems(CompressTab.LEVELS)
        self.batch_level_combo.setCurrentIndex(1)
        compress_form.addRow("Compression level:", self.batch_level_combo)
        compress_page.setLayout(compress_form)

        watermark_page = QWidget()
        watermark_form = QFormLayout()
        self.batch_wm_text = QLineEdit()
        self.batch_wm_text.setPlaceholderText("e.g. DRAFT")
        watermark_form.addRow("Watermark text:", self.batch_wm_text)
        self.batch_wm_size = QSpinBox()
        self.batch_wm_size.setRange(10, 120)
        self.batch_wm_size.setValue(45)
        watermark_form.addRow("Font size:", self.batch_wm_size)
        watermark_page.setLayout(watermark_form)

        protect_page = QWidget()
        protect_form = QFormLayout()
        self.batch_password = QLineEdit()
        self.batch_password.setEchoMode(QLineEdit.Password)
        self.batch_password.setPlaceholderText("Password to apply to every file")
        protect_form.addRow("Password:", self.batch_password)
        protect_page.setLayout(protect_form)

        self.param_stack.addWidget(compress_page)
        self.param_stack.addWidget(watermark_page)
        self.param_stack.addWidget(protect_page)
        layout.addWidget(self.param_stack)

        out_row = QHBoxLayout()
        out_btn = QPushButton("📁 Select Output Folder")
        out_btn.clicked.connect(self.select_output_folder)
        self.out_label = QLabel("No output folder selected")
        out_row.addWidget(out_btn)
        out_row.addWidget(self.out_label, 1)
        layout.addLayout(out_row)

        run_row = QHBoxLayout()
        self.run_btn = QPushButton("▶ Run Batch")
        self.run_btn.setObjectName("primaryBtn")
        self.run_btn.clicked.connect(self.run_batch)
        self.cancel_btn = QPushButton("⏹ Cancel")
        self.cancel_btn.clicked.connect(self.cancel_batch)
        self.cancel_btn.setEnabled(False)
        run_row.addWidget(self.run_btn)
        run_row.addWidget(self.cancel_btn)
        layout.addLayout(run_row)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        layout.addWidget(self.progress_bar)

        self.log_list = QListWidget()
        self.log_list.setMaximumHeight(140)
        layout.addWidget(self.log_list)

        self.setLayout(layout)

    def on_operation_changed(self, idx):
        self.param_stack.setCurrentIndex(idx)

    def select_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select folder containing PDFs")
        if not folder:
            return
        self.files = [
            os.path.join(folder, f) for f in sorted(os.listdir(folder))
            if f.lower().endswith(".pdf")
        ]
        self.file_list.clear()
        for f in self.files:
            self.file_list.addItem(os.path.basename(f))
        self.folder_label.setText(f"{folder}  ({len(self.files)} PDFs found)")

    def select_output_folder(self):
        folder = QFileDialog.getExistingDirectory(self, "Select output folder")
        if folder:
            self.out_dir = folder
            self.out_label.setText(folder)

    def run_batch(self):
        if not self.files:
            QMessageBox.warning(self, "No files", "Please select a folder containing PDFs first.")
            return
        if not self.out_dir:
            QMessageBox.warning(self, "No output folder", "Please select an output folder.")
            return

        op_idx = self.operation_combo.currentIndex()
        if op_idx == 0:
            operation = "compress"
            params = {"level": self.batch_level_combo.currentIndex()}
        elif op_idx == 1:
            text = self.batch_wm_text.text().strip()
            if not text:
                QMessageBox.warning(self, "No text", "Please enter watermark text.")
                return
            operation = "watermark"
            params = {"text": text, "fontsize": self.batch_wm_size.value()}
        else:
            password = self.batch_password.text()
            if not password:
                QMessageBox.warning(self, "No password", "Please enter a password.")
                return
            operation = "protect"
            params = {"password": password}

        self.log_list.clear()
        self.progress_bar.setValue(0)
        self.progress_bar.setMaximum(len(self.files))
        self.run_btn.setEnabled(False)
        self.cancel_btn.setEnabled(True)

        self.worker = BatchWorker(self.files, operation, params, self.out_dir)
        self.worker.progress.connect(self.on_progress)
        self.worker.file_result.connect(self.on_file_result)
        self.worker.finished_all.connect(self.on_finished)
        self.worker.start()

    def cancel_batch(self):
        if self.worker:
            self.worker.cancel()
            self.cancel_btn.setEnabled(False)

    def on_progress(self, current, total):
        self.progress_bar.setValue(current)

    def on_file_result(self, filename, success, message):
        icon = "✅" if success else "❌"
        self.log_list.addItem(f"{icon} {filename} - {message}")

    def on_finished(self):
        self.run_btn.setEnabled(True)
        self.cancel_btn.setEnabled(False)
        QMessageBox.information(self, "Batch complete", f"Processed {len(self.files)} file(s).\nOutput: {self.out_dir}")


# ============================================================
# TAB 10 - COMPARE PDFs
# ============================================================

class CompareTab(QWidget):
    def __init__(self):
        super().__init__()
        self.path_a = None
        self.path_b = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Compare PDFs", "🔍"))
        layout.addWidget(hint_label(
            "Compare two versions of a document (e.g. draft vs final). "
            "Produces a Word report of text changes and/or a visual side-by-side PDF "
            "with differences boxed in red."
        ))

        row_a = QHBoxLayout()
        btn_a = QPushButton("📂 Select File A (Before)")
        btn_a.clicked.connect(self.select_a)
        self.label_a = QLabel("No file selected")
        row_a.addWidget(btn_a)
        row_a.addWidget(self.label_a, 1)
        layout.addLayout(row_a)

        row_b = QHBoxLayout()
        btn_b = QPushButton("📂 Select File B (After)")
        btn_b.clicked.connect(self.select_b)
        self.label_b = QLabel("No file selected")
        row_b.addWidget(btn_b)
        row_b.addWidget(self.label_b, 1)
        layout.addLayout(row_b)

        self.text_diff_check = QCheckBox("Generate text diff report (Word .docx)")
        self.text_diff_check.setChecked(True)
        layout.addWidget(self.text_diff_check)

        self.visual_diff_check = QCheckBox("Generate visual side-by-side diff (PDF)")
        self.visual_diff_check.setChecked(True)
        layout.addWidget(self.visual_diff_check)

        compare_btn = QPushButton("🔍 Compare && Save")
        compare_btn.setObjectName("primaryBtn")
        compare_btn.clicked.connect(self.compare)
        layout.addWidget(compare_btn)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    def select_a(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select File A (Before)", "", "PDF Files (*.pdf)")
        if path:
            self.path_a = path
            self.label_a.setText(os.path.basename(path))

    def select_b(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select File B (After)", "", "PDF Files (*.pdf)")
        if path:
            self.path_b = path
            self.label_b.setText(os.path.basename(path))

    def compare(self):
        if not self.path_a or not self.path_b:
            QMessageBox.warning(self, "Missing files", "Please select both File A and File B.")
            return
        if not self.text_diff_check.isChecked() and not self.visual_diff_check.isChecked():
            QMessageBox.warning(self, "Nothing selected", "Please select at least one output type.")
            return

        out_dir = QFileDialog.getExistingDirectory(self, "Select output folder")
        if not out_dir:
            return

        results = []
        try:
            if self.text_diff_check.isChecked():
                docx_path = os.path.join(out_dir, "text_diff_report.docx")
                has_diff = generate_text_diff_docx(self.path_a, self.path_b, docx_path)
                results.append(f"📝 Text diff report: {docx_path} ({'differences found' if has_diff else 'no differences'})")

            if self.visual_diff_check.isChecked():
                pdf_path = os.path.join(out_dir, "visual_diff.pdf")
                has_diff = generate_visual_diff_pdf(self.path_a, self.path_b, pdf_path)
                results.append(f"🖼 Visual diff PDF: {pdf_path} ({'differences found' if has_diff else 'no differences'})")

            self.status_label.setText("✅ Comparison complete:\n" + "\n".join(results))
            QMessageBox.information(self, "Success", "Comparison complete:\n\n" + "\n".join(results))
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to compare PDFs:\n{e}")


# ============================================================
# TAB 11 - DIGITAL SIGNATURE (image stamp + PKI)
# ============================================================

POSITION_OPTIONS = ["Bottom Right", "Bottom Left", "Top Right", "Top Left"]


class SignTab(QWidget):
    def __init__(self):
        super().__init__()
        self.stamp_pdf_path = None
        self.stamp_image_path = None
        self.pki_pdf_path = None
        self.pki_cert_path = None

        layout = QVBoxLayout()
        layout.addWidget(section_title("Digital Signature", "🖊️"))

        # --- Image e-sign stamp ---
        stamp_box = QGroupBox("E-Sign Stamp (image, for internal drafts)")
        stamp_layout = QVBoxLayout()

        srow1 = QHBoxLayout()
        sbtn1 = QPushButton("📂 Select PDF")
        sbtn1.clicked.connect(self.select_stamp_pdf)
        self.stamp_pdf_label = QLabel("No file selected")
        srow1.addWidget(sbtn1)
        srow1.addWidget(self.stamp_pdf_label, 1)
        stamp_layout.addLayout(srow1)

        srow2 = QHBoxLayout()
        sbtn2 = QPushButton("🖼 Select Signature Image")
        sbtn2.clicked.connect(self.select_stamp_image)
        self.stamp_image_label = QLabel("No image selected")
        srow2.addWidget(sbtn2)
        srow2.addWidget(self.stamp_image_label, 1)
        stamp_layout.addLayout(srow2)

        stamp_form = QFormLayout()
        self.stamp_pages_input = QLineEdit()
        self.stamp_pages_input.setPlaceholderText("e.g. 1 or 'all'")
        stamp_form.addRow("Pages:", self.stamp_pages_input)

        self.stamp_position_combo = QComboBox()
        self.stamp_position_combo.addItems(POSITION_OPTIONS)
        stamp_form.addRow("Position:", self.stamp_position_combo)

        self.stamp_width_spin = QSpinBox()
        self.stamp_width_spin.setRange(50, 400)
        self.stamp_width_spin.setValue(150)
        stamp_form.addRow("Width (pt):", self.stamp_width_spin)
        stamp_layout.addLayout(stamp_form)

        stamp_btn = QPushButton("🖊️ Stamp && Save")
        stamp_btn.setObjectName("primaryBtn")
        stamp_btn.clicked.connect(self.apply_stamp)
        stamp_layout.addWidget(stamp_btn)

        stamp_box.setLayout(stamp_layout)
        layout.addWidget(stamp_box)

        layout.addWidget(divider())

        # --- PKI signature ---
        pki_box = QGroupBox("PKI Digital Signature (.pfx / .p12 certificate file - final signed reports)")
        pki_layout = QVBoxLayout()
        pki_layout.addWidget(hint_label(
            "Signs with a software certificate FILE only. This does not support USB DSC "
            "tokens, which need vendor middleware (e.g. emSigner) to sign."
        ))

        prow1 = QHBoxLayout()
        pbtn1 = QPushButton("📂 Select PDF")
        pbtn1.clicked.connect(self.select_pki_pdf)
        self.pki_pdf_label = QLabel("No file selected")
        prow1.addWidget(pbtn1)
        prow1.addWidget(self.pki_pdf_label, 1)
        pki_layout.addLayout(prow1)

        prow2 = QHBoxLayout()
        pbtn2 = QPushButton("🔑 Select Certificate (.pfx/.p12)")
        pbtn2.clicked.connect(self.select_pki_cert)
        self.pki_cert_label = QLabel("No certificate selected")
        prow2.addWidget(pbtn2)
        prow2.addWidget(self.pki_cert_label, 1)
        pki_layout.addLayout(prow2)

        pki_form = QFormLayout()
        self.pki_cert_password = QLineEdit()
        self.pki_cert_password.setEchoMode(QLineEdit.Password)
        self.pki_cert_password.setPlaceholderText("Certificate password")
        pki_form.addRow("Certificate password:", self.pki_cert_password)

        self.pki_reason = QLineEdit()
        self.pki_reason.setPlaceholderText("e.g. Approved / Audited")
        pki_form.addRow("Reason (optional):", self.pki_reason)

        self.pki_location = QLineEdit()
        self.pki_location.setPlaceholderText("e.g. Delhi")
        pki_form.addRow("Location (optional):", self.pki_location)

        self.pki_page_spin = QSpinBox()
        self.pki_page_spin.setRange(1, 9999)
        self.pki_page_spin.setValue(1)
        pki_form.addRow("Page number:", self.pki_page_spin)

        self.pki_position_combo = QComboBox()
        self.pki_position_combo.addItems(POSITION_OPTIONS)
        pki_form.addRow("Position:", self.pki_position_combo)

        pki_layout.addLayout(pki_form)

        pki_btn = QPushButton("🖊️ Sign && Save")
        pki_btn.setObjectName("primaryBtn")
        pki_btn.clicked.connect(self.apply_pki_signature)
        pki_layout.addWidget(pki_btn)

        pki_box.setLayout(pki_layout)
        layout.addWidget(pki_box)

        self.status_label = hint_label("")
        layout.addWidget(self.status_label)
        layout.addStretch()
        self.setLayout(layout)

    # --- stamp handlers ---
    def select_stamp_pdf(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.stamp_pdf_path = path
            self.stamp_pdf_label.setText(os.path.basename(path))

    def select_stamp_image(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Signature Image", "", "Images (*.png *.jpg *.jpeg)")
        if path:
            self.stamp_image_path = path
            self.stamp_image_label.setText(os.path.basename(path))

    def apply_stamp(self):
        if not self.stamp_pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        if not self.stamp_image_path:
            QMessageBox.warning(self, "No image", "Please select a signature image.")
            return
        try:
            reader = PdfReader(self.stamp_pdf_path)
            total_pages = len(reader.pages)
            pages_set = parse_page_set(self.stamp_pages_input.text() or "1", total_pages)
        except ValueError as e:
            QMessageBox.warning(self, "Invalid input", str(e))
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Signed PDF", "signed.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return
        try:
            op_stamp_signature(
                self.stamp_pdf_path, save_path, self.stamp_image_path,
                pages_set, self.stamp_position_combo.currentText(),
                self.stamp_width_spin.value(),
            )
            self.status_label.setText(f"✅ Signature stamped on {len(pages_set)} page(s) → {save_path}")
            QMessageBox.information(self, "Success", f"Signed PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to stamp signature:\n{e}")

    # --- PKI handlers ---
    def select_pki_pdf(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select PDF", "", "PDF Files (*.pdf)")
        if path:
            self.pki_pdf_path = path
            self.pki_pdf_label.setText(os.path.basename(path))

    def select_pki_cert(self):
        path, _ = QFileDialog.getOpenFileName(self, "Select Certificate", "", "Certificate Files (*.pfx *.p12)")
        if path:
            self.pki_cert_path = path
            self.pki_cert_label.setText(os.path.basename(path))

    def apply_pki_signature(self):
        if not self.pki_pdf_path:
            QMessageBox.warning(self, "No file", "Please select a PDF first.")
            return
        if not self.pki_cert_path:
            QMessageBox.warning(self, "No certificate", "Please select a .pfx/.p12 certificate file.")
            return
        cert_password = self.pki_cert_password.text()
        if not cert_password:
            QMessageBox.warning(self, "No password", "Please enter the certificate password.")
            return

        save_path, _ = QFileDialog.getSaveFileName(self, "Save Signed PDF", "digitally_signed.pdf", "PDF Files (*.pdf)")
        if not save_path:
            return

        try:
            op_pki_sign(
                self.pki_pdf_path, save_path, self.pki_cert_path, cert_password,
                self.pki_reason.text().strip(), self.pki_location.text().strip(),
                self.pki_page_spin.value() - 1, self.pki_position_combo.currentText(),
            )
            self.status_label.setText(f"✅ Digitally signed PDF saved → {save_path}")
            QMessageBox.information(self, "Success", f"Digitally signed PDF saved:\n{save_path}")
        except Exception as e:
            QMessageBox.critical(
                self, "Signing failed",
                f"Failed to apply digital signature:\n{e}\n\n"
                "Common causes: wrong certificate password, or an unsupported certificate file format."
            )


# ============================================================
# MAIN WINDOW
# ============================================================

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF Studio Pro")
        self.resize(980, 720)

        header = QLabel("📄  PDF Studio Pro")
        header.setObjectName("appHeader")
        header.setAlignment(Qt.AlignCenter)

        tabs = QTabWidget()
        tabs.addTab(MergeTab(), "🔗 Merge")
        tabs.addTab(SplitTab(), "✂️ Split")
        tabs.addTab(ExtractTab(), "📑 Extract")
        tabs.addTab(CompressTab(), "🗜️ Compress")
        tabs.addTab(PasswordTab(), "🔒 Protect")
        tabs.addTab(WatermarkTab(), "💧 Watermark")
        tabs.addTab(OrganizeTab(), "🔄 Rotate/Delete")
        tabs.addTab(PdfToExcelTab(), "📊 PDF→Excel")
        tabs.addTab(BatchTab(), "📦 Batch")
        tabs.addTab(CompareTab(), "🔍 Compare")
        tabs.addTab(SignTab(), "🖊️ Sign")

        central = QWidget()
        central_layout = QVBoxLayout()
        central_layout.setContentsMargins(0, 0, 0, 0)
        central_layout.setSpacing(0)
        central_layout.addWidget(header)
        central_layout.addWidget(tabs)
        central.setLayout(central_layout)
        self.setCentralWidget(central)


# ============================================================
# STYLESHEET (vibrant, professional look)
# ============================================================

STYLESHEET = """
QMainWindow {
    background-color: #f4f6fb;
}

#appHeader {
    font-size: 22px;
    font-weight: bold;
    color: white;
    padding: 16px;
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #6C63FF, stop:0.5 #8E54E9, stop:1 #4776E6);
}

QTabWidget::pane {
    border: none;
    background-color: #ffffff;
}

QTabBar::tab {
    background: #e8eaf6;
    color: #333333;
    padding: 8px 12px;
    margin: 2px;
    border-radius: 8px;
    font-weight: 600;
    font-size: 12px;
}

QTabBar::tab:selected {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #6C63FF, stop:1 #4776E6);
    color: white;
}

QTabBar::tab:hover {
    background: #c9cdf0;
}

QWidget {
    font-family: "Segoe UI", Arial, sans-serif;
    font-size: 13px;
    color: #2b2b3a;
}

#tabTitle {
    font-size: 18px;
    font-weight: bold;
    color: #4b3fbf;
    padding: 6px 0 2px 0;
}

#hintLabel {
    color: #6b6b7d;
    padding-bottom: 6px;
}

#fileLabel {
    color: #333;
    font-style: italic;
}

QGroupBox {
    border: 2px solid #dcd6ff;
    border-radius: 10px;
    margin-top: 12px;
    padding: 10px;
    font-weight: bold;
    color: #4b3fbf;
}

QGroupBox::title {
    subcontrol-origin: margin;
    left: 10px;
    padding: 0 6px;
}

QPushButton {
    background-color: #eef0fb;
    color: #333;
    border: 1px solid #c9cdf0;
    border-radius: 8px;
    padding: 8px 14px;
    font-weight: 600;
}

QPushButton:hover {
    background-color: #dfe3fb;
}

QPushButton:pressed {
    background-color: #c9cdf0;
}

QPushButton:disabled {
    background-color: #f0f0f0;
    color: #aaaaaa;
    border: 1px solid #e0e0e0;
}

#primaryBtn {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #6C63FF, stop:1 #4776E6);
    color: white;
    border: none;
    padding: 10px 18px;
    font-size: 14px;
}

#primaryBtn:hover {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #7d75ff, stop:1 #5a86f0);
}

QLineEdit, QComboBox, QSpinBox {
    border: 1px solid #cfd3f0;
    border-radius: 6px;
    padding: 6px;
    background: white;
}

QLineEdit:focus, QComboBox:focus, QSpinBox:focus {
    border: 1px solid #6C63FF;
}

QListWidget {
    border: 1px solid #dcd6ff;
    border-radius: 8px;
    background: #fafaff;
}

QProgressBar {
    border: 1px solid #dcd6ff;
    border-radius: 6px;
    text-align: center;
    background: #fafaff;
}

QProgressBar::chunk {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #6C63FF, stop:1 #4776E6);
    border-radius: 6px;
}

#divider {
    color: #dcd6ff;
    max-height: 1px;
}
"""


def main():
    app = QApplication(sys.argv)
    app.setStyleSheet(STYLESHEET)
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
