
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import random
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

DATA = {
"Mathematics": {
"Chapter 1 - Real Numbers":[("Fill in the Blank","The decimal expansion of a rational number is either terminating or ________.","non-terminating recurring","Easy",1),("MCQ","Which of the following is irrational?","A. 2/3   B. √2   C. 0.25   D. 5","Medium",1),("Short Answer","Use Euclid's division algorithm to find the HCF of 135 and 225.","45","Medium",3),("Long Answer","Prove that √5 is irrational.","Assume √5 = p/q in lowest terms and derive a contradiction.","Hard",5)],
"Chapter 2 - Polynomials":[("Fill in the Blank","The zeroes of a polynomial are the values of x for which the polynomial becomes ________.","zero","Easy",1),("MCQ","If α and β are zeroes of x² - 5x + 6, then α+β is:","A. 6   B. 5   C. -5   D. -6","Easy",1),("Short Answer","Find the zeroes of x² - 7x + 12.","3 and 4","Medium",3),("Long Answer","Explain the relationship between zeroes and coefficients of a quadratic polynomial.","For ax²+bx+c, α+β=-b/a and αβ=c/a.","Hard",5)],
"Chapter 3 - Pair of Linear Equations":[("Fill in the Blank","Two linear equations in two variables have a unique solution when their graphs are ________.","intersecting","Easy",1),("MCQ","The pair x+y=5 and x-y=1 has solution:","A. (2,3)   B. (3,2)   C. (4,1)   D. (1,4)","Easy",1),("Short Answer","Solve 2x+y=7 and x-y=2.","x=3, y=1","Medium",3),("Long Answer","Solve a pair of linear equations by elimination and explain the steps.","Eliminate one variable, solve, then back-substitute.","Hard",5)]},
"Science":{
"Chapter 1 - Chemical Reactions and Equations":[("Fill in the Blank","A reaction in which a substance gains oxygen is called ________.","oxidation","Easy",1),("MCQ","Zn + CuSO4 → ZnSO4 + Cu is a:","A. Combination   B. Displacement   C. Decomposition   D. Neutralisation","Easy",1),("Short Answer","Why is magnesium ribbon cleaned before burning?","To remove the magnesium oxide layer.","Medium",3),("Long Answer","Explain combination, decomposition, displacement and double-displacement reactions with examples.","Define each type and give a balanced equation.","Hard",5)],
"Chapter 2 - Acids, Bases and Salts":[("Fill in the Blank","Acids turn blue litmus paper ________.","red","Easy",1),("MCQ","The pH of a neutral solution is approximately:","A. 0   B. 5   C. 7   D. 14","Easy",1),("Short Answer","What is a neutralisation reaction? Give one example.","An acid reacts with a base to form salt and water.","Medium",3),("Long Answer","Explain the importance of pH in everyday life.","Discuss digestion, tooth decay and soil treatment.","Hard",5)],
"Chapter 3 - Metals and Non-metals":[("Fill in the Blank","Metals generally form ________ ions by losing electrons.","positive","Easy",1),("MCQ","Which metal is liquid at room temperature?","A. Iron   B. Mercury   C. Copper   D. Aluminium","Easy",1),("Short Answer","Why do ionic compounds have high melting points?","Strong electrostatic forces hold ions together.","Medium",3),("Long Answer","Explain the formation of ionic compounds using electron transfer.","Describe electron loss/gain and ion formation using NaCl.","Hard",5)]},
"Social Science":{
"Chapter 1 - The Rise of Nationalism in Europe":[("Fill in the Blank","The French Revolution began in ________.","1789","Easy",1),("MCQ","Who hosted the Congress of Vienna?","A. Metternich   B. Napoleon   C. Garibaldi   D. Bismarck","Easy",1),("Short Answer","What did liberal nationalism mean in nineteenth-century Europe?","It emphasised freedom, equality before law and representative government.","Medium",3),("Long Answer","Explain the process of German unification.","Discuss Prussia, Bismarck, wars and the German Empire.","Hard",5)],
"Chapter 2 - Nationalism in India":[("Fill in the Blank","The Non-Cooperation Movement was adopted in ________.","1920","Easy",1),("MCQ","The Salt March ended at:","A. Dandi   B. Delhi   C. Lahore   D. Bombay","Easy",1),("Short Answer","Why did Gandhi launch the Civil Disobedience Movement?","To challenge colonial laws and demand swaraj.","Medium",3),("Long Answer","Describe the major features of the Civil Disobedience Movement.","Discuss the salt march, participation and limitations.","Hard",5)],
"Chapter 3 - The Making of a Global World":[("Fill in the Blank","The movement of people, goods and capital across countries is associated with ________.","globalisation","Easy",1),("MCQ","Indentured labour was especially associated with:","A. Contract labour migration   B. Factory ownership   C. Military service   D. Land reform","Medium",1),("Short Answer","What was the impact of technology on nineteenth-century global trade?","Steamships, railways and communication reduced time and transport costs.","Medium",3),("Long Answer","Explain how the nineteenth century created a global economy.","Discuss trade, migration, capital, technology and colonial expansion.","Hard",5)]}
}

def generate(subject, chapters, types, difficulty, total):
    pool=[]
    for ch in chapters:
        for q in DATA[subject][ch]:
            if q[0] in types and (difficulty=="Mixed" or q[3]==difficulty):
                pool.append((ch,)+q)
    random.shuffle(pool)
    result=[]; marks=0
    for q in pool:
        if marks+q[5] <= total:
            result.append(q); marks += q[5]
        if marks==total: break
    return result, marks

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("📚 CBSE Exam Generator - Version 1")
        self.geometry("1050x700")
        self.configure(bg="#eef5ff")
        self.paper=[]
        ttk.Style().configure("TButton", font=("Segoe UI", 11))
        tk.Label(self,text="📚 CBSE EXAM PAPER GENERATOR",font=("Segoe UI",22,"bold"),bg="#eef5ff",fg="#17365d").pack(pady=15)
        frm=tk.Frame(self,bg="#eef5ff"); frm.pack(fill="x",padx=25)
        self.subject=tk.StringVar(value="Mathematics")
        self.chapter=tk.StringVar()
        self.difficulty=tk.StringVar(value="Mixed")
        self.marks=tk.StringVar(value="40")
        self.types={x:tk.BooleanVar(value=True) for x in ["Fill in the Blank","MCQ","Short Answer","Long Answer"]}
        self.cb_subject=ttk.Combobox(frm,textvariable=self.subject,values=list(DATA),state="readonly",width=30); self.cb_subject.grid(row=0,column=0,padx=5,pady=5)
        self.cb_subject.bind("<<ComboboxSelected>>",self.update_chapters)
        self.lb=tk.Listbox(frm,selectmode=tk.MULTIPLE,height=4,width=38)
        self.lb.grid(row=0,column=1,padx=5,pady=5)
        ttk.Combobox(frm,textvariable=self.difficulty,values=["Mixed","Easy","Medium","Hard"],state="readonly",width=15).grid(row=0,column=2,padx=5)
        ttk.Combobox(frm,textvariable=self.marks,values=["10","20","30","40","50"],state="readonly",width=10).grid(row=0,column=3,padx=5)
        for i,(name,var) in enumerate(self.types.items()):
            tk.Checkbutton(frm,text=name,variable=var,bg="#eef5ff",font=("Segoe UI",10)).grid(row=1,column=i,pady=8)
        ttk.Button(frm,text="📝 Generate Paper",command=self.make).grid(row=2,column=0,pady=8)
        ttk.Button(frm,text="📄 Save PDF",command=self.pdf).grid(row=2,column=1,pady=8)
        ttk.Button(frm,text="📝 Save Word",command=self.docx).grid(row=2,column=2,pady=8)
        ttk.Button(frm,text="📚 Official NCERT",command=lambda: __import__("webbrowser").open("https://ncert.nic.in/textbook.php")).grid(row=2,column=3,pady=8)
        self.text=tk.Text(self,font=("Segoe UI",11),wrap="word"); self.text.pack(fill="both",expand=True,padx=25,pady=10)
        self.update_chapters()
    def update_chapters(self,event=None):
        self.lb.delete(0,"end")
        for ch in DATA[self.subject.get()]: self.lb.insert("end",ch)
    def make(self):
        chapters=[self.lb.get(i) for i in self.lb.curselection()]
        if not chapters: chapters=list(DATA[self.subject.get()])[:1]
        types=[k for k,v in self.types.items() if v.get()]
        self.paper, achieved=generate(self.subject.get(),chapters,types,self.difficulty.get(),int(self.marks.get()))
        self.text.delete("1.0","end")
        self.text.insert("end",f"CBSE CLASS 10 - {self.subject.get()}\nMaximum Marks: {achieved}\n\n")
        for i,q in enumerate(self.paper,1):
            self.text.insert("end",f"{i}. [{q[1]} | {q[4]} | {q[5]} mark(s)] {q[2]}\n\n")
        if achieved<int(self.marks.get()):
            self.text.insert("end","NOTE: Sample database does not contain enough questions for the selected marks.\n")
    def pdf(self):
        if not self.paper: messagebox.showwarning("Generate first","Please generate a paper first."); return
        path=filedialog.asksaveasfilename(defaultextension=".pdf",filetypes=[("PDF","*.pdf")])
        if not path:return
        doc=SimpleDocTemplate(path,pagesize=A4,rightMargin=40,leftMargin=40,topMargin=40,bottomMargin=40)
        styles=getSampleStyleSheet(); story=[Paragraph(f"CBSE CLASS 10 - {self.subject.get()}",styles["Title"]),Spacer(1,12)]
        for i,q in enumerate(self.paper,1):
            story += [Paragraph(f"{i}. {q[2]} [{q[5]} mark(s)]",styles["BodyText"]),Spacer(1,8)]
        doc.build(story); messagebox.showinfo("Saved","Question paper PDF saved.")
    def docx(self):
        if not self.paper: messagebox.showwarning("Generate first","Please generate a paper first."); return
        path=filedialog.asksaveasfilename(defaultextension=".docx",filetypes=[("Word","*.docx")])
        if not path:return
        d=Document(); d.add_heading(f"CBSE CLASS 10 - {self.subject.get()}",0)
        for i,q in enumerate(self.paper,1): d.add_paragraph(f"{i}. {q[2]} [{q[5]} mark(s)]")
        d.save(path); messagebox.showinfo("Saved","Word paper saved.")

if __name__=="__main__":
    App().mainloop()
