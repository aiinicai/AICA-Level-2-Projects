
import os, re, sys, math
from pathlib import Path
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

import pandas as pd

APP_TITLE = "Ledger Reconciliation & Reporting Tool"

DATE_COLS = ["date","transaction date","voucher date","doc date"]
REF_COLS = ["voucher","voucher no","voucher number","invoice","invoice no","invoice number","document","document no","reference","reference no","ref no"]
DESC_COLS = ["particulars","description","narration","details","remarks"]
DEBIT_COLS = ["debit","dr","debit amount"]
CREDIT_COLS = ["credit","cr","credit amount"]
AMOUNT_COLS = ["amount","transaction amount","value"]
BAL_COLS = ["balance","running balance","closing balance"]

def norm(s):
    return re.sub(r"[^a-z0-9]+", " ", str(s).lower()).strip()

def pick_col(cols, candidates):
    nmap = {norm(c): c for c in cols}
    for cand in candidates:
        if cand in nmap: return nmap[cand]
    for c in cols:
        nc = norm(c)
        if any(cand in nc for cand in candidates):
            return c
    return None

def parse_amount(x):
    if pd.isna(x) or str(x).strip()=="":
        return 0.0
    s = str(x).strip().replace(",", "")
    neg = ("(" in s and ")" in s) or s.endswith(" Cr") or s.endswith(" CR")
    s = re.sub(r"[^0-9.\-]", "", s)
    try: v = float(s) if s else 0.0
    except: return 0.0
    return -abs(v) if neg else v

def read_pdf_ledger(path):
    """Extract the common Tally-style ledger table used in the supplied PDFs."""
    import pdfplumber
    rows = []
    last_date = None
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                for r in table:
                    if not r or len(r) < 7:
                        continue
                    date, drcr, particulars, vtype, vno, credit, debit = r[:7]
                    if not any([date, drcr, particulars, vtype, vno, credit, debit]):
                        continue

                    # Carry forward the date for continuation lines.
                    if date and re.match(r"^\d{1,2}-[A-Za-z]{3}-\d{2}$", str(date).strip()):
                        last_date = str(date).strip()
                    elif not last_date:
                        continue

                    particulars = str(particulars or "").strip()
                    # Ignore report totals, brought forward/carried over rows.
                    if not particulars or any(x in particulars.lower() for x in
                        ["brought forward", "carried over", "closing balance", "opening balance"]):
                        continue

                    rows.append({
                        "Date": last_date,
                        "DrCr": str(drcr or "").strip(),
                        "Particulars": particulars,
                        "Vch Type": str(vtype or "").strip(),
                        "Vch No": str(vno or "").strip(),
                        "Credit": credit,
                        "Debit": debit
                    })
    if not rows:
        raise ValueError("No readable ledger transactions found in PDF. OCR/manual review may be required.")
    return pd.DataFrame(rows)

def read_any(path):
    ext = Path(path).suffix.lower()
    if ext in [".xlsx",".xls",".xlsm",".csv"]:
        if ext == ".csv":
            return pd.read_csv(path)
        sheets = pd.read_excel(path, sheet_name=None)
        frames = []
        for name, df in sheets.items():
            if not df.empty:
                df["__source_sheet"] = name
                frames.append(df)
        return pd.concat(frames, ignore_index=True) if frames else pd.DataFrame()
    if ext == ".pdf":
        return read_pdf_ledger(path)
    if ext in [".jpg",".jpeg",".png",".bmp",".tiff"]:
        import pytesseract
        from PIL import Image
        txt = pytesseract.image_to_string(Image.open(path))
        return pd.DataFrame({"OCR Text":[txt]})
    raise ValueError("Unsupported file format")

def normalize_ledger(df, source):
    if df.empty:
        raise ValueError(f"{source}: no data found")
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]

    # Special handling for Tally-style PDF extraction.
    if "DrCr" in df.columns and "Particulars" in df.columns:
        out = pd.DataFrame(index=df.index)
        out["Date"] = pd.to_datetime(df["Date"], format="%d-%b-%y", errors="coerce")
        out["Reference"] = (
            df.get("Vch Type", "").fillna("").astype(str) + " " +
            df.get("Vch No", "").fillna("").astype(str)
        ).str.strip()
        out["Description"] = df["Particulars"].fillna("").astype(str)
        drcr = df["DrCr"].fillna("").astype(str).str.upper().str.strip()
        debit_vals = df.get("Debit", pd.Series(0, index=df.index)).map(parse_amount)
        credit_vals = df.get("Credit", pd.Series(0, index=df.index)).map(parse_amount)
        gross = debit_vals + credit_vals
        out["Debit"] = gross.where(drcr.eq("DR"), 0.0)
        out["Credit"] = gross.where(drcr.eq("CR"), 0.0)
        # Signed amount is ledger-side debit positive, credit negative.
        out["SignedAmount"] = out["Debit"] - out["Credit"]
        out["RunningBalance"] = float("nan")
        out["Source"] = source
        out["Original Row"] = df.index + 1
        return out

    datec = pick_col(df.columns, DATE_COLS)
    refc = pick_col(df.columns, REF_COLS)
    descc = pick_col(df.columns, DESC_COLS)
    debitc = pick_col(df.columns, DEBIT_COLS)
    creditc = pick_col(df.columns, CREDIT_COLS)
    amountc = pick_col(df.columns, AMOUNT_COLS)
    balc = pick_col(df.columns, BAL_COLS)

    out = pd.DataFrame(index=df.index)
    out["Date"] = pd.to_datetime(df[datec], errors="coerce", dayfirst=True) if datec else pd.NaT
    out["Reference"] = df[refc].fillna("").astype(str) if refc else ""
    out["Description"] = df[descc].fillna("").astype(str) if descc else ""
    if debitc or creditc:
        dr = df[debitc].map(parse_amount) if debitc else 0
        cr = df[creditc].map(parse_amount) if creditc else 0
        out["Debit"] = dr
        out["Credit"] = cr
        out["SignedAmount"] = out["Debit"] - out["Credit"]
    elif amountc:
        vals = df[amountc].map(parse_amount)
        out["SignedAmount"] = vals
        out["Debit"] = vals.clip(lower=0)
        out["Credit"] = (-vals.clip(upper=0))
    else:
        raise ValueError(f"{source}: Debit/Credit or Amount column could not be identified.")
    out["RunningBalance"] = df[balc].map(parse_amount) if balc else float("nan")
    out["Source"] = source
    out["Original Row"] = df.index + 2
    return out

def cutoff_filter(df, cutoff):
    if not cutoff: return df
    c=pd.to_datetime(cutoff, dayfirst=True, errors="coerce")
    if pd.isna(c): return df
    return df[(df["Date"].isna()) | (df["Date"]<=c)].copy()

def closings(df):
    rb=df["RunningBalance"].dropna()
    if len(rb): return float(rb.iloc[-1])
    return float(df["SignedAmount"].sum())

def opening(df):
    rb=df["RunningBalance"].dropna()
    if len(rb):
        first=rb.iloc[0]
        return float(first-df.loc[rb.index[0],"SignedAmount"])
    return 0.0

def similarity(a,b):
    a=norm(a); b=norm(b)
    if not a or not b: return 0
    from difflib import SequenceMatcher
    return SequenceMatcher(None,a,b).ratio()

def reconcile(company, party, tolerance=0.01):
    used_p = set()
    matched = []
    variances = []

    for ci, c in company.iterrows():
        candidates = []
        for pi, p in party.iterrows():
            if pi in used_p:
                continue

            # Reciprocal party ledgers normally have opposite Dr/Cr directions.
            amount_ok = abs(abs(c.SignedAmount) - abs(p.SignedAmount)) <= tolerance
            if not amount_ok:
                continue

            opposite_sign = (c.SignedAmount * p.SignedAmount < 0)
            date_diff = 999 if pd.isna(c.Date) or pd.isna(p.Date) else abs((c.Date - p.Date).days)
            desc_score = similarity(c.Description, p.Description)
            ref_score = 1 if c.Reference and p.Reference and norm(c.Reference) == norm(p.Reference) else 0

            # Amount + reciprocal direction + close date are primary evidence.
            score = (100 if opposite_sign else 20) + max(0, 30 - min(date_diff, 30)) + desc_score * 10 + ref_score * 10
            candidates.append((score, pi, p, date_diff, desc_score, opposite_sign))

        candidates.sort(reverse=True, key=lambda x: x[0])

        if candidates:
            score, pi, p, dd, ds, opposite = candidates[0]
            if dd <= 7 and opposite:
                used_p.add(pi)
                status = "Exact Match" if dd == 0 else "Timing Difference - Matched"
                matched.append([
                    c.Date, c.Reference, c.SignedAmount,
                    p.Date, p.Reference, p.SignedAmount,
                    status, round(score, 2)
                ])
                continue

        variances.append({
            "Date": c.Date, "Reference": c.Reference,
            "Company Amount": c.SignedAmount, "Party Amount": 0.0,
            "Difference": c.SignedAmount,
            "Variance Type": "Only in Company Books / No supported reciprocal match",
            "Reason": "No transaction of the same amount with reciprocal debit/credit effect was found within the matching window.",
            "Action Required By": "Clarification Required from Both Parties",
            "Explanation": "Review supporting voucher, bank reference and timing. Do not assume responsibility without evidence.",
            "Supporting Details": f"Company extracted transaction row {c['Original Row']}"
        })

    for pi, p in party.iterrows():
        if pi not in used_p:
            variances.append({
                "Date": p.Date, "Reference": p.Reference,
                "Company Amount": 0.0, "Party Amount": p.SignedAmount,
                "Difference": -p.SignedAmount,
                "Variance Type": "Only in Party Books / No supported reciprocal match",
                "Reason": "No transaction of the same amount with reciprocal debit/credit effect was found within the matching window.",
                "Action Required By": "Clarification Required from Both Parties",
                "Explanation": "Review supporting voucher, bank reference and timing. Do not assume responsibility without evidence.",
                "Supporting Details": f"Party extracted transaction row {p['Original Row']}"
            })

    vdf = pd.DataFrame(variances)
    if not vdf.empty:
        vdf.insert(0, "Sr. No.", range(1, len(vdf) + 1))
    mdf = pd.DataFrame(matched, columns=[
        "Company Date","Company Reference","Company Amount",
        "Party Date","Party Reference","Party Amount",
        "Match Status","Match Score"
    ])
    return mdf, vdf

def safe_excel(path, sheets):
    with pd.ExcelWriter(path, engine="openpyxl") as w:
        for name, df in sheets.items():
            df.to_excel(w, sheet_name=name[:31], index=False)
        wb=w.book
        for ws in wb.worksheets:
            ws.freeze_panes="A2"
            for col in ws.columns:
                width=min(45,max(12,max(len(str(c.value or "")) for c in col)+2))
                ws.column_dimensions[col[0].column_letter].width=width

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE); self.geometry("850x500")
        self.company=tk.StringVar(); self.party=tk.StringVar(); self.cutoff=tk.StringVar()
        self.outdir=tk.StringVar(value=str(Path.cwd()/"Reconciliation_Output"))
        self.build()
    def build(self):
        frm=ttk.Frame(self,padding=18); frm.pack(fill="both",expand=True)
        ttk.Label(frm,text="Ledger Reconciliation & Reporting Tool",font=("Segoe UI",18,"bold")).grid(row=0,column=0,columnspan=3,sticky="w",pady=(0,18))
        self.row(frm,1,"Company Books",self.company,True)
        self.row(frm,2,"Party Books",self.party,True)
        ttk.Label(frm,text="Reconciliation Date / Cut-off Date").grid(row=3,column=0,sticky="w",pady=8)
        ttk.Entry(frm,textvariable=self.cutoff,width=55).grid(row=3,column=1,sticky="ew")
        ttk.Label(frm,text="Format: DD-MM-YYYY or DD/MM/YYYY").grid(row=3,column=2,sticky="w")
        self.row(frm,4,"Output Folder",self.outdir,False)
        ttk.Button(frm,text="Run Reconciliation & Generate Reports",command=self.run).grid(row=5,column=0,columnspan=3,pady=24,sticky="ew")
        self.log=tk.Text(frm,height=14); self.log.grid(row=6,column=0,columnspan=3,sticky="nsew")
        frm.columnconfigure(1,weight=1); frm.rowconfigure(6,weight=1)
    def row(self,frm,r,label,var,file):
        ttk.Label(frm,text=label).grid(row=r,column=0,sticky="w",pady=8)
        ttk.Entry(frm,textvariable=var,width=55).grid(row=r,column=1,sticky="ew")
        ttk.Button(frm,text="Browse",command=lambda:self.browse(var,file)).grid(row=r,column=2,padx=6)
    def browse(self,var,file):
        p=filedialog.askopenfilename(filetypes=[("Supported","*.xlsx *.xls *.xlsm *.csv *.pdf *.jpg *.jpeg *.png *.bmp *.tiff"),("All","*.*")]) if file else filedialog.askdirectory()
        if p: var.set(p)
    def write(self,s):
        self.log.insert("end",s+"\n"); self.log.see("end"); self.update()
    def run(self):
        try:
            if not self.company.get() or not self.party.get(): raise ValueError("Please select both statements.")
            out=Path(self.outdir.get()); out.mkdir(parents=True,exist_ok=True)
            self.log.delete("1.0","end"); self.write("Reading Company Books...")
            c=normalize_ledger(read_any(self.company.get()),"Company")
            self.write("Reading Party Books...")
            p=normalize_ledger(read_any(self.party.get()),"Party")
            c=cutoff_filter(c,self.cutoff.get()); p=cutoff_filter(p,self.cutoff.get())
            self.write(f"Company rows read: {len(c)} | Party rows read: {len(p)}")
            m,v=reconcile(c,p)
            co,po=opening(c),opening(p); cc,pc=closings(c),closings(p)
            ov=po-co
            residual=pc-(cc+ov)
            total_company=float(v.loc[v["Action Required By"].eq("Action Required by Company"),"Difference"].abs().sum()) if not v.empty else 0
            total_party=float(v.loc[v["Action Required By"].eq("Action Required by Party"),"Difference"].abs().sum()) if not v.empty else 0
            status="FULLY RECONCILED" if abs(residual)<0.01 and v.empty else ("RECONCILED SUBJECT TO ACTION" if abs(residual)<0.01 else "PARTIALLY RECONCILED — CLARIFICATION REQUIRED")
            summary=pd.DataFrame([
                ["Reconciliation Date",self.cutoff.get()],
                ["Company Opening Balance",co],["Party Opening Balance",po],["Opening Variance",ov],
                ["Company Closing Balance",cc],["Party Closing Balance",pc],
                ["Total Variances",len(v)],["Amount requiring Company action",total_company],
                ["Amount requiring Party action",total_party],["Residual Unexplained Difference",residual],
                ["Reconciliation Status",status]
            ],columns=["Particulars","Amount / Status"])
            company_action=v[v["Action Required By"].eq("Action Required by Company")] if not v.empty else v
            party_action=v[v["Action Required By"].eq("Action Required by Party")] if not v.empty else v
            clarify=v[v["Action Required By"].str.contains("Clarification",na=False)] if not v.empty else v
            xlsx=out/"Reconciliation_Report.xlsx"
            safe_excel(xlsx,{"Summary":summary,"Company Books":c,"Party Books":p,"Matched Transactions":m,"Company Action":company_action,"Party Action":party_action,"Clarification Required":clarify,"Reconciliation Calculation":summary,"Detailed Variances":v})
            email=f"""Subject: Reconciliation Statement as at {self.cutoff.get()}

Dear Sir/Madam,

Please find below the summary of the reconciliation carried out between the Company Books and Party Books as at {self.cutoff.get()}.

Company closing balance: {cc:,.2f} Dr/(Cr.)
Party closing balance: {pc:,.2f} Dr/(Cr.)
Opening variance: {ov:,.2f}
Number of identified variances: {len(v)}
Amount requiring Company action: {total_company:,.2f}
Amount requiring Party action: {total_party:,.2f}
Residual unexplained difference: {residual:,.2f}

We request you to review the transactions classified as requiring Party action and provide/confirm the relevant invoices, vouchers, payment references or other supporting documents. Items for which the available evidence is insufficient should be treated as clarification items until supporting records are exchanged.

Current status: {status}

This reconciliation is based solely on the records supplied and does not make unsupported conclusions regarding responsibility.

Regards,
Accounts Team
"""
            (out/"Draft_Email.txt").write_text(email,encoding="utf-8")
            try:
                from docx import Document
                doc=Document(); doc.add_heading("Reconciliation Report",0)
                doc.add_heading("Executive Summary",1)
                for _,r in summary.iterrows(): doc.add_paragraph(f"{r.iloc[0]}: {r.iloc[1]}")
                doc.add_heading("Reconciliation Statement",1)
                doc.add_paragraph(f"Closing as per Company Books: {cc:,.2f} Dr/(Cr.)")
                doc.add_paragraph(f"Adjustment for Opening Variance: {ov:,.2f}")
                doc.add_paragraph(f"Residual unexplained difference: {residual:,.2f}")
                doc.add_paragraph(f"Closing as per Party Books: {pc:,.2f} Dr/(Cr.)")
                doc.add_heading("Detailed Variance Report",1)
                if not v.empty:
                    table=doc.add_table(rows=1, cols=min(8,len(v.columns))); table.style="Table Grid"
                    cols=list(v.columns)[:8]
                    for i,col in enumerate(cols): table.rows[0].cells[i].text=col
                    for _,row in v.iterrows():
                        cells=table.add_row().cells
                        for i,col in enumerate(cols): cells[i].text=str(row[col])
                doc.add_heading("Final Reconciliation Status",1); doc.add_paragraph(status)
                doc.save(out/"Reconciliation_Report.docx")
            except Exception as e:
                self.write(f"Word report skipped: {e}")
            self.write(f"\nCompleted.\nOutput folder: {out}")
            messagebox.showinfo("Completed",f"Reports generated in:\n{out}")
        except Exception as e:
            messagebox.showerror("Error",str(e)); self.write("ERROR: "+str(e))

if __name__=="__main__":
    App().mainloop()
