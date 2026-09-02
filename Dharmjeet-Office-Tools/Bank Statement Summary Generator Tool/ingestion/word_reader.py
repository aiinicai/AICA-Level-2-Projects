"""Word (.docx) Statement Reader."""

import os
import io
import docx
import pandas as pd
from typing import Optional, List

from normalization.schema_mapper import (
    normalize_dataframe, load_bank_templates, detect_bank_from_header_text, extract_account_number
)
from ingestion.excel_reader import find_table_header_row

def read_word_statement(
    file_path_or_bytes: any,
    filename: str = ""
) -> pd.DataFrame:
    """
    Read tables and text from a Word (.docx) document and return a normalized DataFrame.
    """
    fname = filename or (file_path_or_bytes if isinstance(file_path_or_bytes, str) else "statement.docx")
    templates = load_bank_templates()
    
    try:
        if isinstance(file_path_or_bytes, (bytes, bytearray)):
            doc = docx.Document(io.BytesIO(file_path_or_bytes))
        else:
            doc = docx.Document(file_path_or_bytes)
    except Exception as e:
        print(f"Error reading Word document: {e}")
        return pd.DataFrame()

    # Extract paragraph text for metadata
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    metadata_text = " ".join(paragraphs[:15])
    
    bank_name, bank_tmpl = detect_bank_from_header_text(metadata_text, templates)
    acc_no = extract_account_number(metadata_text)

    # Extract all tables
    all_table_rows = []
    for table in doc.tables:
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            if any(row_data):
                all_table_rows.append(row_data)

    if not all_table_rows:
        return pd.DataFrame()

    df_raw = pd.DataFrame(all_table_rows)
    header_idx, headers = find_table_header_row(df_raw)
    
    df_data = df_raw.iloc[header_idx + 1:].copy()
    df_data.columns = headers
    df_data = df_data.dropna(how="all")

    return normalize_dataframe(
        df_data,
        source_file=fname,
        source_bank=bank_name,
        account_number=acc_no,
        bank_template=bank_tmpl
    )
