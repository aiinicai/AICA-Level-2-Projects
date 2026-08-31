"""Word (.docx) Report Generator for CA/ITR Practice.
Produces a narrative-style audit and tax analysis report.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import pandas as pd
from typing import Dict, Any, Optional

NAVY_HEX = "1F4E78"
TEAL_HEX = "2E75B6"
LIGHT_GRAY_HEX = "F2F2F2"
RED_HEX = "C00000"

def set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set inner cell padding in twips."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_styled_heading(doc, text: str, level: int = 1):
    """Add styled heading with custom color and spacing."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def create_table_from_dataframe(
    doc,
    df: pd.DataFrame,
    header_color_hex: str = NAVY_HEX,
    col_widths: Optional[list] = None
):
    """Create a formatted table from a pandas DataFrame."""
    if df is None or df.empty:
        p = doc.add_paragraph("No data available for this section.")
        p.runs[0].italic = True
        return None

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, col_name in enumerate(df.columns):
        hdr_cells[idx].text = str(col_name)
        set_cell_background(hdr_cells[idx], header_color_hex)
        set_cell_margins(hdr_cells[idx], top=120, bottom=120)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.bold = True
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for r_num, (r_idx, row) in enumerate(df.iterrows()):
        row_cells = table.rows[r_num + 1].cells
        bg_color = LIGHT_GRAY_HEX if r_num % 2 == 1 else "FFFFFF"
        for c_idx, col_name in enumerate(df.columns):
            val = row[col_name]
            # Format numbers
            if isinstance(val, (int, float)):
                if "amount" in col_name.lower() or "receipt" in col_name.lower() or "payment" in col_name.lower() or "inr" in col_name.lower():
                    val_str = f"₹{val:,.2f}"
                elif "%" in col_name.lower() or "share" in col_name.lower():
                    val_str = f"{val:.2f}%"
                else:
                    val_str = f"{val:,}"
                align = WD_ALIGN_PARAGRAPH.RIGHT
            elif isinstance(val, list):
                val_str = ", ".join(str(v) for v in val)
                align = WD_ALIGN_PARAGRAPH.LEFT
            else:
                val_str = str(val) if pd.notna(val) else ""
                align = WD_ALIGN_PARAGRAPH.LEFT

            row_cells[c_idx].text = val_str
            set_cell_background(row_cells[c_idx], bg_color)
            set_cell_margins(row_cells[c_idx], top=80, bottom=80)
            p = row_cells[c_idx].paragraphs[0]
            p.alignment = align
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def export_word_report(
    df: pd.DataFrame,
    output_path: str,
    client_name: str = "Client"
) -> str:
    """Generate professional Word (.docx) report."""
    from analysis.summaries import (
        generate_executive_summary, generate_month_wise_summary,
        generate_nature_wise_summary, generate_party_wise_summary,
        generate_top_and_extrema_transactions, generate_cash_summary
    )
    from analysis.red_flags import detect_red_flags
    from analysis.presumptive_tax import analyze_presumptive_tax
    from analysis.reconciliation import validate_running_balances

    exec_sum = generate_executive_summary(df)
    m_sum = generate_month_wise_summary(df)
    nat_cr, nat_dr = generate_nature_wise_summary(df)
    pty_cr, pty_dr = generate_party_wise_summary(df)
    top_ext = generate_top_and_extrema_transactions(df)
    cash_sum = generate_cash_summary(df)
    df_flagged, red_flag_sum = detect_red_flags(df, client_name=client_name)
    presump_sum = analyze_presumptive_tax(df)
    df_recon, recon_sum = validate_running_balances(df)

    doc = docx.Document()
    
    # Set standard margins (0.75 inch)
    for s in doc.sections:
        s.top_margin = Inches(0.75)
        s.bottom_margin = Inches(0.75)
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)

    # 1. Header / Letterhead
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    run_firm = title_p.add_run("DHARMJEET & ASSOCIATES\n")
    run_firm.bold = True
    run_firm.font.size = Pt(16)
    run_firm.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    
    run_sub = title_p.add_run("CHARTERED ACCOUNTANTS | TAX & AUDIT PRACTICE\n")
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Report Title
    rep_title = doc.add_paragraph()
    rep_title.paragraph_format.space_before = Pt(6)
    rep_title.paragraph_format.space_after = Pt(2)
    r_t = rep_title.add_run(f"BANK STATEMENT AUDIT & SUMMARY REPORT — {client_name.upper()}")
    r_t.bold = True
    r_t.font.size = Pt(13)
    r_t.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_after = Pt(10)
    meta_p.add_run(f"Period Covered: {exec_sum['start_date']} to {exec_sum['end_date']}  |  Total Transactions: {exec_sum['total_transactions']:,}  |  Accounts: {exec_sum['accounts_count']}")
    meta_p.runs[0].font.size = Pt(9.5)
    meta_p.runs[0].italic = True

    # 2. Executive Summary Section
    add_styled_heading(doc, "1. Executive Financial Summary", level=1)
    
    exec_table_df = pd.DataFrame([
        {"Metric": "Total Credits (Receipts)", "Value": f"₹{exec_sum['total_credits']:,.2f}", "Count": f"{exec_sum['credit_count']:,} Entries"},
        {"Metric": "Total Debits (Payments)", "Value": f"₹{exec_sum['total_debits']:,.2f}", "Count": f"{exec_sum['debit_count']:,} Entries"},
        {"Metric": "Net Cash Flow / Movement", "Value": f"₹{exec_sum['net_movement']:,.2f}", "Count": "—"},
        {"Metric": "Total Cash Deposits", "Value": f"₹{cash_sum['total_cash_deposits']:,.2f}", "Count": f"{cash_sum['cash_deposit_count']:,} Entries"},
        {"Metric": "Total Cash Withdrawals", "Value": f"₹{cash_sum['total_cash_withdrawals']:,.2f}", "Count": f"{cash_sum['cash_withdrawal_count']:,} Entries"},
        {"Metric": "Reconciliation Status", "Value": recon_sum["status"], "Count": f"{recon_sum['discrepancies_found']} Discrepancies"}
    ])
    create_table_from_dataframe(doc, exec_table_df, header_color_hex=NAVY_HEX)

    # 3. Presumptive Taxation Assessment (Sec 44AD / 44ADA)
    add_styled_heading(doc, "2. Section 44AD / 44ADA Presumptive Tax & Audit Check", level=1)
    p_p = doc.add_paragraph()
    p_p.add_run(
        f"• Gross Receipts / Turnover: ₹{presump_sum['total_turnover']:,.2f}\n"
        f"• Digital Receipts Proportion: {presump_sum['digital_percentage']:.2f}% (Digital: ₹{presump_sum['digital_receipts']:,.2f} | Cash: ₹{presump_sum['cash_receipts']:,.2f})\n"
        f"• Section 44AD Assessment: Applicable Limit is ₹{presump_sum['sec_44ad_limit_applicable']/10000000:.1f} Cr. Status: {'ELIGIBLE (No Audit Under 44AB on Turnover)' if presump_sum['sec_44ad_eligible'] else 'TURNOVER EXCEEDED (Tax Audit Required)'}\n"
        f"• Minimum 44AD Deemed Profit: ₹{presump_sum['min_presumptive_income_44ad']:,.2f} (6% on digital + 8% on cash receipts)\n"
        f"• Section 44ADA Assessment: Applicable Limit is ₹{presump_sum['sec_44ada_limit_applicable']/100000:.1f} Lakhs. Status: {'ELIGIBLE' if presump_sum['sec_44ada_eligible'] else 'LIMIT EXCEEDED (Audit Required)'}"
    )

    # 4. Month-wise & FY Quarter Breakdown
    add_styled_heading(doc, "3. Month-wise & FY Quarter Summary", level=1)
    create_table_from_dataframe(doc, m_sum, header_color_hex=NAVY_HEX)

    # 5. Nature-wise Summary
    add_styled_heading(doc, "4. Nature-wise Classification Breakdown", level=1)
    add_styled_heading(doc, "4.1 Receipts Classification", level=2)
    create_table_from_dataframe(doc, nat_cr, header_color_hex=TEAL_HEX)
    
    add_styled_heading(doc, "4.2 Payments Classification", level=2)
    create_table_from_dataframe(doc, nat_dr, header_color_hex=NAVY_HEX)

    # 6. Party-wise Summary
    add_styled_heading(doc, "5. Top Counterparty Summary", level=1)
    add_styled_heading(doc, "5.1 Top Counterparties (Receipts)", level=2)
    create_table_from_dataframe(doc, pty_cr.head(15), header_color_hex=TEAL_HEX)

    add_styled_heading(doc, "5.2 Top Counterparties (Payments)", level=2)
    create_table_from_dataframe(doc, pty_dr.head(15), header_color_hex=NAVY_HEX)

    # 7. Red Flag & Scrutiny Signals Annexure
    add_styled_heading(doc, "6. Red Flags & Tax Scrutiny Signals Annexure", level=1)
    flag_intro = doc.add_paragraph()
    flag_intro.add_run(
        f"A total of {red_flag_sum['total_flagged_transactions']} transactions amounting to ₹{red_flag_sum['total_flagged_amount']:,.2f} "
        f"have been identified for CA review based on SFT Rule 114E limits, Section 269SS/269T cash loan rules, structuring patterns, round figures, and accommodation entry signals."
    )

    flagged_df = df_flagged[df_flagged["is_flagged"]].copy() if "is_flagged" in df_flagged.columns else pd.DataFrame()
    if not flagged_df.empty:
        display_cols = ["transaction_date", "counterparty_name", "nature", "mode", "debit_amount", "credit_amount", "flag_reasons"]
        flagged_display = flagged_df[[c for c in display_cols if c in flagged_df.columns]].copy().rename(columns={
            "transaction_date": "Date",
            "counterparty_name": "Party",
            "nature": "Nature",
            "mode": "Mode",
            "debit_amount": "Debit (Dr)",
            "credit_amount": "Credit (Cr)",
            "flag_reasons": "Scrutiny Reasons"
        })
        create_table_from_dataframe(doc, flagged_display, header_color_hex=RED_HEX)
    else:
        doc.add_paragraph("No scrutiny red flags detected under standard thresholds.")

    # Save document
    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    return output_path
