"""DOCX adapter over the node tree. Build Prompt v2 §11.1.

A thin adapter, exactly like `html.py`: it decides Word formatting, never
wording. §19 forbids duplicating document prose between the two renderers,
and `tests/test_no_hardcoded_text.py` enforces it by failing on any long
string literal in this package.

The prototype had no footer, no page numbers and 0.75 inch margins.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.render.base import (
    Bullet,
    Document,
    Heading,
    Letterhead,
    Node,
    PageBreak,
    Para,
    Signature,
    SubPara,
    Table,
)

logger = logging.getLogger("auditcraft")

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(10)

MARGIN_TOP = Inches(1)
MARGIN_RIGHT = Inches(1)
MARGIN_BOTTOM = Inches(1)
MARGIN_LEFT = Inches(1.25)  # gutter

SUBPARA_INDENT = Inches(0.4)
HANGING_INDENT = Inches(0.4)


def _field(paragraph: Any, instruction: str) -> None:
    """Insert a Word field code (PAGE, NUMPAGES).

    Written as a real field rather than a literal so the number is correct
    in a document Word repaginates.
    """
    run = paragraph.add_run()

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


@dataclass(frozen=True, slots=True)
class LetterheadBlock:
    """The letterhead a document goes out on --- the firm's or the company's.

    Named `LetterheadBlock` rather than `Letterhead` because `Letterhead` in
    this module is already a document NODE, the addressee block inside an MRL or
    engagement letter. Two different things one word apart.

    **Which party is not a formatting choice.** A Management Representation
    Letter is written by the company to the auditor and a Board's Report is
    issued by the directors; the firm's letterhead on either would make the
    auditor appear to have written the client's own representations. The
    document declares its issuer in the manifest and the caller supplies the
    matching party.

    Nothing is hard-coded: the firm's values come from Admin -> Firm & Partners
    and the company's from its profile. An empty letterhead prints nothing
    rather than a placeholder.
    """

    name: str = ""

    # "Chartered Accountants" under a firm. Empty for a company --- printing it
    # over a client's name would be a false description of the client.
    subtitle: str = ""

    # Address, then the registration identifier: FRN for a firm, CIN for a
    # company. Never labelled "FRN" on a company, which is the sort of detail
    # that survives review because nobody reads a letterhead twice.
    lines: tuple[str, ...] = ()

    # Two logo fields because they are genuinely two things: a preview served
    # over HTTP cannot use a filesystem path, and python-docx cannot use a URL.
    logo_path: str = ""
    logo_url: str = ""

    @property
    def is_empty(self) -> bool:
        return not any((self.name, self.subtitle, self.lines, self.logo_path, self.logo_url))


def _letterhead(docx: Any, section: Any, head: LetterheadBlock) -> None:
    """Put the firm's identity in the FIRST PAGE header only.

    Word repeats a section header on every page, which is right for a page
    number and wrong for a letterhead: printed stationery carries the firm's
    name once, on the sheet the report starts on, and continuation sheets are
    plain. A letterhead on page four reads as four separate letters.

    `different_first_page_header_footer` gives the section a header of its own
    for page one; the ordinary header is left empty, so nothing repeats.
    """
    section.different_first_page_header_footer = True
    paragraph = section.first_page_header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    logo = Path(head.logo_path) if head.logo_path else None
    if logo is not None and logo.is_file():
        # A missing or unreadable logo must not break the export -- a draft is
        # wanted most when the file is half-configured.
        try:
            paragraph.add_run().add_picture(str(logo), height=Inches(0.5))
            paragraph.add_run().add_break()
        except Exception:
            logger.warning("Could not place firm logo %s on the letterhead", logo)

    if head.name:
        run = paragraph.add_run(head.name)
        run.bold = True
        run.font.size = Pt(12)

    if head.subtitle:
        # The firm name holds the name only; every signature block adds
        # "Chartered Accountants" on its own line, so it goes here too rather
        # than being expected inside the stored name. A company has no subtitle.
        second = paragraph.add_run()
        second.add_break()
        second.add_text(head.subtitle)
        second.font.size = Pt(8)

    trailing = "  |  ".join(bit for bit in head.lines if bit)
    if trailing:
        third = paragraph.add_run()
        third.add_break()
        third.add_text(trailing)
        third.font.size = Pt(8)


def _draft_banner(docx: Any, fy_code: str) -> None:
    """Say, on the page, that this is not an issued document.

    A draft printed on the firm's letterhead is indistinguishable from a signed
    report once it is on a desk. This is the only thing preventing that, so it
    is a paragraph in the body rather than a watermark that a copy-paste drops.
    """
    banner = docx.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = banner.add_run(
        f"DRAFT FOR DISCUSSION -- NOT AN ISSUED DOCUMENT{'  ·  ' + fy_code if fy_code else ''}"
    )
    run.bold = True
    run.font.size = Pt(9)


def _configure(docx: Any, footer_text: str, header_text: str) -> None:
    style = docx.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in docx.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = MARGIN_TOP
        section.right_margin = MARGIN_RIGHT
        section.bottom_margin = MARGIN_BOTTOM
        section.left_margin = MARGIN_LEFT

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run(f"{footer_text}    ")
        # "Page X of Y" — two fields, not a typed number.
        footer.add_run("Page ")
        _field(footer, " PAGE ")
        footer.add_run(" of ")
        _field(footer, " NUMPAGES ")

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.add_run(header_text).italic = True


def _add_table(docx: Any, node: Table) -> None:
    table = docx.add_table(rows=1, cols=len(node.headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, node.headers, strict=True):
        cell.text = text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in node.rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row, strict=True):
            cell.text = text
    docx.add_paragraph()


def _add_node(docx: Any, node: Node) -> None:
    match node:
        case Heading(text=text, level=level):
            paragraph = docx.add_paragraph()
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if level <= 1 else WD_ALIGN_PARAGRAPH.LEFT
            )
            run = paragraph.add_run(text.upper() if level <= 1 else text)
            run.bold = True
            run.font.size = Pt(12 if level <= 1 else 11)

        case Para(text=text, number=number):
            paragraph = docx.add_paragraph()
            if number:
                paragraph.paragraph_format.left_indent = HANGING_INDENT
                paragraph.paragraph_format.first_line_indent = -HANGING_INDENT
                paragraph.add_run(f"{number}\t")
            paragraph.add_run(text)

        case SubPara(text=text, number=number):
            paragraph = docx.add_paragraph()
            paragraph.paragraph_format.left_indent = SUBPARA_INDENT + HANGING_INDENT
            paragraph.paragraph_format.first_line_indent = -HANGING_INDENT
            if number:
                paragraph.add_run(f"{number}\t")
            paragraph.add_run(text)

        case Bullet(text=text):
            docx.add_paragraph(text, style="List Bullet")

        case Table():
            _add_table(docx, node)

        case Signature(lines=lines):
            for line in lines:
                paragraph = docx.add_paragraph(line)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_after = Pt(0)
                # §11.1 — a signature block must not split across pages.
                paragraph.paragraph_format.keep_with_next = True

        case Letterhead(lines=lines):
            for index, line in enumerate(lines):
                paragraph = docx.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(line)
                run.bold = index == 0
                if index == 0:
                    run.font.size = Pt(15)
            docx.add_paragraph()

        case PageBreak():
            docx.add_page_break()

        case _:  # pragma: no cover
            raise TypeError(f"no DOCX adapter for {type(node).__name__}")


def render(
    document: Document,
    path: Path,
    *,
    client_name: str = "",
    fy_code: str = "",
    generated_at: str = "",
    letterhead: LetterheadBlock | None = None,
    draft: bool = False,
) -> Path:
    """Write the node tree to a .docx and return the path.

    `letterhead` replaces the plain document-title header with the firm's own,
    and `draft` stamps the page so a work-in-progress cannot be mistaken for an
    issued report. Both are for the data-collection preview export; a generated
    document uses neither.
    """
    docx = DocxDocument()
    footer_bits = [bit for bit in (client_name, fy_code) if bit]
    _configure(docx, "  ·  ".join(footer_bits), document.title)

    if letterhead is not None and not letterhead.is_empty:
        for section in docx.sections:
            # Both are cleared: the ordinary header carried the document title
            # before a letterhead was asked for, and leaving it would print the
            # title on every page below a letterhead that appears on one.
            section.header.paragraphs[0].clear()
            section.different_first_page_header_footer = True
            section.first_page_header.paragraphs[0].clear()
            _letterhead(docx, section, letterhead)

    if draft:
        _draft_banner(docx, fy_code)

    for node in document.nodes:
        _add_node(docx, node)

    # NO TOOL STAMP ON THE DOCUMENT. Removed on the partner's instruction,
    # 19 August 2026: a statutory report goes out over the firm's name and the
    # signing partner's membership number, and the software that typeset it is
    # not a party to it. It was printing
    # "Generated by AuditCraft · Template <version> · <timestamp>" at the foot
    # of every page.
    #
    # **Nothing is lost from the audit trail.** The template version, the
    # generation timestamp, the actor and the SHA-256 of the content all live on
    # the `document_instance` row, which is what §18.7 reprints from and what
    # makes a reprint byte-identical. They were never load-bearing on the page —
    # they were a footer.

    path.parent.mkdir(parents=True, exist_ok=True)
    docx.save(str(path))
    return path
