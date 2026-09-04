"""Tests for Word and Excel exporters (§10)."""
import os
from pathlib import Path
import pytest
from docx import Document
import openpyxl

from src.core.excel_parser import parse_workbook
from src.core.components import map_workbook_components
from src.core.derivations import extract_period_financials
from src.core.assumptions import resolve_principal_repayment, build_assumptions_registry
from src.core.calculator import compute_ratios
from src.core.variance_engine import populate_reasons_for_results
from src.core.integrity import run_integrity_checks
from src.exporters.word_exporter import export_ratios_to_word
from src.exporters.excel_exporter import export_ratios_to_excel


def test_word_and_excel_exports(tmp_path, sample_cy_path, sample_py_path):
    if not sample_cy_path or not sample_py_path:
        pytest.skip("Sample files not present")
        
    cy_res = parse_workbook(sample_cy_path)
    py_res = parse_workbook(sample_py_path)
    
    cy_map, _ = map_workbook_components(cy_res, "CY")
    py_map, _ = map_workbook_components(py_res, "PY")
    
    cy_closing = extract_period_financials(cy_map, "reporting")
    cy_opening = extract_period_financials(cy_map, "comparative")
    py_closing = extract_period_financials(py_map, "reporting")
    py_opening = extract_period_financials(py_map, "comparative")
    
    pr_result = resolve_principal_repayment(cy_closing, cy_opening, py_closing, py_opening)
    assumptions = build_assumptions_registry(pr_result=pr_result, closing_cy=cy_closing, closing_py=py_closing)
    
    res = compute_ratios(cy_closing, cy_opening, py_closing, py_opening, assumptions, 25.0)
    populate_reasons_for_results(res.schedule_iii_ratios, cy_closing, cy_opening, py_closing, py_opening, units=cy_res.units)
    ic_results = run_integrity_checks(cy_res, py_res, cy_map, py_map, cy_closing, cy_opening, py_closing, py_opening)
    
    word_file = str(tmp_path / "test_export.docx")
    excel_file = str(tmp_path / "test_export.xlsx")
    
    export_ratios_to_word(
        file_path=word_file,
        client_name="Test Enterprise LLP",
        fy_end_date="31 March 2026",
        units="Lacs",
        result_set=res,
        assumptions=assumptions,
        integrity_results=ic_results
    )
    
    export_ratios_to_excel(
        file_path=excel_file,
        client_name="Test Enterprise LLP",
        fy_end_date="31 March 2026",
        units="Lacs",
        result_set=res,
        assumptions=assumptions,
        integrity_results=ic_results
    )
    
    assert os.path.exists(word_file)
    assert os.path.exists(excel_file)
    
    # Verify Word document content
    doc = Document(word_file)
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "Test Enterprise LLP" in full_text
    assert "Analytical Ratios" in full_text
    assert len(doc.tables) >= 1
    
    # Verify Excel document content
    wb = openpyxl.load_workbook(excel_file, data_only=False)
    assert "Schedule III Ratios" in wb.sheetnames
    assert "Assumptions & Disclosures" in wb.sheetnames
    ws = wb["Schedule III Ratios"]
    assert ws["A1"].value == "Test Enterprise LLP"
    assert ws["G7"].value.startswith("=")
