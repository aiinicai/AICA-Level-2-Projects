"""Word exporter generating audit-ready .docx document per Schedule III specification (§10)."""
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from src.core.calculator import SingleRatioResult, CalculationResultSet
from src.core.assumptions import AssumptionItem
from src.core.integrity import IntegrityCheckResult


def set_cell_background(cell, hex_color: str):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set internal cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_table_borders(table, color="D6DEE7", sz="4"):
    """Set thin grey grid borders on a table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:left w:val="none"/>'
        f'<w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:right w:val="none"/>'
        f'<w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'<w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def make_row_header(row):
    """Make table row repeat across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def make_row_cant_split(row):
    """Prevent table row from splitting across pages."""
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))


def export_ratios_to_word(
    file_path: str,
    client_name: str,
    fy_end_date: str,
    units: str,
    result_set: CalculationResultSet,
    assumptions: Dict[str, AssumptionItem],
    integrity_results: List[IntegrityCheckResult],
    note_number: str = "",
    include_additional_ratios: bool = False
) -> None:
    """Generate professional Schedule III analytical ratios note in Word (.docx) format."""
    doc = Document()
    
    # 1. Page Setup: A4 Landscape, 2cm margins (~0.787 inches)
    section = doc.sections[0]
    section.page_width = Inches(11.69)   # A4 landscape width (297 mm)
    section.page_height = Inches(8.27)   # A4 landscape height (210 mm)
    section.top_margin = Inches(0.787)
    section.bottom_margin = Inches(0.787)
    section.left_margin = Inches(0.787)
    section.right_margin = Inches(0.787)
    
    # Header & Footer setup
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = f"Generated on {datetime.now().strftime('%d-%m-%Y')} | Schedule III Ratio Analysis"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.style.font.size = Pt(8.5)
    footer_p.style.font.color.rgb = RGBColor(91, 107, 127)

    # 2. Document Title (Centred, Bold, 14pt)
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(client_name)
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(7, 55, 99)  # Deep Navy
    
    # 3. Sub-title
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run(f"Notes forming part of the Financial Statements for the year ended {fy_end_date}")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(26, 35, 48)
    
    # 4. Note Heading & Units
    p_note = doc.add_paragraph()
    note_txt = f"Note {note_number}: Analytical Ratios" if note_number else "Note: Analytical Ratios"
    run_note = p_note.add_run(note_txt)
    run_note.bold = True
    run_note.font.size = Pt(12)
    run_note.font.color.rgb = RGBColor(11, 79, 140)  # Primary Blue
    
    p_units = doc.add_paragraph()
    p_units.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_units.paragraph_format.space_after = Pt(4)
    run_units = p_units.add_run(f"(Rs. in {units})")
    run_units.italic = True
    run_units.font.size = Pt(9.5)
    run_units.font.color.rgb = RGBColor(91, 107, 127)

    # 5. Main Ratios Table
    # Columns: S. No., Ratio, Numerator, Denominator, Current Year, Previous Year, % Variance, Reason
    headers = [
        "S. No.", "Ratio", "Numerator", "Denominator",
        f"Current Year\n({result_set.cy_label})",
        f"Previous Year\n({result_set.py_label})",
        "% Variance", "Reason for variance\n(where 25% or more)"
    ]
    
    col_widths = [Inches(0.6), Inches(1.6), Inches(1.5), Inches(1.5), Inches(0.9), Inches(0.9), Inches(0.9), Inches(2.2)]
    
    table = doc.add_table(rows=len(result_set.schedule_iii_ratios) + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Format Header Row
    header_row = table.rows[0]
    make_row_header(header_row)
    for c_idx, cell in enumerate(header_row.cells):
        cell.width = col_widths[c_idx]
        set_cell_background(cell, "0B4F8C")  # Primary Blue #0B4F8C
        set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx in (4, 5, 6) else (WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
        run = p.add_run(headers[c_idx])
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    # Populate Data Rows
    for r_idx, r_data in enumerate(result_set.schedule_iii_ratios):
        row = table.rows[r_idx + 1]
        make_row_cant_split(row)
        bg_hex = "F2F7FC" if r_idx % 2 == 1 else "FFFFFF"
        
        row_values = [
            str(r_data.id),
            r_data.name,
            r_data.numerator_desc,
            r_data.denominator_desc,
            r_data.value_cy_formatted,
            r_data.value_py_formatted,
            r_data.variance_pct_formatted,
            r_data.reason_final if r_data.reason_final else "Minor variance"
        ]
        
        for c_idx, cell in enumerate(row.cells):
            cell.width = col_widths[c_idx]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx in (4, 5, 6) else (WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            run = p.add_run(row_values[c_idx])
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(26, 35, 48)
            
            # Format flagged variance in bold
            if c_idx == 6 and r_data.is_flagged:
                run.bold = True
                if r_data.variance_pct and r_data.variance_pct > 0:
                    run.font.color.rgb = RGBColor(30, 142, 90)  # Green #1E8E5A
                else:
                    run.font.color.rgb = RGBColor(192, 57, 43)  # Red #C0392B

    # 6. Footnotes and Statutory Statement
    doc.add_paragraph()
    p_stat = doc.add_paragraph()
    p_stat.paragraph_format.space_before = Pt(6)
    p_stat.paragraph_format.space_after = Pt(2)
    run_stat = p_stat.add_run("Note: Ratios have been computed in accordance with the requirements of Schedule III to the Companies Act, 2013.")
    run_stat.italic = True
    run_stat.font.size = Pt(8.5)
    run_stat.font.color.rgb = RGBColor(91, 107, 127)

    # Footnotes for Not Meaningful ratios
    for r in result_set.schedule_iii_ratios:
        if r.footnote:
            p_fn = doc.add_paragraph()
            p_fn.paragraph_format.space_after = Pt(2)
            run_fn = p_fn.add_run(f"* {r.name}: {r.footnote}")
            run_fn.italic = True
            run_fn.font.size = Pt(8)
            run_fn.font.color.rgb = RGBColor(91, 107, 127)

    doc.save(file_path)
