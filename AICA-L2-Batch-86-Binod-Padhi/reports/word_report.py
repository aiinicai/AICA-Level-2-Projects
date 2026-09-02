"""
reports/word_report.py
Generates the Word (.docx) property analysis report using python-docx.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import DISCLAIMER


def _add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def _kv_table(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Light Grid Accent 1"
    for k, v in pairs:
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = str(v)
    return table


def generate_word_report(path, property_input, city_name, locality_name, valuation_output):
    result = valuation_output["result"]
    doc = Document()

    title = doc.add_heading("India Property Valuation Report", level=0)

    _add_heading(doc, "1. Property Details")
    _kv_table(doc, [
        ("Location", f"{locality_name}, {city_name}"),
        ("Property Type", f"{property_input.bhk} BHK {property_input.property_type}"),
        ("Carpet Area", f"{property_input.carpet_area:.0f} sqft"),
        ("Built-up Area", f"{property_input.builtup_area:.0f} sqft"),
        ("Furnishing", property_input.furnishing),
        ("Age of Property", f"{property_input.age_years:.1f} years"),
        ("New / Resale", property_input.new_or_resale),
        ("Asking Price", f"₹{property_input.asking_price:,.0f}"),
        ("Expected Monthly Rent", f"₹{property_input.expected_rent:,.0f}"),
    ])

    _add_heading(doc, "2. Estimated Market Rent")
    _kv_table(doc, [
        ("Low (25th percentile)", f"₹{result.market_rent_low:,.0f}"),
        ("Median", f"₹{result.market_rent_median:,.0f}"),
        ("High (75th percentile)", f"₹{result.market_rent_high:,.0f}"),
        ("Comparable rental observations", valuation_output["rent_stats"]["count"]),
    ])

    _add_heading(doc, "3. Estimated Fair Value")
    _kv_table(doc, [
        ("Comparable Market Value", f"₹{result.comparable_value:,.0f}"),
        ("Rental Capitalization Value", f"₹{result.rental_cap_value:,.0f}"),
        ("Adjusted Value", f"₹{result.adjusted_value:,.0f}"),
        ("Estimated Fair Value Range", f"₹{result.fair_value_low:,.0f} – ₹{result.fair_value_high:,.0f}"),
    ])

    _add_heading(doc, "4. Rental Yield & Price-to-Rent")
    _kv_table(doc, [
        ("Gross Rental Yield", f"{result.gross_yield:.2f}%"),
        ("Net Rental Yield", f"{result.net_yield:.2f}%"),
        ("Price-to-Rent Ratio", f"{result.price_to_rent:.1f} years"),
    ])

    _add_heading(doc, "5. Valuation Verdict")
    verdict_para = doc.add_paragraph()
    run = verdict_para.add_run(result.verdict)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph(f"Estimated Premium/Discount vs Fair Value: {result.premium_pct:+.1f}%")
    doc.add_paragraph(f"Suggested Negotiation Range: ₹{result.fair_value_low:,.0f} – ₹{result.fair_value_high:,.0f}")

    _add_heading(doc, "6. Investment Score & Confidence")
    _kv_table(doc, [
        ("Property Investment Score", f"{result.investment_score:.0f} / 100 ({result.investment_score_label})"),
        ("Valuation Confidence", f"{result.confidence_pct:.0f}%"),
        ("Comparable properties used", result.n_comparables),
        ("Rental observations used", result.n_rental_obs),
        ("Independent sources", result.n_sources),
    ])

    _add_heading(doc, "7. Methodology")
    doc.add_paragraph(result.methodology_notes)
    doc.add_paragraph(f"Data sources: {', '.join(valuation_output['sources']) or 'None available — import market data.'}")

    _add_heading(doc, "8. Disclaimer")
    disc = doc.add_paragraph(DISCLAIMER)
    for run in disc.runs:
        run.italic = True
        run.font.size = Pt(9)

    doc.save(path)
    return path
