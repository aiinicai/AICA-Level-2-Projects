import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Set page margins (1 inch all sides)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling colors
    NAVY = RGBColor(11, 31, 58)      # #0B1F3A
    BLUE = RGBColor(31, 111, 235)    # #1F6FEB
    DARK_GRAY = RGBColor(51, 65, 85) # #334155
    BLACK = RGBColor(15, 23, 42)     # #0F172A

    # --- TITLE ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("H P M S & ASSOCIATES")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = NAVY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("CA Firm Practice Management System — Project Documentation & Prompts Reference")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = BLUE

    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta_p.add_run("Address: A-27, COMMERCIAL MARKET, GOVINDPURAM, GHAZIABAD, UP-201013\nContact: 7290009815, 7290009816")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10.5)
    run_meta.font.italic = True
    run_meta.font.color.rgb = DARK_GRAY

    doc.add_paragraph() # Spacer

    # Helper function for headings
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(18)
        r.font.bold = True
        r.font.color.rgb = NAVY
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        r = h.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(14)
        r.font.bold = True
        r.font.color.rgb = BLUE
        return h

    def add_body(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            rb = p.add_run(bold_prefix)
            rb.font.name = "Calibri"
            rb.font.size = Pt(11)
            rb.font.bold = True
            rb.font.color.rgb = BLACK
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(11)
        r.font.color.rgb = DARK_GRAY
        return p

    def add_code_block(text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, "F8FAFC")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        r = p.add_run(text)
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.font.color.rgb = DARK_GRAY

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_h1("1. Executive Summary & System Overview")
    add_body("H P M S & Associates Practice Management System is a streamlined, professional desktop/web application built for Chartered Accountant (CA) firms. It manages the end-to-end office operational cycle:")
    add_body("Client → Task → Delegation → Due Date → Work Status → Remarks → Completion → Billing → Payment → Outstanding Collection", bold_prefix="Core Workflow: ")
    add_body("The application operates locally using Python, Streamlit, SQLite, Pandas, and Plotly, requiring zero paid APIs, AI API keys, or complex machine learning dependencies.")

    # --- SECTION 2: SYSTEM ARCHITECTURE & FEATURES ---
    add_h1("2. Core Features & Technology Stack")
    
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Component / Module", "Description & Technology"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_background(cell, "0B1F3A")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Arial"
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    data = [
        ("Technology Stack", "Python 3.9+, Streamlit Web UI, SQLite (hpms.db), Pandas DataFrames, Plotly Express."),
        ("Security & Auth", "PBKDF2-HMAC-SHA256 password hashing (100,000 iterations + random salt). Admin auto-creation for 1st user. Strict pre-authorised email signup control."),
        ("Work & Task Tracker", "Preloaded statutory CA task types (GST, Income Tax, Audit, ROC, Accounts). Task delegation, priority, 8 workflow statuses, progress %, and dated activity remarks trail."),
        ("Billing & Invoicing", "Auto-calculated GST @ 18%, Total Bill Amount calculation, printable Tax Invoice generation with CA Emblem & Authorised Signatory, direct WhatsApp & Email sharing provisions.")
    ]

    for row_idx, (col1, col2) in enumerate(data, start=1):
        c1 = table.cell(row_idx, 0)
        c2 = table.cell(row_idx, 1)
        if row_idx % 2 == 1:
            set_cell_background(c1, "F1F5F9")
            set_cell_background(c2, "F1F5F9")
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(col1)
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(10)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(col2)
        r2.font.name = "Calibri"
        r2.font.size = Pt(10)

    doc.add_paragraph()

    # --- SECTION 3: COMPILED USER PROMPTS ---
    add_h1("3. Compiled User Prompts Reference Manual")
    add_body("Below is the complete chronological compilation of all user prompts and instructions provided during the system development:")

    prompts = [
        ("Prompt 1: Initial System Requirements & Specification Document",
"""# ROLE
Act as an experienced Python developer and business-process consultant familiar with the day-to-day working of small Chartered Accountant firms in India.

Build a simple, professional and fully working web application for:
H P M S & Associates
CA Firm Practice Management – Task, Client, Billing & Collection Tracker

Primary objective:
Client -> Task -> Delegation -> Due Date -> Work Status -> Remarks -> Completion -> Billing -> Payment -> Outstanding Collection

Key Requirements:
- Python, Streamlit, SQLite, Pandas, Plotly (simple charts).
- Runs locally, saves data permanently in SQLite (hpms.db).
- First user automatically becomes Admin (Partner role). Public signup stops afterwards.
- Only employees added by Admin in Employee Master can register.
- Password hashing (PBKDF2-HMAC-SHA256, no plain text).
- Access control enforced in SQL queries (WHERE t.assigned_to = ? for staff).
- Client Master: Code, Name, PAN, GSTIN, Contact Person, Mobile, Email, Client Type, Active/Inactive.
- Task Master: Preloaded GST, Income Tax, Audit, ROC, Accounts, Other tasks + Admin option '+ Add New Task Type'.
- Task Delegation: Client, Task, Assignee, Priority, Assignment Date, Due Date, Financial Year, Instructions.
- Task Status: 8 workflow statuses + progress %. Auto-100% on Completed. Dated remark history in task_updates.
- Admin Task Tracker: Filters by Client, Employee, Task, Status, Priority, Due Date, Quick filters.
- Status Colors: Overdue (Red), Urgent/Due Today (Orange), Waiting (Yellow), In Progress (Blue), Completed (Green).
- Client 360° View: Combined work & billing position + summary metrics.
- Billing Module: Related task, Bill Number, Bill Date, Fees, GST, Other Charges, Total Bill Amount (auto-calculated), Payment Due Date, Remarks. Highlight 'Completed - Billing Pending'.
- Payments & Outstanding: Multiple part-payments per bill. Auto-calculated balance and Payment Status (Unpaid / Partially Paid / Paid).
- Collection Follow-up: Dated follow-up remarks history.
- Admin Dashboard & Business Exceptions: Work & Financial KPI cards, 4 exceptions (Overdue Work, Completed Not Billed, Payment Overdue, Waiting for Client), Team-wise & Client-wise tables.
- Demo Data Generator: 8 employees, 10 clients, 30 tasks, 10 bills, part/full payments."""),

        ("Prompt 2: Batch File (.bat) Creation Request",
"""ISKI BAT FILE BANA DO

Purpose: Create Windows batch files (run.bat for direct Streamlit startup & run_menu.bat for interactive control menu)."""),

        ("Prompt 3: Streamlit Height Error Fix",
"""ye error ara hai
[Attached screenshots showing StreamlitInvalidHeightError on st.dataframe call]

Purpose: Fix Streamlit height validation error when height=None was passed into st.dataframe()."""),

        ("Prompt 4: Auto 18% GST & Live Total Calculation",
"""in customer bill total not tacking amount and give option to add gst % @18 %

Purpose: Add GST Rate selection defaulting to 18% (with 12%, 5%, 0%, and Custom) and live dynamic Total Bill Amount box updating as user types."""),

        ("Prompt 5: Client Invoice Preview & Send Provision",
"""aap isi mai bill karte time .client ko bill bhejne ka provsion daal do. HPMS & ASSOCIATES -AUTHORISED SIGNATORY

Purpose: Add 'Print / Send Invoice' tab with HTML/PDF Invoice download, direct WhatsApp share link, Email compose link, and 'HPMS & ASSOCIATES - AUTHORISED SIGNATORY' branding."""),

        ("Prompt 6: Invoice Customization (Address, Contacts & CA India Logo)",
"""dont add unpaid status in customer invoice preview.. add hpms address- A-27, COMMERCIAL MARKET, GOVINDPURAM GHAZIABAD, UP-201013...CONTACT NO . 7290009815, 7290009816... ADD CA INDIA LOGO ALSO

Purpose: 
1. Remove Unpaid/Received/Outstanding payment status overview from customer Tax Invoice preview.
2. Add firm address: A-27, COMMERCIAL MARKET, GOVINDPURAM, GHAZIABAD, UP-201013.
3. Add contact numbers: 7290009815, 7290009816.
4. Embed official CA India Logo SVG emblem badge in Invoice Header."""),

        ("Prompt 7: Prompts Compilation Request",
"""aap mujhe sare prompts ko compile karke de sakte ho mere , i want to save them for my references

Purpose: Compile all prompts and customization instructions into a reference manual.""")
    ]

    for title, content in prompts:
        add_h2(title)
        add_code_block(content)
        doc.add_paragraph()

    # Save document
    doc_path = "H_P_M_S_Project_Prompts_And_Documentation.docx"
    doc.save(doc_path)
    print(f"Document created successfully: {doc_path}")

if __name__ == "__main__":
    create_document()
