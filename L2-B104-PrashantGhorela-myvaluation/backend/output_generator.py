
from pathlib import Path
from typing import Optional, Any, Dict, List

from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================
# HOUSE STYLE
# Based on the user's prior valuation working/report:
# - dark navy section bars
# - white background / restrained formatting
# - gold subsection headings in Word
# - compact professional tables
# - annexure-driven valuation report
# - blue/green input conventions in Excel
# ============================================================

NAVY = "203E6B"
NAVY_DARK = "173052"
GOLD = "B28A16"
LIGHT_GREY = "F2F2F2"
MID_GREY = "D9D9D9"
WHITE = "FFFFFF"
BLACK = "000000"
BLUE_INPUT = "0000FF"
GREEN_LINK = "008000"
RED = "FF0000"
YELLOW = "FFFF00"
LIGHT_GREEN = "C6E0B4"
LIGHT_BLUE = "D9E2F3"

THIN = Side(style="thin", color="BFBFBF")
MEDIUM = Side(style="medium", color=NAVY)

MONEY_FMT = '#,##0.00;[Red](#,##0.00);-'
SHARE_FMT = '#,##0;[Red](#,##0);-'
PERCENT_FMT = '0.00%'
MULTIPLE_FMT = '0.00x'
PER_SHARE_FMT = '"Rs. " #,##0.00'


# ============================================================
# GENERIC HELPERS
# ============================================================

def _f(value: Any, default: float = 0.0) -> float:
    try:
        if value in (None, ""):
            return default
        return float(value)
    except Exception:
        return default


def _fmt(value: Any, digits: int = 2) -> str:
    try:
        return f"{float(value):,.{digits}f}"
    except Exception:
        return "" if value is None else str(value)


def _fmt_int(value: Any) -> str:
    try:
        return f"{float(value):,.0f}"
    except Exception:
        return "" if value is None else str(value)


def _fmt_pct(value: Any) -> str:
    try:
        return f"{float(value):,.2f}%"
    except Exception:
        return "" if value is None else str(value)


def _yes_no(value: Any) -> str:
    return "Yes" if bool(value) else "No"


def _rows(financial_analysis: Dict, bucket: str) -> List[Dict]:
    result = financial_analysis.get(bucket, []) or []
    return [r for r in result if isinstance(r, dict) and r.get("metrics")]


def _all_source_documents(assignment: Dict) -> List[str]:
    docs = assignment.get("documents", {}) or {}
    out = []

    for category, payload in docs.items():
        if not isinstance(payload, dict):
            continue

        for item in payload.get("files", []) or []:
            if isinstance(item, dict):
                name = (
                    item.get("original_name")
                    or item.get("stored_name")
                    or ""
                )
            else:
                name = str(item)

            if name:
                out.append(f"{category}: {name}")

    return out


# ============================================================
# EXCEL STYLE HELPERS
# ============================================================

def _hide_gridlines(ws):
    ws.sheet_view.showGridLines = False


def _navy_bar(ws, row, text, start_col=2, end_col=8):
    ws.merge_cells(
        start_row=row,
        start_column=start_col,
        end_row=row,
        end_column=end_col,
    )

    c = ws.cell(row, start_col, text)
    c.fill = PatternFill("solid", fgColor=NAVY)
    c.font = Font(name="Arial", size=11, bold=True, color=WHITE)
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[row].height = 22


def _sheet_title(ws, company_name, subtitle="", end_col=8):
    _hide_gridlines(ws)

    ws.merge_cells(start_row=2, start_column=2, end_row=2, end_column=end_col)
    c = ws.cell(2, 2, str(company_name or "").upper())
    c.fill = PatternFill("solid", fgColor=NAVY)
    c.font = Font(name="Arial", size=14, bold=True, color=WHITE)
    c.alignment = Alignment(horizontal="left", vertical="center")
    ws.row_dimensions[2].height = 24

    if subtitle:
        ws.merge_cells(start_row=3, start_column=2, end_row=3, end_column=end_col)
        s = ws.cell(3, 2, subtitle)
        s.font = Font(name="Arial", size=10, italic=True, color="666666")
        s.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        ws.row_dimensions[3].height = 22


def _header(ws, row, labels, start_col=2):
    for i, label in enumerate(labels, start=start_col):
        c = ws.cell(row, i, label)
        c.fill = PatternFill("solid", fgColor=NAVY)
        c.font = Font(name="Arial", size=10, bold=True, color=WHITE)
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        c.border = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)

    ws.row_dimensions[row].height = 28


def _border_range(ws, min_row, max_row, min_col, max_col, zebra=False):
    for r in range(min_row, max_row + 1):
        for c in range(min_col, max_col + 1):
            cell = ws.cell(r, c)
            cell.border = Border(top=THIN, bottom=THIN, left=THIN, right=THIN)
            cell.alignment = Alignment(vertical="center", wrap_text=True)

            if zebra and (r - min_row) % 2 == 1:
                cell.fill = PatternFill("solid", fgColor=LIGHT_GREY)


def _set_value_style(cell, hardcoded=False, linked=False, formula=False):
    color = BLACK
    if hardcoded:
        color = BLUE_INPUT
    elif linked:
        color = GREEN_LINK
    elif formula:
        color = BLACK

    cell.font = Font(name="Arial", size=10, color=color)


def _format_cols(ws, money_cols=None, share_cols=None, pct_cols=None, multiple_cols=None, start=1, end=200):
    money_cols = money_cols or []
    share_cols = share_cols or []
    pct_cols = pct_cols or []
    multiple_cols = multiple_cols or []

    for col in money_cols:
        for row in range(start, end + 1):
            ws.cell(row, col).number_format = MONEY_FMT

    for col in share_cols:
        for row in range(start, end + 1):
            ws.cell(row, col).number_format = SHARE_FMT

    for col in pct_cols:
        for row in range(start, end + 1):
            ws.cell(row, col).number_format = PERCENT_FMT

    for col in multiple_cols:
        for row in range(start, end + 1):
            ws.cell(row, col).number_format = MULTIPLE_FMT


def _widths(ws, mapping):
    for col, width in mapping.items():
        ws.column_dimensions[col].width = width


def _print_setup(ws, orientation="portrait"):
    ws.page_setup.orientation = orientation
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 0

    ws.page_margins.left = 0.35
    ws.page_margins.right = 0.35
    ws.page_margins.top = 0.45
    ws.page_margins.bottom = 0.45

    ws.oddFooter.left.text = "Valuation Working - Private & Confidential"
    ws.oddFooter.right.text = "Registered Valuer Review Required"


# ============================================================
# EXCEL SHEETS
# ============================================================

def _excel_cover(wb, assignment, report_status):
    ws = wb.active
    ws.title = "Cover"
    _hide_gridlines(ws)

    ws.column_dimensions["A"].width = 3
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 42

    ws["B3"] = "VALUATION WORKING"
    ws["B3"].font = Font(name="Arial", size=22, bold=True, color=NAVY)

    ws["B5"] = assignment.get("company_name", "")
    ws["B5"].font = Font(name="Arial", size=17, bold=True, color=BLACK)

    rows = [
        ("Assignment ID", assignment.get("assignment_id", "")),
        ("Valuation Date", assignment.get("valuation_date", "")),
        ("Engagement Date", assignment.get("engagement_date", "")),
        ("Report Date", assignment.get("report_date", "")),
        ("Purpose", assignment.get("purpose", "")),
        ("Security", assignment.get("security_type", "")),
        ("Applicable Provision", assignment.get("applicable_provision", "")),
        ("Status", report_status),
    ]

    r = 8
    for label, value in rows:
        ws.cell(r, 2, label).font = Font(name="Arial", bold=True, color=NAVY)
        ws.cell(r, 3, value)
        ws.cell(r, 2).border = Border(bottom=THIN)
        ws.cell(r, 3).border = Border(bottom=THIN)
        ws.cell(r, 3).alignment = Alignment(wrap_text=True)
        r += 2

    ws.merge_cells("B27:C30")
    note = ws["B27"]
    note.value = (
        "PRIVATE & CONFIDENTIAL\n"
        "This working is generated for valuation analysis and remains subject "
        "to review and professional judgement of the Registered Valuer."
    )
    note.fill = PatternFill("solid", fgColor=NAVY)
    note.font = Font(name="Arial", size=10, bold=True, color=WHITE)
    note.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    _print_setup(ws)


def _excel_assumptions(wb, assignment, valuation_inputs, wacc_analysis, review_summary):
    ws = wb.create_sheet("Assumptions")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        "Key valuation assumptions and control inputs",
        end_col=5,
    )

    _navy_bar(ws, 5, "VALUATION ASSUMPTIONS", 2, 5)
    _header(ws, 6, ["Particulars", "Value", "Basis / Source", "Review Note"], 2)

    sources = wacc_analysis.get("sources", {}) or {}

    items = [
        ("Tax Rate", valuation_inputs.get("tax_rate_percent", 0) / 100, "Valuer input", ""),
        ("WACC", valuation_inputs.get("wacc_percent", 0) / 100, "Approved / selected WACC", ""),
        ("Terminal Growth Rate", valuation_inputs.get("terminal_growth_percent", 0) / 100, "Valuer input", ""),
        ("Risk-free Rate", _f(wacc_analysis.get("risk_free_rate_percent")) / 100, sources.get("risk_free_rate", ""), ""),
        ("Equity Risk Premium", _f(wacc_analysis.get("equity_risk_premium_percent")) / 100, sources.get("equity_risk_premium", ""), ""),
        ("Beta", wacc_analysis.get("beta"), sources.get("beta", ""), ""),
        ("Company-specific Risk Premium", _f(wacc_analysis.get("company_specific_risk_premium_percent")) / 100, "Valuer judgement", ""),
        ("Pre-tax Cost of Debt", _f(wacc_analysis.get("pre_tax_cost_of_debt_percent")) / 100, sources.get("cost_of_debt", ""), ""),
        ("Equity Weight", _f(wacc_analysis.get("equity_weight_percent")) / 100, "Capital structure assumption", ""),
        ("Debt Weight", _f(wacc_analysis.get("debt_weight_percent")) / 100, "Capital structure assumption", ""),
        ("Cash (Rs. Lakhs)", valuation_inputs.get("cash", 0), "Valuation-date balance sheet / override", ""),
        ("Debt (Rs. Lakhs)", valuation_inputs.get("debt", 0), "Valuation-date balance sheet / override", ""),
        ("Non-operating Assets (Rs. Lakhs)", valuation_inputs.get("non_operating_assets", 0), "Valuer input", ""),
        ("Fully Diluted Shares", valuation_inputs.get("diluted_shares", 0), "Capital structure working", ""),
        ("Final Report Ready", _yes_no(review_summary.get("final_report_ready")), "System review control", ""),
    ]

    row = 7
    for label, value, basis, note in items:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        ws.cell(row, 4, basis)
        ws.cell(row, 5, note)

        _set_value_style(ws.cell(row, 3), hardcoded=True)
        row += 1

    _border_range(ws, 7, row - 1, 2, 5, zebra=True)

    for r in range(7, 17):
        ws.cell(r, 3).number_format = PERCENT_FMT

    ws.cell(12, 3).number_format = MULTIPLE_FMT

    for r in [17, 18, 19]:
        ws.cell(r, 3).number_format = MONEY_FMT

    ws.cell(20, 3).number_format = SHARE_FMT

    _widths(ws, {"A": 3, "B": 35, "C": 20, "D": 45, "E": 32})
    _print_setup(ws, "landscape")


def _excel_company_profile(wb, assignment):
    ws = wb.create_sheet("Company Profile")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        "Corporate and transaction profile",
        end_col=4,
    )

    _navy_bar(ws, 5, "COMPANY / ASSIGNMENT PROFILE", 2, 4)

    items = [
        ("Company Name", assignment.get("company_name", "")),
        ("CIN", assignment.get("cin", "")),
        ("PAN", assignment.get("pan", "")),
        ("Valuation Date", assignment.get("valuation_date", "")),
        ("Purpose", assignment.get("purpose", "")),
        ("Security / Subject Interest", assignment.get("security_type", "")),
        ("Applicable Provision", assignment.get("applicable_provision", "")),
        ("Transaction Details", assignment.get("transaction_details", "")),
    ]

    row = 6
    for label, value in items:
        ws.cell(row, 2, label).font = Font(name="Arial", bold=True, color=NAVY)
        ws.cell(row, 3, value)
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=4)
        ws.cell(row, 3).alignment = Alignment(wrap_text=True, vertical="top")
        ws.cell(row, 2).border = Border(bottom=THIN)
        ws.cell(row, 3).border = Border(bottom=THIN)
        row += 2

    _widths(ws, {"A": 3, "B": 30, "C": 48, "D": 25})
    _print_setup(ws)


def _excel_financials(wb, assignment, financial_analysis):
    company = assignment.get("company_name", "")
    hist = _rows(financial_analysis, "historical")
    prov = _rows(financial_analysis, "provisional")
    proj = _rows(financial_analysis, "projected")

    # Valuation-date balance sheet
    ws = wb.create_sheet("BS - Valuation Date")
    _sheet_title(ws, company, "Valuation-date / provisional balance sheet summary (Rs. Lakhs)", 5)
    _navy_bar(ws, 5, "BALANCE SHEET SUMMARY", 2, 5)

    _header(ws, 6, ["Particulars", "Amount", "Source / Period", "Review"], 2)

    valuation_row = None
    for r in prov:
        period = str(r.get("period", ""))
        if not period.startswith("FY") and not period.startswith("STUB"):
            valuation_row = r
            break
    if valuation_row is None and prov:
        valuation_row = prov[-1]

    metrics = (valuation_row or {}).get("metrics", {}) or {}
    period_label = (valuation_row or {}).get("display") or (valuation_row or {}).get("period", "")

    bs_items = [
        ("Total Assets", metrics.get("total_assets")),
        ("Current Assets", metrics.get("current_assets")),
        ("Cash & Cash Equivalents", metrics.get("cash")),
        ("Inventory", metrics.get("inventory")),
        ("Trade Receivables", metrics.get("trade_receivables")),
        ("Net Worth / Shareholders' Funds", metrics.get("net_worth")),
        ("Total Debt", metrics.get("total_debt")),
        ("Current Liabilities", metrics.get("current_liabilities")),
        ("Trade Payables", metrics.get("trade_payables")),
        ("Capital Employed", metrics.get("capital_employed")),
    ]

    row = 7
    for label, value in bs_items:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        ws.cell(row, 4, period_label)
        ws.cell(row, 5, "")
        _set_value_style(ws.cell(row, 3), linked=True)
        row += 1

    _border_range(ws, 7, row - 1, 2, 5, zebra=True)
    _format_cols(ws, money_cols=[3], start=7, end=row - 1)
    _widths(ws, {"A": 3, "B": 34, "C": 18, "D": 24, "E": 35})
    _print_setup(ws)

    # Historical + projected P&L
    ws = wb.create_sheet("Projected Financials")
    _sheet_title(
        ws,
        company,
        "Historical and projected financial statements (Rs. Lakhs)",
        end_col=10,
    )

    all_rows = hist + proj
    periods = [
        (r.get("display") or r.get("period", ""))
        for r in all_rows
    ]

    _navy_bar(ws, 6, "PROFIT & LOSS SUMMARY", 2, max(3, 2 + len(periods)))
    headers = ["Particulars"] + periods
    _header(ws, 7, headers, 2)

    metrics_to_show = [
        ("Revenue", "revenue"),
        ("EBITDA", "ebitda"),
        ("EBITDA Margin", "ebitda_margin"),
        ("EBIT", "ebit"),
        ("PAT", "pat"),
        ("PAT Margin", "pat_margin"),
        ("Depreciation", "depreciation"),
        ("Capital Expenditure", "capex"),
        ("Change in Working Capital", "change_working_capital"),
    ]

    row = 8
    for label, key in metrics_to_show:
        ws.cell(row, 2, label)
        for idx, data_row in enumerate(all_rows, start=3):
            value = (data_row.get("metrics", {}) or {}).get(key)
            ws.cell(row, idx, value)
            _set_value_style(ws.cell(row, idx), linked=True)
        row += 1

    _border_range(ws, 8, row - 1, 2, 2 + len(periods), zebra=False)

    for r in range(8, row):
        if "Margin" in str(ws.cell(r, 2).value):
            for c in range(3, 3 + len(periods)):
                ws.cell(r, c).number_format = '0.00"%"'
        else:
            for c in range(3, 3 + len(periods)):
                ws.cell(r, c).number_format = MONEY_FMT

    row += 2
    _navy_bar(ws, row, "BALANCE SHEET / CAPITAL SUMMARY", 2, max(3, 2 + len(periods)))
    row += 1
    _header(ws, row, headers, 2)
    row += 1

    bs_metrics = [
        ("Net Worth", "net_worth"),
        ("Total Debt", "total_debt"),
        ("Debt / Equity", "debt_equity"),
        ("Current Ratio", "current_ratio"),
        ("Interest Coverage", "interest_coverage"),
        ("ROCE", "roce"),
        ("ROE", "roe"),
    ]

    start_bs = row
    for label, key in bs_metrics:
        ws.cell(row, 2, label)
        for idx, data_row in enumerate(all_rows, start=3):
            value = (data_row.get("metrics", {}) or {}).get(key)
            ws.cell(row, idx, value)
            _set_value_style(ws.cell(row, idx), linked=True)
        row += 1

    _border_range(ws, start_bs, row - 1, 2, 2 + len(periods), zebra=False)

    for r in range(start_bs, row):
        label = str(ws.cell(r, 2).value)
        if label in {"ROCE", "ROE"}:
            for c in range(3, 3 + len(periods)):
                ws.cell(r, c).number_format = '0.00"%"'
        elif label in {"Debt / Equity", "Current Ratio", "Interest Coverage"}:
            for c in range(3, 3 + len(periods)):
                ws.cell(r, c).number_format = '0.00x'
        else:
            for c in range(3, 3 + len(periods)):
                ws.cell(r, c).number_format = MONEY_FMT

    _widths(ws, {"A": 3, "B": 38})
    for c in range(3, 3 + len(periods)):
        ws.column_dimensions[get_column_letter(c)].width = 16

    _print_setup(ws, "landscape")


def _excel_wacc(wb, assignment, valuation_inputs, wacc_analysis):
    ws = wb.create_sheet("WACC")
    _sheet_title(ws, assignment.get("company_name", ""), "Weighted Average Cost of Capital - Working", 4)

    sources = wacc_analysis.get("sources", {}) or {}

    sections = [
        (
            "COST OF EQUITY (CAPM)",
            [
                ("Risk-free Rate (Rf)", _f(wacc_analysis.get("risk_free_rate_percent")) / 100, PERCENT_FMT, sources.get("risk_free_rate", "")),
                ("Equity Risk Premium (ERP)", _f(wacc_analysis.get("equity_risk_premium_percent")) / 100, PERCENT_FMT, sources.get("equity_risk_premium", "")),
                ("Levered Beta (Beta)", wacc_analysis.get("beta"), MULTIPLE_FMT, sources.get("beta", "")),
                ("Cost of Equity: Ke = Rf + Beta x ERP + CSRP", _f(wacc_analysis.get("cost_of_equity_percent")) / 100, PERCENT_FMT, "Calculated"),
            ],
        ),
        (
            "COST OF DEBT",
            [
                ("Pre-tax Cost of Debt (Kd)", _f(wacc_analysis.get("pre_tax_cost_of_debt_percent")) / 100, PERCENT_FMT, sources.get("cost_of_debt", "")),
                ("Corporate Tax Rate (t)", _f(wacc_analysis.get("tax_rate_percent", valuation_inputs.get("tax_rate_percent"))) / 100, PERCENT_FMT, "Applicable / selected rate"),
                ("Post-tax Cost of Debt: Kd x (1-t)", _f(wacc_analysis.get("after_tax_cost_of_debt_percent")) / 100, PERCENT_FMT, "Calculated"),
            ],
        ),
        (
            "CAPITAL STRUCTURE",
            [
                ("Weight of Debt: D/(D+E)", _f(wacc_analysis.get("debt_weight_percent")) / 100, PERCENT_FMT, "Selected capital structure"),
                ("Weight of Equity: E/(D+E)", _f(wacc_analysis.get("equity_weight_percent")) / 100, PERCENT_FMT, "Selected capital structure"),
            ],
        ),
    ]

    row = 5
    for title, items in sections:
        _navy_bar(ws, row, title, 2, 4)
        row += 1
        for label, value, numfmt, source in items:
            ws.cell(row, 2, label)
            ws.cell(row, 3, value)
            ws.cell(row, 4, source)
            ws.cell(row, 3).number_format = numfmt
            _set_value_style(ws.cell(row, 3), hardcoded=("Calculated" not in source))
            row += 1
        row += 1

    _navy_bar(ws, row, "WACC CONCLUSION", 2, 4)
    row += 1

    ws.cell(row, 2, "WACC = Ke x We + Kd(1-t) x Wd")
    ws.cell(row, 3, _f(valuation_inputs.get("wacc_percent")) / 100)
    ws.cell(row, 3).number_format = PERCENT_FMT
    ws.cell(row, 3).fill = PatternFill("solid", fgColor=YELLOW)
    ws.cell(row, 3).font = Font(name="Arial", bold=True, color=BLACK)

    row += 2
    ws.cell(row, 2, "Terminal Growth Rate (g)")
    ws.cell(row, 3, _f(valuation_inputs.get("terminal_growth_percent")) / 100)
    ws.cell(row, 3).number_format = PERCENT_FMT
    _set_value_style(ws.cell(row, 3), hardcoded=True)

    row += 1
    ws.cell(row, 2, "Check: WACC > g ?")
    ws.cell(row, 3, "OK - WACC exceeds g" if _f(valuation_inputs.get("wacc_percent")) > _f(valuation_inputs.get("terminal_growth_percent")) else "CHECK REQUIRED")

    _widths(ws, {"A": 3, "B": 48, "C": 18, "D": 55})
    _print_setup(ws)


def _excel_dcf(wb, assignment, dcf_result, valuation_inputs):
    ws = wb.create_sheet("DCF Valuation")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        "Income Approach - Discounted Cash Flow (FCFF) Method | All figures in Rs. Lakhs unless stated otherwise",
        8,
    )

    fcff = dcf_result.get("fcff", []) or []
    years = [str(r.get("year", "")) for r in fcff]

    _header(ws, 6, ["Particulars"] + years, 2)

    items = [
        ("EBIT (per Projected Financials)", "ebit"),
        ("Less: Tax on EBIT", "tax"),
        ("NOPAT (EBIT less tax)", "nopat"),
        ("Add: Depreciation & Amortisation", "depreciation"),
        ("Less: Capital Expenditure", "capex"),
        ("Less: Increase in Net Working Capital", "change_working_capital"),
        ("Free Cash Flow to Firm (FCFF)", "fcff"),
        ("", None),
        ("Discount Factor @ WACC", "discount_factor"),
        ("Present Value of FCFF", "pv_fcff"),
    ]

    row = 7
    for label, key in items:
        ws.cell(row, 2, label)

        if key is None:
            row += 1
            continue

        for idx, item in enumerate(fcff, start=3):
            if key == "tax":
                value = _f(item.get("ebit")) - _f(item.get("nopat"))
            else:
                value = item.get(key)

            ws.cell(row, idx, value)

            if key in {"ebit", "depreciation"}:
                _set_value_style(ws.cell(row, idx), linked=True)

        if label in {"NOPAT (EBIT less tax)", "Free Cash Flow to Firm (FCFF)", "Present Value of FCFF"}:
            ws.cell(row, 2).font = Font(name="Arial", bold=True, color=BLACK)

        row += 1

    max_col = 2 + max(1, len(years))
    _border_range(ws, 7, row - 1, 2, max_col, zebra=False)

    for r in range(7, row):
        label = str(ws.cell(r, 2).value or "")
        for c in range(3, max_col + 1):
            if "Discount Factor" in label:
                ws.cell(r, c).number_format = "0.0000"
            else:
                ws.cell(r, c).number_format = MONEY_FMT

    row += 1
    _navy_bar(ws, row, "TERMINAL VALUE", 2, 5)
    row += 1

    terminal_rows = [
        ("FCFF - Final Forecast Year", (fcff[-1].get("fcff") if fcff else 0), MONEY_FMT),
        ("Terminal Growth Rate (g)", _f(valuation_inputs.get("terminal_growth_percent")) / 100, PERCENT_FMT),
        ("Terminal Value", dcf_result.get("terminal_value", 0), MONEY_FMT),
        ("Present Value of Terminal Value", dcf_result.get("pv_terminal_value", 0), MONEY_FMT),
    ]

    for label, value, fmt in terminal_rows:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        ws.cell(row, 3).number_format = fmt
        row += 1

    row += 1
    _navy_bar(ws, row, "ENTERPRISE VALUE TO EQUITY VALUE BRIDGE", 2, 5)
    row += 1

    bridge = [
        ("Sum of PV of explicit period FCFF", dcf_result.get("pv_explicit_fcff", 0), MONEY_FMT),
        ("Add: PV of Terminal Value", dcf_result.get("pv_terminal_value", 0), MONEY_FMT),
        ("Enterprise Value (Operating Business)", dcf_result.get("enterprise_value", 0), MONEY_FMT),
        ("Add: Cash & Cash Equivalents", valuation_inputs.get("cash", 0), MONEY_FMT),
        ("Add: Non-Operating Assets", valuation_inputs.get("non_operating_assets", 0), MONEY_FMT),
        ("Less: Debt", valuation_inputs.get("debt", 0), MONEY_FMT),
        ("Equity Value under DCF", dcf_result.get("equity_value", 0), MONEY_FMT),
        ("No. of Fully Diluted Equity Shares", valuation_inputs.get("diluted_shares", 0), SHARE_FMT),
        ("Value per Equity Share under DCF (Rs.)", dcf_result.get("value_per_share", 0), PER_SHARE_FMT),
    ]

    for label, value, fmt in bridge:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        ws.cell(row, 3).number_format = fmt

        if label == "Equity Value under DCF":
            ws.cell(row, 3).fill = PatternFill("solid", fgColor=YELLOW)
            ws.cell(row, 3).font = Font(name="Arial", bold=True)
        elif label == "Value per Equity Share under DCF (Rs.)":
            ws.cell(row, 3).fill = PatternFill("solid", fgColor=LIGHT_GREEN)
            ws.cell(row, 3).font = Font(name="Arial", bold=True)
        elif label == "Enterprise Value (Operating Business)":
            ws.cell(row, 3).fill = PatternFill("solid", fgColor=LIGHT_BLUE)
            ws.cell(row, 3).font = Font(name="Arial", bold=True)

        row += 1

    _widths(ws, {"A": 3, "B": 58})
    for c in range(3, max_col + 1):
        ws.column_dimensions[get_column_letter(c)].width = 18

    _print_setup(ws, "landscape")


def _excel_sensitivity(wb, assignment, dcf_result):
    ws = wb.create_sheet("DCF Sensitivity")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        "DCF Sensitivity - Value per Equity Share (Rs.)",
        8,
    )

    sensitivity = dcf_result.get("sensitivity", {}) or {}
    waccs = sensitivity.get("wacc_values_percent", []) or []

    headers = ["Terminal Growth / WACC"] + [f"{_f(x):.2f}%" for x in waccs]
    _header(ws, 6, headers, 2)

    base_wacc = sensitivity.get("base_wacc_percent")
    base_g = sensitivity.get("base_terminal_growth_percent")

    row = 7
    for sens_row in sensitivity.get("rows", []) or []:
        g = sens_row.get("terminal_growth")
        ws.cell(row, 2, _f(g) / 100)
        ws.cell(row, 2).number_format = PERCENT_FMT

        for idx, cell_data in enumerate(sens_row.get("values", []) or [], start=3):
            value = cell_data.get("value_per_share")
            ws.cell(row, idx, value)
            ws.cell(row, idx).number_format = PER_SHARE_FMT

            if (
                g is not None
                and base_g is not None
                and abs(_f(g) - _f(base_g)) < 0.001
                and cell_data.get("wacc") is not None
                and base_wacc is not None
                and abs(_f(cell_data.get("wacc")) - _f(base_wacc)) < 0.001
            ):
                ws.cell(row, idx).fill = PatternFill("solid", fgColor=YELLOW)
                ws.cell(row, idx).font = Font(name="Arial", bold=True)

        row += 1

    _border_range(ws, 7, row - 1, 2, 2 + len(waccs), zebra=True)
    _widths(ws, {"A": 3, "B": 26})
    for c in range(3, 3 + len(waccs)):
        ws.column_dimensions[get_column_letter(c)].width = 18

    _print_setup(ws, "landscape")


def _excel_nav(wb, assignment, nav_result, valuation_inputs):
    ws = wb.create_sheet("NAV Valuation")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        "Asset Approach - Net Asset Value Method | All figures in Rs. Lakhs unless stated otherwise",
        5,
    )

    _navy_bar(ws, 5, "NAV WORKING", 2, 5)
    _header(ws, 6, ["Particulars", "Amount", "Basis / Comment", "Review"], 2)

    rows = [
        ("Total / Adjusted Assets", valuation_inputs.get("adjusted_assets", 0), "Valuation-date balance sheet / valuer adjustment", ""),
        ("Less: Total / Adjusted Liabilities", valuation_inputs.get("adjusted_liabilities", 0), "Valuation-date balance sheet / valuer adjustment", ""),
        ("Net Assets available to Equity Shareholders", nav_result.get("equity_value", 0), "Assets less liabilities", ""),
        ("Fully Diluted Equity Shares", valuation_inputs.get("diluted_shares", 0), "Capital structure working", ""),
        ("Value per Equity Share under NAV (Rs.)", nav_result.get("value_per_share", 0), "NAV / diluted shares", ""),
    ]

    row = 7
    for label, value, basis, review in rows:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        ws.cell(row, 4, basis)
        ws.cell(row, 5, review)

        if label == "Value per Equity Share under NAV (Rs.)":
            ws.cell(row, 3).number_format = PER_SHARE_FMT
            ws.cell(row, 3).fill = PatternFill("solid", fgColor=LIGHT_GREEN)
            ws.cell(row, 3).font = Font(name="Arial", bold=True)
        elif "Shares" in label:
            ws.cell(row, 3).number_format = SHARE_FMT
        else:
            ws.cell(row, 3).number_format = MONEY_FMT

        row += 1

    _border_range(ws, 7, row - 1, 2, 5, zebra=True)
    _widths(ws, {"A": 3, "B": 45, "C": 20, "D": 48, "E": 28})
    _print_setup(ws)


def _excel_cap_structure(wb, assignment, capital_structure):
    ws = wb.create_sheet("Capital Structure")
    _sheet_title(ws, assignment.get("company_name", ""), "Fully diluted capital structure", 7)

    _navy_bar(ws, 5, "FULLY DILUTED SHARE RECONCILIATION", 2, 7)

    items = [
        ("Basic Equity Shares", capital_structure.get("basic_equity_shares")),
        ("CCPS Outstanding", capital_structure.get("ccps_outstanding")),
        ("Equity from CCPS", capital_structure.get("equity_from_ccps")),
        ("Equity from Warrants", capital_structure.get("equity_from_warrants")),
        ("Equity from ESOPs", capital_structure.get("equity_from_esops")),
        ("Fully Diluted Shares", capital_structure.get("fully_diluted_shares")),
        ("Future Exercise Proceeds (Rs. Lakhs)", capital_structure.get("future_cash_receivable_on_exercise")),
        ("Validation Status", capital_structure.get("status", "")),
    ]

    row = 6
    for label, value in items:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        if "Proceeds" in label:
            ws.cell(row, 3).number_format = MONEY_FMT
        elif "Shares" in label or "Outstanding" in label or "Equity from" in label:
            ws.cell(row, 3).number_format = SHARE_FMT
        row += 1

    holders = capital_structure.get("holders", []) or []
    if holders:
        row += 1
        _navy_bar(ws, row, "HOLDER-WISE FULLY DILUTED SHAREHOLDING", 2, 8)
        row += 1
        _header(ws, row, ["Holder", "Basic Shares", "CCPS", "Warrants", "ESOPs", "Fully Diluted", "FD %"], 2)
        row += 1
        start = row

        for holder in holders:
            values = [
                holder.get("holder", ""),
                holder.get("basic_equity_shares"),
                holder.get("equity_from_ccps"),
                holder.get("equity_from_warrants"),
                holder.get("equity_from_esops"),
                holder.get("fully_diluted_shares"),
                holder.get("fully_diluted_percentage"),
            ]
            for idx, value in enumerate(values, start=2):
                ws.cell(row, idx, value)
            row += 1

        _border_range(ws, start, row - 1, 2, 8, zebra=True)
        for c in range(3, 8):
            for r in range(start, row):
                ws.cell(r, c).number_format = SHARE_FMT
        for r in range(start, row):
            ws.cell(r, 8).number_format = '0.00"%"'

    _widths(ws, {"A": 3, "B": 34, "C": 18, "D": 18, "E": 18, "F": 18, "G": 18, "H": 16})
    _print_setup(ws, "landscape")


def _excel_summary(wb, assignment, dcf_result, nav_result, weightage_result, valuation_inputs, report_status):
    ws = wb.create_sheet("Valuation Summary")
    _sheet_title(
        ws,
        assignment.get("company_name", ""),
        f"Valuation Date: {assignment.get('valuation_date', '')}   |   Fair value per equity share",
        7,
    )

    _navy_bar(ws, 5, "VALUATION SUMMARY - FAIR VALUE PER EQUITY SHARE", 2, 7)
    _header(
        ws,
        6,
        ["Method", "Approach", "Indicated Value per Share (Rs.)", "Weight Assigned", "Weighted Value (Rs.)", "Comment"],
        2,
    )

    methods = weightage_result.get("methods", []) or []
    row = 7

    for method in methods:
        name = str(method.get("method", ""))
        approach = "Income Approach" if "DCF" in name.upper() else ("Asset Approach" if "NAV" in name.upper() else "Other")
        per_share = None

        if "DCF" in name.upper():
            per_share = dcf_result.get("value_per_share")
        elif "NAV" in name.upper():
            per_share = nav_result.get("value_per_share")
        else:
            value = _f(method.get("value"))
            shares = _f(valuation_inputs.get("diluted_shares"))
            per_share = (value * 100000 / shares) if shares > 0 else 0

        weight = _f(method.get("weight")) / 100
        weighted_per_share = _f(per_share) * weight

        ws.cell(row, 2, name)
        ws.cell(row, 3, approach)
        ws.cell(row, 4, per_share)
        ws.cell(row, 5, weight)
        ws.cell(row, 6, weighted_per_share)
        ws.cell(row, 7, "")

        ws.cell(row, 4).number_format = PER_SHARE_FMT
        ws.cell(row, 5).number_format = PERCENT_FMT
        ws.cell(row, 6).number_format = PER_SHARE_FMT

        _set_value_style(ws.cell(row, 4), linked=True)
        _set_value_style(ws.cell(row, 5), hardcoded=True)
        row += 1

    ws.cell(row, 2, "Total")
    ws.cell(row, 5, f"=SUM(E7:E{row-1})")
    ws.cell(row, 6, f"=SUM(F7:F{row-1})")
    ws.cell(row, 2).font = Font(name="Arial", bold=True)
    ws.cell(row, 5).font = Font(name="Arial", bold=True)
    ws.cell(row, 6).font = Font(name="Arial", bold=True)
    ws.cell(row, 5).number_format = PERCENT_FMT
    ws.cell(row, 6).number_format = PER_SHARE_FMT

    _border_range(ws, 7, row, 2, 7, zebra=True)

    row += 2
    ws.cell(row, 2, "Fair Value per Equity Share (Rs.)")
    ws.cell(row, 2).font = Font(name="Arial", size=12, bold=True)
    ws.cell(row, 4, weightage_result.get("value_per_share", 0))
    ws.cell(row, 4).number_format = PER_SHARE_FMT
    ws.cell(row, 4).fill = PatternFill("solid", fgColor=YELLOW)
    ws.cell(row, 4).font = Font(name="Arial", size=13, bold=True, italic=True)

    row += 3
    _navy_bar(ws, row, "BASIS FOR WEIGHTS ASSIGNED", 2, 7)
    row += 1
    ws.merge_cells(start_row=row, start_column=2, end_row=row + 6, end_column=7)
    text = (
        "The weighting of valuation methods is a matter of professional judgement. "
        "The generated working reflects the weights selected in the application. "
        "The final report should document the rationale for the relative reliability "
        "of each method having regard to the Company's circumstances, quality of data, "
        "purpose of valuation and valuation standards."
    )
    ws.cell(row, 2, text)
    ws.cell(row, 2).font = Font(name="Arial", size=10, italic=True, color="666666")
    ws.cell(row, 2).alignment = Alignment(wrap_text=True, vertical="top")

    row += 8
    _navy_bar(ws, row, "STATUS / CAVEATS", 2, 7)
    row += 1
    ws.merge_cells(start_row=row, start_column=2, end_row=row + 3, end_column=7)
    ws.cell(row, 2, f"System output status: {report_status}. Final issuance remains subject to Registered Valuer review.")
    ws.cell(row, 2).font = Font(name="Arial", size=10, color=RED)
    ws.cell(row, 2).alignment = Alignment(wrap_text=True, vertical="top")

    _widths(ws, {"A": 3, "B": 38, "C": 24, "D": 22, "E": 18, "F": 22, "G": 35})
    _print_setup(ws, "landscape")


def _excel_review(wb, assignment, review_summary, financial_analysis):
    ws = wb.create_sheet("Review & Controls")
    _sheet_title(ws, assignment.get("company_name", ""), "Data validation and Registered Valuer review status", 8)

    _navy_bar(ws, 5, "REVIEW SUMMARY", 2, 8)

    summary = [
        ("Total Review Items", review_summary.get("total", 0)),
        ("Open Review Items", review_summary.get("unresolved", 0)),
        ("Material Open Items", review_summary.get("unresolved_material", 0)),
        ("Failed Financial Cross-checks", review_summary.get("failed_cross_checks", 0)),
        ("Failed Capital Checks", review_summary.get("failed_capital_checks", 0)),
        ("Data Ready for Valuation", _yes_no(review_summary.get("data_ready_for_valuation"))),
        ("Final Report Ready", _yes_no(review_summary.get("final_report_ready"))),
    ]

    row = 6
    for label, value in summary:
        ws.cell(row, 2, label)
        ws.cell(row, 3, value)
        row += 1

    review_items = financial_analysis.get("review_required", []) or []
    if review_items:
        row += 1
        _navy_bar(ws, row, "REVIEW ITEMS", 2, 9)
        row += 1
        _header(ws, row, ["Status", "Material", "File", "Sheet", "Row", "Source Label", "Suggested Field", "Reason"], 2)
        row += 1
        start = row

        for item in review_items:
            values = [
                item.get("review_status", "pending"),
                _yes_no(item.get("material")),
                item.get("file_name", ""),
                item.get("sheet", ""),
                item.get("row"),
                item.get("source_label", ""),
                item.get("canonical_field", ""),
                item.get("reason", ""),
            ]
            for idx, value in enumerate(values, start=2):
                ws.cell(row, idx, value)
            row += 1

        _border_range(ws, start, row - 1, 2, 9, zebra=True)

    _widths(ws, {"A": 3, "B": 20, "C": 15, "D": 28, "E": 18, "F": 10, "G": 30, "H": 25, "I": 45})
    _print_setup(ws, "landscape")


# ============================================================
# MAIN EXCEL GENERATOR
# ============================================================

def generate_excel_working(
    output_path: Path,
    assignment: dict,
    dcf_result: dict,
    nav_result: dict,
    weightage_result: dict,
    valuation_inputs: dict,
    wacc_analysis: Optional[dict] = None,
    capital_structure: Optional[dict] = None,
    review_summary: Optional[dict] = None,
    report_status: str = "DRAFT_ONLY_REVIEW_PENDING",
    financial_analysis: Optional[dict] = None,
):
    wacc_analysis = wacc_analysis or {}
    capital_structure = capital_structure or {}
    review_summary = review_summary or {}
    financial_analysis = financial_analysis or {}

    wb = Workbook()

    _excel_cover(wb, assignment, report_status)
    _excel_assumptions(wb, assignment, valuation_inputs, wacc_analysis, review_summary)
    _excel_company_profile(wb, assignment)
    _excel_financials(wb, assignment, financial_analysis)
    _excel_wacc(wb, assignment, valuation_inputs, wacc_analysis)
    _excel_dcf(wb, assignment, dcf_result, valuation_inputs)
    _excel_sensitivity(wb, assignment, dcf_result)
    _excel_nav(wb, assignment, nav_result, valuation_inputs)
    _excel_cap_structure(wb, assignment, capital_structure)
    _excel_summary(wb, assignment, dcf_result, nav_result, weightage_result, valuation_inputs, report_status)
    _excel_review(wb, assignment, review_summary, financial_analysis)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(output_path)
    return output_path


# ============================================================
# WORD STYLE HELPERS
# ============================================================

def _shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def _setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in [
        ("Heading 1", 13.5, NAVY),
        ("Heading 2", 10.5, GOLD),
    ]:
        st = doc.styles[style_name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    if "Report Note" not in doc.styles:
        note = doc.styles.add_style("Report Note", WD_STYLE_TYPE.PARAGRAPH)
        note.font.name = "Arial"
        note.font.size = Pt(8.5)
        note.font.italic = True
        note.font.color.rgb = RGBColor.from_string("666666")


def _set_header_footer(doc, assignment):
    section = doc.sections[0]

    h = section.header.paragraphs[0]
    h.text = ""
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left = h.add_run("Private & Confidential")
    left.font.name = "Arial"
    left.font.size = Pt(7.5)
    left.font.italic = True
    left.font.color.rgb = RGBColor.from_string("666666")

    h.add_run(" " * 16)

    right = h.add_run(
        f"Report on Fair Value of Equity Shares - {assignment.get('company_name', '')}"
    )
    right.font.name = "Arial"
    right.font.size = Pt(7.5)
    right.font.italic = True
    right.font.color.rgb = RGBColor.from_string("666666")

    f = section.footer.paragraphs[0]
    f.text = ""
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = f.add_run("Registered Valuer - Securities or Financial Assets")
    run.font.name = "Arial"
    run.font.size = Pt(7.5)
    run.font.color.rgb = RGBColor.from_string("666666")


def _table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = str(text)
        _shade(hdr[i], NAVY)
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)

    _repeat_header(table.rows[0])

    for ridx, row in enumerate(rows):
        cells = table.add_row().cells

        for i, value in enumerate(row):
            cells[i].text = "" if value is None else str(value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if ridx % 2 == 1:
                _shade(cells[i], LIGHT_GREY)

            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)

    return table


def _kv(doc, rows):
    t = _table(doc, ["Particulars", "Value / Details"], rows)

    for row in t.rows[1:]:
        _shade(row.cells[0], LIGHT_GREY)
        for p in row.cells[0].paragraphs:
            for r in p.runs:
                r.font.bold = True

    return t


def _note(doc, text):
    p = doc.add_paragraph(style="Report Note")
    p.add_run(text)
    return p


def _title_para(doc, text, size=19, color=NAVY):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(str(item))


# ============================================================
# WORD REPORT
# ============================================================

def generate_word_report(
    output_path: Path,
    assignment: dict,
    dcf_result: dict,
    nav_result: dict,
    weightage_result: dict,
    valuation_inputs: dict,
    wacc_analysis: Optional[dict] = None,
    capital_structure: Optional[dict] = None,
    review_summary: Optional[dict] = None,
    report_status: str = "DRAFT_ONLY_REVIEW_PENDING",
    financial_analysis: Optional[dict] = None,
):
    wacc_analysis = wacc_analysis or {}
    capital_structure = capital_structure or {}
    review_summary = review_summary or {}
    financial_analysis = financial_analysis or {}

    doc = Document()
    _setup_doc(doc)
    _set_header_footer(doc, assignment)

    company = assignment.get("company_name", "")
    valuation_date = assignment.get("valuation_date", "")
    report_date = assignment.get("report_date", "")

    # --------------------------------------------------------
    # COVER
    # --------------------------------------------------------
    doc.add_paragraph("")
    _title_para(doc, "VALUATION REPORT", 21)
    doc.add_paragraph("")
    _title_para(doc, company, 15, BLACK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        str(
            assignment.get("purpose", "")
            or "Valuation of Equity Shares"
        )
    )
    r.font.name = "Arial"
    r.font.size = Pt(10.5)

    doc.add_paragraph("")

    _kv(
        doc,
        [
            ["Valuation Date", valuation_date],
            ["Report Date", report_date],
            ["Security / Subject Interest", assignment.get("security_type", "")],
            ["Applicable Provision", assignment.get("applicable_provision", "")],
            ["System Status", report_status],
        ],
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STRICTLY PRIVATE AND CONFIDENTIAL")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    _note(
        doc,
        "Prepared solely for the stated purpose. This system-generated report is a draft "
        "and must be reviewed, modified as necessary and approved by the Registered Valuer before issuance."
    )

    doc.add_page_break()

    # --------------------------------------------------------
    # CONTENTS
    # --------------------------------------------------------
    doc.add_heading("Contents", level=1)

    contents = [
        "1. Executive Summary",
        "2. Purpose and Scope of the Assignment",
        "3. Sources of Information",
        "4. Background of the Company",
        "5. Regulatory Framework and Basis of Valuation",
        "6. Valuation Approaches and Methodology",
        "7. Key Assumptions and Rationale",
        "8. Valuation Summary, Weightage and Recommended Fair Value",
        "9. Limiting Conditions and Disclaimer",
        "10. Valuer Review / Declaration",
        "ANNEXURE A - Income Approach - Discounted Cash Flow Working",
        "ANNEXURE B - Asset Approach - Net Asset Value Working",
        "ANNEXURE C - Weighted Average Cost of Capital Working",
        "ANNEXURE D - Historical and Projected Financial Statements",
        "ANNEXURE E - Summary of Key Assumptions and Review Items",
    ]

    for item in contents:
        doc.add_paragraph(item)

    doc.add_page_break()

    # --------------------------------------------------------
    # 1. EXECUTIVE SUMMARY
    # --------------------------------------------------------
    doc.add_heading("1. Executive Summary", level=1)

    doc.add_paragraph(
        f"This draft report sets out the valuation analysis of {company} as at "
        f"{valuation_date} ('the Valuation Date') for the purpose stated in Section 2 below. "
        "The valuation conclusion has been derived using the valuation methods selected in "
        "the application and is subject to the assumptions, qualifications, source data and "
        "limiting conditions set out in this Report and its Annexures."
    )

    _kv(
        doc,
        [
            ["DCF Equity Value (Rs. Lakhs)", _fmt(dcf_result.get("equity_value"))],
            ["DCF Value per Equity Share (Rs.)", _fmt(dcf_result.get("value_per_share"))],
            ["NAV Equity Value (Rs. Lakhs)", _fmt(nav_result.get("equity_value"))],
            ["NAV Value per Equity Share (Rs.)", _fmt(nav_result.get("value_per_share"))],
            ["Concluded Equity Value (Rs. Lakhs)", _fmt(weightage_result.get("concluded_value"))],
            ["Recommended Fair Value per Equity Share (Rs.)", _fmt(weightage_result.get("value_per_share"))],
            ["Fully Diluted Equity Shares", _fmt_int(valuation_inputs.get("diluted_shares"))],
        ],
    )

    doc.add_paragraph(
        "This Executive Summary must be read together with the whole of this Report, "
        "including the detailed workings and assumptions in the Annexures."
    )

    # --------------------------------------------------------
    # 2. PURPOSE / SCOPE
    # --------------------------------------------------------
    doc.add_heading("2. Purpose and Scope of the Assignment", level=1)

    doc.add_heading("2.1 Purpose of Valuation", level=2)
    doc.add_paragraph(str(assignment.get("purpose", "")))

    doc.add_heading("2.2 Transaction Background", level=2)
    transaction = assignment.get("transaction_details", "")
    doc.add_paragraph(
        str(transaction)
        if transaction
        else "The transaction background should be completed from the engagement information before final issuance."
    )

    doc.add_heading("2.3 Basis and Standard of Value", level=2)
    doc.add_paragraph(
        "The valuation has been prepared on the basis recorded in the engagement and "
        "subject to the applicable valuation standards and regulatory requirements. "
        "The final signed report should expressly state the adopted basis of value and "
        "premise of value after Registered Valuer review."
    )

    # --------------------------------------------------------
    # 3. SOURCES
    # --------------------------------------------------------
    doc.add_heading("3. Sources of Information", level=1)

    docs = _all_source_documents(assignment)
    if docs:
        _add_bullets(doc, docs)
    else:
        _add_bullets(
            doc,
            [
                "Provisional financial statements as at the Valuation Date.",
                "Historical audited financial statements made available for the assignment.",
                "Projected financial statements / business plan.",
                "Capital structure and fully diluted shareholding information.",
                "Management explanations and transaction information.",
                "Market / WACC data recorded in the valuation working.",
            ],
        )

    doc.add_paragraph(
        "The information referred to above has been relied upon without independent audit "
        "or verification unless expressly stated otherwise in the final report."
    )

    # --------------------------------------------------------
    # 4. BACKGROUND
    # --------------------------------------------------------
    doc.add_heading("4. Background of the Company", level=1)

    _kv(
        doc,
        [
            ["Name of the Company", company],
            ["CIN", assignment.get("cin", "")],
            ["PAN", assignment.get("pan", "")],
            ["Security / Subject Interest", assignment.get("security_type", "")],
            ["Valuation Date", valuation_date],
        ],
    )

    doc.add_paragraph(
        "The detailed corporate background, nature of business, history, operations and "
        "recent corporate actions should be completed from the company profile and engagement "
        "documents before final issuance. The system does not invent information that has not "
        "been supplied or extracted."
    )

    # --------------------------------------------------------
    # 5. REGULATORY
    # --------------------------------------------------------
    doc.add_heading("5. Regulatory Framework and Basis of Valuation", level=1)

    doc.add_heading("5.1 Applicable Provision", level=2)
    doc.add_paragraph(str(assignment.get("applicable_provision", "")))

    doc.add_heading("5.2 Valuation Standards", level=2)
    doc.add_paragraph(
        "The final report should identify the valuation standards and statutory / regulatory "
        "provisions applicable to the actual purpose of valuation. The methodology in this draft "
        "has been structured around generally recognised Income and Asset approaches."
    )

    # --------------------------------------------------------
    # 6. METHODS
    # --------------------------------------------------------
    doc.add_heading("6. Valuation Approaches and Methodology", level=1)

    doc.add_heading("6.1 Income Approach - Discounted Cash Flow Method", level=2)
    doc.add_paragraph(
        "Under the Discounted Cash Flow ('DCF') Method, the enterprise value is estimated as "
        "the present value of projected Free Cash Flow to Firm ('FCFF') discounted at the "
        "Weighted Average Cost of Capital ('WACC'), together with the present value of a terminal "
        "value representing cash flows beyond the explicit forecast period."
    )

    doc.add_heading("6.2 Asset Approach - Net Asset Value Method", level=2)
    doc.add_paragraph(
        "Under the Net Asset Value ('NAV') Method, the equity value is determined by reference "
        "to the adjusted value of assets less adjusted liabilities. The figures automatically "
        "populated by the system are starting values and require review for fair-value and "
        "off-balance-sheet adjustments."
    )

    # --------------------------------------------------------
    # 7. ASSUMPTIONS
    # --------------------------------------------------------
    doc.add_heading("7. Key Assumptions and Rationale", level=1)

    assumptions = [
        f"The valuation date is {valuation_date}.",
        f"The WACC used in the DCF is {_fmt_pct(valuation_inputs.get('wacc_percent'))}.",
        f"The terminal growth rate is {_fmt_pct(valuation_inputs.get('terminal_growth_percent'))}.",
        f"The tax rate applied in the DCF is {_fmt_pct(valuation_inputs.get('tax_rate_percent'))}.",
        f"Fully diluted equity shares considered are {_fmt_int(valuation_inputs.get('diluted_shares'))}.",
        "Financial projections have been relied upon as management estimates and remain subject to uncertainty.",
        "Extracted / automatically mapped information remains subject to Registered Valuer review.",
    ]

    _add_bullets(doc, assumptions)

    # --------------------------------------------------------
    # 8. SUMMARY / WEIGHTING
    # --------------------------------------------------------
    doc.add_heading("8. Valuation Summary, Weightage and Recommended Fair Value", level=1)

    method_rows = []
    for method in weightage_result.get("methods", []) or []:
        name = str(method.get("method", ""))
        approach = "Income Approach" if "DCF" in name.upper() else ("Asset Approach" if "NAV" in name.upper() else "Other")
        if "DCF" in name.upper():
            per_share = dcf_result.get("value_per_share")
        elif "NAV" in name.upper():
            per_share = nav_result.get("value_per_share")
        else:
            value = _f(method.get("value"))
            shares = _f(valuation_inputs.get("diluted_shares"))
            per_share = value * 100000 / shares if shares else 0

        weight = _f(method.get("weight"))
        method_rows.append([
            name,
            approach,
            _fmt(per_share),
            _fmt_pct(weight),
            _fmt(_f(per_share) * weight / 100),
        ])

    _table(
        doc,
        ["Method", "Approach", "Indicated Value per Share (Rs.)", "Weight Assigned", "Weighted Value (Rs.)"],
        method_rows,
    )

    p = doc.add_paragraph()
    r = p.add_run(
        "Total Weighted Fair Value per Equity Share: Rs. "
        + _fmt(weightage_result.get("value_per_share"))
    )
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)

    doc.add_heading("8.1 Weightage - Rationale for Weights Assigned", level=2)
    doc.add_paragraph(
        "The weights assigned to the valuation methods are a matter of professional judgement. "
        "The final report should record the specific rationale for the weights selected, having "
        "regard to the reliability of the projections, relevance of asset values, nature of the "
        "business, purpose of valuation and other facts and circumstances."
    )

    doc.add_heading("8.2 Recommended Fair Value", level=2)
    doc.add_paragraph(
        "Based on the analysis set out in this draft Report and subject to the assumptions, "
        "scope limitations and qualifications stated herein, the concluded value per equity "
        "share is set out below:"
    )

    _kv(
        doc,
        [
            ["Recommended Fair Value per Equity Share (Rs.)", _fmt(weightage_result.get("value_per_share"))],
            ["Fully Diluted Equity Shares", _fmt_int(valuation_inputs.get("diluted_shares"))],
            ["Total Equity Value (Rs. Lakhs)", _fmt(weightage_result.get("concluded_value"))],
        ],
    )

    # --------------------------------------------------------
    # 9. LIMITING CONDITIONS
    # --------------------------------------------------------
    doc.add_heading("9. Limiting Conditions and Disclaimer", level=1)

    limitations = [
        "This Report is restricted to the stated purpose and should not be used for another purpose without appropriate review.",
        "The Valuer relies upon financial and other information supplied by management and available sources without independent audit unless expressly stated otherwise.",
        "Projected financial information is inherently uncertain and actual results may differ materially from projections.",
        "The valuation reflects conditions as at the Valuation Date and does not automatically incorporate subsequent events.",
        "Valuation inherently involves judgement and different valuers may reasonably arrive at different conclusions.",
        "This Report does not constitute legal, tax, investment or accounting advice.",
        "The system-generated document is a draft and cannot be issued without Registered Valuer review and approval.",
    ]

    _add_bullets(doc, limitations)

    # --------------------------------------------------------
    # 10. DECLARATION
    # --------------------------------------------------------
    doc.add_heading("10. Valuer Review / Declaration", level=1)

    doc.add_paragraph(
        "The following declaration section is to be completed and confirmed by the Registered "
        "Valuer after review of the final report, supporting working papers, source information, "
        "assumptions, calculations and regulatory requirements."
    )

    _add_bullets(
        doc,
        [
            "The valuation has been carried out impartially and to the best of the Valuer's knowledge and belief.",
            "The Valuer is not disqualified from acting for the assignment and has considered conflicts of interest.",
            "The valuation methodology and assumptions are considered appropriate for the stated purpose.",
            "The fee for the engagement is not contingent upon the concluded value.",
        ],
    )

    doc.add_paragraph("")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Registered Valuer - Securities or Financial Assets")
    doc.add_paragraph("Registration No.: _______________________________")
    doc.add_paragraph("Place: __________________  Date: __________________")

    doc.add_page_break()

    # ========================================================
    # ANNEXURE A - DCF
    # ========================================================
    doc.add_heading("ANNEXURE A", level=1)
    doc.add_heading("Income Approach - Discounted Cash Flow (FCFF) Method", level=1)

    doc.add_heading("A.1 Basis", level=2)
    doc.add_paragraph(
        "FCFF has been derived from the projected financial information used in the valuation "
        "working. A terminal value has been determined using the Gordon Growth Model."
    )

    doc.add_heading("A.2 Free Cash Flow to Firm - Working", level=2)
    _note(doc, "All figures in Rs. Lakhs, unless otherwise stated.")

    fcff_rows = []
    for item in dcf_result.get("fcff", []) or []:
        fcff_rows.append([
            item.get("year"),
            _fmt(item.get("ebit")),
            _fmt(_f(item.get("ebit")) - _f(item.get("nopat"))),
            _fmt(item.get("nopat")),
            _fmt(item.get("depreciation")),
            _fmt(item.get("capex")),
            _fmt(item.get("change_working_capital")),
            _fmt(item.get("fcff")),
        ])

    _table(
        doc,
        ["Year", "EBIT", "Tax", "NOPAT", "Dep.", "Capex", "Increase in NWC", "FCFF"],
        fcff_rows,
    )

    doc.add_heading("A.3 Discounting and Terminal Value", level=2)

    discount_rows = []
    for item in dcf_result.get("fcff", []) or []:
        discount_rows.append([
            item.get("year"),
            _fmt(item.get("discount_factor"), 4),
            _fmt(item.get("pv_fcff")),
        ])

    _table(doc, ["Year", "Discount Factor", "Present Value of FCFF"], discount_rows)

    _kv(
        doc,
        [
            ["Terminal Growth Rate", _fmt_pct(valuation_inputs.get("terminal_growth_percent"))],
            ["Terminal Value (Rs. Lakhs)", _fmt(dcf_result.get("terminal_value"))],
            ["Present Value of Terminal Value (Rs. Lakhs)", _fmt(dcf_result.get("pv_terminal_value"))],
        ],
    )

    doc.add_heading("A.4 Enterprise Value to Equity Value Bridge", level=2)
    _kv(
        doc,
        [
            ["Sum of PV of explicit period FCFF", _fmt(dcf_result.get("pv_explicit_fcff"))],
            ["Add: PV of Terminal Value", _fmt(dcf_result.get("pv_terminal_value"))],
            ["Enterprise Value", _fmt(dcf_result.get("enterprise_value"))],
            ["Add: Cash & Cash Equivalents", _fmt(valuation_inputs.get("cash"))],
            ["Add: Non-Operating Assets", _fmt(valuation_inputs.get("non_operating_assets"))],
            ["Less: Debt", _fmt(valuation_inputs.get("debt"))],
            ["Equity Value under DCF (Rs. Lakhs)", _fmt(dcf_result.get("equity_value"))],
            ["Fully Diluted Equity Shares", _fmt_int(valuation_inputs.get("diluted_shares"))],
            ["Value per Equity Share under DCF (Rs.)", _fmt(dcf_result.get("value_per_share"))],
        ],
    )

    # ========================================================
    # ANNEXURE B - NAV
    # ========================================================
    doc.add_page_break()
    doc.add_heading("ANNEXURE B", level=1)
    doc.add_heading("Asset Approach - Net Asset Value Method", level=1)

    doc.add_heading("B.1 Basis", level=2)
    doc.add_paragraph(
        "NAV has been computed from the adjusted / considered assets and liabilities recorded "
        "in the valuation working. The automatically populated values are starting figures and "
        "require valuer review for fair-value and other adjustments."
    )

    doc.add_heading("B.2 Working", level=2)
    _kv(
        doc,
        [
            ["Adjusted / Considered Assets (Rs. Lakhs)", _fmt(valuation_inputs.get("adjusted_assets"))],
            ["Less: Adjusted / Considered Liabilities (Rs. Lakhs)", _fmt(valuation_inputs.get("adjusted_liabilities"))],
            ["Net Assets available to Equity Shareholders (Rs. Lakhs)", _fmt(nav_result.get("equity_value"))],
            ["Fully Diluted Equity Shares", _fmt_int(valuation_inputs.get("diluted_shares"))],
            ["Value per Equity Share under NAV (Rs.)", _fmt(nav_result.get("value_per_share"))],
        ],
    )

    # ========================================================
    # ANNEXURE C - WACC
    # ========================================================
    doc.add_page_break()
    doc.add_heading("ANNEXURE C", level=1)
    doc.add_heading("Weighted Average Cost of Capital - Working", level=1)

    sources = wacc_analysis.get("sources", {}) or {}

    doc.add_heading("C.1 Cost of Equity", level=2)
    _table(
        doc,
        ["Particulars", "Value", "Basis / Source"],
        [
            ["Risk-free Rate", _fmt_pct(wacc_analysis.get("risk_free_rate_percent")), sources.get("risk_free_rate", "")],
            ["Equity Risk Premium", _fmt_pct(wacc_analysis.get("equity_risk_premium_percent")), sources.get("equity_risk_premium", "")],
            ["Levered Beta", _fmt(wacc_analysis.get("beta")), sources.get("beta", "")],
            ["Company-specific Risk Premium", _fmt_pct(wacc_analysis.get("company_specific_risk_premium_percent")), "Valuer judgement"],
            ["Cost of Equity", _fmt_pct(wacc_analysis.get("cost_of_equity_percent")), "Rf + Beta x ERP + CSRP"],
        ],
    )

    doc.add_heading("C.2 Cost of Debt", level=2)
    _table(
        doc,
        ["Particulars", "Value", "Basis / Source"],
        [
            ["Pre-tax Cost of Debt", _fmt_pct(wacc_analysis.get("pre_tax_cost_of_debt_percent")), sources.get("cost_of_debt", "")],
            ["Corporate Tax Rate", _fmt_pct(wacc_analysis.get("tax_rate_percent", valuation_inputs.get("tax_rate_percent"))), "Selected / applicable rate"],
            ["Post-tax Cost of Debt", _fmt_pct(wacc_analysis.get("after_tax_cost_of_debt_percent")), "Kd x (1-t)"],
        ],
    )

    doc.add_heading("C.3 Capital Structure and WACC", level=2)
    _kv(
        doc,
        [
            ["Weight of Debt", _fmt_pct(wacc_analysis.get("debt_weight_percent"))],
            ["Weight of Equity", _fmt_pct(wacc_analysis.get("equity_weight_percent"))],
            ["WACC", _fmt_pct(valuation_inputs.get("wacc_percent"))],
            ["Terminal Growth Rate", _fmt_pct(valuation_inputs.get("terminal_growth_percent"))],
        ],
    )

    # ========================================================
    # ANNEXURE D - FINANCIALS
    # ========================================================
    doc.add_page_break()
    doc.add_heading("ANNEXURE D", level=1)
    doc.add_heading("Historical and Projected Financial Statements", level=1)

    hist = _rows(financial_analysis, "historical")
    proj = _rows(financial_analysis, "projected")

    if hist or proj:
        rows = []
        for r in hist + proj:
            m = r.get("metrics", {}) or {}
            rows.append([
                r.get("display") or r.get("period", ""),
                _fmt(m.get("revenue")),
                _fmt(m.get("ebitda")),
                _fmt_pct(m.get("ebitda_margin")),
                _fmt(m.get("ebit")),
                _fmt(m.get("pat")),
                _fmt(m.get("capex")),
                _fmt(m.get("change_working_capital")),
            ])

        _table(
            doc,
            ["Period", "Revenue", "EBITDA", "EBITDA %", "EBIT", "PAT", "Capex", "Change in WC"],
            rows,
        )
    else:
        doc.add_paragraph("Financial analysis data was not available in the generated output.")

    # ========================================================
    # ANNEXURE E - ASSUMPTIONS / REVIEW
    # ========================================================
    doc.add_page_break()
    doc.add_heading("ANNEXURE E", level=1)
    doc.add_heading("Summary of Key Assumptions and Review Items", level=1)

    review_items = [
        f"WACC used: {_fmt_pct(valuation_inputs.get('wacc_percent'))}.",
        f"Terminal growth rate: {_fmt_pct(valuation_inputs.get('terminal_growth_percent'))}.",
        f"Tax rate: {_fmt_pct(valuation_inputs.get('tax_rate_percent'))}.",
        f"Fully diluted shares: {_fmt_int(valuation_inputs.get('diluted_shares'))}.",
        f"Open review items: {review_summary.get('unresolved', 0)}.",
        f"Material open review items: {review_summary.get('unresolved_material', 0)}.",
        f"Failed financial cross-checks: {review_summary.get('failed_cross_checks', 0)}.",
        f"Failed capital checks: {review_summary.get('failed_capital_checks', 0)}.",
        f"Final report ready: {_yes_no(review_summary.get('final_report_ready'))}.",
    ]

    _add_bullets(doc, review_items)

    _note(
        doc,
        "Any assumption, source, extracted field or validation item flagged for review must be "
        "resolved or expressly addressed by the Registered Valuer before final issuance."
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path
