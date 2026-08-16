#!/usr/bin/env python3
"""Desktop entry point for The 45-Day Clock.

FastAPI serves a UTF-8 local application on a random loopback port. pywebview
hosts that application in a native Windows window. A fresh secret token is
required for every API request and exists only for the current process.
"""

from __future__ import annotations

import argparse
import faulthandler
import hashlib
import io
import json
import re
import secrets
import socket
import threading
import time
import traceback
import os
import sys
import tempfile
from dataclasses import asdict
from copy import deepcopy
from datetime import date, datetime
from decimal import Decimal
from pathlib import Path
from typing import Any, Optional
from urllib.request import urlopen

from fastapi import FastAPI, File, Header, HTTPException, Request, UploadFile
from fastapi.responses import HTMLResponse, Response
from pydantic import BaseModel, Field

from clock45.classify import (
    EVIDENCE_RANK,
    GATE_ACTIVITY,
    GATE_CLASS,
    GATE_REGISTRATION,
    GATE_TIMING,
    SRC_ASSUMED,
    SRC_CERTIFICATE,
    SRC_CLIENT_FLAG,
    SRC_DECLARATION,
    UdyamRecord,
    assess_coverage,
)
from clock45.demo_data import build_demo_dataset
from clock45.engine import (
    ACC_GRN_DATE,
    ACC_INVOICE_DATE,
    ACC_INVOICE_PLUS,
    ACC_POLICY_TEXT,
    PaymentLine,
    PurchaseLine,
    action_list,
    exclusion_register,
    interest_only_register,
    run_assessment,
)
from clock45.ingest import (
    InvoiceSupplement,
    IngestError,
    ManualEntryGrid,
    RowProblem,
    VendorImportData,
    _parse_rows,
    import_tally_xml,
    suggest_mapping,
    suggest_mapping_details,
)
from clock45.rules import assess_invoice, fy_bounds, resolve_credit_period
from clock45.store import Store, StoreError
from clock45.license import LicenceError, LicenceManager, TRIAL_LINE_LIMIT
from clock45.udyam import UdyamParseError, parse_udyam_certificate


ROOT = Path(__file__).resolve().parent
WEB = ROOT / "web"
_STARTUP_LOG_HANDLE = None
if os.environ.get("CLOCK45_STARTUP_LOG"):
    _STARTUP_LOG_HANDLE = open(os.environ["CLOCK45_STARTUP_LOG"], "a", encoding="utf-8")
    faulthandler.enable(file=_STARTUP_LOG_HANDLE)
    faulthandler.dump_traceback_later(15, repeat=True, file=_STARTUP_LOG_HANDLE)


def _startup_marker(message: str) -> None:
    if _STARTUP_LOG_HANDLE:
        _STARTUP_LOG_HANDLE.write(f"STARTUP: {message}\n")
        _STARTUP_LOG_HANDLE.flush()


def _write_local_crash_log(store: Optional[Store], heading: str) -> None:
    """Record an unexpected desktop failure only inside the chosen data folder."""
    if store is None:
        return
    try:
        target = store.folder / "clock45-crash.log"
        with target.open("a", encoding="utf-8") as handle:
            handle.write(f"\n[{datetime.now().isoformat(timespec='seconds')}] {heading}\n")
            handle.write(traceback.format_exc())
    except Exception:
        pass
GATE_LABELS = {
    GATE_CLASS: "Enterprise class",
    GATE_ACTIVITY: "Trader activity",
    GATE_REGISTRATION: "Udyam registration",
    GATE_TIMING: "Registration timing",
}
SOURCE_LABELS = {
    SRC_CERTIFICATE: "Udyam certificate",
    SRC_DECLARATION: "Vendor declaration",
    SRC_CLIENT_FLAG: "Client ERP flag",
    SRC_ASSUMED: "ASSUMED",
}


def money(value: Any) -> str:
    """Indian digit grouping used by every API response and export."""
    amount = Decimal(value or 0).quantize(Decimal("0.01"))
    negative = amount < 0
    whole, fraction = f"{abs(amount):.2f}".split(".")
    if len(whole) > 3:
        head, tail = whole[:-3], whole[-3:]
        groups = []
        while len(head) > 2:
            groups.insert(0, head[-2:])
            head = head[:-2]
        if head:
            groups.insert(0, head)
        whole = ",".join(groups + [tail])
    suffix = "" if fraction == "00" else "." + fraction
    return f"{'-' if negative else ''}Rs {whole}{suffix}"


def iso(value: Optional[date | datetime]) -> Optional[str]:
    return value.isoformat() if value else None


def _jsonable(value: Any) -> Any:
    if isinstance(value, Decimal):
        return str(value)
    if isinstance(value, (date, datetime)):
        return value.isoformat()
    if isinstance(value, dict):
        return {key: _jsonable(item) for key, item in value.items()}
    if isinstance(value, (list, tuple)):
        return [_jsonable(item) for item in value]
    return value


class MappingRequest(BaseModel):
    mapping: dict[str, str]
    record_type: str = "purchase"


class AnalysisStartRequest(BaseModel):
    entity_name: str
    entity_pan: str = ""
    fy: str


class FirmProfileRequest(BaseModel):
    firm_name: str
    frn: str = ""
    address: str = ""
    email: str = ""
    phone: str = ""
    preparer: str = ""
    reviewer: str = ""
    document_status: str = "DRAFT"
    tax_rate_pct: str = "25.17"
    tax_rate_basis: str = "Illustrative effective rate; confirm for this client"


class ManualRequest(BaseModel):
    text: str
    record_type: str = "purchase"
    expected_total: Optional[str] = None


class AssumptionRequest(BaseModel):
    policy: str
    plus_days: int = 0
    payments_confirmed: bool = False


class BulkRequest(BaseModel):
    vendor_ids: list[str]
    source: str


class VendorUpdateRequest(BaseModel):
    udyam_no: str = ""
    enterprise_class: str = ""
    nic_code: str = ""
    activity_label: str = ""
    registration_date: str = ""
    source: str = SRC_ASSUMED
    confirmed_by: str = "Desktop user"


class CertificateReviewRequest(BaseModel):
    evidence_id: str
    udyam_no: str
    enterprise_name: str = ""
    enterprise_class: str
    major_activity: str = ""
    pan: str = ""
    organisation_type: str = ""
    incorporation_date: str = ""
    commencement_date: str = ""
    registration_date: str = ""
    registered_address: str = ""
    nic_code: str
    confirmed_by: str
    classification_history: list[dict[str, Any]] = Field(default_factory=list)


class DesktopState:
    def __init__(self, store: Optional[Store] = None) -> None:
        self.store: Optional[Store] = None
        self.client_id: Optional[str] = None
        self.analysis_id: Optional[str] = None
        self.completed_run_id: Optional[str] = None
        self.entity_pan = ""
        self.stage = "home"
        self.demo_mode = False
        self.entity_name = ""
        self.fy = "2025-26"
        self.purchases: list[PurchaseLine] = []
        self.payments: list[PaymentLine] = []
        self.invoice_supplements: dict[str, InvoiceSupplement] = {}
        self.vendor_import_data: dict[str, VendorImportData] = {}
        self.udyam: dict[str, UdyamRecord] = {}
        self.run = None
        self.runs: list[dict[str, Any]] = []
        self.control_totals: Optional[dict[str, Any]] = None
        self.confirmations: list[dict[str, Any]] = []
        self.pending_rows: list[dict[str, Any]] = []
        self.pending_columns: list[str] = []
        self.pending_fingerprint = ""
        self.pending_source = ""
        self.pending_source_hash = ""
        self.pending_certificate_reviews: dict[str, dict[str, Any]] = {}
        self.assumptions_confirmed = False
        self.payment_information_confirmed = False
        if store:
            self.attach_store(store)

    def attach_store(self, store: Store) -> None:
        old_store = self.store
        self.store = store
        if old_store is not None and old_store is not store:
            old_store.close()
        self.reload_from_store()

    def reload_from_store(self) -> None:
        if not self.store:
            return
        saved = self.store.load_latest_analysis()
        self._refresh_runs()
        if saved is None:
            self.reset_analysis()
            return
        self.client_id = saved.client_id
        self.analysis_id = saved.analysis_id
        self.completed_run_id = saved.completed_run_id
        self.entity_name = saved.entity_name
        self.entity_pan = saved.entity_pan
        self.fy = saved.fy
        self.stage = saved.stage
        self.control_totals = saved.control_totals
        self.purchases = saved.purchases
        self.payments = saved.payments
        self.invoice_supplements = saved.invoice_supplements
        self.udyam = saved.udyam
        self.assumptions_confirmed = False
        self.payment_information_confirmed = False
        self.run = None
        if saved.completed_run_id:
            run, purchases, payments, udyam = self.store.load_completed_run(saved.completed_run_id)
            self.run = run
            self.purchases = purchases
            self.payments = payments
            self.udyam = udyam
            self.assumptions_confirmed = True
            self.payment_information_confirmed = bool(
                run.control_totals.get("payment_information_confirmed", True)
            )
            self.stage = "results"

    def storage_status(self) -> dict[str, Any]:
        return {
            "configured": self.store is not None,
            "folder": str(self.store.folder) if self.store else "",
            "database": str(self.store.path) if self.store else "",
            "database_display": "clock45.sqlite3 in the selected local folder" if self.store else "",
            "integrity": self.store.integrity_check() if self.store else None,
        }

    def firm_profile(self) -> dict[str, str]:
        defaults = {
            "firm_name": "",
            "frn": "",
            "address": "",
            "email": "",
            "phone": "",
            "preparer": "",
            "reviewer": "",
            "document_status": "DRAFT",
            "tax_rate_pct": "25.17",
            "tax_rate_basis": "Legacy illustrative rate: 22% base × 10% surcharge × 4% cess; confirm for this client",
        }
        if not self.store:
            return defaults
        return {
            key: self.store.get_setting(f"firm_profile:{key}", value) or value
            for key, value in defaults.items()
        }

    def save_firm_profile(self, values: dict[str, str]) -> dict[str, str]:
        if not self.store:
            raise StoreError("Choose the local data folder before saving the firm profile")
        firm_name = values.get("firm_name", "").strip()
        if not firm_name:
            raise ValueError("Enter the CA firm name")
        status = values.get("document_status", "DRAFT").strip().upper()
        if status not in {"DRAFT", "REVIEWED", "FINAL"}:
            raise ValueError("Document status must be Draft, Reviewed or Final")
        cleaned = {
            key: str(value or "").strip()
            for key, value in values.items()
        }
        cleaned["firm_name"] = firm_name
        cleaned["document_status"] = status
        try:
            tax_rate = Decimal(cleaned.get("tax_rate_pct", "25.17"))
        except Exception as exc:
            raise ValueError("Enter the tax-impact rate as a percentage, for example 25.17") from exc
        if tax_rate < 0 or tax_rate > 100:
            raise ValueError("Tax-impact rate must be between 0 and 100")
        cleaned["tax_rate_pct"] = str(tax_rate.quantize(Decimal("0.01")))
        if not cleaned.get("tax_rate_basis", "").strip():
            raise ValueError("Explain the basis of the Action List tax-impact rate")
        email = cleaned.get("email", "")
        if email and not re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", email):
            raise ValueError("Enter a valid professional email address")
        phone = re.sub(r"[\s()+-]", "", cleaned.get("phone", ""))
        if phone and (not phone.isdigit() or not 7 <= len(phone) <= 15):
            raise ValueError("Telephone must contain 7 to 15 digits")
        frn = cleaned.get("frn", "").upper()
        if frn and not re.fullmatch(r"\d{6}[A-Z]", frn):
            raise ValueError("Firm Registration Number must look like 012345N")
        cleaned["frn"] = frn
        if status == "FINAL" and (not frn or not cleaned.get("preparer") or not cleaned.get("reviewer")):
            raise ValueError("Final documents require the FRN, Prepared by and Reviewed by fields")
        for key, value in cleaned.items():
            self.store.set_setting(f"firm_profile:{key}", value)
        return self.firm_profile()

    def start_analysis(self, entity_name: str, entity_pan: str, fy: str) -> None:
        name = entity_name.strip()
        pan = entity_pan.strip().upper()
        financial_year = fy.strip()
        if not name:
            raise ValueError("Enter the client entity name")
        if pan and not re.fullmatch(r"[A-Z]{5}[0-9]{4}[A-Z]", pan):
            raise ValueError("Client PAN must contain five letters, four digits and one final letter")
        if len(financial_year) != 7 or financial_year[4] != "-":
            raise ValueError("Financial year must be written like 2025-26")
        try:
            fy_bounds(financial_year)
        except Exception as exc:
            raise ValueError("Enter a valid financial year such as 2025-26") from exc
        self.reset_analysis()
        self.entity_name = name
        self.entity_pan = pan
        self.fy = financial_year
        self._ensure_analysis()

    def licence_status(self) -> dict[str, Any]:
        if not self.store:
            return {
                "mode": "SETUP_REQUIRED", "firm_name": "", "seat_count": 0,
                "expiry_date": None, "days_remaining": 30,
                "line_limit": TRIAL_LINE_LIMIT, "can_analyse": False,
                "message": "Choose a local data folder to start the 30-day trial.",
            }
        return LicenceManager(self.store).status().as_dict()

    def startup_status(self) -> dict[str, Any]:
        first_run_complete = bool(self.store and self.store.get_setting("welcome_complete") == "yes")
        return {
            "needs_welcome": not first_run_complete,
            "storage": self.storage_status(),
            "licence": self.licence_status(),
        }

    def complete_welcome(self) -> None:
        if not self.store:
            raise LicenceError("Choose a local data folder before continuing")
        self.store.set_setting("welcome_complete", "yes")

    def home_data(self) -> dict[str, Any]:
        if self.store:
            entities = [
                {
                    "name": item["entity_name"],
                    "fy": item["fy"] or "Not started",
                    "status": "Completed" if item["stage"] == "results" else "Analysis in progress",
                }
                for item in self.store.list_client_workspaces()
            ]
        elif self.entity_name:
            entities = [{"name": self.entity_name, "fy": self.fy, "status": "Current session only"}]
        else:
            entities = []
        return {
            "entities": entities,
            "runs": self.runs,
            "storage": self.storage_status(),
            "active_stage": self.stage,
            "active_entity": self.entity_name,
            "active_fy": self.fy,
        }

    def _refresh_runs(self) -> None:
        if not self.store:
            return
        self.runs = [
            {
                "run_id": row["run_id"],
                "entity": row["entity_name"],
                "fy": row["fy"],
                "disallowance": Decimal(row["disallowance_total"]),
                "disallowance_display": money(row["disallowance_total"]),
                "run_date": datetime.fromisoformat(row["run_at"]).strftime("%d-%b-%Y %H:%M"),
                "run_hash": row["run_hash"],
            }
            for row in self.store.list_completed_runs()
        ]

    def _ensure_analysis(self) -> None:
        if not self.store:
            return
        self.client_id = self.store.get_or_create_client(self.entity_name, self.entity_pan)
        if not self.analysis_id:
            self.analysis_id = self.store.start_or_resume_analysis(self.client_id, self.fy)

    def _persist_ledger(self, source: str, source_type: str = "application") -> None:
        if not self.store:
            return
        self._ensure_analysis()
        self.store.save_analysis(
            self.analysis_id,
            stage="vendors",
            purchases=self.purchases,
            payments=self.payments,
            control_totals=self.control_totals,
            source_label=source,
            source_type=source_type,
            invoice_supplements=self.invoice_supplements,
        )
        self.stage = "vendors"

    def persist_vendor(self, vendor_id: str, changed_by: str = "Desktop user") -> None:
        if not self.store or not self.client_id or vendor_id not in self.udyam:
            return
        name = next(
            (line.vendor_name_as_written for line in self.purchases if line.vendor_id == vendor_id),
            vendor_id,
        )
        pan_gstin = vendor_id if not vendor_id.startswith("NAME:") else ""
        self.store.upsert_vendor(
            self.client_id, self.udyam[vendor_id], vendor_name=name,
            pan_gstin=pan_gstin, changed_by=changed_by,
        )

    def export_context(self):
        from clock45.export import ExportContext

        if self.run is None or not self.assumptions_confirmed:
            raise ValueError("Confirm the assumptions and run the analysis before exporting")
        profile = self.firm_profile()
        vendor_metadata = {}
        evidence_documents = []
        if self.store and self.client_id:
            for vendor_id in self.udyam:
                vendor_metadata[vendor_id] = self.store.load_vendor_metadata(
                    self.client_id, vendor_id
                )
                for summary in self.store.list_vendor_evidence(self.client_id, vendor_id):
                    item = self.store.get_vendor_evidence(summary["evidence_id"])
                    item["bytes"] = len(item["content"])
                    evidence_documents.append(item)
        return ExportContext(
            run=self.run,
            purchases=self.purchases,
            payments=self.payments,
            udyam=self.udyam,
            entity_pan=self.entity_pan or "Not provided",
            preparer=profile["preparer"],
            reviewer=profile["reviewer"],
            firm_name=profile["firm_name"] or "FIRM PROFILE NOT CONFIGURED",
            firm_frn=profile["frn"],
            firm_address=profile["address"],
            firm_email=profile["email"],
            firm_phone=profile["phone"],
            document_status=profile["document_status"],
            tax_rate=Decimal(profile["tax_rate_pct"]) / Decimal("100"),
            tax_rate_basis=profile["tax_rate_basis"],
            vendor_metadata=vendor_metadata,
            evidence_documents=tuple(evidence_documents),
        )

    def reset_analysis(self) -> None:
        self.client_id = None
        self.analysis_id = None
        self.completed_run_id = None
        self.entity_pan = ""
        self.stage = "home"
        self.demo_mode = False
        self.entity_name = ""
        self.purchases = []
        self.payments = []
        self.invoice_supplements = {}
        self.vendor_import_data = {}
        self.udyam = {}
        self.run = None
        self.control_totals = None
        self.confirmations = []
        self.pending_rows = []
        self.pending_columns = []
        self.pending_fingerprint = ""
        self.pending_source = ""
        self.pending_source_hash = ""
        self.pending_certificate_reviews = {}
        self.assumptions_confirmed = False
        self.payment_information_confirmed = False

    def load_demo(self) -> dict[str, Any]:
        demo = build_demo_dataset()
        self.entity_name = demo["entity_name"]
        self.fy = demo["fy"]
        self.purchases = demo["purchases"]
        self.payments = demo["payments"]
        self.invoice_supplements = {}
        self.udyam = demo["udyam"]
        total = sum((line.amount for line in self.purchases), Decimal("0.00"))
        demo_payload = {
            "purchases": [asdict(line) for line in self.purchases],
            "payments": [asdict(line) for line in self.payments],
            "udyam": {key: asdict(value) for key, value in sorted(self.udyam.items())},
        }
        demo_sha256 = hashlib.sha256(
            json.dumps(demo_payload, default=str, sort_keys=True).encode("utf-8")
        ).hexdigest()
        self.control_totals = {
            "lines_read": len(self.purchases),
            "total_value_read": total,
            "total_value_accounted_for": total,
            "in_period_lines": len(self.purchases),
            "in_period_value": total,
            "out_of_period_lines": 0,
            "out_of_period_value": Decimal("0.00"),
            "ties": True,
            "source": "Deterministic demonstration ledger",
            "source_file_name": "Built-in synthetic demonstration ledger v1",
            "source_file_sha256": demo_sha256,
            "imported_at": datetime.now().isoformat(timespec="seconds"),
        }
        self.assumptions_confirmed = False
        self.payment_information_confirmed = True
        self.demo_mode = True
        self.run = None
        self.completed_run_id = None
        self.analysis_id = None
        self._ensure_analysis()
        if self.store and self.client_id:
            names = {}
            for line in self.purchases:
                names.setdefault(line.vendor_id, line.vendor_name_as_written)
            for vendor_id, record in self.udyam.items():
                self.store.upsert_vendor(
                    self.client_id, record, vendor_name=names.get(vendor_id, vendor_id),
                    pan_gstin=vendor_id if not vendor_id.startswith("NAME:") else "",
                    changed_by="Demonstration loader",
                )
            self._persist_ledger("Deterministic demonstration ledger", "demo")
        return self.control_totals

    def add_import(self, result, *, replace: bool = False, source_hash: str = "") -> None:
        prospective_count = len(result.purchases) + (0 if replace else len(self.purchases))
        if self.store:
            LicenceManager(self.store).require_analysis(prospective_count)
        elif prospective_count > TRIAL_LINE_LIMIT:
            raise LicenceError("Choose a local data folder before importing a ledger")
        prospective_purchases = list(result.purchases) if replace else [*self.purchases, *result.purchases]
        invoice_index = {line.invoice_id: line for line in prospective_purchases}
        prospective_payments = list(result.payments) if replace else [*self.payments, *result.payments]
        if prospective_payments and not invoice_index:
            raise IngestError("Import refused: load the purchase invoices before loading payments.")
        payment_problems = []
        paid_by_invoice: dict[str, Decimal] = {}
        for position, payment in enumerate(prospective_payments, 2):
            purchase = invoice_index.get(payment.invoice_id)
            reasons = []
            if purchase is None:
                reasons.append(
                    f"payment refers to invoice {payment.invoice_id!r}, which is not in the purchase ledger"
                )
            elif payment.payment_date < purchase.invoice_date:
                reasons.append(
                    f"payment date {payment.payment_date:%d-%m-%Y} is before invoice date "
                    f"{purchase.invoice_date:%d-%m-%Y}"
                )
            paid_by_invoice[payment.invoice_id] = (
                paid_by_invoice.get(payment.invoice_id, Decimal("0.00")) + payment.amount
            )
            if reasons:
                payment_problems.append(RowProblem(position, tuple(reasons), {
                    "invoice_number": payment.invoice_id,
                    "payment_date": payment.payment_date.isoformat(),
                    "amount": str(payment.amount),
                }))
        for invoice_id, total_paid in paid_by_invoice.items():
            purchase = invoice_index.get(invoice_id)
            if purchase and total_paid > purchase.amount:
                payment_problems.append(RowProblem(0, (
                    f"total payments {total_paid} exceed invoice amount {purchase.amount} for {invoice_id!r}",
                ), {"invoice_number": invoice_id, "total_paid": str(total_paid)}))
        if payment_problems:
            raise IngestError(
                f"Import refused: {len(payment_problems)} payment validation issue(s) must be corrected.",
                row_problems=payment_problems,
            )
        self.demo_mode = False
        if replace:
            self.purchases = []
            self.payments = []
        self.purchases.extend(result.purchases)
        self.payments.extend(result.payments)
        self.invoice_supplements.update(result.invoice_supplements)
        self.vendor_import_data.update(result.vendor_data)
        for line in result.purchases:
            self.udyam.setdefault(line.vendor_id, UdyamRecord(vendor_id=line.vendor_id))
        for vendor_id, imported in result.vendor_data.items():
            record = self.udyam.setdefault(vendor_id, UdyamRecord(vendor_id=vendor_id))
            # Imported values are helpful prefill only. ASSUMED remains in force
            # until a human confirms a declaration or certificate.
            record.udyam_no = record.udyam_no or imported.udyam_no or None
            record.enterprise_class = record.enterprise_class or imported.enterprise_class or None
            record.nic_code = record.nic_code or imported.nic_code or None
            record.activity_label = record.activity_label or imported.major_activity or None
        self.entity_name = self.entity_name or "Imported client"
        totals = result.control_totals
        import_totals = {
            "lines_read": totals.lines_read,
            "total_value_read": totals.total_value_read,
            "total_value_accounted_for": totals.total_value_accounted_for,
            "ties": totals.ties,
            "source": result.source,
            "source_file_name": Path(result.source).name if result.source else "",
            "source_file_sha256": source_hash,
            "imported_at": datetime.now().isoformat(timespec="seconds"),
            "unmapped_columns": list(result.unmapped_columns),
        }
        if self.purchases:
            fy_start, fy_end = fy_bounds(self.fy)
            in_period = [line for line in self.purchases if fy_start <= line.invoice_date <= fy_end]
            out_of_period = [line for line in self.purchases if line not in in_period]
            import_totals.update({
                "selected_fy": self.fy,
                "in_period_lines": len(in_period),
                "in_period_value": sum((line.amount for line in in_period), Decimal("0")),
                "out_of_period_lines": len(out_of_period),
                "out_of_period_value": sum((line.amount for line in out_of_period), Decimal("0")),
                "scope_ok": bool(in_period),
            })
        else:
            import_totals["scope_ok"] = request_scope_ok = bool(self.control_totals and self.control_totals.get("scope_ok"))
            if request_scope_ok and self.control_totals:
                for key in (
                    "selected_fy", "in_period_lines", "in_period_value",
                    "out_of_period_lines", "out_of_period_value",
                ):
                    import_totals[key] = self.control_totals.get(key)
        import_totals["can_continue"] = bool(import_totals["ties"] and import_totals.get("scope_ok"))
        self.control_totals = import_totals
        self.confirmations = [_jsonable(asdict(flag)) for flag in result.confirmations]
        self.run = None
        self.assumptions_confirmed = False
        self.payment_information_confirmed = False
        self.completed_run_id = None
        self._ensure_analysis()
        if self.store and self.client_id:
            persisted = self.store.load_vendor_master(self.client_id)
            for vendor_id, record in persisted.items():
                if vendor_id in self.udyam:
                    self.udyam[vendor_id] = record
            for vendor_id in self.udyam:
                self.persist_vendor(vendor_id, "Ledger import")
                imported = result.vendor_data.get(vendor_id)
                if imported:
                    self.store.upsert_vendor_metadata(
                        self.client_id, vendor_id, pan=imported.pan, gstin=imported.gstin,
                        contact=imported.contact,
                        registration_status=imported.registration_status,
                        verification_source=imported.verification_source or "IMPORTED_LEDGER",
                    )
            self._persist_ledger(result.source, "ledger")

    def vendor_rows(self) -> list[dict[str, Any]]:
        first_supply: dict[str, date] = {}
        names: dict[str, str] = {}
        ids: dict[str, str] = {}
        for line in self.purchases:
            first_supply[line.vendor_id] = min(
                first_supply.get(line.vendor_id, line.invoice_date), line.invoice_date
            )
            names.setdefault(line.vendor_id, line.vendor_name_as_written)
            ids.setdefault(line.vendor_id, line.vendor_id)
        rows = []
        for vendor_id, supply_date in first_supply.items():
            record = self.udyam.get(vendor_id) or UdyamRecord(vendor_id=vendor_id)
            history = (
                self.store.classification_history(self.client_id, vendor_id)
                if self.store and self.client_id else []
            )
            relevant = next(
                (item for item in history if item["classification_year"] == self.fy), None
            )
            effective_record = deepcopy(record)
            if relevant:
                effective_record.enterprise_class = relevant["enterprise_class"]
            coverage = assess_coverage(effective_record, supply_date)
            metadata = (
                self.store.load_vendor_metadata(self.client_id, vendor_id)
                if self.store and self.client_id else {}
            )
            credit_days = sorted({
                line.agreement_days for line in self.purchases
                if line.vendor_id == vendor_id and line.agreement_days is not None
            })
            evidence = (
                self.store.list_vendor_evidence(self.client_id, vendor_id)
                if self.store and self.client_id else []
            )
            if record.source == SRC_ASSUMED:
                validation_status = "REVIEW_REQUIRED"
                required_action = "Confirm evidence and classification"
            elif not record.udyam_no:
                validation_status = "UNREGISTERED"
                required_action = "Retain declaration/evidence for non-registration"
            elif not relevant and history:
                validation_status = "YEAR_REVIEW_REQUIRED"
                required_action = f"Confirm classification for FY {self.fy}"
            elif record.evidence_file_hash:
                validation_status = "VERIFIED"
                required_action = "No action"
            else:
                validation_status = "EVIDENCE_PENDING"
                required_action = "Upload or retain supporting evidence"
            rows.append({
                "vendor_id": vendor_id,
                "name": names[vendor_id],
                "pan_gstin": metadata.get("gstin") or metadata.get("pan") or (
                    vendor_id if not vendor_id.startswith("NAME:") else "Not in ledger"
                ),
                "pan": metadata.get("pan", ""),
                "gstin": metadata.get("gstin", ""),
                "contact": metadata.get("contact", ""),
                "udyam_no": record.udyam_no or "Not provided",
                "udyam_no_raw": record.udyam_no or "",
                "enterprise_class": effective_record.enterprise_class or "Unconfirmed",
                "enterprise_class_raw": effective_record.enterprise_class or "",
                "classification_history": history,
                "relevant_classification_year": relevant["classification_year"] if relevant else "",
                "nic_code": record.nic_code or "Unconfirmed",
                "nic_code_raw": record.nic_code or "",
                "activity": record.activity_label or "Unconfirmed",
                "activity_raw": record.activity_label or "",
                "registration_date": (
                    record.registration_date.isoformat() if record.registration_date else ""
                ),
                "coverage": "Covered" if coverage.covered else "Excluded",
                "gate": coverage.gate_failed,
                "coverage_reason": coverage.reason,
                "source": record.source,
                "source_label": SOURCE_LABELS.get(record.source, record.source),
                "has_certificate_evidence": bool(record.evidence_file_hash),
                "evidence_count": len(evidence),
                "evidence_strength": coverage.evidence_strength,
                "unconfirmed": coverage.needs_human_confirmation,
                "is_trader": record.is_trader,
                "is_medium": record.enterprise_class == "MEDIUM",
                "is_assumed": record.source == SRC_ASSUMED,
                "confirmed_by": record.confirmed_by or "",
                "confirmed_on": record.confirmed_on.isoformat() if record.confirmed_on else "",
                "credit_days": ", ".join(str(item) for item in credit_days) or "Not available",
                "validation_status": validation_status,
                "required_action": required_action,
            })
        return sorted(rows, key=lambda row: (not row["is_assumed"], row["name"].casefold()))

    def update_vendor(self, vendor_id: str, values: dict[str, Any]) -> dict[str, Any]:
        record = self.udyam.get(vendor_id)
        if record is None:
            raise ValueError("Vendor not found in the current ledger")

        udyam_no = str(values.get("udyam_no", "")).strip().upper()
        enterprise_class = str(values.get("enterprise_class", "")).strip().upper()
        nic_code = re.sub(r"\s+", "", str(values.get("nic_code", "")).strip())
        activity_label = str(values.get("activity_label", "")).strip()
        registration_text = str(values.get("registration_date", "")).strip()
        source = str(values.get("source", SRC_ASSUMED)).strip()
        confirmed_by = str(values.get("confirmed_by", "")).strip() or "Desktop user"

        if udyam_no and not re.fullmatch(r"UDYAM-[A-Z]{2}-\d{2}-\d{7}", udyam_no):
            raise ValueError("Udyam number must look like UDYAM-MH-26-0123456")
        if enterprise_class and enterprise_class not in {"MICRO", "SMALL", "MEDIUM"}:
            raise ValueError("Enterprise class must be Micro, Small or Medium")
        if nic_code and (not nic_code.isdigit() or not 2 <= len(nic_code) <= 5):
            raise ValueError("NIC code must contain between 2 and 5 digits")
        if source not in EVIDENCE_RANK:
            raise ValueError("Choose a recognised evidence source")
        if source == SRC_CERTIFICATE and not record.evidence_file_hash:
            raise ValueError("Upload the Udyam certificate before selecting certificate evidence")
        if udyam_no and not enterprise_class:
            raise ValueError("Select the enterprise class shown on the Udyam certificate")
        if udyam_no and not nic_code:
            raise ValueError("Enter the NIC code shown on the Udyam certificate or declaration")

        registration_date = None
        if registration_text:
            try:
                registration_date = date.fromisoformat(registration_text)
            except ValueError as exc:
                raise ValueError("Registration date must be a valid date") from exc

        record.udyam_no = udyam_no or None
        record.enterprise_class = enterprise_class or None
        record.nic_code = nic_code or None
        record.activity_label = activity_label or None
        record.registration_date = registration_date
        record.source = source
        record.confirmed_by = confirmed_by
        record.confirmed_on = date.today()
        self.persist_vendor(vendor_id, f"Vendor classification confirmed by {confirmed_by}")
        return next(row for row in self.vendor_rows() if row["vendor_id"] == vendor_id)

    def execute(self, policy: str, plus_days: int = 0, payments_confirmed: bool = False) -> None:
        if policy not in ACC_POLICY_TEXT:
            raise ValueError("Unknown acceptance-date policy")
        fy_start, fy_end = fy_bounds(self.fy)
        in_period = [line for line in self.purchases if fy_start <= line.invoice_date <= fy_end]
        if not in_period:
            raise ValueError(
                f"No purchase invoice falls within FY {self.fy}. "
                "Change the financial year or load the correct purchase ledger."
            )
        if not (payments_confirmed or self.payment_information_confirmed):
            raise ValueError(
                "Confirm that payment information is complete before calculating. "
                "This includes confirming that invoices with no payment row were genuinely unpaid."
            )
        assumed = [row for row in self.vendor_rows() if row["is_assumed"]]
        if assumed:
            raise ValueError(
                f"{len(assumed)} vendor(s) still use ASSUMED evidence. "
                "Resolve them before sign-off."
            )
        if not self.demo_mode:
            if not self.store:
                raise LicenceError("Choose a local data folder before running an analysis")
            LicenceManager(self.store).require_analysis(len(self.purchases))
        udyam_for_run = {vendor_id: deepcopy(record) for vendor_id, record in self.udyam.items()}
        if self.store and self.client_id:
            for vendor_id, record in udyam_for_run.items():
                history = self.store.classification_history(self.client_id, vendor_id)
                relevant = next(
                    (item for item in history if item["classification_year"] == self.fy), None
                )
                if relevant:
                    record.enterprise_class = relevant["enterprise_class"]
                elif history:
                    raise ValueError(
                        f"{next((line.vendor_name_as_written for line in self.purchases if line.vendor_id == vendor_id), vendor_id)} "
                        f"has certificate classification history, but no confirmed entry for FY {self.fy}."
                    )
        self.run = run_assessment(
            entity_name=self.entity_name,
            fy=self.fy,
            operator="Desktop user",
            purchases=self.purchases,
            payments=self.payments,
            udyam=udyam_for_run,
            acceptance_policy=policy,
            acceptance_plus_days=plus_days,
        )
        imported = self.control_totals or {}
        self.run.control_totals.update({
            "import_lines_read": imported.get("lines_read", len(self.purchases)),
            "import_value_read": imported.get(
                "total_value_read", sum((line.amount for line in self.purchases), Decimal("0"))
            ),
            "import_value_accounted_for": imported.get(
                "total_value_accounted_for", sum((line.amount for line in self.purchases), Decimal("0"))
            ),
            "import_ties": imported.get("ties", True),
            "source_file_name": imported.get("source_file_name", imported.get("source", "")),
            "source_file_sha256": imported.get("source_file_sha256", ""),
            "imported_at": imported.get("imported_at", ""),
            "payment_information_confirmed": True,
        })
        self.run.control_totals["ties"] = bool(
            self.run.control_totals.get("ties") and self.run.control_totals.get("import_ties")
        )
        self.assumptions_confirmed = True
        self.payment_information_confirmed = True
        self.stage = "results"
        summary = {
            "entity": self.entity_name,
            "fy": self.fy,
            "disallowance": self.run.disallowance_total,
            "disallowance_display": money(self.run.disallowance_total),
            "run_date": self.run.run_at.strftime("%d-%b-%Y %H:%M"),
            "run_hash": self.run.run_hash(),
        }
        if self.store:
            self._ensure_analysis()
            previous = self.store.list_completed_runs(self.client_id)
            supersedes = previous[0]["run_id"] if previous else None
            self.completed_run_id = self.store.save_completed_run(
                self.analysis_id, self.run, self.purchases, self.payments, udyam_for_run,
                entity_pan=self.entity_pan, acceptance_plus_days=plus_days,
                supersedes_run_id=supersedes,
            )
            self._refresh_runs()
        else:
            self.runs.insert(0, summary)

    def project_readiness(self) -> dict[str, Any]:
        vendors = self.vendor_rows() if self.purchases else []
        total_vendors = len(vendors)
        pan_count = sum(bool(row.get("pan") or row.get("gstin")) for row in vendors)
        class_count = sum(bool(row.get("enterprise_class_raw")) for row in vendors)
        missing_certificates = sum(not row.get("has_certificate_evidence") for row in vendors)
        grn_missing = sum(line.grn_date is None for line in self.purchases)
        credit_missing = len({
            line.vendor_id for line in self.purchases if line.agreement_days is None
        })
        issues = []
        if not self.purchases:
            issues.append("Import a purchase ledger")
        if self.control_totals and not self.control_totals.get("can_continue", self.control_totals.get("ties")):
            issues.append("Resolve the import control-total or financial-year scope block")
        assumed = sum(row.get("is_assumed", False) for row in vendors)
        if assumed:
            issues.append(f"Resolve evidence for {assumed} vendor(s)")
        if any(row.get("validation_status") == "YEAR_REVIEW_REQUIRED" for row in vendors):
            issues.append(f"Confirm year-specific Udyam classification for FY {self.fy}")
        if self.purchases and not self.payment_information_confirmed:
            issues.append("Confirm that payment information is complete, including genuinely unpaid invoices")
        return {
            "status": "READY_FOR_REVIEW" if not issues and self.run else (
                "READY_TO_CALCULATE" if not issues else "ISSUES_REQUIRE_ATTENTION"
            ),
            "issue_count": len(issues),
            "issues": issues,
            "checks": [
                {"label": "Ledger imported", "value": bool(self.purchases)},
                {"label": "Column mapping and totals complete", "value": bool(self.control_totals and self.control_totals.get("ties"))},
                {"label": "Vendor PAN/GSTIN coverage", "value": f"{round(100 * pan_count / total_vendors) if total_vendors else 0}%"},
                {"label": "Udyam classification coverage", "value": f"{round(100 * class_count / total_vendors) if total_vendors else 0}%"},
                {"label": "Missing certificate files", "value": missing_certificates},
                {"label": "GRN date missing", "value": f"{grn_missing} invoice(s)"},
                {"label": "Credit terms missing", "value": f"{credit_missing} vendor(s)"},
                {"label": "Payment information confirmed", "value": self.payment_information_confirmed},
                {"label": "Calculation completed", "value": self.run is not None},
            ],
        }


class NativeBridge:
    """Native folder dialogs for the database, backups and finished exports."""

    def __init__(self, state: DesktopState) -> None:
        self.state = state

    @staticmethod
    def _save_bytes(content: bytes, default_name: str, file_type: str) -> dict[str, Any]:
        import webview

        selected = webview.windows[0].create_file_dialog(
            webview.SAVE_DIALOG,
            save_filename=default_name,
            file_types=(file_type,),
        )
        if not selected:
            return {"cancelled": True}
        target_value = selected[0] if isinstance(selected, (list, tuple)) else selected
        target = Path(target_value).expanduser().resolve()
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(content)
        if not target.is_file() or target.stat().st_size == 0:
            raise OSError("The file could not be written")
        return {"cancelled": False, "file": str(target), "bytes": target.stat().st_size}

    def save_input_template(self) -> dict[str, Any]:
        try:
            return self._save_bytes(
                _import_template_bytes(),
                "45-Day-Clock-Import-Template.xlsx",
                "Excel workbook (*.xlsx)",
            )
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def save_vendor_declaration(self, vendor_id: str) -> dict[str, Any]:
        try:
            content, filename = _vendor_declaration_bytes(self.state, vendor_id)
            return self._save_bytes(content, filename, "Word document (*.docx)")
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def save_vendor_evidence(self, evidence_id: str) -> dict[str, Any]:
        if not self.state.store:
            return {"cancelled": False, "error": "Choose the local data folder first."}
        try:
            item = self.state.store.get_vendor_evidence(evidence_id)
            suffix = Path(item["filename"]).suffix.lower()
            description = "PDF document (*.pdf)" if suffix == ".pdf" else "Evidence file (*.*)"
            return self._save_bytes(bytes(item["content"]), item["filename"], description)
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def save_export(self, kind: str) -> dict[str, Any]:
        try:
            content, filename, media = _finished_export_bytes(kind, self.state)
            file_type = "PDF document (*.pdf)" if media == "application/pdf" else "Excel workbook (*.xlsx)"
            result = self._save_bytes(content, filename, file_type)
            if (not result.get("cancelled") and not result.get("error")
                    and self.state.store and self.state.completed_run_id):
                path = result["file"]
                self.state.store.record_export(
                    self.state.completed_run_id, Path(path).parent, [path],
                    exported_by="Desktop user",
                )
            return result
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def choose_and_export(self) -> dict[str, Any]:
        import webview
        from clock45.export import export_all

        try:
            selected = webview.windows[0].create_file_dialog(webview.FOLDER_DIALOG)
            if not selected:
                return {"cancelled": True}
            folder = selected[0] if isinstance(selected, (list, tuple)) else selected
            bundle = export_all(folder, self.state.export_context())
            if self.state.store and self.state.completed_run_id:
                self.state.store.record_export(
                    self.state.completed_run_id, folder, list(bundle.paths()),
                    exported_by="Desktop user",
                )
            return {
                "cancelled": False,
                "folder": str(Path(folder).resolve()),
                "files": [str(path) for path in bundle.paths()],
            }
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def choose_and_export_audit_pack(self) -> dict[str, Any]:
        import webview
        from clock45.export import export_complete_audit_pack

        try:
            selected = webview.windows[0].create_file_dialog(webview.FOLDER_DIALOG)
            if not selected:
                return {"cancelled": True}
            folder = selected[0] if isinstance(selected, (list, tuple)) else selected
            bundle = export_complete_audit_pack(folder, self.state.export_context())
            if self.state.store and self.state.completed_run_id:
                self.state.store.record_export(
                    self.state.completed_run_id, bundle.root, list(bundle.paths()),
                    exported_by="Desktop user - complete audit pack",
                )
            return {
                "cancelled": False,
                "folder": str(bundle.root),
                "files": [str(path) for path in bundle.paths()],
            }
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def choose_data_folder(self) -> dict[str, Any]:
        import webview

        try:
            selected = webview.windows[0].create_file_dialog(webview.FOLDER_DIALOG)
            if not selected:
                return {"cancelled": True}
            folder = selected[0] if isinstance(selected, (list, tuple)) else selected
            self.state.attach_store(Store(folder))
            return {"cancelled": False, **self.state.storage_status()}
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def backup_database(self) -> dict[str, Any]:
        import webview

        if not self.state.store:
            return {"cancelled": False, "error": "Choose a data folder first."}
        try:
            selected = webview.windows[0].create_file_dialog(webview.FOLDER_DIALOG)
            if not selected:
                return {"cancelled": True}
            folder = selected[0] if isinstance(selected, (list, tuple)) else selected
            backup = self.state.store.backup(folder)
            return {"cancelled": False, "file": str(backup)}
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def restore_database(self) -> dict[str, Any]:
        import webview

        if not self.state.store:
            return {"cancelled": False, "error": "Choose the live data folder first."}
        try:
            selected = webview.windows[0].create_file_dialog(
                webview.OPEN_DIALOG, allow_multiple=False,
                file_types=("45-Day Clock backup (*.sqlite3)",),
            )
            if not selected:
                return {"cancelled": True}
            source = selected[0] if isinstance(selected, (list, tuple)) else selected
            self.state.store.restore(source)
            self.state.reload_from_store()
            return {"cancelled": False, "file": str(Path(source).resolve())}
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}

    def choose_licence_file(self) -> dict[str, Any]:
        import webview

        if not self.state.store:
            return {"cancelled": False, "error": "Choose the local data folder first."}
        try:
            selected = webview.windows[0].create_file_dialog(
                webview.OPEN_DIALOG, allow_multiple=False,
                file_types=("45-Day Clock licence (*.json)",),
            )
            if not selected:
                return {"cancelled": True}
            source = selected[0] if isinstance(selected, (list, tuple)) else selected
            status = LicenceManager(self.state.store).install(source)
            return {"cancelled": False, "licence": status.as_dict()}
        except Exception as exc:
            return {"cancelled": False, "error": str(exc)}


def _read_uploaded_table(data: bytes, filename: str):
    import pandas as pd

    suffix = Path(filename).suffix.lower()
    stream = io.BytesIO(data)
    if suffix in {".xlsx", ".xlsm"}:
        return pd.read_excel(stream, dtype=object, keep_default_na=False)
    if suffix == ".csv":
        return pd.read_csv(stream, dtype=object, keep_default_na=False, encoding="utf-8-sig")
    raise IngestError("Choose an .xlsx, .xlsm or .csv file")


def _error_payload(exc: IngestError) -> dict[str, Any]:
    totals = _jsonable(asdict(exc.control_totals)) if exc.control_totals else None
    return {
        "message": str(exc),
        "control_totals": totals,
        "rows": [_jsonable(asdict(problem)) for problem in exc.row_problems],
    }


def _xlsx_bytes(title: str, headers: list[str], rows: list[list[Any]]) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = title[:31]
    sheet.append(headers)
    for row in rows:
        sheet.append([_jsonable(value) for value in row])
    header_fill = PatternFill("solid", fgColor="12233A")
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(vertical="center")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column in sheet.columns:
        width = min(55, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
        sheet.column_dimensions[column[0].column_letter].width = width
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _import_template_bytes() -> bytes:
    from openpyxl import Workbook
    from openpyxl.formatting.rule import FormulaRule
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.worksheet.datavalidation import DataValidation

    workbook = Workbook()
    purchases = workbook.active
    purchases.title = "Purchase Invoices"
    payments = workbook.create_sheet("Payments")
    instructions = workbook.create_sheet("Read Me", 0)
    purchase_headers = [
        "Invoice Number", "Invoice Date", "Vendor Name", "Vendor PAN", "Vendor GSTIN",
        "Invoice Amount", "Udyam Registration Number", "Enterprise Type",
        "Udyam Classification Year", "NIC Code", "Major Activity",
        "Registration Status", "Udyam Verification Source", "GRN Date",
        "Agreement Credit Days", "Agreed Payment Due Date", "Actual Payment Date",
        "Outstanding Amount", "Ledger / Expense Category", "Vendor Contact", "Remarks",
    ]
    payment_headers = ["Invoice Number", "Payment Date", "Amount"]
    instructions.append(["THE 45-DAY CLOCK - IMPORT TEMPLATE"])
    instructions.append(["Purpose", "Prepare purchase invoices and invoice-wise payments for one selected financial year."])
    instructions.append(["Dates", "Use DD-MM-YYYY, DD/MM/YYYY or DD-Mon-YYYY."])
    instructions.append(["Amounts", "Enter positive invoice/payment values. Indian comma formatting is accepted."])
    instructions.append(["Payments", "Use the same invoice number as the Purchase Invoices sheet."])
    instructions.append(["Privacy", "Keep this file in the client-controlled folder. Do not upload real data to an AI service."])
    instructions.append(["Control", "Do not add total rows; the application calculates and reconciles control totals."])
    instructions.append(["Important", "Vendor/MSME values imported from this workbook remain marked for human review until evidence is confirmed."])
    instructions.append(["Required purchase fields", "Invoice Number, Invoice Date, Vendor Name and Invoice Amount. PAN/GSTIN is strongly recommended but may be completed later."])
    instructions.append(["GRN and credit terms", "Optional, but these fields affect the acceptance-date and credit-period arithmetic when selected/available."])
    instructions.append(["Payment dates", "Use the Payments sheet for multiple invoice-wise allocations. If Actual Payment Date is supplied on a purchase row, the application derives the paid amount from Invoice Amount less Outstanding Amount; if Outstanding Amount is blank, full settlement is inferred and flagged for confirmation."])
    purchases.append(purchase_headers)
    payments.append(payment_headers)
    for sheet in (purchases, payments):
        for cell in sheet[1]:
            cell.fill = PatternFill("solid", fgColor="0F766E")
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(wrap_text=True, vertical="center")
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = f"A1:{sheet.cell(1, sheet.max_column).column_letter}1"
        sheet.row_dimensions[1].height = 34
        for column in range(1, sheet.max_column + 1):
            sheet.column_dimensions[sheet.cell(1, column).column_letter].width = 24
        sheet.sheet_view.showGridLines = False
        sheet.auto_filter.ref = f"A1:{sheet.cell(1, sheet.max_column).column_letter}5000"
    purchases.freeze_panes = "D2"
    purchases.column_dimensions["A"].width = 18
    purchases.column_dimensions["B"].width = 15
    purchases.column_dimensions["C"].width = 30
    for letter in ("D", "E", "G", "I", "J", "L", "M"):
        purchases.column_dimensions[letter].width = 23
    for letter in ("F", "R"):
        purchases.column_dimensions[letter].width = 18
    for letter in ("N", "P", "Q"):
        purchases.column_dimensions[letter].width = 20
    purchases.column_dimensions["S"].width = 25
    purchases.column_dimensions["T"].width = 22
    purchases.column_dimensions["U"].width = 38
    for column in ("B", "N", "P", "Q"):
        for row in range(2, 5001):
            purchases[f"{column}{row}"].number_format = "dd-mmm-yyyy"
    for column in ("F", "R"):
        for row in range(2, 5001):
            purchases[f"{column}{row}"].number_format = "#,##,##0.00"
    for column in ("B", "C", "F"):
        purchases[f"{column}1"].fill = PatternFill("solid", fgColor="12233A")
    class_validation = DataValidation(
        type="list", formula1='"Micro,Small,Medium,Unregistered,Not available"',
        allow_blank=True,
    )
    status_validation = DataValidation(
        type="list", formula1='"Registered,Unregistered,Certificate unavailable,Information incomplete,Verification pending"',
        allow_blank=True,
    )
    source_validation = DataValidation(
        type="list", formula1='"Imported ledger,Udyam certificate,Vendor declaration,Client ERP flag,Manual input"',
        allow_blank=True,
    )
    purchases.add_data_validation(class_validation)
    purchases.add_data_validation(status_validation)
    purchases.add_data_validation(source_validation)
    class_validation.add("H2:H5000")
    status_validation.add("L2:L5000")
    source_validation.add("M2:M5000")
    purchases.conditional_formatting.add(
        "A2:U5000",
        FormulaRule(
            formula=['AND(COUNTA($A2:$U2)>0,OR($A2="",$B2="",$C2="",$F2=""))'],
            fill=PatternFill("solid", fgColor="FEF3C7"),
        ),
    )
    instructions.column_dimensions["A"].width = 22
    instructions.column_dimensions["B"].width = 92
    instructions["A1"].fill = PatternFill("solid", fgColor="12233A")
    instructions["A1"].font = Font(color="FFFFFF", bold=True, size=16)
    instructions.merge_cells("A1:B1")
    for row in range(2, instructions.max_row + 1):
        instructions.cell(row, 1).font = Font(bold=True, color="12233A")
        instructions.cell(row, 2).alignment = Alignment(wrap_text=True, vertical="top")
    instructions.sheet_view.showGridLines = False
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _vendor_declaration_bytes(state: DesktopState, vendor_id: str) -> tuple[bytes, str]:
    """Create the actual declaration file used by both API and native Save As."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    row = next((item for item in state.vendor_rows() if item["vendor_id"] == vendor_id), None)
    if row is None:
        raise ValueError("Vendor not found")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("MSME / UDYAM STATUS DECLARATION")
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(18, 35, 58)
    document.add_paragraph(date.today().strftime("%d %B %Y"))
    document.add_paragraph(f"To\n{state.entity_name}")
    document.add_paragraph(f"Financial year under review: {state.fy}")
    document.add_paragraph(f"Vendor / supplier: {row['name']}")
    document.add_paragraph(
        "Please confirm the following information for the above financial year. "
        "This declaration will be retained with the buyer's tax-audit working papers."
    )
    fields = [
        "Udyam Registration Number",
        "Enterprise classification for the above year (Micro / Small / Medium)",
        "Classification year and classification date",
        "Major activity (Manufacturing / Services / Trading)",
        "Principal NIC code and description",
        "Udyam registration date",
        "PAN",
        "GSTIN",
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Particulars"
    table.rows[0].cells[1].text = "Vendor confirmation"
    for label in fields:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = ""
    document.add_paragraph(
        "We confirm that the information above is complete and correct for the stated "
        "period and enclose the relevant Udyam Registration Certificate, including its annexure."
    )
    document.add_paragraph("For: ______________________________________________")
    document.add_paragraph("Authorised signatory: ______________________________")
    document.add_paragraph("Name and designation: ______________________________")
    document.add_paragraph("Date and place: _____________________________________")
    note = document.add_paragraph(
        "Generated by The 45-Day Clock as a request for evidence. The engagement team must review "
        "the response before changing the vendor classification."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(8)
    output = io.BytesIO()
    document.save(output)
    safe_vendor = re.sub(r"[^A-Za-z0-9]+", "-", row["name"]).strip("-") or "Vendor"
    return output.getvalue(), f"{safe_vendor}-{state.fy}-MSME-Declaration.docx"


def _working_paper_pdf(state: DesktopState) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    output = io.BytesIO()
    pdf = canvas.Canvas(output, pagesize=A4)
    width, height = A4
    y = height - 54
    pdf.setTitle("The 45-Day Clock - Working Paper")
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(48, y, "THE 45-DAY CLOCK - WORKING PAPER")
    y -= 30
    pdf.setFont("Helvetica", 10)
    lines = [
        f"Entity: {state.entity_name}",
        f"Tax year: {state.fy}",
        f"Run hash: {state.run.run_hash()}",
        f"Disallowance: {money(state.run.disallowance_total)}",
        f"MSMED s.16 interest (not deductible under s.23): {money(state.run.interest_total)}",
        f"Correctly not disallowed: {money(state.run.excluded_total)}",
        "",
        "Acceptance-date policy:",
        ACC_POLICY_TEXT[state.run.acceptance_policy],
        "",
        "This is a computation aid. Vendor classifications require human confirmation.",
    ]
    for line in lines:
        words = line.split()
        current = ""
        wrapped = []
        for word in words:
            trial = (current + " " + word).strip()
            if pdf.stringWidth(trial, "Helvetica", 10) > width - 96:
                wrapped.append(current)
                current = word
            else:
                current = trial
        wrapped.append(current)
        for item in wrapped:
            pdf.drawString(48, y, item)
            y -= 15
    pdf.save()
    return output.getvalue()


def _finished_export_bytes(kind: str, state: DesktopState) -> tuple[bytes, str, str]:
    """Generate the same finished file used by Export All, never a debug substitute."""
    from clock45.export import (
        build_action_list_workbook,
        build_clause_22_workbook,
        build_exclusion_register_workbook,
        build_working_paper_pdf,
    )

    builders = {
        "clause22": (build_clause_22_workbook, "Clause-22-Workbook.xlsx",
                     "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        "action": (build_action_list_workbook, "31-March-Action-List.xlsx",
                   "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        "exclusions": (build_exclusion_register_workbook, "Exclusion-Register.xlsx",
                       "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"),
        "working-paper": (build_working_paper_pdf, "Working-Paper.pdf", "application/pdf"),
    }
    if kind not in builders:
        raise ValueError("Unknown export")
    builder, suffix, media = builders[kind]
    safe_entity = re.sub(r"[^A-Za-z0-9]+", "-", state.entity_name).strip("-") or "Client"
    filename = f"{safe_entity}-{state.fy}-{suffix}"
    with tempfile.TemporaryDirectory(prefix="clock45-export-") as temporary:
        path = Path(temporary) / filename
        builder(path, state.export_context())
        content = path.read_bytes()
    if not content:
        raise ValueError("The export was generated with no content")
    return content, filename, media


def create_application(state: Optional[DesktopState] = None, token: Optional[str] = None) -> FastAPI:
    state = state or DesktopState()
    token = token or secrets.token_urlsafe(32)
    application = FastAPI(title="The 45-Day Clock", docs_url=None, redoc_url=None)
    application.state.desktop = state
    application.state.session_token = token

    @application.middleware("http")
    async def secure_api(request: Request, call_next):
        if request.url.path.startswith("/api/"):
            supplied = request.headers.get("X-Session-Token") or request.query_params.get("token")
            if not supplied or not secrets.compare_digest(supplied, token):
                return Response("Session token required", status_code=403, media_type="text/plain")
        response = await call_next(request)
        response.headers["Cache-Control"] = "no-store"
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Content-Security-Policy"] = (
            "default-src 'self'; style-src 'self'; script-src 'self'; "
            "img-src 'self' data:; connect-src 'self'; frame-ancestors 'none'"
        )
        return response

    @application.get("/", response_class=HTMLResponse)
    def index():
        html = (WEB / "index.html").read_text(encoding="utf-8")
        return HTMLResponse(html.replace("__SESSION_TOKEN__", token), media_type="text/html; charset=utf-8")

    @application.get("/assets/app.css")
    def styles():
        return Response((WEB / "app.css").read_text(encoding="utf-8"), media_type="text/css; charset=utf-8")

    @application.get("/assets/app.js")
    def scripts():
        return Response((WEB / "app.js").read_text(encoding="utf-8"), media_type="application/javascript; charset=utf-8")

    @application.get("/assets/storage.css")
    def storage_styles():
        return Response(
            (WEB / "storage.css").read_text(encoding="utf-8"),
            media_type="text/css; charset=utf-8",
        )

    @application.get("/assets/welcome.css")
    def welcome_styles():
        return Response(
            (WEB / "welcome.css").read_text(encoding="utf-8"),
            media_type="text/css; charset=utf-8",
        )

    @application.get("/assets/capstone.css")
    def capstone_styles():
        return Response(
            (WEB / "capstone.css").read_text(encoding="utf-8"),
            media_type="text/css; charset=utf-8",
        )

    @application.get("/api/home")
    def home():
        return _jsonable({**state.home_data(), "readiness": state.project_readiness()})

    @application.post("/api/analysis/start")
    def start_analysis(request: AnalysisStartRequest):
        try:
            state.start_analysis(request.entity_name, request.entity_pan, request.fy)
        except (ValueError, StoreError) as exc:
            raise HTTPException(422, str(exc)) from exc
        return _jsonable({"entity": state.entity_name, "pan": state.entity_pan, "fy": state.fy})

    @application.get("/api/firm-profile")
    def firm_profile():
        return state.firm_profile()

    @application.post("/api/firm-profile")
    def save_firm_profile(request: FirmProfileRequest):
        try:
            return state.save_firm_profile(request.model_dump())
        except (ValueError, StoreError) as exc:
            raise HTTPException(422, str(exc)) from exc

    @application.get("/api/import/template")
    def import_template():
        headers = {
            "Content-Disposition": 'attachment; filename="45-Day-Clock-Import-Template.xlsx"'
        }
        return Response(
            _import_template_bytes(),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers=headers,
        )

    @application.get("/api/storage/status")
    def storage_status():
        return state.storage_status()

    @application.get("/api/startup")
    def startup():
        return state.startup_status()

    @application.post("/api/startup/complete")
    def startup_complete():
        try:
            state.complete_welcome()
        except LicenceError as exc:
            raise HTTPException(409, str(exc)) from exc
        return {"ok": True, "licence": state.licence_status()}

    @application.post("/api/analysis/reset")
    def reset():
        state.reset_analysis()
        return {"ok": True}

    @application.post("/api/demo/load")
    def load_demo():
        totals = state.load_demo()
        return _jsonable({"entity": state.entity_name, "fy": state.fy, "control_totals": totals})

    @application.post("/api/import/file")
    async def import_file(record_type: str = "purchase", file: UploadFile = File(...)):
        data = await file.read()
        filename = file.filename or "ledger"
        try:
            if filename.lower().endswith(".xml"):
                result = import_tally_xml(data)
                state.add_import(result, replace=True, source_hash=hashlib.sha256(data).hexdigest())
                return _jsonable({"complete": True, "control_totals": state.control_totals,
                                  "confirmations": state.confirmations})
            table = _read_uploaded_table(data, filename)
            state.pending_rows = [row.to_dict() for _, row in table.iterrows()]
            state.pending_columns = [str(column) for column in table.columns]
            state.pending_fingerprint = "|".join(column.strip().casefold() for column in state.pending_columns)
            state.pending_source = filename
            state.pending_source_hash = hashlib.sha256(data).hexdigest()
            if not state.entity_name:
                state.entity_name = "Imported client"
            state._ensure_analysis()
            remembered = None
            if state.store and state.client_id:
                remembered = state.store.load_column_mapping(
                    state.client_id, record_type, state.pending_fingerprint
                )
            return {
                "complete": False,
                "columns": state.pending_columns,
                "suggested_mapping": remembered or suggest_mapping_details(
                    state.pending_columns, record_type
                )[0],
                "mapping_confidence": (
                    {field: "REMEMBERED" for field in remembered}
                    if remembered else suggest_mapping_details(state.pending_columns, record_type)[1]
                ),
                "mapping_remembered": remembered is not None,
                "record_type": record_type,
            }
        except LicenceError as exc:
            raise HTTPException(409, str(exc)) from exc
        except IngestError as exc:
            raise HTTPException(422, _error_payload(exc)) from exc
        except Exception as exc:
            raise HTTPException(422, f"Could not read {filename}: {exc}") from exc

    @application.post("/api/import/map")
    def map_columns(request: MappingRequest):
        if not state.pending_rows:
            raise HTTPException(409, "No file is waiting for column mapping")
        try:
            result = _parse_rows(
                ((number, row) for number, row in enumerate(state.pending_rows, 2)),
                request.mapping,
                record_type=request.record_type,
                source=state.pending_source,
                source_columns=state.pending_columns,
            )
            state.add_import(
                result,
                replace=request.record_type == "purchase",
                source_hash=state.pending_source_hash,
            )
            if state.store and state.client_id:
                state.store.save_column_mapping(
                    state.client_id, request.record_type, state.pending_fingerprint, request.mapping
                )
            return _jsonable({
                "control_totals": state.control_totals,
                "unmapped_columns": list(result.unmapped_columns),
                "confirmations": state.confirmations,
            })
        except LicenceError as exc:
            raise HTTPException(409, str(exc)) from exc
        except IngestError as exc:
            raise HTTPException(422, _error_payload(exc)) from exc

    @application.post("/api/import/manual")
    def manual(request: ManualRequest):
        try:
            grid = ManualEntryGrid(request.record_type)
            mapping = grid.paste_from_excel(request.text)
            result = grid.import_rows(mapping, expected_total=request.expected_total)
            state.add_import(
                result,
                replace=request.record_type == "purchase",
                source_hash=hashlib.sha256(request.text.encode("utf-8")).hexdigest(),
            )
            return _jsonable({
                "control_totals": state.control_totals,
                "confirmations": state.confirmations,
            })
        except LicenceError as exc:
            raise HTTPException(409, str(exc)) from exc
        except IngestError as exc:
            raise HTTPException(422, _error_payload(exc)) from exc

    @application.get("/api/vendors")
    def vendors(filter: str = "all"):
        all_rows = state.vendor_rows()
        assumed_count = sum(row["is_assumed"] for row in all_rows)
        rows = all_rows
        if filter == "unconfirmed":
            rows = [row for row in rows if row["unconfirmed"]]
        elif filter == "traders":
            rows = [row for row in rows if row["is_trader"]]
        elif filter == "medium":
            rows = [row for row in rows if row["is_medium"]]
        elif filter == "covered":
            rows = [row for row in rows if row["coverage"] == "Covered"]
        return {"vendors": rows, "assumed_count": assumed_count, "total_count": len(all_rows)}

    @application.get("/api/readiness")
    def readiness():
        return _jsonable(state.project_readiness())

    @application.get("/api/vendors/{vendor_id}")
    def vendor(vendor_id: str):
        row = next((item for item in state.vendor_rows() if item["vendor_id"] == vendor_id), None)
        if row is None:
            raise HTTPException(404, "Vendor not found")
        return _jsonable(row)

    @application.put("/api/vendors/{vendor_id}")
    def update_vendor(vendor_id: str, request: VendorUpdateRequest):
        try:
            return _jsonable(state.update_vendor(vendor_id, request.model_dump()))
        except (ValueError, StoreError) as exc:
            raise HTTPException(422, str(exc)) from exc

    @application.post("/api/vendors/bulk")
    def bulk(request: BulkRequest):
        if request.source not in EVIDENCE_RANK:
            raise HTTPException(422, "Unknown evidence source")
        for vendor_id in request.vendor_ids:
            record = state.udyam.get(vendor_id)
            if record:
                record.source = request.source
                state.persist_vendor(vendor_id, "Desktop user - bulk evidence change")
        return {"updated": len(request.vendor_ids)}

    @application.post("/api/vendors/{vendor_id}/certificate")
    async def certificate(vendor_id: str, file: UploadFile = File(...)):
        record = state.udyam.get(vendor_id)
        if record is None:
            raise HTTPException(404, "Vendor not found")
        content = await file.read()
        if not content:
            raise HTTPException(422, "The selected certificate is empty")
        if not state.store or not state.client_id:
            raise HTTPException(
                409, "Choose the local data folder before uploading evidence so the original can be preserved."
            )
        try:
            evidence_hash = state.store.add_vendor_evidence(
                state.client_id, vendor_id, filename=file.filename or "Udyam certificate",
                media_type=file.content_type or "application/octet-stream", content=content,
                added_by="Desktop user",
            )
        except StoreError as exc:
            raise HTTPException(409, str(exc)) from exc
        evidence = state.store.find_vendor_evidence(state.client_id, vendor_id, evidence_hash)
        if not evidence:
            raise HTTPException(500, "The certificate was stored but could not be reopened")
        try:
            parsed = parse_udyam_certificate(
                content,
                filename=file.filename or "certificate.pdf",
                media_type=file.content_type or "application/octet-stream",
            )
            extracted = parsed.as_dict()
            relevant = parsed.class_for_year(state.fy)
            extracted["project_year_class"] = (
                relevant.enterprise_class if relevant else parsed.enterprise_class
            )
            extracted["project_year_match"] = bool(relevant)
            extracted["nic_code"] = (
                parsed.nic_activities[0].nic_code if parsed.nic_activities else ""
            )
            parse_status = "EXTRACTED_REVIEW_REQUIRED"
            parse_message = "Values were extracted locally. Review every field before saving."
        except UdyamParseError as exc:
            if exc.code not in {"OCR_REQUIRED"}:
                raise HTTPException(422, {
                    "message": str(exc), "code": exc.code,
                    "evidence_retained": True, "evidence_id": evidence["evidence_id"],
                }) from exc
            extracted = {
                "classification_history": [], "nic_code": "", "warnings": [str(exc)],
                "field_status": {}, "project_year_class": "", "project_year_match": False,
            }
            parse_status = exc.code
            parse_message = str(exc)
        current_row = next(item for item in state.vendor_rows() if item["vendor_id"] == vendor_id)
        comparisons = {
            "udyam_no": (current_row.get("udyam_no_raw", ""), extracted.get("udyam_no", "")),
            "enterprise_class": (
                current_row.get("enterprise_class_raw", ""), extracted.get("project_year_class", "")
            ),
            "nic_code": (current_row.get("nic_code_raw", ""), extracted.get("nic_code", "")),
            "major_activity": (current_row.get("activity_raw", ""), extracted.get("major_activity", "")),
            "pan": (current_row.get("pan", ""), extracted.get("pan", "")),
            "enterprise_name": (current_row.get("name", ""), extracted.get("enterprise_name", "")),
        }
        conflicts = {
            key: {"current": old, "extracted": new}
            for key, (old, new) in comparisons.items()
            if old and new and str(old).strip().casefold() != str(new).strip().casefold()
        }
        state.pending_certificate_reviews[evidence["evidence_id"]] = {
            "vendor_id": vendor_id, "parsed": extracted, "conflicts": conflicts,
            "hash": evidence_hash,
        }
        return _jsonable({
            "updated": False,
            "evidence_id": evidence["evidence_id"],
            "hash": evidence_hash,
            "parse_status": parse_status,
            "message": parse_message,
            "extracted": extracted,
            "conflicts": conflicts,
            "requires_confirmation": True,
        })

    @application.post("/api/vendors/{vendor_id}/certificate/confirm")
    def confirm_certificate(vendor_id: str, request: CertificateReviewRequest):
        if not state.store or not state.client_id:
            raise HTTPException(409, "Choose the local data folder first")
        pending = state.pending_certificate_reviews.get(request.evidence_id)
        if not pending or pending["vendor_id"] != vendor_id:
            raise HTTPException(409, "Upload the certificate again before confirming it")
        evidence = state.store.get_vendor_evidence(request.evidence_id)
        if evidence["client_id"] != state.client_id or evidence["vendor_id"] != vendor_id:
            raise HTTPException(403, "Evidence does not belong to this vendor")
        history = []
        for item in request.classification_history:
            year = str(item.get("classification_year", "")).strip()
            enterprise_class = str(item.get("enterprise_class", "")).strip().upper()
            classification_date = str(item.get("classification_date", "")).strip()
            if not re.fullmatch(r"20\d{2}-\d{2}", year):
                raise HTTPException(422, f"Classification year {year!r} must look like 2025-26")
            if enterprise_class not in {"MICRO", "SMALL", "MEDIUM"}:
                raise HTTPException(422, f"Invalid enterprise class for {year}")
            if classification_date:
                try:
                    date.fromisoformat(classification_date)
                except ValueError as exc:
                    raise HTTPException(422, f"Invalid classification date for {year}") from exc
            history.append({
                "classification_year": year,
                "enterprise_class": enterprise_class,
                "classification_date": classification_date,
            })
        relevant = next((item for item in history if item["classification_year"] == state.fy), None)
        if history and not relevant:
            raise HTTPException(
                422, f"The certificate history has no classification for project FY {state.fy}. "
                "Add or confirm the applicable year before saving."
            )
        record = state.udyam.get(vendor_id)
        previous_hash = record.evidence_file_hash
        record.evidence_file_hash = evidence["sha256"]
        try:
            state.update_vendor(vendor_id, {
                "udyam_no": request.udyam_no,
                "enterprise_class": relevant["enterprise_class"] if relevant else request.enterprise_class,
                "nic_code": request.nic_code,
                "activity_label": request.major_activity,
                "registration_date": request.registration_date,
                "source": SRC_CERTIFICATE,
                "confirmed_by": request.confirmed_by,
            })
        except (ValueError, StoreError) as exc:
            record.evidence_file_hash = previous_hash
            raise HTTPException(422, str(exc)) from exc
        confirmed = request.model_dump()
        state.store.record_evidence_review(
            evidence_id=request.evidence_id,
            client_id=state.client_id,
            vendor_id=vendor_id,
            parsed=pending["parsed"],
            confirmed=confirmed,
            conflicts=pending["conflicts"],
            confirmed_by=request.confirmed_by,
            classification_history=history,
        )
        state.store.upsert_vendor_metadata(
            state.client_id, vendor_id,
            pan=request.pan, organisation_type=request.organisation_type,
            incorporation_date=request.incorporation_date,
            commencement_date=request.commencement_date,
            registered_address=request.registered_address,
            registration_status="REGISTERED", verification_source="UDYAM_CERTIFICATE",
        )
        state.pending_certificate_reviews.pop(request.evidence_id, None)
        return _jsonable({
            "saved": True,
            "vendor": next(item for item in state.vendor_rows() if item["vendor_id"] == vendor_id),
        })

    @application.get("/api/vendors/{vendor_id}/evidence")
    def vendor_evidence(vendor_id: str):
        if not state.store or not state.client_id:
            return {"evidence": []}
        return {"evidence": state.store.list_vendor_evidence(state.client_id, vendor_id)}

    @application.get("/api/vendors/{vendor_id}/declaration")
    def declaration(vendor_id: str):
        try:
            content, filename = _vendor_declaration_bytes(state, vendor_id)
        except ValueError as exc:
            raise HTTPException(404, str(exc)) from exc
        headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
        return Response(
            content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers=headers,
        )

    @application.get("/api/assumptions")
    def assumptions():
        return {"policies": [
            {"id": ACC_INVOICE_DATE, "label": "Invoice date", "wording": ACC_POLICY_TEXT[ACC_INVOICE_DATE]},
            {"id": ACC_GRN_DATE, "label": "GRN date where available", "wording": ACC_POLICY_TEXT[ACC_GRN_DATE]},
            {"id": ACC_INVOICE_PLUS, "label": "Invoice date plus N days", "wording": ACC_POLICY_TEXT[ACC_INVOICE_PLUS]},
        ]}

    @application.post("/api/assumptions/confirm")
    def confirm_assumptions(request: AssumptionRequest):
        try:
            if state.store and state.analysis_id:
                state.store.update_analysis_stage(
                    state.analysis_id, "assumptions", acceptance_policy=request.policy,
                    acceptance_plus_days=request.plus_days,
                )
            state.execute(request.policy, request.plus_days, request.payments_confirmed)
        except ValueError as exc:
            raise HTTPException(409, str(exc)) from exc
        return {"ok": True, "run_hash": state.run.run_hash()}

    def require_run():
        if state.run is None or not state.assumptions_confirmed:
            raise HTTPException(409, "Confirm the assumptions before viewing results")
        return state.run

    @application.get("/api/results")
    def results():
        run = require_run()
        _, year_end = fy_bounds(state.fy)
        purchases = {line.invoice_id: line for line in state.purchases}
        payments_by_invoice: dict[str, list[PaymentLine]] = {}
        for payment in state.payments:
            payments_by_invoice.setdefault(payment.invoice_id, []).append(payment)
        paid_within_year = Decimal("0.00")
        outstanding_at_year_end = Decimal("0.00")
        for finding in run.findings:
            purchase = purchases.get(finding.invoice_id)
            if not purchase:
                continue
            paid = min(
                purchase.amount,
                sum(
                    (payment.amount for payment in payments_by_invoice.get(finding.invoice_id, [])
                     if payment.payment_date <= year_end),
                    Decimal("0.00"),
                ),
            )
            paid_within_year += paid
            outstanding_at_year_end += max(Decimal("0.00"), purchase.amount - paid)
        original_population = sum((finding.amount for finding in run.findings), Decimal("0.00"))
        excluded_population = sum(
            (finding.amount for finding in run.findings if finding.gate_failed), Decimal("0.00")
        )
        return _jsonable({
            "section": run.statute["section"],
            "act": run.statute["act"],
            "disallowance": run.disallowance_total,
            "disallowance_display": money(run.disallowance_total),
            "interest": run.interest_total,
            "interest_display": money(run.interest_total),
            "excluded": run.excluded_total,
            "excluded_display": money(run.excluded_total),
            "action_list": action_list(run, top_n=10),
            "interest_only": interest_only_register(run),
            "reconciliation": {
                "original_population": original_population,
                "original_population_display": money(original_population),
                "paid_within_year": paid_within_year,
                "paid_within_year_display": money(paid_within_year),
                "outstanding_at_year_end": outstanding_at_year_end,
                "outstanding_at_year_end_display": money(outstanding_at_year_end),
                "excluded_population": excluded_population,
                "excluded_population_display": money(excluded_population),
                "population_considered": original_population - excluded_population,
                "population_considered_display": money(original_population - excluded_population),
            },
            "formula_notes": {
                "disallowance": "Sum of unpaid balance at 31 March for covered invoices whose statutory due date has expired.",
                "interest": "Invoice-wise MSMED s.16 compound interest with monthly rests and dated RBI Bank Rate segments.",
                "excluded": "Counterfactual disallowance prevented by the first failed vendor coverage gate.",
            },
        })

    @application.get("/api/findings")
    def findings(kind: str = "disallowance"):
        run = require_run()
        if kind == "interest":
            selected = [finding for finding in run.findings if finding.interest > 0]
        elif kind == "excluded":
            selected = [finding for finding in run.findings if finding.gate_failed]
        else:
            selected = [finding for finding in run.findings if finding.disallowance > 0]
        return {"findings": [_jsonable({
            "invoice_id": f.invoice_id,
            "vendor": f.vendor_name,
            "invoice_date": f.invoice_date,
            "amount": f.amount,
            "amount_display": money(f.amount),
            "disallowance_display": money(f.disallowance),
            "interest_display": money(f.interest),
            "status": f.status,
            "due_date": f.due_date,
            "reason": f.reason,
            "treatment": f"Excluded - {f.gate_failed}" if f.gate_failed else f.status,
        }) for f in selected]}

    @application.get("/api/findings/{invoice_id:path}")
    def finding_detail(invoice_id: str):
        run = require_run()
        finding = next((item for item in run.findings if item.invoice_id == invoice_id), None)
        purchase = next((item for item in state.purchases if item.invoice_id == invoice_id), None)
        if finding is None or purchase is None:
            raise HTTPException(404, "Invoice not found")
        payments = [item for item in state.payments if item.invoice_id == invoice_id]
        credit = resolve_credit_period(purchase.agreement_days)
        verdict = assess_invoice(
            amount=purchase.amount,
            acceptance_date=finding.acceptance_date,
            agreement_days=purchase.agreement_days,
            payments=[(payment.payment_date, payment.amount) for payment in payments],
            fy=state.fy,
        )
        return _jsonable({
            "invoice_id": invoice_id,
            "vendor": finding.vendor_name,
            "invoice_date": finding.invoice_date,
            "acceptance_date": finding.acceptance_date,
            "amount": finding.amount,
            "amount_display": money(finding.amount),
            "credit_days": credit.days,
            "credit_basis": credit.basis,
            "ceiling_applied": credit.ceiling_applied,
            "due_date": verdict.due_date,
            "appointed_day": verdict.appointed_day,
            "payments": [{"date": p.payment_date, "amount": p.amount,
                          "amount_display": money(p.amount)} for p in payments],
            "interest_segments": verdict.interest_segments,
            "reason": finding.reason,
            "status": finding.status,
            "disallowance_display": money(finding.disallowance),
            "interest_display": money(finding.interest),
        })

    @application.get("/api/exclusions")
    def exclusions():
        run = require_run()
        rows = exclusion_register(run)
        groups: dict[str, dict[str, Any]] = {}
        for row in rows:
            gate = row["gate"]
            group = groups.setdefault(gate, {
                "gate": gate,
                "label": GATE_LABELS.get(gate, gate),
                "reason": row["reason"],
                "total": Decimal("0.00"),
                "rows": [],
            })
            group["total"] += row["wrongly_disallowable"]
            group["rows"].append({**row, "amount_display": money(row["wrongly_disallowable"])})
        for group in groups.values():
            group["total_display"] = money(group["total"])
        gross_original = sum((finding.amount for finding in run.findings), Decimal("0.00"))
        gross_excluded = sum(
            (finding.amount for finding in run.findings if finding.gate_failed), Decimal("0.00")
        )
        return _jsonable({
            "groups": list(groups.values()),
            "grand_total_display": money(run.excluded_total),
            "reconciliation": {
                "original_population": gross_original,
                "original_population_display": money(gross_original),
                "excluded_population": gross_excluded,
                "excluded_population_display": money(gross_excluded),
                "considered_population": gross_original - gross_excluded,
                "considered_population_display": money(gross_original - gross_excluded),
                "invoice_count": len(run.findings),
                "excluded_invoice_count": sum(bool(f.gate_failed) for f in run.findings),
            },
        })

    @application.get("/api/exports/preview")
    def export_preview():
        run = require_run()
        return _jsonable({
            "run_hash": run.run_hash(),
            "files": [
                {"id": "clause22", "name": "Clause 22 Workbook.xlsx", "description": f"Invoice-level disallowance under section {run.statute['section']}", "rows": sum(f.disallowance > 0 for f in run.findings)},
                {"id": "action", "name": "31 March Action List.xlsx", "description": "Vendors ranked by payment required before year end", "rows": len(action_list(run))},
                {"id": "exclusions", "name": "Exclusion Register.xlsx", "description": "Defensible non-disallowance by failed coverage gate", "rows": len(exclusion_register(run))},
                {"id": "working-paper", "name": "Working Paper.pdf", "description": "Run controls, assumptions and headline conclusions", "rows": len(run.findings)},
            ],
        })

    @application.get("/api/exports/{kind}")
    def export(kind: str):
        require_run()
        try:
            content, filename, media = _finished_export_bytes(kind, state)
        except ValueError as exc:
            raise HTTPException(404, str(exc)) from exc
        return Response(content, media_type=media, headers={"Content-Disposition": f'attachment; filename="{filename}"'})

    return application


def _free_port() -> int:
    with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def _wait_until_ready(url: str, timeout: float = 10.0) -> None:
    deadline = time.monotonic() + timeout
    while time.monotonic() < deadline:
        try:
            with urlopen(url, timeout=0.5) as response:
                if response.status == 200:
                    return
        except Exception:
            time.sleep(0.05)
    raise RuntimeError("The local interface did not start")


def main() -> None:
    _startup_marker("entered main")
    parser = argparse.ArgumentParser(description="The 45-Day Clock desktop application")
    parser.add_argument("--no-window", action="store_true", help="Run the local server without pywebview")
    parser.add_argument("--port", type=int, default=0, help="Testing only; default is a random free port")
    parser.add_argument(
        "--data-folder", type=str, default="",
        help="Use an explicitly chosen local folder for the SQLite database",
    )
    args = parser.parse_args()
    _startup_marker("arguments parsed")

    import uvicorn
    _startup_marker("uvicorn imported")

    token = secrets.token_urlsafe(32)
    port = args.port or _free_port()
    store = Store(args.data_folder) if args.data_folder else None
    _startup_marker("store opened")
    application = create_application(DesktopState(store), token=token)
    _startup_marker("FastAPI application created")
    _startup_marker("creating uvicorn configuration")
    config = uvicorn.Config(
        application, host="127.0.0.1", port=port, log_level="warning",
        access_log=False, log_config=None,
    )
    _startup_marker("uvicorn configuration created")
    server = uvicorn.Server(config)
    _startup_marker("uvicorn server created")
    thread = threading.Thread(target=server.run, name="clock45-local-server", daemon=True)
    _startup_marker("server thread object created")
    thread.start()
    _startup_marker("server thread started")
    local_url = f"http://127.0.0.1:{port}/"
    _wait_until_ready(local_url)
    _startup_marker("local server ready")
    if sys.stdout:
        print(f"The 45-Day Clock is running locally at {local_url}", flush=True)

    if args.no_window:
        try:
            while thread.is_alive():
                time.sleep(0.25)
        except KeyboardInterrupt:
            server.should_exit = True
        finally:
            if application.state.desktop.store:
                application.state.desktop.store.close()
        return

    try:
        import webview
    except ImportError as exc:
        server.should_exit = True
        raise RuntimeError("pywebview is not installed. Install requirements.txt first.") from exc
    webview.create_window(
        "The 45-Day Clock",
        local_url,
        js_api=NativeBridge(application.state.desktop),
        width=1366,
        height=768,
        min_size=(1100, 680),
        background_color="#F7F8FA",
        text_select=True,
    )
    try:
        webview.start(debug=False)
    except BaseException:
        _write_local_crash_log(application.state.desktop.store, "Desktop window stopped unexpectedly")
        raise
    finally:
        server.should_exit = True
        thread.join(timeout=3)
        if application.state.desktop.store:
            application.state.desktop.store.close()


if __name__ == "__main__":
    main()
