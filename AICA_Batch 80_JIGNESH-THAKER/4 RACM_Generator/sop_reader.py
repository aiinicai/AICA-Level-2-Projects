"""
SOP Reader — extracts plain text from a client's SOP file,
whether it's a Word document (.docx) or a PDF (.pdf).
"""

import os
from docx import Document
from pypdf import PdfReader


def read_sop(file_path):
    """
    Takes a file path, detects whether it's .docx or .pdf,
    and returns the full text content as a single string.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Could not find file: {file_path}")

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".docx":
        return _read_docx(file_path)
    elif extension == ".pdf":
        return _read_pdf(file_path)
    else:
        raise ValueError(
            f"Unsupported file type '{extension}'. Please provide a .docx or .pdf file."
        )


def _read_docx(file_path):
    doc = Document(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also pull text out of any tables in the SOP (many SOPs use tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())

    return "\n".join(paragraphs)


def _read_pdf(file_path):
    reader = PdfReader(file_path)
    text_chunks = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


# Quick self-test: only runs if you execute THIS file directly
if __name__ == "__main__":
    raw_input_path = input("Paste the full path to a test SOP file (.docx or .pdf): ")
    # Remove surrounding quotes AND any accidental leading/trailing spaces
    test_path = raw_input_path.strip().strip('"').strip()

    print(f"\nChecking this exact path:\n{test_path}\n")

    if not os.path.exists(test_path):
        print("File not found. Let's check the folder contents instead...")
        folder = os.path.dirname(test_path)
        if os.path.exists(folder):
            print(f"\nFiles actually found in '{folder}':")
            for f in os.listdir(folder):
                print(f"  - {f}")
        else:
            print(f"Even the folder doesn't exist: {folder}")
    else:
        content = read_sop(test_path)
        print(f"\nExtracted {len(content)} characters.\n")
        print("First 500 characters preview:\n")
        print(content[:500])